from __future__ import annotations

from dataclasses import replace
from datetime import timedelta
from concurrent.futures import ThreadPoolExecutor
from io import BytesIO
import json
from pathlib import Path
import secrets
import subprocess
import threading
from tempfile import TemporaryDirectory
import time
from unittest import mock

from django.contrib.auth import get_user_model
from django.core import mail
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TransactionTestCase, override_settings
from django.urls import reverse
from django.utils import timezone
from PIL import Image

from src import openai_ai
from src import web_actions as actions
from src.youtube_tools import SubtitleCue
from .legal_documents import legal_document_content
from . import views
from .localization import app_messages, language_options, translate
from .models import AccountProfile, CommunityPurchase, CommunityWork, DesignerAsset, DesignerProject, JobEventRecord, JobOutputRecord, JobRecord, LearningArticle, MagicLoginToken, MusicEditorAsset, MusicEditorProject, VideoEditorAsset, VideoEditorProject, WorkspaceShare
from .views import _job_record_params


class SiteSmokeTests(TransactionTestCase):
    def test_core_site_pages_render(self) -> None:
        routes = [
            reverse("studio:landing"),
            reverse("studio:index"),
            reverse("studio:dashboard_detail", args=["files"]),
            reverse("studio:design_project_list"),
            reverse("studio:cherryx_pay"),
            reverse("studio:video_project_list"),
            reverse("studio:designer"),
            reverse("studio:video_editor"),
        ]

        for route in routes:
            with self.subTest(route=route):
                response = self.client.get(route)
                self.assertEqual(response.status_code, 200)

    def test_workspace_design_switch_opens_project_list(self) -> None:
        response = self.client.get(reverse("studio:index"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'href="/app/design-projects/"')
        self.assertContains(response, 'href="/app/cherryx-pay/"')
        self.assertContains(response, "data-designer-launch")

    def test_magic_login_token_logs_user_in_once(self) -> None:
        user = get_user_model().objects.create_user(username="magic@example.com", email="magic@example.com", password="Strong123")
        token = MagicLoginToken.create_for_user(user, timezone.now() + timedelta(minutes=30))

        response = self.client.get(reverse("studio:magic_login", args=[token.token]))

        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.url, reverse("studio:index"))
        token.refresh_from_db()
        self.assertIsNotNone(token.used_at)
        self.assertIn("_auth_user_id", self.client.session)

        self.client.logout()
        response = self.client.get(reverse("studio:magic_login", args=[token.token]))

        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.url, reverse("studio:login"))
        self.assertNotIn("_auth_user_id", self.client.session)

    def test_legal_pages_use_telegram_stars_without_old_requisites(self) -> None:
        forbidden = (
            "WayForPay",
            "3795908055",
            "Tax ID",
            "ФОП",
            "ИНН",
            "ІПН",
            "РНОКПП",
            "Individual entrepreneur",
            "Entrepreneur individuel",
            "Einzelunternehmer",
        )
        routes = [
            reverse("studio:legal_info"),
            reverse("studio:legal_terms"),
            reverse("studio:legal_refund"),
            reverse("studio:legal_contacts"),
        ]

        for route in routes:
            with self.subTest(route=route):
                response = self.client.get(route)
                self.assertEqual(response.status_code, 200)
                content = response.content.decode("utf-8")
                self.assertIn("Telegram Stars", content)
                self.assertIn("cherryxdigital@gmail.com", content)
                for marker in forbidden:
                    self.assertNotIn(marker, content)

    def test_active_legal_documents_exist_for_all_site_languages(self) -> None:
        forbidden = ("WayForPay", "3795908055", "ФОП", "ИНН", "ІПН", "РНОКПП", "Tax ID", "Individual entrepreneur")
        languages = ("ru", "uk", "en", "fr", "de", "es", "it", "ka", "hy")

        for language in languages:
            for document_type in ("terms", "refund", "contacts"):
                with self.subTest(language=language, document_type=document_type):
                    document = legal_document_content(document_type, language)
                    serialized = json.dumps(document, ensure_ascii=False)
                    self.assertTrue(document["sections"])
                    self.assertIn("Telegram Stars", serialized)
                    for marker in forbidden:
                        self.assertNotIn(marker, serialized)

    def test_learning_and_community_pages_render_with_cherryx_purchase(self) -> None:
        seller = get_user_model().objects.create_user(username="seller@example.com", email="seller@example.com", password="pass12345")
        buyer = get_user_model().objects.create_user(username="buyer@example.com", email="buyer@example.com", password="pass12345")
        AccountProfile.objects.create(user=seller, cherryx_balance=0)
        AccountProfile.objects.create(user=buyer, cherryx_balance=200)
        article = LearningArticle.objects.create(
            title="First CherryX lesson",
            excerpt="A practical guide for creators.",
            body="Open workspace.\nPublish better.",
            status=LearningArticle.STATUS_PUBLISHED,
        )
        video = CommunityWork.objects.create(
            owner=seller,
            title="Paid video template",
            kind=CommunityWork.KIND_VIDEO,
            excerpt="A paid public work.",
            access=CommunityWork.ACCESS_PAID,
            price_cherryx=120,
            status=CommunityWork.STATUS_PUBLISHED,
        )
        image = CommunityWork.objects.create(
            owner=seller,
            title="Free cover idea",
            kind=CommunityWork.KIND_IMAGE,
            access=CommunityWork.ACCESS_FREE,
            status=CommunityWork.STATUS_PUBLISHED,
        )
        text = CommunityWork.objects.create(
            owner=seller,
            title="Caption pack",
            kind=CommunityWork.KIND_TEXT,
            access=CommunityWork.ACCESS_FREE,
            status=CommunityWork.STATUS_PUBLISHED,
        )
        paid_tmp = TemporaryDirectory()
        paid_media_root = Path(paid_tmp.name)
        paid_image_bytes = BytesIO()
        Image.new("RGB", (80, 80), "#e11d48").save(paid_image_bytes, "PNG")
        with override_settings(MEDIA_ROOT=paid_media_root):
            paid_image = CommunityWork.objects.create(
                owner=seller,
                title="Paid source image",
                kind=CommunityWork.KIND_IMAGE,
                access=CommunityWork.ACCESS_PAID,
                price_cherryx=30,
                status=CommunityWork.STATUS_PUBLISHED,
                media_file=SimpleUploadedFile("paid.png", paid_image_bytes.getvalue(), content_type="image/png"),
            )

            feed = self.client.get(reverse("studio:community_images"))
            self.assertContains(feed, "Paid source image")
            self.assertNotContains(feed, paid_image.media_file.url)
            self.assertContains(feed, reverse("studio:community_work_preview", args=[paid_image.slug]))

            detail = self.client.get(reverse("studio:community_work_detail", args=[paid_image.slug]))
            self.assertContains(detail, reverse("studio:community_work_preview", args=[paid_image.slug]))
            self.assertNotContains(detail, paid_image.media_file.url)

            preview = self.client.get(reverse("studio:community_work_preview", args=[paid_image.slug]))
            self.assertEqual(preview.status_code, 200)
            self.assertEqual(preview["Content-Type"], "image/webp")
            self.assertEqual(preview["Cache-Control"], "no-store, private")

            locked = self.client.get(reverse("studio:community_work_download", args=[paid_image.slug]))
            self.assertEqual(locked.status_code, 302)
            self.assertIn(reverse("studio:login"), locked.url)

            direct_private = self.client.get(paid_image.media_file.url)
            self.assertEqual(direct_private.status_code, 404)

        for route, expected in (
            (reverse("studio:learn_index"), article.title),
            (reverse("studio:learn_article", args=[article.slug]), "Open workspace."),
            (reverse("studio:community_videos"), video.title),
            (reverse("studio:community_images"), image.title),
            (reverse("studio:community_texts"), text.title),
            (reverse("studio:community_work_detail", args=[video.slug]), "120 CherryX"),
        ):
            with self.subTest(route=route):
                response = self.client.get(route)
                self.assertEqual(response.status_code, 200)
                self.assertContains(response, expected)

        self.client.force_login(buyer)
        response = self.client.post(reverse("studio:community_purchase", args=[video.slug]))

        self.assertEqual(response.status_code, 302)
        self.assertTrue(CommunityPurchase.objects.filter(work=video, buyer=buyer).exists())
        buyer.studio_profile.refresh_from_db()
        seller.studio_profile.refresh_from_db()
        self.assertEqual(buyer.studio_profile.cherryx_balance, 80)
        self.assertEqual(seller.studio_profile.cherryx_balance, 120)

        with override_settings(MEDIA_ROOT=paid_media_root):
            response = self.client.post(reverse("studio:community_purchase", args=[paid_image.slug]))
            self.assertEqual(response.status_code, 302)
            download = self.client.get(reverse("studio:community_work_download", args=[paid_image.slug]))
            self.assertEqual(download.status_code, 200)
            download.close()
            paid_image.refresh_from_db()
            self.assertEqual(paid_image.download_count, 1)
            buyer.studio_profile.refresh_from_db()
            seller.studio_profile.refresh_from_db()
            self.assertEqual(buyer.studio_profile.cherryx_balance, 50)
            self.assertEqual(seller.studio_profile.cherryx_balance, 150)
        paid_tmp.cleanup()

        publish = self.client.post(
            reverse("studio:community_publish"),
            {
                "kind": CommunityWork.KIND_TEXT,
                "title": "Buyer public text",
                "excerpt": "Shared from the account.",
                "body": "Ready text pack.",
                "access": CommunityWork.ACCESS_FREE,
                "price_cherryx": 0,
            },
        )

        self.assertEqual(publish.status_code, 302)
        published = CommunityWork.objects.get(title="Buyer public text")
        self.assertEqual(published.owner, buyer)
        self.assertEqual(published.status, CommunityWork.STATUS_PUBLISHED)
        self.assertContains(self.client.get(reverse("studio:community_texts")), "Buyer public text")

        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            media_root = root / "media"
            source_image = root / "result.png"
            Image.new("RGB", (80, 80), "#2563eb").save(source_image, "PNG")
            job = JobRecord.objects.create(
                owner=buyer,
                job_id="publishjob001",
                kind="cover",
                title="Generated cover",
                status="completed",
                progress=100,
                output_count=1,
                params_json=json.dumps({"action": "cover"}),
            )
            JobOutputRecord.objects.create(job=job, label="Cover", path=str(source_image), media_type="image/png", size=source_image.stat().st_size)
            with override_settings(MEDIA_ROOT=media_root):
                response = self.client.post(
                    reverse("studio:community_publish"),
                    {
                        "source": "job",
                        "source_id": job.job_id,
                        "kind": CommunityWork.KIND_IMAGE,
                        "title": "Published generated cover",
                        "excerpt": "Copied from job output.",
                        "body": "",
                        "access": CommunityWork.ACCESS_FREE,
                        "price_cherryx": 0,
                    },
                )

            self.assertEqual(response.status_code, 302)
            published_from_job = CommunityWork.objects.get(title="Published generated cover")
            self.assertEqual(published_from_job.source_job, job)
            self.assertTrue(published_from_job.media_file.name.endswith(".webp"))

    def test_designer_page_exposes_mobile_palette_and_drawer_hooks(self) -> None:
        response = self.client.get(reverse("studio:designer"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "data-designer-mobile-palette")
        self.assertContains(response, "data-designer-mobile-drawer")

    def test_files_stats_ajax_partial_contains_controls(self) -> None:
        response = self.client.get(
            reverse("studio:dashboard_detail", args=["files"]),
            {"type": "image", "q": "cover"},
            HTTP_X_REQUESTED_WITH="XMLHttpRequest",
            HTTP_ACCEPT="application/json",
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertIn("html", payload)
        self.assertIn("stats-type-picker", payload["html"])
        self.assertIn("stats-files-table-head", payload["html"])

    def test_localized_ui_text_repairs_mojibake_at_output_boundary(self) -> None:
        self.assertEqual(translate("pricing", "ru"), "Тарифы")
        self.assertEqual(translate("login", "uk"), "Увійти")
        self.assertEqual(translate("edit_design", "uk"), "Редагувати дизайн")
        self.assertEqual(translate("canvas", "ru"), "Холст")
        self.assertEqual(translate("local_draft", "uk"), "Локальна чернетка")
        self.assertEqual(translate("ai_clip_planner", "ru"), "План клипов")
        self.assertEqual(app_messages("ru")["canvas"], "Холст")
        self.assertEqual(app_messages("en")["ai_fallback"], "AI fallback")
        options = {item["code"]: item for item in language_options("ru")}
        self.assertEqual(options["ru"]["native"], "Русский")
        self.assertEqual(options["fr"]["native"], "Français")

    def test_subtitle_style_picker_and_detail_pages_render(self) -> None:
        response = self.client.get(reverse("studio:index"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "data-subtitle-style-picker")
        self.assertContains(response, "target=\"_blank\"")
        for style in actions.SUBTITLE_STYLE_CHOICES:
            detail = self.client.get(reverse("studio:subtitle_style_detail", args=[style[0]]))
            self.assertEqual(detail.status_code, 200)
            self.assertContains(detail, "subtitle-style-demo-line")

    def test_originality_tool_renders_and_accepts_text(self) -> None:
        response = self.client.get(reverse("studio:index"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "data-originality-form")
        self.assertContains(response, reverse("studio:start_originality"))

        repeated = "Данная работа в современном мире показывает важную роль образовательных технологий согласно статистике."
        text = f"{repeated} {repeated} Автор добавляет собственный пример и сравнивает несколько подходов к проверке текста."
        result = self.client.post(reverse("studio:start_originality"), {"text": text})

        self.assertEqual(result.status_code, 200)
        payload = result.json()["analysis"]
        self.assertIn("overall", payload)
        self.assertEqual(len(payload["metrics"]), 6)
        self.assertTrue(any(segment["severity"] != "none" for segment in payload["segments"]))
        self.assertIn("job", result.json())
        self.assertEqual(result.json()["job"]["kind"], "originality")
        self.assertEqual(len(result.json()["job"]["outputs"]), 2)
        self.assertEqual(payload["check"]["mode"], "local")
        self.assertEqual(payload["check"]["price_cherryx"], 5)

        send = self.client.post(reverse("studio:send_originality_report", args=[result.json()["job"]["id"]]), {"email": "teacher@example.com"})
        self.assertEqual(send.status_code, 200)
        self.assertEqual(send.json()["ok"], True)
        self.assertEqual(len(mail.outbox), 1)
        self.assertIn("teacher@example.com", mail.outbox[0].to)

    def test_originality_web_mode_prepares_share_report(self) -> None:
        text = "It is important to note that academic work needs sources and clear structure. " * 4
        response = self.client.post(reverse("studio:start_originality"), {"text": text, "mode": "web"})

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["analysis"]["check"]["mode"], "web")
        self.assertEqual(payload["analysis"]["check"]["price_cherryx"], 25)
        self.assertEqual(payload["analysis"]["check"]["web_queries_limit"], 5)
        self.assertTrue(payload["analysis"]["check"]["web_probes"])
        record = JobRecord.objects.get(job_id=payload["job"]["id"])
        token = _job_record_params(record)["originality"]["share_token"]
        shared = self.client.get(reverse("studio:originality_shared_report", args=[token]))
        self.assertEqual(shared.status_code, 200)
        self.assertContains(shared, "Verified by CherryX Originality")

    def test_originality_tool_reads_docx_upload(self) -> None:
        from docx import Document

        document = Document()
        document.add_paragraph("It is important to note that academic work needs sources and clear structure.")
        document.add_paragraph("It is important to note that academic work needs sources and clear structure.")
        buffer = BytesIO()
        document.save(buffer)
        upload = SimpleUploadedFile(
            "paper.docx",
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        response = self.client.post(reverse("studio:start_originality"), {"file": upload})

        self.assertEqual(response.status_code, 200)
        payload = response.json()["analysis"]
        self.assertGreater(payload["source"]["words"], 10)
        self.assertTrue(any(metric["key"] == "ai_risk" for metric in payload["metrics"]))
        self.assertEqual(response.json()["job"]["status"], "completed")


class QueryOptimizationTests(TransactionTestCase):
    def setUp(self) -> None:
        self.user = get_user_model().objects.create_user("perf@example.com", email="perf@example.com", password="pass12345")
        self.client.force_login(self.user)

    def test_video_project_list_payload_uses_prefetched_assets(self) -> None:
        for index in range(50):
            project = VideoEditorProject.objects.create(
                owner=self.user,
                title=f"Video {index}",
                state_json={"clips": [{"id": f"clip-{index}", "duration": 5}], "tracks": []},
                asset_count=1,
                clip_count=1,
                duration_seconds=5,
            )
            VideoEditorAsset.objects.create(project=project, kind="image", file_path=f"/tmp/{index}.jpg", media_type="image/jpeg", size=10, original_name=f"{index}.jpg")

        with self.assertNumQueries(2):
            queryset = _video_project_queryset_for_test(self.user.id).prefetch_related("assets")
            rows = views._attach_access_roles(list(queryset[:50]), WorkspaceShare.RESOURCE_VIDEO, self.user.id, "")
            payloads = [views._video_project_payload(project, include_state=False, owner_id=self.user.id) for project in rows]

        self.assertEqual(len(payloads), 50)
        self.assertTrue(all(item["asset_count"] == 1 for item in payloads))

    def test_dashboard_output_url_builder_batches_job_records(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            job_dicts = []
            for index in range(25):
                path = root / f"cover-{index}.png"
                Image.new("RGB", (16, 16), "#2563eb").save(path)
                record = JobRecord.objects.create(
                    owner=self.user,
                    job_id=f"job{index:04d}",
                    kind="cover",
                    title=f"Cover {index}",
                    status="completed",
                    progress=100,
                    params_json=json.dumps({"design_projects": {str(path.resolve()): 1000 + index}}),
                )
                JobOutputRecord.objects.create(job=record, label="PNG-cover", path=str(path), media_type="image/png", size=path.stat().st_size)
            records = JobRecord.objects.filter(owner=self.user).prefetch_related("outputs").order_by("id")
            job_dicts = [actions._serialize_job_record(record) for record in records]

            with self.assertNumQueries(2):
                prepared = views._attach_output_urls_many(job_dicts)

        self.assertEqual(len(prepared), 25)
        self.assertTrue(all(job["outputs"][0]["can_edit_design"] for job in prepared))

    def test_account_stats_uses_aggregates(self) -> None:
        records = [
            JobRecord.objects.create(owner=self.user, job_id=f"stats{index}", kind="cover", title="Stats", status="completed", progress=100)
            for index in range(3)
        ]
        for index, record in enumerate(records, start=1):
            JobOutputRecord.objects.create(job=record, label="Output", path=f"/tmp/stats-{index}.png", media_type="image/png", size=index * 100)

        with self.assertNumQueries(2):
            stats = actions.get_account_stats(self.user.id, "")

        self.assertEqual(stats["total_jobs"], 3)
        self.assertEqual(stats["output_count"], 3)
        self.assertEqual(stats["total_output_size"], 600)

    def test_dashboard_files_reads_paginated_outputs_from_database(self) -> None:
        for index in range(205):
            record = JobRecord.objects.create(
                owner=self.user,
                job_id=f"filedash{index:04d}",
                kind="cover",
                title=f"Archive job {index}",
                status="completed",
                progress=100,
            )
            JobOutputRecord.objects.create(
                job=record,
                label=f"Archive {index}",
                path=f"/tmp/archive-{index}.png",
                media_type="image/png",
                size=100 + index,
            )

        response = self.client.get(reverse("studio:dashboard_detail", args=["files"]), {"page": "11"})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context["total_outputs"], 205)
        self.assertEqual(response.context["page"], 11)
        self.assertEqual(response.context["pages"], 11)
        self.assertContains(response, "Archive 0")
        self.assertNotContains(response, "Archive 204")

    def test_dashboard_jobs_are_paginated_from_database(self) -> None:
        for index in range(205):
            JobRecord.objects.create(
                owner=self.user,
                job_id=f"jobdash{index:04d}",
                kind="cover",
                title=f"Dashboard job {index}",
                status="completed",
                progress=100,
            )

        response = self.client.get(reverse("studio:dashboard_detail", args=["all"]), {"page": "11"})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context["total_jobs"], 205)
        self.assertEqual(response.context["page"], 11)
        self.assertEqual(response.context["pages"], 11)
        self.assertContains(response, "Dashboard job 0")
        self.assertNotContains(response, "Dashboard job 204")

    def test_queued_job_can_be_cancelled(self) -> None:
        record = JobRecord.objects.create(
            owner=self.user,
            job_id="cancelqueued01",
            kind="cover",
            title="Queued job",
            status="queued",
            progress=1,
            params_json=json.dumps({"action": "cover"}),
        )

        response = self.client.post(reverse("studio:cancel_job", args=[record.job_id]), HTTP_X_REQUESTED_WITH="XMLHttpRequest")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        record.refresh_from_db()
        self.assertEqual(payload["job"]["status"], "cancelled")
        self.assertEqual(record.status, "cancelled")
        self.assertTrue(record.events.filter(status="cancelled").exists())

    def test_running_job_can_be_stopped(self) -> None:
        record = JobRecord.objects.create(
            owner=self.user,
            job_id="stoprunning01",
            kind="cover",
            title="Running job",
            status="running",
            progress=40,
            params_json=json.dumps({"action": "cover"}),
        )

        response = self.client.post(reverse("studio:cancel_job", args=[record.job_id]), HTTP_X_REQUESTED_WITH="XMLHttpRequest")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        record.refresh_from_db()
        self.assertEqual(payload["job"]["status"], "cancelled")
        self.assertEqual(record.status, "cancelled")
        self.assertTrue(record.events.filter(status="cancelled").exists())

    def test_dashboard_job_actions_follow_status(self) -> None:
        JobRecord.objects.create(owner=self.user, job_id="queuedaction01", kind="cover", title="Queued action", status="queued", progress=1, params_json=json.dumps({"action": "cover"}))
        JobRecord.objects.create(owner=self.user, job_id="runningaction1", kind="cover", title="Running action", status="running", progress=40, params_json=json.dumps({"action": "cover"}))
        completed = JobRecord.objects.create(owner=self.user, job_id="completeaction", kind="cover", title="Complete action", status="completed", progress=100, params_json=json.dumps({"action": "cover"}))
        JobOutputRecord.objects.create(job=completed, label="Output", path="/tmp/complete-action.png", media_type="image/png", size=10)
        JobRecord.objects.create(owner=self.user, job_id="failedaction01", kind="cover", title="Failed action", status="failed", progress=100, error="Render failed", params_json=json.dumps({"action": "cover"}))
        JobRecord.objects.create(owner=self.user, job_id="cancelaction01", kind="cover", title="Cancel action", status="cancelled", progress=100, params_json=json.dumps({"action": "cover"}))

        response = self.client.get(reverse("studio:dashboard_detail", args=["all"]))

        self.assertContains(response, 'action="/api/jobs/queuedaction01/cancel/"')
        self.assertContains(response, 'action="/api/jobs/runningaction1/cancel/"')
        self.assertContains(response, 'action="/api/jobs/failedaction01/repeat/"')
        self.assertContains(response, 'action="/api/jobs/cancelaction01/repeat/"')
        self.assertContains(response, 'href="/download/completeaction/all/"')

    def test_dashboard_tasks_ajax_returns_task_panel(self) -> None:
        JobRecord.objects.create(owner=self.user, job_id="ajaxtask0001", kind="cover", title="Ajax task", status="completed", progress=100)

        response = self.client.get(reverse("studio:dashboard_detail", args=["all"]), HTTP_X_REQUESTED_WITH="XMLHttpRequest")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertIn('data-section="all"', payload["html"])
        self.assertIn("Ajax task", payload["html"])
        self.assertNotIn("stats-files-panel", payload["html"])

    def test_ajax_delete_job_removes_database_rows_and_media_files(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            output_root = root / "outputs" / "django"
            storage_root = root / "storage" / "django"
            job_id = "delmedia0001"
            output_file = output_root / job_id / "cover" / "result.png"
            source_file = storage_root / job_id / "uploads" / "source.png"
            output_file.parent.mkdir(parents=True)
            source_file.parent.mkdir(parents=True)
            output_file.write_bytes(b"generated")
            source_file.write_bytes(b"source")

            original_settings = actions.settings
            try:
                actions.settings = replace(actions.settings, storage_dir=root / "storage")
                with mock.patch.object(actions, "WEB_OUTPUT_ROOT", output_root), mock.patch.object(actions, "WEB_STORAGE_ROOT", storage_root):
                    record = JobRecord.objects.create(
                        owner=self.user,
                        job_id=job_id,
                        kind="cover",
                        title="Delete media",
                        status="completed",
                        progress=100,
                        params_json=json.dumps({"source": str(source_file)}),
                    )
                    JobOutputRecord.objects.create(job=record, label="Result", path=str(output_file), media_type="image/png", size=output_file.stat().st_size)
                    JobEventRecord.objects.create(job=record, status="completed", progress=100, message="Ready")

                    response = self.client.post(reverse("studio:delete_job", args=[job_id]), HTTP_X_REQUESTED_WITH="XMLHttpRequest")
            finally:
                actions.settings = original_settings

            self.assertEqual(response.status_code, 200)
            self.assertTrue(response.json()["ok"])
            self.assertFalse(JobRecord.objects.filter(job_id=job_id).exists())
            self.assertFalse(JobOutputRecord.objects.filter(job__job_id=job_id).exists())
            self.assertFalse(JobEventRecord.objects.filter(job__job_id=job_id).exists())
            self.assertFalse(output_file.exists())
            self.assertFalse(source_file.exists())
            self.assertFalse((output_root / job_id).exists())
            self.assertFalse((storage_root / job_id).exists())

    def test_dashboard_tasks_hide_ai_fallback_summary(self) -> None:
        JobRecord.objects.create(
            owner=self.user,
            job_id="aitasksummary1",
            kind="youtube",
            title="AI task",
            status="failed",
            progress=100,
            params_json=json.dumps(
                {
                    "ai": {
                        "clip_planner": {
                            "status": "fallback",
                            "fallback_reason": "OpenAI is not configured",
                            "selected_outputs": [],
                        }
                    }
                }
            ),
        )

        response = self.client.get(reverse("studio:dashboard_detail", args=["all"]))

        self.assertContains(response, "AI task")
        self.assertNotContains(response, "AI fallback")
        self.assertNotContains(response, "Clip planner")
        self.assertNotContains(response, "OpenAI is not configured")

    def test_share_roles_are_attached_in_one_batch_for_mixed_resources(self) -> None:
        owner = get_user_model().objects.create_user("owner-perf@example.com", email="owner-perf@example.com", password="pass12345")
        design = DesignerProject.objects.create(owner=owner, title="Design")
        video = VideoEditorProject.objects.create(owner=owner, title="Video")
        music = MusicEditorProject.objects.create(owner=owner, title="Music")
        expires_at = timezone.now() + timezone.timedelta(days=7)
        for resource_type, resource in (
            (WorkspaceShare.RESOURCE_DESIGN, design),
            (WorkspaceShare.RESOURCE_VIDEO, video),
            (WorkspaceShare.RESOURCE_MUSIC, music),
        ):
            WorkspaceShare.objects.create(
                owner=owner,
                invited_user=self.user,
                resource_type=resource_type,
                resource_id=resource.id,
                email=self.user.email,
                role=WorkspaceShare.ROLE_EDITOR,
                token=secrets.token_urlsafe(12),
                status=WorkspaceShare.STATUS_ACCEPTED,
                expires_at=expires_at,
            )

        with self.assertNumQueries(1):
            views._attach_access_roles([design], WorkspaceShare.RESOURCE_DESIGN, self.user.id, "")
        with self.assertNumQueries(1):
            views._attach_access_roles([video], WorkspaceShare.RESOURCE_VIDEO, self.user.id, "")
        with self.assertNumQueries(1):
            views._attach_access_roles([music], WorkspaceShare.RESOURCE_MUSIC, self.user.id, "")

        self.assertEqual(design._access_role, WorkspaceShare.ROLE_EDITOR)
        self.assertEqual(video._access_role, WorkspaceShare.ROLE_EDITOR)
        self.assertEqual(music._access_role, WorkspaceShare.ROLE_EDITOR)

    def test_project_list_api_supports_pagination_search_and_sort_metadata(self) -> None:
        for index in range(5):
            VideoEditorProject.objects.create(owner=self.user, title=f"Launch cut {index}", asset_count=index)
        VideoEditorProject.objects.create(owner=self.user, title="Archived draft")

        response = self.client.get(
            reverse("studio:video_projects"),
            {"q": "Launch", "page": "2", "per_page": "2", "sort": "title"},
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["page"], 2)
        self.assertEqual(payload["per_page"], 2)
        self.assertTrue(payload["has_more"])
        self.assertEqual(payload["total"], 5)
        self.assertEqual(payload["query"], "Launch")
        self.assertEqual(payload["sort"], "title")
        self.assertEqual(len(payload["projects"]), 2)
        self.assertTrue(all("Launch" in project["title"] for project in payload["projects"]))

    def test_project_list_page_renders_search_controls_and_filtered_results(self) -> None:
        VideoEditorProject.objects.create(owner=self.user, title="Launch teaser")
        VideoEditorProject.objects.create(owner=self.user, title="Archived draft")

        response = self.client.get(reverse("studio:video_project_list"), {"q": "Launch", "sort": "title"})

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'class="project-filter-bar"')
        self.assertContains(response, "Launch teaser")
        self.assertNotContains(response, "Archived draft")

    def test_project_list_page_wires_live_pagination_hooks(self) -> None:
        VideoEditorProject.objects.create(owner=self.user, title="Launch teaser")
        VideoEditorProject.objects.create(owner=self.user, title="Launch final")

        response = self.client.get(reverse("studio:video_project_list"), {"q": "Launch", "per_page": "1"})

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "data-project-list")
        self.assertContains(response, 'data-project-page="next"')
        self.assertContains(response, "studio/project_lists.js")

    def test_video_editor_exposes_expanded_export_presets(self) -> None:
        response = self.client.get(reverse("studio:video_editor"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'data-export-preset="feed"')
        self.assertContains(response, 'data-export-preset="pinterest"')
        self.assertContains(response, 'data-export-preset="cinema"')
        self.assertContains(response, 'data-aspect="21/9"')

    def test_video_export_size_respects_wide_and_portrait_aspects(self) -> None:
        self.assertEqual(views._video_export_size({"aspect": "4 / 5"}, "1080p"), (1080, 1350))
        self.assertEqual(views._video_export_size({"aspect": "3 / 4"}, "720p"), (720, 960))
        self.assertEqual(views._video_export_size({"aspect": "21 / 9"}, "1080p"), (2520, 1080))
        self.assertEqual(views._video_export_size({"aspect": "4 / 3"}, "720p"), (960, 720))

    def test_all_subtitle_languages_have_picker_flags(self) -> None:
        languages = views._localized_subtitle_languages("ru")
        missing = [item["code"] for item in languages if item["code"] != "auto" and not item.get("flag")]

        self.assertEqual(missing, [])


class JobSchedulerTests(TransactionTestCase):
    def setUp(self) -> None:
        self.user = get_user_model().objects.create_user("queue@example.com", email="queue@example.com", password="pass12345")
        self.original_settings = actions.settings
        self.original_executor = actions._executor
        actions.settings = replace(actions.settings, job_max_workers=20, account_concurrent_jobs=10)
        actions._executor = ThreadPoolExecutor(max_workers=20)
        with actions._lock:
            actions._jobs.clear()
            actions._pending_jobs.clear()
            actions._running_job_ids.clear()
            actions._running_by_account.clear()
        self.addCleanup(self.cleanup_scheduler)

    def cleanup_scheduler(self) -> None:
        with actions._lock:
            actions._pending_jobs.clear()
            actions._running_job_ids.clear()
            actions._running_by_account.clear()
            actions._jobs.clear()
        actions._executor.shutdown(wait=True, cancel_futures=True)
        actions._executor = self.original_executor
        actions.settings = self.original_settings

    def test_scheduler_limits_one_account_to_ten_running_jobs(self) -> None:
        release = threading.Event()
        started: list[str] = []
        started_lock = threading.Lock()
        ten_started = threading.Event()

        def worker(job: actions.WebJob) -> None:
            with started_lock:
                started.append(job.id)
                if len(started) == 10:
                    ten_started.set()
            release.wait(timeout=5)

        for index in range(12):
            actions._submit_job("test", f"Job {index}", worker, {"action": "test"}, self.user.id, "")

        self.assertTrue(ten_started.wait(timeout=5))
        with actions._lock:
            running = [job for job in actions._jobs.values() if job.status == "running"]
            queued = [job for job in actions._jobs.values() if job.status == "queued"]
        self.assertEqual(len(running), 10)
        self.assertEqual(len(queued), 2)

        release.set()
        deadline = time.time() + 5
        while time.time() < deadline:
            with actions._lock:
                completed = [job for job in actions._jobs.values() if job.status == "completed"]
            if len(completed) == 12:
                break
            time.sleep(0.05)
        self.assertEqual(len(completed), 12)


def _video_project_queryset_for_test(owner_id: int):
    return VideoEditorProject.objects.filter(owner_id=owner_id)


class WebOpenAIUpgradeTests(TransactionTestCase):
    def test_ai_disabled_clip_planner_falls_back_to_local_starts(self) -> None:
        job = actions.WebJob(id="aitest01", kind="youtube", title="AI test", params={})

        with mock.patch.object(actions.openai_ai, "is_openai_ready", return_value=False):
            starts = actions._ai_improve_clip_starts(job, "Demo", 120, [10, 40, 80], 3)

        self.assertEqual(starts, [10, 40, 80])
        self.assertEqual(job.params["ai"]["clip_planner"]["status"], "fallback")
        self.assertEqual(job.params["ai"]["clip_planner"]["fallback_reason"], "OpenAI is not configured")
        self.assertEqual(job.params["ai"]["clip_planner"]["selected_outputs"], [])

    def test_ai_clip_planner_uses_valid_returned_starts(self) -> None:
        job = actions.WebJob(id="aitest02", kind="youtube", title="AI test", params={})

        with (
            mock.patch.object(actions.openai_ai, "is_openai_ready", return_value=True),
            mock.patch.object(actions.openai_ai, "plan_clip_moments", return_value={"starts": [80, 999, 10], "model": "test-model"}),
        ):
            starts = actions._ai_improve_clip_starts(job, "Demo", 120, [10, 40, 80], 2)

        self.assertEqual(starts, [10, 80])
        self.assertEqual(job.params["ai"]["clip_planner"]["status"], "used")

    def test_ai_cover_failure_keeps_local_cover_only(self) -> None:
        job = actions.WebJob(id="aitest03", kind="cover", title="Cover", params={})
        with TemporaryDirectory() as tmp:
            reference = Path(tmp) / "cover.png"
            Image.new("RGB", (1280, 720), "#2563eb").save(reference)

            with (
                mock.patch.object(actions.openai_ai, "is_openai_ready", return_value=True),
                mock.patch.object(actions.openai_ai, "generate_cover_prompt", side_effect=RuntimeError("API unavailable")),
            ):
                output = actions._maybe_add_ai_cover(job, reference, Path(tmp) / "ai", "Cover")

        self.assertIsNone(output)
        self.assertEqual(job.outputs, [])
        self.assertEqual(job.params["ai"]["cover"]["status"], "fallback")

    def test_openai_word_transcription_maps_to_subtitle_cues(self) -> None:
        cues = openai_ai._transcription_to_cues(
            {
                "words": [
                    {"word": "Hello", "start": 0.0, "end": 0.4},
                    {"word": "world", "start": 0.45, "end": 0.8},
                ]
            }
        )

        self.assertEqual(cues, [SubtitleCue(start=0.0, end=0.8, text="Hello world")])


class EditableCoverDesignTests(TransactionTestCase):
    def setUp(self) -> None:
        self.temp_dir = TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)
        self.original_settings = views.settings
        views.settings = replace(views.settings, storage_dir=Path(self.temp_dir.name))
        self.addCleanup(lambda: setattr(views, "settings", self.original_settings))
        session = self.client.session
        session["editable_cover_test"] = "1"
        session.save()
        self.guest_key = session.session_key

    def make_job_output(self, *, kind: str = "cover", label: str = "PNG-cover", media_type: str = "image/png", name: str = "cover_business_cover.png") -> tuple[JobRecord, Path]:
        path = Path(self.temp_dir.name) / name
        if media_type.startswith("image/"):
            Image.new("RGB", (1280, 720), "#163c66").save(path)
        else:
            path.write_text("not an image", encoding="utf-8")
        record = JobRecord.objects.create(job_id=secrets.token_hex(6), kind=kind, title="Smoke Cover", status="completed", progress=100, guest_key=self.guest_key, params_json="{}")
        JobOutputRecord.objects.create(job=record, label=label, path=str(path), media_type=media_type, size=path.stat().st_size)
        return record, path

    def test_cover_output_creates_editable_design_once(self) -> None:
        record, _path = self.make_job_output()
        url = reverse("studio:edit_output_design", args=[record.job_id, 0])

        response = self.client.post(url)
        self.assertEqual(response.status_code, 200)
        project_id = response.json()["project"]["id"]
        project = DesignerProject.objects.get(id=project_id)
        self.assertEqual(project.assets.count(), 1)
        self.assertSetEqual({item.get("type") for item in project.state_json["objects"]}, {"frame", "image", "shape", "text"})

        repeat = self.client.post(url)
        self.assertEqual(repeat.status_code, 200)
        self.assertEqual(repeat.json()["project"]["id"], project_id)
        self.assertEqual(DesignerProject.objects.count(), 1)

    def test_non_image_output_is_not_editable_design(self) -> None:
        record, _path = self.make_job_output(kind="subtitles", label="ASS subtitles", media_type="text/plain", name="captions.ass")
        job = views._attach_output_urls(actions.get_job(record.job_id, None, self.guest_key))

        self.assertFalse(job["outputs"][0].get("can_edit_design"))
        response = self.client.post(reverse("studio:edit_output_design", args=[record.job_id, 0]))
        self.assertEqual(response.status_code, 404)


class VideoProjectDeleteTests(TransactionTestCase):
    def setUp(self) -> None:
        self.temp_dir = TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)
        self.original_settings = views.settings
        views.settings = replace(views.settings, storage_dir=Path(self.temp_dir.name))
        self.addCleanup(lambda: setattr(views, "settings", self.original_settings))
        self.user = get_user_model().objects.create_user("owner@example.com", password="pass12345")
        self.other_user = get_user_model().objects.create_user("other@example.com", password="pass12345")
        self.client.force_login(self.user)

    def make_project(self, owner=None, title: str = "Project") -> tuple[VideoEditorProject, Path, Path]:
        project = VideoEditorProject.objects.create(owner=owner or self.user, title=title, state_json={"title": title})
        media_dir = views._video_project_media_dir(project)
        media_dir.mkdir(parents=True, exist_ok=True)
        file_path = media_dir / "clip.mp4"
        thumb_path = media_dir / "thumb.jpg"
        file_path.write_bytes(b"video")
        thumb_path.write_bytes(b"thumb")
        VideoEditorAsset.objects.create(
            project=project,
            kind="video",
            file_path=str(file_path),
            thumbnail_path=str(thumb_path),
            size=file_path.stat().st_size,
            original_name="clip.mp4",
        )
        project.storage_bytes = views._video_project_storage_bytes(project)
        project.save(update_fields=["storage_bytes"])
        return project, media_dir, file_path

    def test_single_delete_removes_project_and_media_directory(self) -> None:
        project, media_dir, file_path = self.make_project()

        response = self.client.post(reverse("studio:delete_video_project", args=[project.id]), data="{}", content_type="application/json")

        self.assertEqual(response.status_code, 200)
        self.assertFalse(VideoEditorProject.objects.filter(id=project.id).exists())
        self.assertFalse(file_path.exists())
        self.assertFalse(media_dir.exists())

    def test_bulk_delete_removes_media_for_owned_projects_only(self) -> None:
        project_a, media_dir_a, file_path_a = self.make_project(title="A")
        project_b, media_dir_b, file_path_b = self.make_project(title="B")
        other_project, other_media_dir, other_file_path = self.make_project(owner=self.other_user, title="Other")

        response = self.client.post(
            reverse("studio:delete_video_projects"),
            data=json.dumps({"ids": [project_a.id, project_b.id, other_project.id]}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["deleted_count"], 2)
        self.assertCountEqual(payload["deleted_ids"], [project_a.id, project_b.id])
        self.assertFalse(VideoEditorProject.objects.filter(id__in=[project_a.id, project_b.id]).exists())
        self.assertTrue(VideoEditorProject.objects.filter(id=other_project.id).exists())
        self.assertFalse(file_path_a.exists())
        self.assertFalse(file_path_b.exists())
        self.assertFalse(media_dir_a.exists())
        self.assertFalse(media_dir_b.exists())
        self.assertTrue(other_file_path.exists())
        self.assertTrue(other_media_dir.exists())

    def test_rename_updates_title_without_clearing_project_state(self) -> None:
        project, _, _ = self.make_project(title="Original")
        project.state_json = {"title": "Original", "clips": [{"id": "clip-1"}], "tracks": [{"id": "track-1"}]}
        project.save(update_fields=["state_json"])

        response = self.client.post(
            reverse("studio:rename_video_project", args=[project.id]),
            data=json.dumps({"title": "Fresh cut"}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        project.refresh_from_db()
        self.assertEqual(project.title, "Fresh cut")
        self.assertEqual(project.state_json["clips"], [{"id": "clip-1"}])
        self.assertEqual(project.state_json["tracks"], [{"id": "track-1"}])

    def test_asset_delete_removes_files_and_referencing_clips(self) -> None:
        project, media_dir, file_path = self.make_project(title="Asset delete")
        asset = project.assets.first()
        project.state_json = {"clips": [{"id": "clip-1", "assetId": asset.id}, {"id": "clip-2", "assetId": 999}]}
        project.save(update_fields=["state_json"])

        response = self.client.post(reverse("studio:delete_video_project_asset", args=[project.id, asset.id]))

        self.assertEqual(response.status_code, 200)
        project.refresh_from_db()
        self.assertFalse(file_path.exists())
        self.assertEqual(project.state_json["clips"], [{"id": "clip-2", "assetId": 999}])
        self.assertTrue(media_dir.exists())

    def test_asset_rename_updates_owned_asset_display_name_only(self) -> None:
        project, _, file_path = self.make_project(title="Rename asset")
        asset = project.assets.first()

        response = self.client.post(
            reverse("studio:rename_video_project_asset", args=[project.id, asset.id]),
            data=json.dumps({"name": "Fresh name.mp4"}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        asset.refresh_from_db()
        self.assertEqual(asset.original_name, "Fresh name.mp4")
        self.assertEqual(Path(asset.file_path), file_path)
        self.assertEqual(response.json()["asset"]["name"], "Fresh name.mp4")

    def test_asset_rename_rejects_foreign_project(self) -> None:
        other_project, _, _ = self.make_project(owner=self.other_user, title="Other asset")
        asset = other_project.assets.first()

        response = self.client.post(
            reverse("studio:rename_video_project_asset", args=[other_project.id, asset.id]),
            data=json.dumps({"name": "Nope.mp4"}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 404)

    def test_export_endpoint_rejects_foreign_project(self) -> None:
        other_project, _, _ = self.make_project(owner=self.other_user, title="Other export")

        response = self.client.post(
            reverse("studio:start_video_project_export", args=[other_project.id]),
            data=json.dumps({"quality": "720p"}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 404)

    def test_export_job_creates_mp4_for_empty_project(self) -> None:
        project = VideoEditorProject.objects.create(owner=self.user, title="Exportable", state_json={"aspect": "1 / 1", "clips": []})

        response = self.client.post(
            reverse("studio:start_video_project_export", args=[project.id]),
            data=json.dumps({"quality": "720p"}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        job = response.json()["job"]
        deadline = time.time() + 20
        while job["status"] not in {"done", "failed"} and time.time() < deadline:
            time.sleep(0.4)
            job = self.client.get(reverse("studio:video_project_export_status", args=[project.id, job["id"]])).json()["job"]
        self.assertEqual(job["status"], "done", job.get("error"))
        self.assertTrue(JobRecord.objects.filter(job_id=job["id"], kind="video_export", status="completed").exists())
        self.assertTrue(job["download_url"])
        download = self.client.get(job["download_url"])
        self.assertEqual(download.status_code, 200)
        download.close()


class DesignerProjectTests(TransactionTestCase):
    def setUp(self) -> None:
        self.temp_dir = TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)
        self.original_settings = views.settings
        views.settings = replace(views.settings, storage_dir=Path(self.temp_dir.name))
        self.addCleanup(lambda: setattr(views, "settings", self.original_settings))
        self.user = get_user_model().objects.create_user("designer@example.com", password="pass12345")
        self.other_user = get_user_model().objects.create_user("designer-other@example.com", password="pass12345")
        self.client.force_login(self.user)

    def image_upload(self, name: str = "poster.png") -> SimpleUploadedFile:
        image = Image.new("RGB", (64, 48), "#2563eb")
        buffer = BytesIO()
        image.save(buffer, "PNG")
        return SimpleUploadedFile(name, buffer.getvalue(), content_type="image/png")

    def make_project(self, owner=None, title: str = "Design") -> DesignerProject:
        project = DesignerProject.objects.create(owner=owner or self.user, title=title, state_json={"title": title, "objects": [], "vectors": []})
        media_dir = views._design_project_media_dir(project)
        media_dir.mkdir(parents=True, exist_ok=True)
        file_path = media_dir / "image.jpg"
        file_path.write_bytes(b"image")
        DesignerAsset.objects.create(project=project, kind="image", file_path=str(file_path), media_type="image/jpeg", size=file_path.stat().st_size, original_name="image.jpg")
        project.storage_bytes = views._design_project_storage_bytes(project)
        project.save(update_fields=["storage_bytes"])
        return project

    def test_create_save_rename_and_detail_for_owned_project(self) -> None:
        response = self.client.post(
            reverse("studio:create_design_project"),
            data=json.dumps({"title": "Brand board", "state": {"title": "Brand board", "objects": [{"id": "a"}], "vectors": []}}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        project_id = response.json()["project"]["id"]
        response = self.client.post(
            reverse("studio:save_design_project", args=[project_id]),
            data=json.dumps({"title": "Brand board v2", "state": {"title": "Brand board v2", "objects": [{"id": "b"}], "vectors": []}}),
            content_type="application/json",
        )
        self.assertEqual(response.status_code, 200)
        project = DesignerProject.objects.get(id=project_id)
        self.assertEqual(project.title, "Brand board v2")
        self.assertEqual(project.state_json["objects"], [{"id": "b"}])

        response = self.client.post(
            reverse("studio:rename_design_project", args=[project_id]),
            data=json.dumps({"title": "Launch cover"}),
            content_type="application/json",
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["project"]["title"], "Launch cover")

    def test_design_project_payload_focuses_latest_visual_layer(self) -> None:
        project = DesignerProject.objects.create(
            owner=self.user,
            title="Focused preview",
            state_json={
                "objects": [
                    {"id": "first", "x": 100, "y": 100, "w": 300, "h": 180},
                    {"id": "latest", "x": 7600, "y": 4900, "w": 600, "h": 420},
                ],
                "strokes": [
                    {"id": "older-stroke", "points": [{"x": 120, "y": 100}, {"x": 180, "y": 140}]}
                ],
                "vectors": [],
            },
        )

        payload = views._design_project_payload(project, include_state=False, owner_id=self.user.id)

        self.assertEqual(payload["preview_focus"], {"x": 88, "y": 80})

    def test_design_project_list_renders_zoomed_preview_focus(self) -> None:
        project = DesignerProject.objects.create(
            owner=self.user,
            title="Preview card",
            state_json={"objects": [{"id": "layer", "x": 4200, "y": 2800, "w": 900, "h": 600}], "vectors": []},
        )
        media_dir = views._design_project_media_dir(project)
        media_dir.mkdir(parents=True, exist_ok=True)
        preview_path = media_dir / "preview.jpg"
        preview_path.write_bytes(b"preview")
        project.preview_path = str(preview_path)
        project.save(update_fields=["preview_path"])

        response = self.client.get(reverse("studio:design_project_list"))
        html = response.content.decode()

        self.assertEqual(response.status_code, 200)
        self.assertIn("design-project-preview-frame", html)
        self.assertIn("design-project-preview-image", html)
        self.assertIn("--preview-x: 52%; --preview-y: 48%;", html)

    def test_foreign_design_project_is_not_accessible(self) -> None:
        other_project = self.make_project(owner=self.other_user)

        response = self.client.get(reverse("studio:design_project_detail", args=[other_project.id]))

        self.assertEqual(response.status_code, 404)

    def test_upload_image_optimizes_and_counts_storage(self) -> None:
        project = DesignerProject.objects.create(owner=self.user, title="Upload", state_json={"objects": []})

        response = self.client.post(reverse("studio:upload_design_project_asset", args=[project.id]), data={"file": self.image_upload()})

        self.assertEqual(response.status_code, 200)
        asset = DesignerAsset.objects.get(project=project)
        self.assertTrue(Path(asset.file_path).exists())
        project.refresh_from_db()
        self.assertGreater(project.storage_bytes, asset.size)
        self.assertEqual(response.json()["asset"]["kind"], "image")

    def test_delete_design_project_removes_media_directory(self) -> None:
        project = self.make_project()
        media_dir = views._design_project_media_dir(project)

        response = self.client.post(reverse("studio:delete_design_project", args=[project.id]), data="{}", content_type="application/json")

        self.assertEqual(response.status_code, 200)
        self.assertFalse(DesignerProject.objects.filter(id=project.id).exists())
        self.assertFalse(media_dir.exists())

    def test_bulk_delete_ignores_foreign_projects(self) -> None:
        project = self.make_project(title="A")
        other_project = self.make_project(owner=self.other_user, title="Other")

        response = self.client.post(
            reverse("studio:delete_design_projects"),
            data=json.dumps({"ids": [project.id, other_project.id]}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["deleted_ids"], [project.id])
        self.assertFalse(DesignerProject.objects.filter(id=project.id).exists())
        self.assertTrue(DesignerProject.objects.filter(id=other_project.id).exists())

    def test_storage_quota_includes_design_project_bytes(self) -> None:
        project = self.make_project()
        stats = views._storage_quota(self.client.get(reverse("studio:index")).wsgi_request, {"total_output_size": 0})

        self.assertIn("total_output_size_text", stats)
        self.assertGreaterEqual(stats["storage_percent"], 0)
        self.assertLess(stats["storage_available"], stats["storage_limit"])
        self.assertGreater(project.storage_bytes, 0)

    def test_export_list_rejects_foreign_project_and_lists_owned_job(self) -> None:
        project = VideoEditorProject.objects.create(owner=self.user, title="Export list", state_json={"aspect": "1 / 1", "clips": []})
        other_project = VideoEditorProject.objects.create(owner=self.other_user, title="Other list", state_json={"clips": []})
        JobRecord.objects.create(owner=self.user, job_id="owned-export", kind="video_export", title="Export", params_json=json.dumps({"project_id": project.id}), status="completed", progress=100)

        response = self.client.get(reverse("studio:list_video_project_exports", args=[project.id]))
        foreign = self.client.get(reverse("studio:list_video_project_exports", args=[other_project.id]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(foreign.status_code, 404)
        self.assertEqual(response.json()["jobs"][0]["id"], "owned-export")

    def test_cancel_export_marks_job_cancelled(self) -> None:
        project = VideoEditorProject.objects.create(owner=self.user, title="Cancel export", state_json={"clips": []})
        JobRecord.objects.create(owner=self.user, job_id="cancel-me", kind="video_export", title="Export", params_json=json.dumps({"project_id": project.id}), status="queued", progress=2)
        views._video_export_jobs["cancel-me"] = {"id": "cancel-me", "project_id": project.id, "owner_id": self.user.id, "guest_key": "", "status": "queued", "path": "", "quality": "720p"}

        response = self.client.post(reverse("studio:cancel_video_project_export", args=[project.id, "cancel-me"]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["job"]["status"], "cancelled")
        self.assertEqual(JobRecord.objects.get(job_id="cancel-me").status, "cancelled")

    def test_upload_image_returns_optimized_visual_asset(self) -> None:
        project = VideoEditorProject.objects.create(owner=self.user, title="Image upload", state_json={"clips": []})
        image_io = BytesIO()
        Image.new("RGB", (3200, 1800), "#4f46e5").save(image_io, "JPEG", quality=95)
        upload = SimpleUploadedFile("poster.jpg", image_io.getvalue(), content_type="image/jpeg")

        response = self.client.post(reverse("studio:upload_video_project_asset", args=[project.id]), {"file": upload, "kind": "image"})

        self.assertEqual(response.status_code, 200)
        asset_payload = response.json()["asset"]
        self.assertEqual(asset_payload["kind"], "image")
        self.assertEqual(asset_payload["visual_kind"], "image")
        self.assertTrue(asset_payload["is_previewable_image"])
        asset = VideoEditorAsset.objects.get(id=asset_payload["id"])
        self.assertIn(asset.media_type, {"image/jpeg", "image/webp"})
        self.assertTrue(Path(asset.file_path).exists())
        with Image.open(asset.file_path) as saved:
            self.assertLessEqual(max(saved.size), 2560)

    def test_upload_pdf_returns_visual_file_payload_and_preview(self) -> None:
        project = VideoEditorProject.objects.create(owner=self.user, title="PDF upload", state_json={"clips": []})
        upload = SimpleUploadedFile("brief.pdf", b"%PDF-1.4\n%test\n", content_type="application/pdf")

        response = self.client.post(reverse("studio:upload_video_project_asset", args=[project.id]), {"file": upload, "kind": "image"})

        self.assertEqual(response.status_code, 200)
        asset_payload = response.json()["asset"]
        self.assertEqual(asset_payload["kind"], "image")
        self.assertEqual(asset_payload["visual_kind"], "pdf")
        self.assertFalse(asset_payload["is_previewable_image"])
        preview = self.client.get(asset_payload["preview_url"])
        self.assertEqual(preview.status_code, 200)
        preview.close()

    def test_waveform_endpoint_creates_compact_json_for_audio_asset(self) -> None:
        project = VideoEditorProject.objects.create(owner=self.user, title="Waveform", state_json={"clips": []})
        media_dir = views._video_project_media_dir(project)
        media_dir.mkdir(parents=True, exist_ok=True)
        audio_path = media_dir / "tone.mp3"
        audio_path.write_bytes(bytes(range(256)) * 8)
        asset = VideoEditorAsset.objects.create(project=project, kind="audio", file_path=str(audio_path), media_type="audio/mpeg", size=audio_path.stat().st_size, original_name="tone.mp3")

        response = self.client.post(reverse("studio:video_project_asset_waveform", args=[project.id, asset.id]))

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["count"], 96)
        self.assertEqual(len(payload["samples"]), 96)

    def test_cover_endpoint_creates_protected_image_output(self) -> None:
        project = VideoEditorProject.objects.create(owner=self.user, title="Cover", state_json={"aspect": "1 / 1", "clips": []})

        response = self.client.post(
            reverse("studio:export_video_project_cover", args=[project.id]),
            data=json.dumps({"time": 0}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        job = response.json()["job"]
        deadline = time.time() + 8
        while job["status"] not in {"done", "failed"} and time.time() < deadline:
            time.sleep(0.2)
            job = self.client.get(reverse("studio:video_project_export_status", args=[project.id, job["id"]])).json()["job"]
        self.assertEqual(job["status"], "done")
        download = self.client.get(job["download_url"])
        self.assertEqual(download.status_code, 200)
        self.assertEqual(download["Content-Type"], "image/jpeg")
        download.close()

    def test_export_job_creates_mp4_for_image_and_pdf_timeline(self) -> None:
        project = VideoEditorProject.objects.create(owner=self.user, title="Visual export", state_json={"aspect": "1 / 1", "clips": []})
        media_dir = views._video_project_media_dir(project)
        media_dir.mkdir(parents=True, exist_ok=True)
        image_path = media_dir / "poster.jpg"
        Image.new("RGB", (720, 720), "#10b981").save(image_path, "JPEG", quality=90)
        pdf_path = media_dir / "brief.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n%test\n")
        image_asset = VideoEditorAsset.objects.create(project=project, kind="image", file_path=str(image_path), media_type="image/jpeg", size=image_path.stat().st_size, original_name="poster.jpg")
        pdf_asset = VideoEditorAsset.objects.create(project=project, kind="image", file_path=str(pdf_path), media_type="application/pdf", size=pdf_path.stat().st_size, original_name="brief.pdf")
        project.state_json = {
            "aspect": "1 / 1",
            "clips": [
                {"id": "image-1", "type": "image", "assetId": image_asset.id, "start": 0, "duration": 1.2, "x": 50, "y": 50},
                {"id": "pdf-1", "type": "image", "assetId": pdf_asset.id, "start": 0.4, "duration": 1.0, "x": 50, "y": 70},
            ],
        }
        project.save(update_fields=["state_json"])

        response = self.client.post(
            reverse("studio:start_video_project_export", args=[project.id]),
            data=json.dumps({"quality": "720p"}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        job = response.json()["job"]
        deadline = time.time() + 20
        while job["status"] not in {"done", "failed"} and time.time() < deadline:
            time.sleep(0.4)
            job = self.client.get(reverse("studio:video_project_export_status", args=[project.id, job["id"]])).json()["job"]
        self.assertEqual(job["status"], "done", job.get("error"))
        download = self.client.get(job["download_url"])
        self.assertEqual(download.status_code, 200)
        download.close()

    def test_export_job_creates_mp4_for_video_with_image_and_text_overlays(self) -> None:
        project = VideoEditorProject.objects.create(owner=self.user, title="Overlay export", state_json={"aspect": "16 / 9", "clips": []})
        media_dir = views._video_project_media_dir(project)
        media_dir.mkdir(parents=True, exist_ok=True)
        video_path = media_dir / "base.mp4"
        image_path = media_dir / "overlay.jpg"
        Image.new("RGB", (320, 240), "#f59e0b").save(image_path, "JPEG", quality=90)
        completed = subprocess.run(
            [
                views.ffmpeg_path(),
                "-y",
                "-f",
                "lavfi",
                "-i",
                "color=c=blue:s=320x180:d=1.2:r=24",
                "-pix_fmt",
                "yuv420p",
                str(video_path),
            ],
            capture_output=True,
            text=True,
            timeout=20,
        )
        self.assertEqual(completed.returncode, 0, completed.stderr)
        video_asset = VideoEditorAsset.objects.create(project=project, kind="video", file_path=str(video_path), media_type="video/mp4", size=video_path.stat().st_size, original_name="base.mp4", duration=1.2)
        image_asset = VideoEditorAsset.objects.create(project=project, kind="image", file_path=str(image_path), media_type="image/jpeg", size=image_path.stat().st_size, original_name="overlay.jpg")
        project.state_json = {
            "aspect": "16 / 9",
            "clips": [
                {"id": "video-1", "type": "video", "assetId": video_asset.id, "start": 0, "duration": 1.2, "sourceStart": 0},
                {"id": "image-1", "type": "image", "assetId": image_asset.id, "start": 0.1, "duration": 0.8, "x": 55, "y": 45, "scale": 30},
                {"id": "text-1", "type": "text", "start": 0.2, "duration": 0.7, "text": "Hello", "y": 72, "style": {"size": 28}},
            ],
        }
        project.save(update_fields=["state_json"])

        response = self.client.post(
            reverse("studio:start_video_project_export", args=[project.id]),
            data=json.dumps({"quality": "720p"}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        job = response.json()["job"]
        deadline = time.time() + 20
        while job["status"] not in {"done", "failed"} and time.time() < deadline:
            time.sleep(0.4)
            job = self.client.get(reverse("studio:video_project_export_status", args=[project.id, job["id"]])).json()["job"]
        self.assertEqual(job["status"], "done", job.get("error"))
        download = self.client.get(job["download_url"])
        self.assertEqual(download.status_code, 200)
        download.close()


@override_settings(EMAIL_BACKEND="django.core.mail.backends.locmem.EmailBackend")
class WorkspaceSharingTests(TransactionTestCase):
    def setUp(self) -> None:
        self.temp_dir = TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)
        self.original_settings = views.settings
        views.settings = replace(views.settings, storage_dir=Path(self.temp_dir.name))
        self.addCleanup(lambda: setattr(views, "settings", self.original_settings))
        mail.outbox = []
        self.owner = get_user_model().objects.create_user("owner-share@example.com", email="owner-share@example.com", password="pass12345", first_name="Owner")
        self.viewer = get_user_model().objects.create_user("viewer-share@example.com", email="viewer-share@example.com", password="pass12345")
        self.editor = get_user_model().objects.create_user("editor-share@example.com", email="editor-share@example.com", password="pass12345")
        self.other = get_user_model().objects.create_user("other-share@example.com", email="other-share@example.com", password="pass12345")
        self.design = DesignerProject.objects.create(owner=self.owner, title="Shared design", state_json={"objects": [], "vectors": []})
        self.video = VideoEditorProject.objects.create(owner=self.owner, title="Shared video", state_json={"clips": []})
        self.client.force_login(self.owner)

    def make_share(self, resource_type: str, resource, user=None, email: str = "", role: str = WorkspaceShare.ROLE_VIEWER, status: str = WorkspaceShare.STATUS_ACCEPTED, expires_at=None) -> WorkspaceShare:
        email = email or getattr(user, "email", "")
        return WorkspaceShare.objects.create(
            owner=self.owner,
            invited_user=user if status == WorkspaceShare.STATUS_ACCEPTED else None,
            resource_type=resource_type,
            resource_id=resource.id,
            email=email.lower(),
            role=role,
            status=status,
            token=secrets.token_urlsafe(12),
            expires_at=expires_at or timezone.now() + timezone.timedelta(days=14),
        )

    def image_upload(self, name: str = "share.png") -> SimpleUploadedFile:
        image_io = BytesIO()
        Image.new("RGB", (64, 64), "#2563eb").save(image_io, "PNG")
        return SimpleUploadedFile(name, image_io.getvalue(), content_type="image/png")

    def test_owner_invites_existing_and_new_email(self) -> None:
        response = self.client.post(
            reverse("studio:workspace_shares"),
            data=json.dumps({"resource_type": WorkspaceShare.RESOURCE_DESIGN, "resource_id": self.design.id, "email": self.viewer.email, "role": WorkspaceShare.ROLE_EDITOR}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        share = WorkspaceShare.objects.get(id=response.json()["share"]["id"])
        self.assertEqual(share.invited_user, self.viewer)
        self.assertEqual(share.status, WorkspaceShare.STATUS_PENDING)
        self.assertEqual(share.role, WorkspaceShare.ROLE_EDITOR)
        self.assertEqual(len(mail.outbox), 1)
        self.assertIn(share.token, mail.outbox[0].body)

        response = self.client.post(
            reverse("studio:workspace_shares"),
            data=json.dumps({"resource_type": WorkspaceShare.RESOURCE_VIDEO, "resource_id": self.video.id, "email": "new-person@example.com", "role": WorkspaceShare.ROLE_VIEWER}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        share = WorkspaceShare.objects.get(id=response.json()["share"]["id"])
        self.assertIsNone(share.invited_user)
        self.assertEqual(len(mail.outbox), 2)

    def test_login_and_register_auto_accept_pending_invite(self) -> None:
        login_share = self.make_share(WorkspaceShare.RESOURCE_DESIGN, self.design, user=self.viewer, status=WorkspaceShare.STATUS_PENDING)
        self.client.logout()

        response = self.client.post(
            reverse("studio:login"),
            data={"email": self.viewer.email, "password": "pass12345", "next": reverse("studio:workspace_invite", args=[login_share.token])},
        )

        self.assertEqual(response.status_code, 302)
        login_share.refresh_from_db()
        self.assertEqual(login_share.status, WorkspaceShare.STATUS_ACCEPTED)
        self.assertEqual(login_share.invited_user, self.viewer)

        register_share = self.make_share(WorkspaceShare.RESOURCE_VIDEO, self.video, email="brand-new-share@example.com", status=WorkspaceShare.STATUS_PENDING)
        self.client.logout()
        response = self.client.post(
            reverse("studio:register"),
            data={
                "name": "Brand New",
                "email": "brand-new-share@example.com",
                "password": "pass12345",
                "password_confirm": "pass12345",
                "next": reverse("studio:workspace_invite", args=[register_share.token]),
            },
        )

        self.assertEqual(response.status_code, 302)
        register_share.refresh_from_db()
        self.assertEqual(register_share.status, WorkspaceShare.STATUS_ACCEPTED)
        self.assertEqual(register_share.invited_user.email, "brand-new-share@example.com")

    def test_viewer_can_read_but_cannot_mutate_or_share(self) -> None:
        self.make_share(WorkspaceShare.RESOURCE_DESIGN, self.design, self.viewer, role=WorkspaceShare.ROLE_VIEWER)
        self.client.force_login(self.viewer)

        detail = self.client.get(reverse("studio:design_project_detail", args=[self.design.id]))
        save = self.client.post(reverse("studio:save_design_project", args=[self.design.id]), data=json.dumps({"state": {"objects": [{"id": "x"}]}}), content_type="application/json")
        upload = self.client.post(reverse("studio:upload_design_project_asset", args=[self.design.id]), data={"file": self.image_upload()})
        delete = self.client.post(reverse("studio:delete_design_project", args=[self.design.id]), data="{}", content_type="application/json")
        share = self.client.post(
            reverse("studio:workspace_shares"),
            data=json.dumps({"resource_type": WorkspaceShare.RESOURCE_DESIGN, "resource_id": self.design.id, "email": self.other.email, "role": WorkspaceShare.ROLE_VIEWER}),
            content_type="application/json",
        )

        self.assertEqual(detail.status_code, 200)
        self.assertFalse(detail.json()["project"]["is_owner"])
        self.assertFalse(detail.json()["project"]["can_edit"])
        self.assertFalse(detail.json()["project"]["can_share"])
        self.assertEqual(save.status_code, 404)
        self.assertEqual(upload.status_code, 404)
        self.assertEqual(delete.status_code, 404)
        self.assertEqual(share.status_code, 404)

    def test_editor_can_edit_content_but_cannot_manage_project(self) -> None:
        self.make_share(WorkspaceShare.RESOURCE_DESIGN, self.design, self.editor, role=WorkspaceShare.ROLE_EDITOR)
        self.make_share(WorkspaceShare.RESOURCE_VIDEO, self.video, self.editor, role=WorkspaceShare.ROLE_EDITOR)
        self.client.force_login(self.editor)

        save = self.client.post(reverse("studio:save_design_project", args=[self.design.id]), data=json.dumps({"state": {"objects": [{"id": "editor"}]}}), content_type="application/json")
        upload = self.client.post(reverse("studio:upload_design_project_asset", args=[self.design.id]), data={"file": self.image_upload("editor.png")})
        export = self.client.post(reverse("studio:start_video_project_export", args=[self.video.id]), data=json.dumps({"quality": "720p"}), content_type="application/json")
        delete = self.client.post(reverse("studio:delete_design_project", args=[self.design.id]), data="{}", content_type="application/json")
        share = self.client.post(
            reverse("studio:workspace_shares"),
            data=json.dumps({"resource_type": WorkspaceShare.RESOURCE_VIDEO, "resource_id": self.video.id, "email": self.other.email, "role": WorkspaceShare.ROLE_VIEWER}),
            content_type="application/json",
        )

        self.assertEqual(save.status_code, 200)
        self.assertEqual(upload.status_code, 200)
        self.assertEqual(export.status_code, 200)
        self.assertEqual(delete.status_code, 404)
        self.assertEqual(share.status_code, 404)

    def test_revoked_and_expired_shares_do_not_grant_access(self) -> None:
        self.make_share(WorkspaceShare.RESOURCE_DESIGN, self.design, self.viewer, status=WorkspaceShare.STATUS_REVOKED)
        self.make_share(WorkspaceShare.RESOURCE_VIDEO, self.video, self.viewer, status=WorkspaceShare.STATUS_ACCEPTED, expires_at=timezone.now() - timezone.timedelta(seconds=1))
        self.client.force_login(self.viewer)

        design = self.client.get(reverse("studio:design_project_detail", args=[self.design.id]))
        video = self.client.get(reverse("studio:video_project_detail", args=[self.video.id]))

        self.assertEqual(design.status_code, 404)
        self.assertEqual(video.status_code, 404)

    def test_shared_projects_appear_in_project_lists_with_access_flags(self) -> None:
        self.make_share(WorkspaceShare.RESOURCE_DESIGN, self.design, self.viewer, role=WorkspaceShare.ROLE_VIEWER)
        self.make_share(WorkspaceShare.RESOURCE_VIDEO, self.video, self.viewer, role=WorkspaceShare.ROLE_EDITOR)
        self.client.force_login(self.viewer)

        design_projects = self.client.get(reverse("studio:design_projects")).json()["projects"]
        video_projects = self.client.get(reverse("studio:video_projects")).json()["projects"]

        self.assertEqual(design_projects[0]["id"], self.design.id)
        self.assertEqual(design_projects[0]["access_role"], WorkspaceShare.ROLE_VIEWER)
        self.assertFalse(design_projects[0]["is_owner"])
        self.assertFalse(design_projects[0]["can_edit"])
        self.assertEqual(video_projects[0]["id"], self.video.id)
        self.assertEqual(video_projects[0]["access_role"], WorkspaceShare.ROLE_EDITOR)
        self.assertTrue(video_projects[0]["can_edit"])

    def test_wrong_email_cannot_accept_invite(self) -> None:
        share = self.make_share(WorkspaceShare.RESOURCE_DESIGN, self.design, user=self.viewer, status=WorkspaceShare.STATUS_PENDING)
        self.client.force_login(self.other)

        response = self.client.get(reverse("studio:workspace_invite", args=[share.token]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "This invite is for another email")
        share.refresh_from_db()
        self.assertEqual(share.status, WorkspaceShare.STATUS_PENDING)

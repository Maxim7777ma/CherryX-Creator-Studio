from __future__ import annotations

from base64 import b64decode
from collections import Counter
from concurrent.futures import ThreadPoolExecutor
from io import BytesIO
from pathlib import Path
from urllib.parse import quote, urlencode
import json
import html
import math
import mimetypes
import random
import re
import secrets
import shutil
import statistics
import subprocess
import threading
import time
import uuid
import zipfile

from django.contrib.auth import get_user_model, login, logout, update_session_auth_hash
from django.contrib.auth.decorators import login_required
from django.core.exceptions import ValidationError
from django.core.files import File
from django.core.files.base import ContentFile
from django.core.mail import EmailMultiAlternatives
from django.core.validators import validate_email
from django.db import transaction
from django.db.models import Count, F, Q, Sum
from django.http import FileResponse, Http404, HttpRequest, HttpResponse, JsonResponse, StreamingHttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.template.loader import render_to_string
from django.urls import reverse
from django.utils.html import strip_tags
from django.utils.http import url_has_allowed_host_and_scheme
from django.utils import timezone
from django.views.decorators.csrf import ensure_csrf_cookie
from django.views.decorators.http import require_GET, require_POST, require_http_methods
from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from pypdf import PdfReader

from billing.plans import PLANS, get_plan
from billing.services import active_access_until, cherryx_to_usd_display_approx, cherryx_to_usd_cents_approx, ensure_telegram_link_token, prorated_due_cents, telegram_stars_rate, transfer_guest_workspace, user_has_active_access
from src.config import get_settings
from src.image_tools import clean_base_name, human_size
from src.job_service import job_service as actions
from src import openai_ai
from src.video_tools import ffmpeg_path, inspect_video
from src.youtube_tools import SubtitleUnavailableError, normalize_subtitle_language, transcribe_subtitle_cues
from .forms import AccountSettingsForm, CommunityWorkForm, EmailLoginForm, RegisterForm
from .documentation_content import documentation_content
from .legal_documents import legal_document_content
from .localization import LANGUAGE_OPTIONS, app_messages, cherryx_pay_messages, clean_language, localized_plan, music_messages, repair_mojibake, translate
from .models import AccountProfile, CommunityPurchase, CommunityWork, DesignerAsset, DesignerProject, JobEventRecord, JobOutputRecord, JobRecord, LearningArticle, MagicLoginToken, MusicEditorAsset, MusicEditorProject, VideoEditorAsset, VideoEditorProject, WorkspaceShare
from .wallet import WalletError, create_cherryx_withdrawal_request, recent_wallet_transactions, transfer_cherryx_by_email


settings = get_settings()
_video_export_executor = ThreadPoolExecutor(max_workers=max(1, settings.video_export_workers))
_video_export_jobs: dict[str, dict[str, object]] = {}
_video_export_lock = threading.RLock()
_video_export_processes: dict[str, subprocess.Popen] = {}
RESUME_FIELDS = [
    "name",
    "position",
    "resume_mode",
    "target_role",
    "target_scope",
    "work_format",
    "value_offer",
    "vacancy_text",
    "cover_letter",
    "contact",
    "links",
    "summary",
    "experience",
    "education",
    "skills",
    "achievements",
    "additional",
]
ORIGINALITY_MAX_UPLOAD_BYTES = 8 * 1024 * 1024
ORIGINALITY_MAX_ANALYSIS_CHARS = 120_000
ORIGINALITY_MODES = {
    "local": {
        "label": "Local",
        "price": 5,
        "max_chars": 20_000,
        "max_web_queries": 0,
        "summary": "Fast local AI risk, repetition, sources and structure scan.",
    },
    "web": {
        "label": "Web",
        "price": 25,
        "max_chars": 45_000,
        "max_web_queries": 5,
        "summary": "Prepared for web originality search with selected quote checks.",
    },
    "deep": {
        "label": "Deep Web",
        "price": 60,
        "max_chars": 120_000,
        "max_web_queries": 20,
        "summary": "Prepared for deep web search, source matching and citation review.",
    },
}
ORIGINALITY_VISIBLE_SEGMENTS = 240
ORIGINALITY_WORD_RE = re.compile(r"[\wА-Яа-яЁёІіЇїЄєҐґ’'-]+", re.UNICODE)
ORIGINALITY_SENTENCE_RE = re.compile(r"[^.!?…\n]+(?:[.!?…]+|$)", re.UNICODE)


@require_GET
def landing(request: HttpRequest):
    if request.user.is_authenticated:
        return redirect("studio:index")
    return render(request, "studio/landing.html")


@require_GET
def legal_info(request: HttpRequest):
    return render(request, "studio/legal_info.html")


@require_GET
def legal_terms(request: HttpRequest):
    language = getattr(request, "interface_language", "en")
    return render(request, "studio/legal_document.html", {"document_type": "terms", "legal_doc": legal_document_content("terms", language)})


@require_GET
def legal_refund(request: HttpRequest):
    language = getattr(request, "interface_language", "en")
    return render(request, "studio/legal_document.html", {"document_type": "refund", "legal_doc": legal_document_content("refund", language)})


@require_GET
def legal_contacts(request: HttpRequest):
    language = getattr(request, "interface_language", "en")
    return render(request, "studio/legal_document.html", {"document_type": "contacts", "legal_doc": legal_document_content("contacts", language)})


@require_GET
def documentation(request: HttpRequest):
    language = getattr(request, "interface_language", "en")
    return render(request, "studio/documentation.html", {"docs": documentation_content(language)})


@require_GET
def learn_index(request: HttpRequest):
    articles = LearningArticle.objects.filter(status=LearningArticle.STATUS_PUBLISHED).order_by("-featured", "-published_at", "-created_at")
    return render(
        request,
        "studio/learn_index.html",
        {
            "articles": articles,
            "featured_article": articles.first(),
            "seo_title": "CherryX learning articles",
            "seo_description": "Practical CherryX Creator Studio lessons for video, images, texts, publishing workflow and Telegram Stars payments.",
        },
    )


@require_GET
def learn_article(request: HttpRequest, slug: str):
    article = get_object_or_404(LearningArticle, slug=slug, status=LearningArticle.STATUS_PUBLISHED)
    return render(
        request,
        "studio/learn_article.html",
        {
            "article": article,
            "seo_title": article.seo_title or article.title,
            "seo_description": article.seo_description or article.excerpt or strip_tags(article.body)[:240],
        },
    )


def _community_base_queryset():
    return (
        CommunityWork.objects.filter(status=CommunityWork.STATUS_PUBLISHED)
        .select_related("owner")
        .annotate(total_purchases=Count("purchases"))
    )


def _community_apply_filters(queryset, request: HttpRequest):
    query = (request.GET.get("q") or "").strip()
    access = (request.GET.get("access") or "").strip()
    sort = (request.GET.get("sort") or "new").strip()
    min_price = (request.GET.get("min_price") or "").strip()
    max_price = (request.GET.get("max_price") or "").strip()

    if query:
        queryset = queryset.filter(
            Q(title__icontains=query)
            | Q(excerpt__icontains=query)
            | Q(body__icontains=query)
            | Q(owner__first_name__icontains=query)
            | Q(owner__username__icontains=query)
            | Q(owner__email__icontains=query)
        )
    if access in {CommunityWork.ACCESS_FREE, CommunityWork.ACCESS_PAID}:
        queryset = queryset.filter(access=access)
    if min_price.isdigit():
        queryset = queryset.filter(price_cherryx__gte=int(min_price))
    if max_price.isdigit():
        queryset = queryset.filter(price_cherryx__lte=int(max_price))

    orderings = {
        "popular": ("-featured", "-total_purchases", "-download_count", "-published_at", "-created_at"),
        "price_low": ("price_cherryx", "-featured", "-published_at", "-created_at"),
        "price_high": ("-price_cherryx", "-featured", "-published_at", "-created_at"),
        "new": ("-featured", "-published_at", "-created_at"),
    }
    queryset = queryset.order_by(*orderings.get(sort, orderings["new"]))
    filters = {
        "q": query,
        "access": access,
        "sort": sort if sort in orderings else "new",
        "min_price": min_price,
        "max_price": max_price,
    }
    return queryset, filters


def _community_queryset(kind: str, request: HttpRequest | None = None):
    queryset = _community_base_queryset().filter(kind=kind)
    if request is None:
        return queryset.order_by("-featured", "-published_at", "-created_at")
    return _community_apply_filters(queryset, request)


def _community_kind_meta(kind: str) -> dict[str, str]:
    return {
        CommunityWork.KIND_VIDEO: {
            "title": "CherryX Video Gallery",
            "eyebrow": "Video network",
            "description": "Pinterest-style feed of videos created or shared through CherryX Creator Studio.",
            "icon": "video",
        },
        CommunityWork.KIND_IMAGE: {
            "title": "CherryX Image Gallery",
            "eyebrow": "Image network",
            "description": "A visual feed for covers, designs, previews and images created on the platform or uploaded by users.",
            "icon": "image",
        },
        CommunityWork.KIND_TEXT: {
            "title": "CherryX Text Library",
            "eyebrow": "Text network",
            "description": "Articles, text packs, prompts and written works shared by CherryX users.",
            "icon": "file-text",
        },
        CommunityWork.KIND_MUSIC: {
            "title": "CherryX Music Market",
            "eyebrow": "Music network",
            "description": "Beats, finished tracks and editable CherryX music projects shared by creators.",
            "icon": "music-2",
        },
    }[kind]


def _community_usd_display(cherryx_amount: int) -> str:
    return cherryx_to_usd_display_approx(int(cherryx_amount or 0))


def _community_usd_rate_per_cherryx() -> str:
    return f"{cherryx_to_usd_cents_approx(1) / 100:.6f}"


def _community_work_preview_url(work: CommunityWork, request: HttpRequest) -> str:
    if work.cover_image:
        return work.cover_image.url
    if work.kind == CommunityWork.KIND_IMAGE and work.media_file:
        return work.media_file.url
    if work.status == CommunityWork.STATUS_PUBLISHED and _community_preview_field(work):
        return reverse("studio:community_work_preview", args=[work.slug])
    return ""


def _delete_community_work_files(work: CommunityWork) -> None:
    for field_name in ("cover_image", "media_file"):
        field = getattr(work, field_name, None)
        if field and field.name:
            field.delete(save=False)


def _community_manage_payload(work: CommunityWork, request: HttpRequest) -> dict[str, object]:
    public_path = reverse("studio:community_work_detail", args=[work.slug])
    purchases = int(getattr(work, "total_purchases", None) or work.purchase_count or 0)
    revenue = int(getattr(work, "total_revenue", None) or 0)
    if not revenue and work.is_paid:
        revenue = purchases * int(work.price_cherryx or 0)
    return {
        "work": work,
        "preview_url": _community_work_preview_url(work, request),
        "public_url": request.build_absolute_uri(public_path),
        "public_path": public_path,
        "purchases": purchases,
        "revenue": revenue,
        "is_published": work.status == CommunityWork.STATUS_PUBLISHED,
        "visibility_action": "hide" if work.status == CommunityWork.STATUS_PUBLISHED else "publish",
        "visibility_label": "Hide from public" if work.status == CommunityWork.STATUS_PUBLISHED else "Publish again",
    }


@require_GET
def community_market(request: HttpRequest):
    works, filters = _community_apply_filters(_community_base_queryset(), request)
    featured = works[:18]
    articles = LearningArticle.objects.filter(status=LearningArticle.STATUS_PUBLISHED).order_by("-featured", "-published_at", "-created_at")[:6]
    return render(
        request,
        "studio/community_market.html",
        {
            "kind": "market",
            "works": featured,
            "articles": articles,
            "filters": filters,
            "cherryx_usd_rate": _community_usd_rate_per_cherryx(),
            "seo_title": "CherryX Marketplace",
            "seo_description": "Explore CherryX videos, images, texts and learning articles with free and paid CherryX downloads.",
        },
    )


def _community_source_kind_from_media(media_type: str, path: str) -> str:
    value = (media_type or mimetypes.guess_type(path)[0] or "").lower()
    suffix = Path(path or "").suffix.lower()
    if value.startswith("video/") or suffix in {".mp4", ".mov", ".webm", ".m4v"}:
        return CommunityWork.KIND_VIDEO
    if value.startswith("image/") or suffix in {".png", ".jpg", ".jpeg", ".webp", ".gif"}:
        return CommunityWork.KIND_IMAGE
    return CommunityWork.KIND_TEXT


def _community_clean_source_title(value: str) -> str:
    title = str(value or "").strip()
    title = re.sub(r"^(?:edit[:_\-\s]+)+", "", title, flags=re.IGNORECASE).strip()
    title = re.sub(r"_[0-9]{3,4}p(?:_[a-f0-9]{8,})?$", "", title, flags=re.IGNORECASE)
    title = re.sub(r"_[a-f0-9]{12,}$", "", title, flags=re.IGNORECASE)
    title = re.sub(r"_short_\d+(?:_\d+)?$", "", title, flags=re.IGNORECASE)
    title = title.replace("_", " ")
    title = re.sub(r"\s+", " ", title).strip(" -_")
    return title[:180] or str(value or "").strip()[:180] or "CherryX work"


def _community_source_for_user(user, source: str, source_id: str) -> dict[str, object] | None:
    source = (source or "").strip()
    source_id = (source_id or "").strip()
    if not user or not getattr(user, "is_authenticated", False) or not source_id:
        return None
    if source == "job":
        record = JobRecord.objects.filter(owner=user, job_id=source_id).prefetch_related("outputs").first()
        if not record:
            return None
        output = record.outputs.first()
        path = str(output.path) if output else ""
        kind = _community_source_kind_from_media(output.media_type if output else "", path)
        return {
            "source": source,
            "source_id": source_id,
            "title": _community_clean_source_title(record.title),
            "excerpt": f"Published from CherryX task {record.job_id}.",
            "kind": kind,
            "media_path": path if path and Path(path).exists() else "",
            "cover_path": path if kind == CommunityWork.KIND_IMAGE and path and Path(path).exists() else "",
            "source_job": record,
        }
    if source == "video_project":
        project = VideoEditorProject.objects.filter(owner=user, id=source_id).prefetch_related("assets").first()
        if not project:
            return None
        first_video = next((asset for asset in project.assets.all() if str(asset.kind).lower() == "video" and asset.file_path and Path(asset.file_path).exists()), None)
        first_thumb = next((asset for asset in project.assets.all() if asset.thumbnail_path and Path(asset.thumbnail_path).exists()), None)
        return {
            "source": source,
            "source_id": source_id,
            "title": _community_clean_source_title(project.title),
            "excerpt": "Published from a CherryX video workspace.",
            "kind": CommunityWork.KIND_VIDEO,
            "media_path": first_video.file_path if first_video else "",
            "cover_path": project.thumbnail_path if project.thumbnail_path and Path(project.thumbnail_path).exists() else (first_thumb.thumbnail_path if first_thumb else ""),
            "source_video_project": project,
        }
    if source == "design_project":
        project = DesignerProject.objects.filter(owner=user, id=source_id).first()
        if not project:
            return None
        return {
            "source": source,
            "source_id": source_id,
            "title": _community_clean_source_title(project.title),
            "excerpt": "Published from a CherryX design workspace.",
            "kind": CommunityWork.KIND_IMAGE,
            "media_path": project.preview_path if project.preview_path and Path(project.preview_path).exists() else "",
            "cover_path": project.preview_path if project.preview_path and Path(project.preview_path).exists() else "",
            "source_design_project": project,
        }
    if source == "music_project":
        project = MusicEditorProject.objects.filter(owner=user, id=source_id).prefetch_related("assets").first()
        if not project:
            return None
        first_audio = next((asset for asset in project.assets.all() if asset.kind == "audio" and asset.file_path and Path(asset.file_path).exists()), None)
        return {
            "source": source,
            "source_id": source_id,
            "title": _community_clean_source_title(project.title),
            "excerpt": "Published from a CherryX music workspace.",
            "kind": CommunityWork.KIND_MUSIC,
            "body": "Beat, track or editable CherryX music project. Add usage notes, BPM, license terms or lyrics before publishing.",
            "media_path": first_audio.file_path if first_audio else "",
            "cover_path": "",
            "source_music_project": project,
        }
    return None


def _media_url_for_path(path_text: str) -> str:
    if not path_text:
        return ""
    try:
        media_root = Path(settings.MEDIA_ROOT).resolve()
        path = Path(path_text).resolve()
        relative = path.relative_to(media_root)
    except Exception:
        return ""
    media_url = str(getattr(settings, "MEDIA_URL", "/media/") or "/media/")
    if not media_url.startswith("/"):
        media_url = f"/{media_url}"
    if not media_url.endswith("/"):
        media_url = f"{media_url}/"
    return f"{media_url}{quote(relative.as_posix(), safe='/')}"


def _community_source_publish_info(source_payload: dict[str, object] | None) -> dict[str, object]:
    if not source_payload:
        return {
            "connected": False,
            "title": "",
            "source_label": "Manual upload",
            "media_ready": False,
            "cover_ready": False,
            "media_url": "",
            "cover_url": "",
            "kind": "",
            "steps": [
                {"label": "Choose category", "state": "active"},
                {"label": "Upload media or text", "state": "pending"},
                {"label": "Add preview", "state": "pending"},
                {"label": "Publish from account", "state": "pending"},
            ],
        }
    source_labels = {
        "job": "Workspace task",
        "video_project": "Video editor project",
        "design_project": "Design project",
        "music_project": "Music project",
    }
    media_path = str(source_payload.get("media_path") or "")
    cover_path = str(source_payload.get("cover_path") or "")
    media_ready = bool(media_path)
    cover_ready = bool(cover_path)
    source = str(source_payload.get("source") or "")
    source_id = str(source_payload.get("source_id") or "")
    source_preview_url = f"{reverse('studio:community_publish_source_preview')}?{urlencode({'source': source, 'id': source_id})}" if source and source_id and (cover_ready or media_ready) else ""
    return {
        "connected": True,
        "title": str(source_payload.get("title") or ""),
        "source_label": source_labels.get(str(source_payload.get("source") or ""), "Connected source"),
        "media_ready": media_ready,
        "cover_ready": cover_ready,
        "media_url": _media_url_for_path(media_path),
        "cover_url": _media_url_for_path(cover_path) or source_preview_url,
        "source_preview_url": source_preview_url,
        "kind": str(source_payload.get("kind") or ""),
        "steps": [
            {"label": "Account selected", "state": "done"},
            {"label": "Source attached", "state": "done"},
            {"label": "Media will copy automatically" if media_ready else "Upload main media", "state": "done" if media_ready else "active"},
            {"label": "Preview ready" if cover_ready else "Upload or generate preview", "state": "done" if cover_ready else "active"},
            {"label": "Confirm rights and publish", "state": "pending"},
        ],
    }


def _copy_source_path_to_field(work: CommunityWork, field_name: str, source_path: str, _folder: str) -> None:
    path = Path(source_path or "")
    if not path.exists() or not path.is_file():
        return
    safe_name = f"{uuid.uuid4().hex[:12]}-{path.name[:80]}"
    field = getattr(work, field_name)
    with path.open("rb") as source_file:
        field.save(safe_name, File(source_file), save=False)


def _apply_community_source(work: CommunityWork, source_payload: dict[str, object] | None) -> None:
    if not source_payload:
        return
    if source_payload.get("source_job"):
        work.source_job = source_payload["source_job"]
    if source_payload.get("source_video_project"):
        work.source_video_project = source_payload["source_video_project"]
    if source_payload.get("source_design_project"):
        work.source_design_project = source_payload["source_design_project"]
    if source_payload.get("source_music_project"):
        work.source_music_project = source_payload["source_music_project"]
    if not work.media_file and source_payload.get("media_path"):
        _copy_source_path_to_field(work, "media_file", str(source_payload["media_path"]), "community/private")
    if not work.cover_image and source_payload.get("cover_path"):
        _copy_source_path_to_field(work, "cover_image", str(source_payload["cover_path"]), "community/source")


def _community_cover_bytes_from_image(file_field) -> bytes:
    source = getattr(file_field, "path", None) or file_field
    try:
        if not source:
            return b""
        if source is file_field:
            file_field.open("rb")
        with Image.open(source) as image:
            image = ImageOps.exif_transpose(image).convert("RGB")
            image.thumbnail((1400, 1400), Image.Resampling.LANCZOS)
            output = BytesIO()
            image.save(output, format="WEBP", quality=86, method=6)
            return output.getvalue()
    finally:
        try:
            if source is file_field:
                file_field.close()
        except Exception:
            pass


def _community_cover_bytes_from_video(path: Path) -> bytes:
    if not path.exists() or not path.is_file():
        return b""
    frame_dir = settings.storage_dir / "community_preview_frames"
    frame_dir.mkdir(parents=True, exist_ok=True)
    frame_path = frame_dir / f"{uuid.uuid4().hex}.jpg"
    attempts = (
        ["-ss", "0.5", "-i", str(path)],
        ["-i", str(path)],
    )
    try:
        for input_args in attempts:
            completed = subprocess.run(
                [
                    str(ffmpeg_path()),
                    "-y",
                    *input_args,
                    "-frames:v",
                    "1",
                    "-vf",
                    "thumbnail,scale=1280:-2:force_original_aspect_ratio=decrease",
                    str(frame_path),
                ],
                capture_output=True,
                text=True,
                timeout=min(30, int(getattr(settings, "video_timeout_seconds", 30) or 30)),
            )
            if completed.returncode == 0 and frame_path.exists() and frame_path.stat().st_size:
                with Image.open(frame_path) as image:
                    image = ImageOps.exif_transpose(image).convert("RGB")
                    output = BytesIO()
                    image.save(output, format="WEBP", quality=84, method=6)
                    return output.getvalue()
        return b""
    finally:
        frame_path.unlink(missing_ok=True)


def _community_placeholder_preview(kind: str, title: str) -> bytes:
    palette = {
        CommunityWork.KIND_MUSIC: ((26, 20, 62), (124, 58, 237), "MUSIC"),
        CommunityWork.KIND_TEXT: ((120, 53, 15), (245, 158, 11), "TEXT"),
        CommunityWork.KIND_VIDEO: ((15, 23, 42), (37, 99, 235), "VIDEO"),
        CommunityWork.KIND_IMAGE: ((20, 83, 45), (22, 163, 74), "IMAGE"),
    }
    bg, accent, label = palette.get(kind, palette[CommunityWork.KIND_VIDEO])
    image = Image.new("RGB", (1280, 720), bg)
    overlay = Image.new("RGBA", image.size, (255, 255, 255, 0))
    draw = ImageDraw.Draw(overlay)
    draw.ellipse((820, -120, 1420, 480), fill=(*accent, 96))
    draw.ellipse((-160, 400, 420, 980), fill=(255, 255, 255, 34))
    image = Image.alpha_composite(image.convert("RGBA"), overlay)
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arial.ttf", 62)
        label_font = ImageFont.truetype("arial.ttf", 28)
    except OSError:
        title_font = ImageFont.load_default()
        label_font = ImageFont.load_default()
    draw.rounded_rectangle((70, 64, 260, 112), radius=24, fill=(*accent, 235))
    draw.text((92, 76), label, fill=(255, 255, 255, 255), font=label_font)
    words = str(title or "CherryX preview").split()
    lines: list[str] = []
    current = ""
    for word in words:
        probe = f"{current} {word}".strip()
        if len(probe) > 28 and current:
            lines.append(current)
            current = word
        else:
            current = probe
    if current:
        lines.append(current)
    y = 250
    for line in lines[:4]:
        draw.text((76, y), line, fill=(255, 255, 255, 244), font=title_font)
        y += 70
    output = BytesIO()
    image.convert("RGB").save(output, format="WEBP", quality=86, method=6)
    return output.getvalue()


def _community_protected_preview_dir() -> Path:
    path = settings.storage_dir / "community_protected_previews"
    path.mkdir(parents=True, exist_ok=True)
    return path


def _community_protected_video_path(work: CommunityWork) -> Path:
    source = Path(work.media_file.path)
    output = _community_protected_preview_dir() / f"{work.slug}-{int(source.stat().st_mtime)}-preview.mp4"
    if output.exists() and output.stat().st_size:
        return output
    vf = (
        "scale='min(720,iw)':-2:force_original_aspect_ratio=decrease,"
        "pad=720:1280:(ow-iw)/2:(oh-ih)/2:color=#0f172a,"
        "drawtext=text='CHERRYX PREVIEW':x=(w-text_w)/2:y=(h-text_h)/2:"
        "fontsize=54:fontcolor=white@0.36:box=1:boxcolor=black@0.18:boxborderw=18,"
        "drawtext=text='Unlock clean original after payment':x=34:y=h-92:"
        "fontsize=28:fontcolor=white@0.92:box=1:boxcolor=black@0.45:boxborderw=12"
    )
    completed = subprocess.run(
        [
            str(ffmpeg_path()),
            "-y",
            "-t",
            "24",
            "-i",
            str(source),
            "-vf",
            vf,
            "-c:v",
            "libx264",
            "-preset",
            "veryfast",
            "-crf",
            "30",
            "-c:a",
            "aac",
            "-b:a",
            "64k",
            "-movflags",
            "+faststart",
            str(output),
        ],
        capture_output=True,
        text=True,
        timeout=min(60, int(getattr(settings, "video_timeout_seconds", 60) or 60)),
    )
    if completed.returncode != 0 or not output.exists() or not output.stat().st_size:
        output.unlink(missing_ok=True)
        raise Http404("Protected preview unavailable")
    return output


def _community_protected_audio_path(work: CommunityWork) -> Path:
    source = Path(work.media_file.path)
    output = _community_protected_preview_dir() / f"{work.slug}-{int(source.stat().st_mtime)}-preview.mp3"
    if output.exists() and output.stat().st_size:
        return output
    completed = subprocess.run(
        [
            str(ffmpeg_path()),
            "-y",
            "-t",
            "30",
            "-i",
            str(source),
            "-af",
            "volume=0.72,afade=t=out:st=27:d=3",
            "-vn",
            "-c:a",
            "libmp3lame",
            "-b:a",
            "96k",
            str(output),
        ],
        capture_output=True,
        text=True,
        timeout=45,
    )
    if completed.returncode != 0 or not output.exists() or not output.stat().st_size:
        output.unlink(missing_ok=True)
        raise Http404("Protected preview unavailable")
    return output


def _ensure_community_cover_from_media(work: CommunityWork) -> bool:
    if work.cover_image or not work.media_file:
        return False
    try:
        if work.kind == CommunityWork.KIND_IMAGE:
            payload = _community_cover_bytes_from_image(work.media_file)
        elif work.kind == CommunityWork.KIND_VIDEO:
            payload = _community_cover_bytes_from_video(Path(work.media_file.path))
        else:
            payload = b""
    except Exception:
        payload = b""
    if not payload:
        return False
    safe_slug = work.slug or f"work-{work.pk or uuid.uuid4().hex[:8]}"
    work.cover_image.save(f"community/source/{safe_slug}-auto-preview.webp", ContentFile(payload), save=False)
    work.save(update_fields=["cover_image", "updated_at"])
    return True


def _community_work_can_access(work: CommunityWork, user) -> bool:
    if not work.is_paid:
        return True
    if user and getattr(user, "is_authenticated", False):
        if work.owner_id == user.id:
            return True
        return CommunityPurchase.objects.filter(work=work, buyer=user).exists()
    return False


def _community_work_has_purchase_access(work: CommunityWork, user) -> bool:
    if not work.is_paid:
        return True
    if user and getattr(user, "is_authenticated", False):
        return CommunityPurchase.objects.filter(work=work, buyer=user).exists()
    return False


def _community_platform_fee(price: int) -> int:
    return max(1, math.ceil(max(0, int(price or 0)) * 0.01)) if int(price or 0) > 0 else 0


def _community_purchase_total(price: int) -> int:
    base = max(0, int(price or 0))
    return base + _community_platform_fee(base)


def _community_preview_field(work: CommunityWork):
    if work.cover_image:
        return work.cover_image
    if work.kind == CommunityWork.KIND_IMAGE and work.media_file:
        return work.media_file
    return None


def _watermark_community_image(file_field, slug: str) -> bytes:
    try:
        file_field.open("rb")
        with Image.open(file_field) as image:
            image = ImageOps.exif_transpose(image).convert("RGBA")
    except FileNotFoundError as exc:
        raise Http404("Preview not found") from exc
    finally:
        try:
            file_field.close()
        except Exception:
            pass
    image.thumbnail((1400, 1400), Image.Resampling.LANCZOS)
    overlay = Image.new("RGBA", image.size, (255, 255, 255, 0))
    draw = ImageDraw.Draw(overlay)
    base_size = max(28, min(image.size) // 7)
    try:
        font = ImageFont.truetype("arial.ttf", base_size)
    except OSError:
        font = ImageFont.load_default()
    rng = random.Random(f"{slug}:{time.time_ns()}")
    marks = max(10, (image.width * image.height) // 95_000)
    for index in range(marks):
        text = "CX"
        x = rng.randint(-base_size, max(1, image.width - base_size))
        y = rng.randint(-base_size, max(1, image.height - base_size))
        alpha = rng.randint(62, 112)
        color = (255, 255, 255, alpha) if index % 2 else (37, 99, 235, alpha)
        draw.text((x, y), text, fill=color, font=font)
    ribbon = Image.new("RGBA", (image.width, max(42, image.height // 11)), (15, 23, 42, 118))
    image.alpha_composite(ribbon, (0, max(0, image.height - ribbon.height)))
    draw = ImageDraw.Draw(image)
    draw.text((24, max(8, image.height - ribbon.height + 12)), "CX preview - unlock original with CherryX", fill=(255, 255, 255, 220), font=font)
    image = Image.alpha_composite(image, overlay).convert("RGB")
    output = BytesIO()
    image.save(output, format="WEBP", quality=78, method=6)
    return output.getvalue()


@require_GET
def community_feed(request: HttpRequest, kind: str):
    if kind not in {CommunityWork.KIND_VIDEO, CommunityWork.KIND_IMAGE, CommunityWork.KIND_TEXT, CommunityWork.KIND_MUSIC}:
        raise Http404("Unknown community feed")
    meta = _community_kind_meta(kind)
    works, filters = _community_queryset(kind, request)
    return render(
        request,
        "studio/community_feed.html",
        {
            "kind": kind,
            "meta": meta,
            "works": works,
            "filters": filters,
            "cherryx_usd_rate": _community_usd_rate_per_cherryx(),
            "preview_route_name": "studio:community_work_preview",
            "seo_title": meta["title"],
            "seo_description": meta["description"],
        },
    )


@require_GET
def community_work_detail(request: HttpRequest, slug: str):
    work = get_object_or_404(CommunityWork.objects.select_related("owner"), slug=slug, status=CommunityWork.STATUS_PUBLISHED)
    CommunityWork.objects.filter(pk=work.pk).update(view_count=F("view_count") + 1, updated_at=timezone.now())
    work.view_count = int(work.view_count or 0) + 1
    has_purchase = bool(request.user.is_authenticated and CommunityPurchase.objects.filter(work=work, buyer=request.user).exists())
    is_owner = bool(request.user.is_authenticated and work.owner_id == request.user.id)
    can_access = _community_work_can_access(work, request.user)
    can_view_original = _community_work_has_purchase_access(work, request.user)
    platform_fee = _community_platform_fee(work.price_cherryx) if work.is_paid else 0
    purchase_total = _community_purchase_total(work.price_cherryx) if work.is_paid else 0
    same_kind = CommunityWork.objects.filter(status=CommunityWork.STATUS_PUBLISHED, kind=work.kind).exclude(pk=work.pk).select_related("owner")[:10]
    fallback = CommunityWork.objects.filter(status=CommunityWork.STATUS_PUBLISHED).exclude(pk=work.pk).exclude(kind=work.kind).select_related("owner")[:10]
    related_works = list(same_kind)
    seen_related = {item.pk for item in related_works}
    for item in fallback:
        if item.pk not in seen_related:
            related_works.append(item)
        if len(related_works) >= 14:
            break
    return render(
        request,
        "studio/community_detail.html",
        {
            "work": work,
            "related_works": related_works,
            "can_access": can_access,
            "can_view_original": can_view_original,
            "is_owner": is_owner,
            "has_purchase": has_purchase,
            "balance_cherryx": _cherryx_balance(request),
            "platform_fee": platform_fee,
            "purchase_total": purchase_total,
            "protected_preview_url": reverse("studio:community_work_protected_preview", args=[work.slug]) if work.is_paid and work.media_file else "",
            "download_url": reverse("studio:community_work_download", args=[work.slug]) if work.media_file else "",
            "preview_url": reverse("studio:community_work_preview", args=[work.slug]) if _community_preview_field(work) else "",
            "seo_title": work.title,
            "seo_description": work.excerpt or strip_tags(work.body)[:240] or "CherryX public community work.",
        },
    )


@login_required
@require_http_methods(["GET", "POST"])
def community_publish(request: HttpRequest):
    source = request.POST.get("source") if request.method == "POST" else request.GET.get("source")
    source_id = request.POST.get("source_id") if request.method == "POST" else request.GET.get("id")
    source_payload = _community_source_for_user(request.user, source or "", source_id or "")
    source_has_media = bool(source_payload and (source_payload.get("media_path") or source_payload.get("cover_path") or source_payload.get("source_music_project")))
    initial = {}
    if source_payload:
        initial = {
            "kind": source_payload.get("kind") or CommunityWork.KIND_TEXT,
            "title": source_payload.get("title") or "",
            "excerpt": source_payload.get("excerpt") or "",
            "body": source_payload.get("body") or "",
            "access": CommunityWork.ACCESS_FREE,
            "price_cherryx": 0,
        }
    form = CommunityWorkForm(request.POST or None, request.FILES or None, initial=initial, source_has_media=source_has_media)
    if request.method == "POST" and form.is_valid():
        work = form.save(commit=False)
        work.owner = request.user
        work.status = CommunityWork.STATUS_PUBLISHED
        _apply_community_source(work, source_payload)
        work.save()
        _ensure_community_cover_from_media(work)
        return redirect("studio:community_work_detail", slug=work.slug)
    publish_account = {
        "name": request.user.get_full_name() or request.user.get_username() or request.user.email,
        "email": request.user.email,
        "avatar_url": _avatar_url(request),
    }
    return render(
        request,
        "studio/community_publish.html",
        {
            "form": form,
            "source_payload": source_payload,
            "source_publish_info": _community_source_publish_info(source_payload),
            "publish_account": publish_account,
            "cherryx_usd_rate": _community_usd_rate_per_cherryx(),
            "telegram_stars_rate": telegram_stars_rate(),
            "cherryx_unit_usd": _community_usd_display(1),
            "seo_title": "Publish to CherryX network",
            "seo_description": "Share a free or paid CherryX video, image, music or text work with the public community feed.",
        },
    )


@login_required
@require_GET
def community_my_works(request: HttpRequest):
    works = (
        CommunityWork.objects.filter(owner=request.user)
        .annotate(total_purchases=Count("purchases"), total_revenue=Sum("purchases__price_cherryx"))
        .order_by("-updated_at", "-created_at")
    )
    items = [_community_manage_payload(work, request) for work in works]
    totals = {
        "works": len(items),
        "published": sum(1 for item in items if item["is_published"]),
        "hidden": sum(1 for item in items if not item["is_published"]),
        "views": sum(int(item["work"].view_count or 0) for item in items),
        "purchases": sum(int(item["purchases"] or 0) for item in items),
        "revenue": sum(int(item["revenue"] or 0) for item in items),
        "downloads": sum(int(item["work"].download_count or 0) for item in items),
    }
    return render(
        request,
        "studio/community_my_works.html",
        {
            "items": items,
            "totals": totals,
            "seo_title": "My community listings",
            "seo_description": "Manage your CherryX community publications, stats, links and visibility.",
        },
    )


@login_required
@require_POST
def community_work_visibility(request: HttpRequest, slug: str):
    work = get_object_or_404(CommunityWork, slug=slug, owner=request.user)
    action = request.POST.get("action") or ""
    if action == "hide":
        work.status = CommunityWork.STATUS_DRAFT
        work.save(update_fields=["status", "updated_at"])
        message = "Listing is hidden from public feeds."
    elif action == "publish":
        work.status = CommunityWork.STATUS_PUBLISHED
        work.save(update_fields=["status", "published_at", "updated_at"])
        message = "Listing is public again."
    else:
        return JsonResponse({"ok": False, "message": "Unknown visibility action."}, status=400)
    payload = _community_manage_payload(work, request)
    return JsonResponse(
        {
            "ok": True,
            "message": message,
            "status": work.status,
            "is_published": payload["is_published"],
            "visibility_action": payload["visibility_action"],
            "visibility_label": payload["visibility_label"],
            "public_url": payload["public_url"],
        }
    )


@login_required
@require_POST
def community_work_delete(request: HttpRequest, slug: str):
    work = get_object_or_404(CommunityWork, slug=slug, owner=request.user)
    work_id = work.pk
    with transaction.atomic():
        _delete_community_work_files(work)
        work.delete()
    return JsonResponse({"ok": True, "id": work_id, "message": "Listing was deleted from the system and database."})


@login_required
@require_GET
def community_publish_source_preview(request: HttpRequest):
    source_payload = _community_source_for_user(request.user, request.GET.get("source") or "", request.GET.get("id") or "")
    if not source_payload:
        raise Http404("Source not found")
    kind = str(source_payload.get("kind") or "")
    title = str(source_payload.get("title") or "CherryX preview")
    cover_path = Path(str(source_payload.get("cover_path") or ""))
    media_path = Path(str(source_payload.get("media_path") or ""))
    if cover_path.exists() and cover_path.is_file():
        content_type = mimetypes.guess_type(cover_path.name)[0] or "image/webp"
        response = FileResponse(cover_path.open("rb"), as_attachment=False, filename=cover_path.name, content_type=content_type)
        response["Cache-Control"] = "private, max-age=120"
        response["X-Content-Type-Options"] = "nosniff"
        return response
    if kind == CommunityWork.KIND_IMAGE and media_path.exists() and media_path.is_file():
        content_type = mimetypes.guess_type(media_path.name)[0] or "image/webp"
        response = FileResponse(media_path.open("rb"), as_attachment=False, filename=media_path.name, content_type=content_type)
        response["Cache-Control"] = "private, max-age=120"
        response["X-Content-Type-Options"] = "nosniff"
        return response
    if kind == CommunityWork.KIND_VIDEO and media_path.exists() and media_path.is_file():
        payload = _community_cover_bytes_from_video(media_path)
        if payload:
            response = HttpResponse(payload, content_type="image/webp")
            response["Cache-Control"] = "private, max-age=120"
            response["X-Content-Type-Options"] = "nosniff"
            return response
    payload = _community_placeholder_preview(kind, title)
    response = HttpResponse(payload, content_type="image/webp")
    response["Cache-Control"] = "private, max-age=120"
    response["X-Content-Type-Options"] = "nosniff"
    return response


@login_required
@require_POST
def community_purchase(request: HttpRequest, slug: str):
    work = get_object_or_404(CommunityWork.objects.select_related("owner"), slug=slug, status=CommunityWork.STATUS_PUBLISHED)
    if not work.is_paid or work.owner_id == request.user.id:
        return redirect("studio:community_work_detail", slug=work.slug)
    with transaction.atomic():
        work = CommunityWork.objects.select_for_update().select_related("owner").get(pk=work.pk)
        purchase, created = CommunityPurchase.objects.get_or_create(
            work=work,
            buyer=request.user,
            defaults={"seller": work.owner, "price_cherryx": work.price_cherryx},
        )
        if created:
            purchase_total = _community_purchase_total(work.price_cherryx)
            buyer_profile, _ = AccountProfile.objects.select_for_update().get_or_create(user=request.user)
            if int(buyer_profile.cherryx_balance or 0) < purchase_total:
                purchase.delete()
                return redirect(f"{reverse('studio:community_work_detail', args=[work.slug])}?payment=insufficient")
            buyer_profile.cherryx_balance = int(buyer_profile.cherryx_balance or 0) - purchase_total
            buyer_profile.save(update_fields=["cherryx_balance", "updated_at"])
            if work.owner_id:
                seller_profile, _ = AccountProfile.objects.select_for_update().get_or_create(user=work.owner)
                seller_profile.cherryx_balance = int(seller_profile.cherryx_balance or 0) + work.price_cherryx
                seller_profile.save(update_fields=["cherryx_balance", "updated_at"])
            work.purchase_count = CommunityPurchase.objects.filter(work=work).count()
            work.save(update_fields=["purchase_count", "updated_at"])
    return redirect("studio:community_work_detail", slug=work.slug)


def _clone_music_project_for_user(source: MusicEditorProject, user, title: str) -> MusicEditorProject:
    target = MusicEditorProject.objects.create(
        owner=user,
        guest_key="",
        title=_clean_project_title(f"{title} - CherryX copy"),
        state_json={},
        clip_count=source.clip_count,
        duration_seconds=source.duration_seconds,
        last_export_status=source.last_export_status,
    )
    target_dir = _music_project_media_dir(target)
    target_dir.mkdir(parents=True, exist_ok=True)
    asset_id_map: dict[str, int] = {}
    for asset in source.assets.all():
        source_path = Path(asset.file_path)
        if not source_path.exists() or not source_path.is_file():
            continue
        target_path = target_dir / f"{uuid.uuid4().hex[:12]}_{source_path.name[:96]}"
        shutil.copy2(source_path, target_path)
        new_asset = MusicEditorAsset.objects.create(
            project=target,
            kind=asset.kind,
            file_path=str(target_path),
            media_type=asset.media_type,
            size=target_path.stat().st_size,
            original_name=asset.original_name,
            duration=asset.duration,
        )
        asset_id_map[str(asset.id)] = new_asset.id
    state = json.loads(json.dumps(source.state_json or {}))
    if isinstance(state, dict):
        for asset_state in state.get("assets", []) if isinstance(state.get("assets"), list) else []:
            old_id = str(asset_state.get("serverId") or asset_state.get("id") or "")
            new_id = asset_id_map.get(old_id)
            if new_id:
                asset_state["serverId"] = new_id
                asset_state["url"] = reverse("studio:music_project_asset_preview", args=[target.id, new_id])
                if str(asset_state.get("id") or "").isdigit():
                    asset_state["id"] = new_id
        for clip in state.get("clips", []) if isinstance(state.get("clips"), list) else []:
            old_id = str(clip.get("assetId") or "")
            new_id = asset_id_map.get(old_id)
            if new_id:
                clip["assetId"] = new_id
    target.state_json = state if isinstance(state, dict) else {}
    _update_music_project_metadata(target)
    target.storage_bytes = _music_project_storage_bytes(target)
    target.save(update_fields=["state_json", "storage_bytes", "asset_count", "clip_count", "duration_seconds", "last_export_status", "updated_at"])
    return target


@login_required
@require_GET
def community_open_music_project(request: HttpRequest, slug: str):
    work = get_object_or_404(
        CommunityWork.objects.select_related("source_music_project").prefetch_related("source_music_project__assets"),
        slug=slug,
        status=CommunityWork.STATUS_PUBLISHED,
        kind=CommunityWork.KIND_MUSIC,
    )
    if not work.source_music_project:
        raise Http404("Music project not found")
    if not _community_work_can_access(work, request.user):
        return redirect(f"{reverse('studio:community_work_detail', args=[work.slug])}?payment=required")
    source = work.source_music_project
    if source.owner_id == request.user.id:
        return redirect(f"{reverse('studio:music_editor')}?{urlencode({'project': source.id})}")
    target = _clone_music_project_for_user(source, request.user, work.title)
    return redirect(f"{reverse('studio:music_editor')}?{urlencode({'project': target.id})}")


@require_GET
def community_work_download(request: HttpRequest, slug: str):
    work = get_object_or_404(CommunityWork, slug=slug, status=CommunityWork.STATUS_PUBLISHED)
    if not work.media_file:
        raise Http404("File not found")
    if work.is_paid and not _community_work_has_purchase_access(work, request.user):
        if not request.user.is_authenticated:
            return redirect(f"{reverse('studio:login')}?{urlencode({'next': request.path})}")
        return redirect(f"{reverse('studio:community_work_detail', args=[work.slug])}?payment=required")
    try:
        file_handle = work.media_file.open("rb")
    except FileNotFoundError as exc:
        raise Http404("File not found") from exc
    CommunityWork.objects.filter(pk=work.pk).update(download_count=F("download_count") + 1, updated_at=timezone.now())
    filename = Path(work.media_file.name).name or f"{work.slug}.bin"
    content_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"
    response = FileResponse(file_handle, content_type=content_type)
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    response["X-Content-Type-Options"] = "nosniff"
    return response


@require_GET
def community_work_preview(request: HttpRequest, slug: str):
    work = get_object_or_404(CommunityWork, slug=slug, status=CommunityWork.STATUS_PUBLISHED)
    field = _community_preview_field(work)
    if not field:
        raise Http404("Preview not found")
    if work.is_paid and not _community_work_has_purchase_access(work, request.user):
        payload = _watermark_community_image(field, work.slug)
        response = HttpResponse(payload, content_type="image/webp")
        response["Cache-Control"] = "no-store, private"
        response["X-Content-Type-Options"] = "nosniff"
        return response
    try:
        file_handle = field.open("rb")
    except FileNotFoundError as exc:
        raise Http404("Preview not found") from exc
    content_type = mimetypes.guess_type(field.name)[0] or "application/octet-stream"
    response = FileResponse(file_handle, content_type=content_type)
    response["Cache-Control"] = "private, max-age=300"
    response["X-Content-Type-Options"] = "nosniff"
    return response


@require_GET
def community_work_protected_preview(request: HttpRequest, slug: str):
    work = get_object_or_404(CommunityWork, slug=slug, status=CommunityWork.STATUS_PUBLISHED)
    if not work.is_paid:
        if work.media_file:
            return redirect("studio:community_work_download", slug=work.slug)
        return redirect("studio:community_work_preview", slug=work.slug)
    if _community_work_has_purchase_access(work, request.user):
        if work.media_file:
            return redirect("studio:community_work_download", slug=work.slug)
        return redirect("studio:community_work_preview", slug=work.slug)
    if work.kind == CommunityWork.KIND_VIDEO and work.media_file:
        path = _community_protected_video_path(work)
        response = _range_file_response(request, path, "video/mp4", path.name)
        response["Cache-Control"] = "no-store, private"
        response["X-Content-Type-Options"] = "nosniff"
        return response
    if work.kind == CommunityWork.KIND_MUSIC and work.media_file:
        path = _community_protected_audio_path(work)
        response = FileResponse(path.open("rb"), as_attachment=False, filename=path.name, content_type="audio/mpeg")
        response["Cache-Control"] = "no-store, private"
        response["X-Content-Type-Options"] = "nosniff"
        return response
    return redirect("studio:community_work_preview", slug=work.slug)


@require_GET
def favicon(request: HttpRequest) -> HttpResponse:
    svg = (
        '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 64 64">'
        '<rect width="64" height="64" rx="16" fill="#0f172a"/>'
        '<path d="M18 38c0 5 4 9 10 9 9 0 18-8 18-22V15h-8v10c0 7-4 14-10 14-2 0-3-1-3-3 0-3 3-5 8-5h5v-7h-6c-9 0-14 6-14 14z" fill="#ff4d6d"/>'
        '<path d="M40 15h10v10H40z" fill="#22c55e"/>'
        "</svg>"
    )
    response = HttpResponse(svg, content_type="image/svg+xml")
    response["Cache-Control"] = "public, max-age=86400"
    return response


@require_http_methods(["GET", "POST"])
def register(request: HttpRequest):
    if request.user.is_authenticated:
        return redirect(_safe_next_url(request) or reverse("studio:index"))
    form = RegisterForm(request.POST or None, language=getattr(request, "interface_language", "en"))
    if request.method == "POST" and form.is_valid():
        guest_key = _guest_key(request)
        user = form.save()
        login(request, user)
        transfer_guest_workspace(guest_key, user)
        _transfer_guest_video_projects(guest_key, user)
        _transfer_guest_design_projects(guest_key, user)
        _accept_pending_workspace_shares(user)
        return redirect(_safe_next_url(request) or reverse("studio:index"))
    return render(request, "studio/auth.html", {**_auth_context(request), "form": form, "mode": "register"})


@require_http_methods(["GET", "POST"])
def login_view(request: HttpRequest):
    if request.user.is_authenticated:
        return redirect(_safe_next_url(request) or reverse("studio:index"))
    form = EmailLoginForm(request.POST or None, language=getattr(request, "interface_language", "en"))
    if request.method == "POST" and form.is_valid():
        guest_key = _guest_key(request)
        login(request, form.cleaned_data["user"])
        transfer_guest_workspace(guest_key, form.cleaned_data["user"])
        _transfer_guest_video_projects(guest_key, form.cleaned_data["user"])
        _transfer_guest_design_projects(guest_key, form.cleaned_data["user"])
        _accept_pending_workspace_shares(form.cleaned_data["user"])
        return redirect(_safe_next_url(request) or reverse("studio:index"))
    return render(request, "studio/auth.html", {**_auth_context(request), "form": form, "mode": "login"})


@require_GET
def magic_login(request: HttpRequest, token: str):
    guest_key = _guest_key(request)
    with transaction.atomic():
        magic_token = MagicLoginToken.objects.select_for_update().select_related("user").filter(token=(token or "").strip()).first()
        if not magic_token or not magic_token.is_usable:
            return redirect("studio:login")
        user = magic_token.user
        magic_token.used_at = timezone.now()
        magic_token.save(update_fields=["used_at"])
    login(request, user)
    transfer_guest_workspace(guest_key, user)
    _transfer_guest_video_projects(guest_key, user)
    _transfer_guest_design_projects(guest_key, user)
    _accept_pending_workspace_shares(user)
    return redirect(_safe_next_url(request) or reverse("studio:index"))


@require_POST
def logout_view(request: HttpRequest):
    logout(request)
    return redirect("studio:landing")


@require_POST
def set_interface_language(request: HttpRequest):
    language = clean_language(request.POST.get("language"))
    request.session["interface_language"] = language
    request.interface_language = language
    if request.user.is_authenticated:
        profile, _ = AccountProfile.objects.get_or_create(user=request.user)
        profile.interface_language = language
        profile.save(update_fields=["interface_language", "updated_at"])
    response = redirect(request.POST.get("next") or request.META.get("HTTP_REFERER") or reverse("studio:landing"))
    response.set_cookie("interface_language", language, max_age=60 * 60 * 24 * 365, samesite="Lax")
    return response


@login_required
@require_http_methods(["GET", "POST"])
def account_settings(request: HttpRequest):
    profile, _ = AccountProfile.objects.get_or_create(user=request.user)
    form = AccountSettingsForm(
        request.POST or None,
        request.FILES or None,
        user=request.user,
        profile=profile,
        language=getattr(request, "interface_language", "en"),
    )
    if request.method == "POST" and form.is_valid():
        password_changed = form.save_profile()
        avatar_crop_data = form.cleaned_data.get("avatar_crop_data")
        avatar = form.cleaned_data.get("avatar_file")
        if avatar_crop_data:
            profile.avatar_path = str(_save_avatar_crop(avatar_crop_data, request.user.id))
            profile.avatar_url = ""
            profile.save(update_fields=["avatar_path", "avatar_url", "updated_at"])
        elif avatar:
            profile.avatar_path = str(_save_avatar_upload(avatar, request.user.id))
            profile.avatar_url = ""
            profile.save(update_fields=["avatar_path", "avatar_url", "updated_at"])
        if password_changed:
            update_session_auth_hash(request, request.user)
        return redirect("studio:account_settings")
    return render(
        request,
        "studio/account_settings.html",
        {
            "form": form,
            "avatar_url": _avatar_url(request),
            "display_name": _display_name(request),
            "accent_color": _accent_color(request),
            "ui_accent_color": _ui_accent_color(request),
            "theme_mode": _theme_mode(request),
            "language_options": _language_options_for_form(form),
        },
    )


@login_required
@require_GET
def account_avatar(request: HttpRequest):
    try:
        avatar_path = request.user.studio_profile.avatar_path
    except AccountProfile.DoesNotExist:
        avatar_path = ""
    if not avatar_path:
        raise Http404("Avatar not found")
    path = Path(avatar_path)
    if not path.exists() or not path.is_file():
        raise Http404("Avatar not found")
    return FileResponse(path.open("rb"), as_attachment=False)


@require_GET
def index(request: HttpRequest):
    owner_id, guest_key = _workspace_identity(request)
    initial_jobs = _attach_output_urls_many(actions.get_recent_jobs(5, owner_id, guest_key))
    account_stats = actions.get_account_stats(owner_id, guest_key)
    account_stats.update(_storage_quota(request, account_stats))
    native_status = actions.native_status()
    has_access = user_has_active_access(request.user)
    language = getattr(request, "interface_language", "en")
    return render(
        request,
        "studio/index.html",
        {
            "image_formats": actions.IMAGE_FORMAT_CHOICES,
            "video_formats": actions.VIDEO_FORMAT_CHOICES,
            "image_modes": _localized_image_modes(language),
            "youtube_modes": _localized_youtube_modes(language),
            "subtitle_styles": _localized_subtitle_styles(request),
            "subtitle_languages": _localized_subtitle_languages(language),
            "resume_templates": _localized_resume_templates(language),
            "max_image_mb": settings.max_image_mb,
            "max_video_mb": settings.max_video_mb,
            "youtube_max_shorts": settings.youtube_max_shorts,
            "initial_jobs": initial_jobs,
            "account_stats": account_stats,
            "has_access": has_access,
            "openai_ready": actions.is_openai_ready(),
            "native_status": native_status,
            "active_until": active_access_until(request.user),
            "is_guest": not request.user.is_authenticated,
            "display_name": _display_name(request),
            "checkout_url": _checkout_url(request),
            "pricing_url": reverse("billing:pricing"),
            "login_url": reverse("studio:login"),
            "settings_url": reverse("studio:account_settings"),
            "designer_url": reverse("studio:designer"),
            "design_projects_url": reverse("studio:design_project_list"),
            "cherryx_pay_url": reverse("studio:cherryx_pay"),
            "video_editor_url": reverse("studio:video_project_list"),
            "music_projects_url": reverse("studio:music_project_list"),
            "cherryx_balance": _cherryx_balance(request),
            "avatar_url": _avatar_url(request),
            "accent_color": _accent_color(request),
            "ui_accent_color": _ui_accent_color(request),
            "theme_mode": _theme_mode(request),
            "subscription_meter": _subscription_meter(request),
            "subscription_panel": _subscription_panel(request, account_stats, language),
            "app_messages": app_messages(language),
            "resume_prefill": _resume_prefill(request),
        },
    )


@require_GET
def subtitle_style_detail(request: HttpRequest, style: str):
    language = getattr(request, "interface_language", "en")
    detail = _subtitle_style_detail(style, language)
    if not detail:
        raise Http404("Subtitle style not found")
    back_url = request.GET.get("next") or f"{reverse('studio:index')}#tab-subtitles"
    if not url_has_allowed_host_and_scheme(back_url, allowed_hosts={request.get_host()}, require_https=request.is_secure()):
        back_url = f"{reverse('studio:index')}#tab-subtitles"
    return render(
        request,
        "studio/subtitle_style_detail.html",
        {
            "style_detail": detail,
            "back_url": back_url,
            "accent_color": _accent_color(request),
            "ui_accent_color": _ui_accent_color(request),
            "theme_mode": _theme_mode(request),
            "app_messages": app_messages(language),
        },
    )


@require_GET
@ensure_csrf_cookie
def designer_mode(request: HttpRequest):
    language = getattr(request, "interface_language", "en")
    owner_id, guest_key = _workspace_identity(request)
    project = None
    project_id = request.GET.get("project", "")
    if project_id.isdigit():
        project = _design_project_queryset(owner_id, guest_key).prefetch_related("assets").filter(id=int(project_id)).first()
    return render(
        request,
        "studio/designer.html",
        {
            "designer_url": reverse("studio:designer"),
            "design_projects_url": reverse("studio:design_project_list"),
            "design_projects_api_url": reverse("studio:design_projects"),
            "current_design_project": _design_project_payload(project, owner_id=owner_id, guest_key=guest_key) if project else None,
            "designer_fullscreen": True,
            "accent_color": _accent_color(request),
            "ui_accent_color": _ui_accent_color(request),
            "theme_mode": _theme_mode(request),
            "app_messages": app_messages(language),
            "subtitle_languages": _localized_subtitle_languages(language),
        },
    )


@require_GET
@ensure_csrf_cookie
def design_project_list(request: HttpRequest):
    owner_id, guest_key = _workspace_identity(request)
    listing = _project_listing_payload(
        request,
        _design_project_queryset(owner_id, guest_key).prefetch_related("assets"),
        WorkspaceShare.RESOURCE_DESIGN,
        owner_id,
        guest_key,
        lambda project: _design_project_payload(project, include_state=False, owner_id=owner_id, guest_key=guest_key),
        per_page_default=120,
    )
    return render(
        request,
        "studio/design_projects.html",
        {
            "design_projects": listing["projects"],
            "project_listing": listing,
            "designer_url": reverse("studio:designer"),
            "design_projects_api_url": reverse("studio:design_projects"),
            "accent_color": _accent_color(request),
            "ui_accent_color": _ui_accent_color(request),
            "theme_mode": _theme_mode(request),
            "app_messages": app_messages(getattr(request, "interface_language", "en")),
            "music_messages_json": json.dumps(music_messages(getattr(request, "interface_language", "en")), ensure_ascii=False),
        },
    )


@ensure_csrf_cookie
@require_GET
def cherryx_pay(request: HttpRequest):
    telegram_link_token = ensure_telegram_link_token(request.user) if request.user.is_authenticated else ""
    profile = getattr(request.user, "studio_profile", None) if request.user.is_authenticated else None
    language = getattr(request, "interface_language", "en")
    wallet_transactions = recent_wallet_transactions(request.user, 8) if request.user.is_authenticated else []
    return render(
        request,
        "studio/cherryx_pay.html",
        {
            "balance": _cherryx_balance(request),
            "withdrawal_available": _cherryx_balance(request),
            "usd_rate": 100,
            "telegram_stars_rate": telegram_stars_rate(),
            "telegram_link_token": telegram_link_token,
            "telegram_user_id": getattr(profile, "telegram_user_id", "") if profile else "",
            "telegram_username": getattr(profile, "telegram_username", "") if profile else "",
            "telegram_first_name": getattr(profile, "telegram_first_name", "") if profile else "",
            "wallet_transactions": wallet_transactions,
            "display_name": _display_name(request),
            "checkout_url": _checkout_url(request),
            "pricing_url": reverse("billing:pricing"),
            "plans": [localized_plan(plan, language) for plan in PLANS],
            "pay_text": cherryx_pay_messages(language),
            "accent_color": _accent_color(request),
            "ui_accent_color": _ui_accent_color(request),
            "theme_mode": _theme_mode(request),
            "app_messages": app_messages(getattr(request, "interface_language", "en")),
        },
    )


def _wallet_transaction_payload(transaction) -> dict[str, object]:
    related = transaction.related_user
    return {
        "id": transaction.id,
        "type": transaction.type,
        "label": transaction.get_type_display(),
        "amount": transaction.amount,
        "balance_after": transaction.balance_after,
        "status": transaction.status,
        "related": (related.email or related.username) if related else "",
        "created_at": transaction.created_at.isoformat(),
    }


@login_required
@require_POST
def cherryx_transfer(request: HttpRequest) -> JsonResponse:
    try:
        result = transfer_cherryx_by_email(
            request.user,
            request.POST.get("email", ""),
            request.POST.get("credits") or request.POST.get("amount") or "0",
        )
    except WalletError as exc:
        return JsonResponse({"ok": False, "error": exc.reason, "message": str(exc)}, status=400)
    return JsonResponse(
        {
            "ok": True,
            "balance": result["balance"],
            "amount": result["amount"],
            "recipient_email": result["recipient_email"],
            "transaction": _wallet_transaction_payload(result["transaction"]),
        }
    )


@login_required
@require_POST
def cherryx_withdrawal_request(request: HttpRequest) -> JsonResponse:
    try:
        result = create_cherryx_withdrawal_request(
            request.user,
            request.POST.get("credits") or request.POST.get("amount") or "0",
        )
    except WalletError as exc:
        return JsonResponse({"ok": False, "error": exc.reason, "message": str(exc)}, status=400)
    return JsonResponse(
        {
            "ok": True,
            "balance": result["balance"],
            "amount": result["amount"],
            "estimated_stars": result["estimated_stars"],
            "withdrawal_id": result["withdrawal"].id,
            "transaction": _wallet_transaction_payload(result["transaction"]),
        }
    )


@login_required
@require_POST
def telegram_link_token(request: HttpRequest) -> JsonResponse:
    profile, _ = AccountProfile.objects.get_or_create(user=request.user)
    profile.telegram_link_token = secrets.token_urlsafe(18)[:32]
    profile.telegram_link_token_created_at = timezone.now()
    profile.save(update_fields=["telegram_link_token", "telegram_link_token_created_at", "updated_at"])
    return JsonResponse({
        "token": profile.telegram_link_token,
        "telegram_user_id": profile.telegram_user_id,
    })


@require_GET
def design_projects(request: HttpRequest) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    listing = _project_listing_payload(
        request,
        _design_project_queryset(owner_id, guest_key).prefetch_related("assets"),
        WorkspaceShare.RESOURCE_DESIGN,
        owner_id,
        guest_key,
        lambda project: _design_project_payload(project, include_state=False, owner_id=owner_id, guest_key=guest_key),
    )
    return JsonResponse(listing)


@require_POST
def create_design_project(request: HttpRequest) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    data = _json_body(request)
    state = data.get("state") if isinstance(data.get("state"), dict) else {}
    title = _clean_project_title(str(data.get("title") or state.get("title") or "New design"))
    project = DesignerProject.objects.create(
        owner=request.user if request.user.is_authenticated else None,
        guest_key="" if owner_id else guest_key,
        title=title,
        state_json=state,
        storage_bytes=_json_size(state),
    )
    preview = str(data.get("preview") or "")
    if preview.startswith("data:image/"):
        if project.preview_path:
            try:
                old_preview = Path(project.preview_path)
                if old_preview.resolve().is_relative_to(settings.storage_dir.resolve()) and old_preview.exists() and old_preview.is_file():
                    old_preview.unlink()
            except Exception:
                pass
        project.preview_path = str(_save_design_project_preview(preview, _design_project_media_dir(project)))
    _update_design_project_metadata(project, assets=[])
    project.storage_bytes = _design_project_storage_bytes(project)
    project.save(update_fields=["preview_path", "storage_bytes", "asset_count", "object_count", "updated_at"])
    return JsonResponse({"project": _design_project_payload(project, owner_id=owner_id, guest_key=guest_key)})


@require_GET
def design_project_detail(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _design_project_queryset(owner_id, guest_key).prefetch_related("assets").filter(id=project_id).first()
    if not project:
        raise Http404("Design project not found")
    return JsonResponse({"project": _design_project_payload(project, owner_id=owner_id, guest_key=guest_key)})


@require_POST
def save_design_project(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _design_project_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Design project not found")
    if not _resource_can_edit(WorkspaceShare.RESOURCE_DESIGN, project, owner_id, guest_key):
        raise Http404("Design project not found")
    data = _json_body(request)
    state = data.get("state") if isinstance(data.get("state"), dict) else {}
    title = _clean_project_title(str(data.get("title") or state.get("title") or project.title))
    project.title = title
    project.state_json = state
    preview = str(data.get("preview") or "")
    if preview.startswith("data:image/"):
        if project.preview_path:
            try:
                old_preview = Path(project.preview_path)
                if old_preview.resolve().is_relative_to(settings.storage_dir.resolve()) and old_preview.exists() and old_preview.is_file():
                    old_preview.unlink()
            except Exception:
                pass
        project.preview_path = str(_save_design_project_preview(preview, _design_project_media_dir(project)))
    _update_design_project_metadata(project)
    project.storage_bytes = _design_project_storage_bytes(project)
    project.save(update_fields=["title", "state_json", "preview_path", "storage_bytes", "asset_count", "object_count", "updated_at"])
    return JsonResponse({"project": _design_project_payload(project, owner_id=owner_id, guest_key=guest_key)})


@require_POST
def rename_design_project(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _design_owner_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Design project not found")
    data = _json_body(request)
    project.title = _clean_project_title(str(data.get("title") or project.title))
    project.save(update_fields=["title", "updated_at"])
    return JsonResponse({"project": _design_project_payload(project, include_state=False, owner_id=owner_id, guest_key=guest_key)})


@require_POST
def duplicate_design_project(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    source = _design_owner_queryset(owner_id, guest_key).prefetch_related("assets").filter(id=project_id).first()
    if not source:
        raise Http404("Design project not found")
    copy = DesignerProject.objects.create(
        owner=source.owner,
        guest_key=source.guest_key,
        title=_clean_project_title(f"Copy of {source.title}"),
        state_json=source.state_json or {},
        preview_path="",
    )
    target_dir = _design_project_media_dir(copy)
    target_dir.mkdir(parents=True, exist_ok=True)
    asset_id_map: dict[str, int] = {}
    for asset in source.assets.all():
        source_path = Path(asset.file_path)
        if not source_path.exists() or not source_path.is_file():
            continue
        target_path = target_dir / f"{uuid.uuid4().hex[:12]}_{source_path.name}"
        shutil.copy2(source_path, target_path)
        new_asset = DesignerAsset.objects.create(project=copy, kind=asset.kind, file_path=str(target_path), media_type=asset.media_type, size=target_path.stat().st_size, original_name=asset.original_name)
        asset_id_map[str(asset.id)] = new_asset.id
    if source.preview_path:
        preview_path = Path(source.preview_path)
        if preview_path.exists() and preview_path.is_file():
            target_preview = target_dir / f"preview_{uuid.uuid4().hex[:12]}.jpg"
            shutil.copy2(preview_path, target_preview)
            copy.preview_path = str(target_preview)
    state = json.loads(json.dumps(source.state_json or {}))
    if isinstance(state, dict):
        for obj in state.get("objects", []) if isinstance(state.get("objects"), list) else []:
            asset_id = str(obj.get("assetId") or "")
            if asset_id in asset_id_map:
                obj["assetId"] = asset_id_map[asset_id]
                obj["src"] = reverse("studio:design_project_asset_preview", args=[copy.id, asset_id_map[asset_id]])
    copy.state_json = state
    _update_design_project_metadata(copy)
    copy.storage_bytes = _design_project_storage_bytes(copy)
    copy.save(update_fields=["state_json", "preview_path", "storage_bytes", "asset_count", "object_count", "updated_at"])
    return JsonResponse({"project": _design_project_payload(copy, owner_id=owner_id, guest_key=guest_key)})


@require_POST
def delete_design_project(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _design_owner_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Design project not found")
    _delete_design_project_media(project)
    deleted, _ = project.delete()
    if not deleted:
        raise Http404("Design project not found")
    stats = actions.get_account_stats(owner_id, guest_key)
    stats.update(_storage_quota(request, stats))
    return JsonResponse({"ok": True, "deleted_ids": [project_id], "account_stats": stats})


@require_POST
def delete_design_projects(request: HttpRequest) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    data = _json_body(request)
    raw_ids = data.get("ids") if isinstance(data.get("ids"), list) else []
    project_ids = []
    for raw_id in raw_ids:
        try:
            project_id = int(raw_id)
        except (TypeError, ValueError):
            continue
        if project_id > 0 and project_id not in project_ids:
            project_ids.append(project_id)
    projects = list(_design_owner_queryset(owner_id, guest_key).filter(id__in=project_ids))
    deleted_ids = []
    for project in projects:
        _delete_design_project_media(project)
        deleted_ids.append(project.id)
    if deleted_ids:
        DesignerProject.objects.filter(id__in=deleted_ids).delete()
    stats = actions.get_account_stats(owner_id, guest_key)
    stats.update(_storage_quota(request, stats))
    return JsonResponse({"ok": True, "deleted_ids": deleted_ids, "deleted_count": len(deleted_ids), "account_stats": stats})


@require_POST
def upload_design_project_asset(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _design_project_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Design project not found")
    if not _resource_can_edit(WorkspaceShare.RESOURCE_DESIGN, project, owner_id, guest_key):
        raise Http404("Design project not found")
    upload = _require_file(request, "file")
    media_type = upload.content_type or mimetypes.guess_type(upload.name or "")[0] or "application/octet-stream"
    if not media_type.startswith("image/"):
        return _error_json(ValueError("Only image assets are supported"), 400)
    if upload.size > settings.max_image_mb * 1024 * 1024:
        return _error_json(ValueError(f"Image limit: {settings.max_image_mb} MB"), 400)
    asset_dir = _design_project_media_dir(project)
    asset_dir.mkdir(parents=True, exist_ok=True)
    path, media_type = _save_optimized_editor_image(upload, asset_dir, clean_base_name(upload.name or "image", "image"))
    asset = DesignerAsset.objects.create(project=project, kind="image", file_path=str(path), media_type=media_type, size=path.stat().st_size, original_name=(upload.name or "image")[:240])
    _update_design_project_metadata(project)
    project.storage_bytes = _design_project_storage_bytes(project)
    project.save(update_fields=["storage_bytes", "asset_count", "object_count", "updated_at"])
    return JsonResponse({"asset": _design_asset_payload(asset), "project": _design_project_payload(project, owner_id=owner_id, guest_key=guest_key)})


@require_GET
def preview_design_project_asset(request: HttpRequest, project_id: int, asset_id: int) -> FileResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _design_project_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Design project not found")
    asset = project.assets.filter(id=asset_id).first()
    if not asset:
        raise Http404("Design asset not found")
    path = Path(asset.file_path)
    if not path.exists() or not path.is_file():
        raise Http404("Design asset file not found")
    return FileResponse(path.open("rb"), as_attachment=False, filename=asset.original_name or path.name, content_type=asset.media_type)


@require_GET
def preview_design_project(request: HttpRequest, project_id: int) -> HttpResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _design_project_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Design preview not found")
    state = project.state_json if isinstance(project.state_json, dict) else {}
    should_render_live = bool(_state_list(state, "vectors")) or not project.preview_path
    if should_render_live:
        rendered = _render_design_project_preview(project)
        if rendered:
            return HttpResponse(rendered.getvalue(), content_type="image/jpeg")
    if project.preview_path:
        path = Path(project.preview_path)
        if path.exists() and path.is_file():
            return FileResponse(path.open("rb"), as_attachment=False, filename=path.name, content_type="image/jpeg")
    rendered = _render_design_project_preview(project)
    if rendered:
        return HttpResponse(rendered.getvalue(), content_type="image/jpeg")
    raise Http404("Design preview not found")


@require_GET
def workspace_invite(request: HttpRequest, token: str):
    share = WorkspaceShare.objects.filter(token=token).select_related("owner", "invited_user").first()
    if not share or share.status == WorkspaceShare.STATUS_REVOKED or share.expires_at <= timezone.now():
        return render(request, "studio/invite.html", {**_invite_context(request, share, "expired"), "status": "expired"})
    resource = _share_resource(share.resource_type, share.resource_id)
    if not resource:
        return render(request, "studio/invite.html", {**_invite_context(request, share, "missing"), "status": "missing"})
    if request.user.is_authenticated:
        if _clean_email(request.user.email) != share.email:
            return render(request, "studio/invite.html", {**_invite_context(request, share, "wrong_email"), "status": "wrong_email"})
        if share.status == WorkspaceShare.STATUS_PENDING:
            share.invited_user = request.user
            share.status = WorkspaceShare.STATUS_ACCEPTED
            share.save(update_fields=["invited_user", "status", "updated_at"])
        return redirect(_share_resource_url(share.resource_type, share.resource_id))
    return render(request, "studio/invite.html", {**_invite_context(request, share, "ready"), "status": "ready"})


@require_http_methods(["GET", "POST"])
def workspace_shares(request: HttpRequest) -> JsonResponse:
    if not request.user.is_authenticated:
        return _error_json(ValueError("Sign in to share projects"), 403)
    if request.method == "GET":
        resource_type = str(request.GET.get("resource_type") or "")
        try:
            resource_id = int(request.GET.get("resource_id") or 0)
        except (TypeError, ValueError):
            resource_id = 0
        resource = _share_resource(resource_type, resource_id)
        if not resource or not _resource_is_owner(resource, request.user.id, ""):
            raise Http404("Project not found")
        shares = WorkspaceShare.objects.filter(owner=request.user, resource_type=resource_type, resource_id=resource_id).exclude(status=WorkspaceShare.STATUS_REVOKED).select_related("owner", "invited_user")
        return JsonResponse({
            "resource": _workspace_share_resource_payload(request, resource_type, resource),
            "shares": [_workspace_share_payload(request, share) for share in shares],
        })

    data = _json_body(request)
    resource_type = str(data.get("resource_type") or "").strip()
    try:
        resource_id = int(data.get("resource_id") or 0)
    except (TypeError, ValueError):
        resource_id = 0
    resource = _share_resource(resource_type, resource_id)
    if resource_type not in SHARE_RESOURCE_TYPES or not resource or not _resource_is_owner(resource, request.user.id, ""):
        raise Http404("Project not found")
    email = _clean_email(str(data.get("email") or ""))
    if not email:
        return _error_json(ValueError("Email is required"), 400)
    if email == _clean_email(request.user.email):
        return _error_json(ValueError("You already own this project"), 400)
    role = str(data.get("role") or WorkspaceShare.ROLE_VIEWER).strip().lower()
    if role not in {WorkspaceShare.ROLE_VIEWER, WorkspaceShare.ROLE_EDITOR}:
        role = WorkspaceShare.ROLE_VIEWER
    invited_user = get_user_model().objects.filter(email__iexact=email).first()
    share, created = WorkspaceShare.objects.get_or_create(
        owner=request.user,
        resource_type=resource_type,
        resource_id=resource_id,
        email=email,
        defaults={
            "invited_user": invited_user,
            "role": role,
            "status": WorkspaceShare.STATUS_PENDING,
            "token": secrets.token_urlsafe(32),
            "expires_at": timezone.now() + timezone.timedelta(days=14),
        },
    )
    if not created:
        share.invited_user = invited_user or share.invited_user
        share.role = role
        share.token = secrets.token_urlsafe(32)
        share.expires_at = timezone.now() + timezone.timedelta(days=14)
        update_fields = ["invited_user", "role", "token", "expires_at", "updated_at"]
        if share.status != WorkspaceShare.STATUS_ACCEPTED:
            share.status = WorkspaceShare.STATUS_PENDING
            update_fields.append("status")
        share.save(update_fields=update_fields)
    _send_workspace_invite(request, share)
    return JsonResponse({
        "resource": _workspace_share_resource_payload(request, resource_type, resource),
        "share": _workspace_share_payload(request, share),
    })


@require_POST
def workspace_share_role(request: HttpRequest, share_id: int) -> JsonResponse:
    if not request.user.is_authenticated:
        return _error_json(ValueError("Sign in to manage sharing"), 403)
    share = WorkspaceShare.objects.filter(id=share_id, owner=request.user).select_related("owner", "invited_user").first()
    if not share:
        raise Http404("Share not found")
    data = _json_body(request)
    role = str(data.get("role") or "").strip().lower()
    if role not in {WorkspaceShare.ROLE_VIEWER, WorkspaceShare.ROLE_EDITOR}:
        return _error_json(ValueError("Invalid role"), 400)
    share.role = role
    share.save(update_fields=["role", "updated_at"])
    return JsonResponse({"share": _workspace_share_payload(request, share)})


@require_POST
def revoke_workspace_share(request: HttpRequest, share_id: int) -> JsonResponse:
    if not request.user.is_authenticated:
        return _error_json(ValueError("Sign in to manage sharing"), 403)
    share = WorkspaceShare.objects.filter(id=share_id, owner=request.user).select_related("owner", "invited_user").first()
    if not share:
        raise Http404("Share not found")
    share.status = WorkspaceShare.STATUS_REVOKED
    share.save(update_fields=["status", "updated_at"])
    return JsonResponse({"ok": True, "share": _workspace_share_payload(request, share)})


@require_GET
@ensure_csrf_cookie
def video_editor(request: HttpRequest):
    language = getattr(request, "interface_language", "en")
    owner_id, guest_key = _workspace_identity(request)
    project = None
    project_id = request.GET.get("project", "")
    if project_id.isdigit():
        project = _video_project_queryset(owner_id, guest_key).prefetch_related("assets").filter(id=int(project_id)).first()
    return render(
        request,
        "studio/video_editor.html",
        {
            "current_video_project": _video_project_payload(project, owner_id=owner_id, guest_key=guest_key) if project else None,
            "video_projects_api_url": reverse("studio:video_projects"),
            "accent_color": _accent_color(request),
            "ui_accent_color": _ui_accent_color(request),
            "theme_mode": _theme_mode(request),
            "app_messages": app_messages(language),
            "subtitle_languages": _localized_subtitle_languages(language),
        },
    )


@require_GET
@ensure_csrf_cookie
def music_editor(request: HttpRequest):
    language = getattr(request, "interface_language", "en")
    owner_id, guest_key = _workspace_identity(request)
    project = None
    project_id = request.GET.get("project", "")
    if project_id.isdigit():
        project = _music_project_queryset(owner_id, guest_key).prefetch_related("assets").filter(id=int(project_id)).first()
    current_music_project = _music_project_payload(project, owner_id=owner_id, guest_key=guest_key) if project else None

    return render(
        request,
        "studio/music_editor.html",
        {
            "current_music_project": current_music_project,
            "current_music_project_json": json.dumps(current_music_project or {}),
            "music_messages_json": json.dumps(music_messages(language), ensure_ascii=False),
            "music_projects_api_url": reverse("studio:music_projects"),
            "accent_color": _accent_color(request),
            "ui_accent_color": _ui_accent_color(request),
            "theme_mode": _theme_mode(request),
            "app_messages": app_messages(language),
        },
    )



@require_GET
def music_project_list(request: HttpRequest):
    owner_id, guest_key = _workspace_identity(request)
    listing = _project_listing_payload(
        request,
        _music_project_queryset(owner_id, guest_key).prefetch_related("assets"),
        WorkspaceShare.RESOURCE_MUSIC,
        owner_id,
        guest_key,
        lambda project: _music_project_payload(project, include_state=False, owner_id=owner_id, guest_key=guest_key),
        per_page_default=120,
    )
    return render(
        request,
        "studio/music_projects.html",
        {
            "music_projects": listing["projects"],
            "project_listing": listing,
            "music_editor_url": reverse("studio:music_editor"),
            "music_messages_json": json.dumps(music_messages(getattr(request, "interface_language", "en")), ensure_ascii=False),
            "music_projects_api_url": reverse("studio:music_projects"),
            "accent_color": _accent_color(request),
            "ui_accent_color": _ui_accent_color(request),
            "theme_mode": _theme_mode(request),
            "app_messages": app_messages(getattr(request, "interface_language", "en")),
        },
    )



@require_GET
def music_projects(request: HttpRequest) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)

    listing = _project_listing_payload(
        request,
        _music_project_queryset(owner_id, guest_key).prefetch_related("assets"),
        WorkspaceShare.RESOURCE_MUSIC,
        owner_id,
        guest_key,
        lambda project: _music_project_payload(project, include_state=False, owner_id=owner_id, guest_key=guest_key),
    )
    return JsonResponse(listing)


@require_POST
def create_music_project(request: HttpRequest) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    data = _json_body(request)

    state = data.get("state") if isinstance(data.get("state"), dict) else {}
    title = _clean_project_title(
        str(data.get("title") or state.get("title") or "New music project")
    )

    project = MusicEditorProject.objects.create(
        owner=request.user if request.user.is_authenticated else None,
        guest_key="" if owner_id else guest_key,
        title=title,
        state_json=state,
        storage_bytes=_json_size(state),
    )
    _update_music_project_metadata(project, assets=[])
    project.save(update_fields=["asset_count", "clip_count", "duration_seconds", "updated_at"])

    return JsonResponse({
        "project": _music_project_payload(project, owner_id=owner_id, guest_key=guest_key)
    })

@require_GET
def music_project_detail(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _music_project_queryset(owner_id, guest_key).prefetch_related("assets").filter(id=project_id).first()
    if not project:
        raise Http404("Music project not found")
    return JsonResponse({"project": _music_project_payload(project, owner_id=owner_id, guest_key=guest_key)})


@require_POST
def save_music_project(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)

    project = _music_project_queryset(owner_id, guest_key).filter(id=project_id).first()

    if not project:
        raise Http404("Music project not found")

    if not _resource_can_edit(WorkspaceShare.RESOURCE_MUSIC, project, owner_id, guest_key):
        raise Http404("Music project not found")

    data = _json_body(request)

    state = data.get("state") if isinstance(data.get("state"), dict) else {}
    title = _clean_project_title(
        str(data.get("title") or state.get("title") or project.title)
    )

    project.title = title
    project.state_json = state
    _update_music_project_metadata(project)
    project.storage_bytes = _music_project_storage_bytes(project)
    project.save(update_fields=["title", "state_json", "storage_bytes", "asset_count", "clip_count", "duration_seconds", "updated_at"])

    return JsonResponse({
        "project": _music_project_payload(project, owner_id=owner_id, guest_key=guest_key)
    })


@require_POST
def rename_music_project(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)

    project = _music_owner_queryset(owner_id, guest_key).filter(id=project_id).first()

    if not project:
        raise Http404("Music project not found")

    data = _json_body(request)

    project.title = _clean_project_title(str(data.get("title") or project.title))
    project.save(update_fields=["title", "updated_at"])

    return JsonResponse({
        "project": _music_project_payload(
            project,
            include_state=False,
            owner_id=owner_id,
            guest_key=guest_key,
        )
    })

@require_POST
def delete_music_project(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)

    project = _music_owner_queryset(owner_id, guest_key).filter(id=project_id).first()

    if not project:
        raise Http404("Music project not found")

    _delete_music_project_media(project)

    deleted, _ = project.delete()

    if not deleted:
        raise Http404("Music project not found")

    stats = actions.get_account_stats(owner_id, guest_key)
    stats.update(_storage_quota(request, stats))

    return JsonResponse({
        "ok": True,
        "deleted_ids": [project_id],
        "account_stats": stats,
    })

@require_POST
def delete_music_projects(request: HttpRequest) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    data = _json_body(request)
    raw_ids = data.get("ids") if isinstance(data.get("ids"), list) else []
    project_ids = []
    for raw_id in raw_ids:
        try:
            project_id = int(raw_id)
        except (TypeError, ValueError):
            continue
        if project_id > 0 and project_id not in project_ids:
            project_ids.append(project_id)
    projects = list(_music_owner_queryset(owner_id, guest_key).filter(id__in=project_ids))
    deleted_ids = []
    for project in projects:
        _delete_music_project_media(project)
        deleted_ids.append(project.id)
    if deleted_ids:
        MusicEditorProject.objects.filter(id__in=deleted_ids).delete()
    stats = actions.get_account_stats(owner_id, guest_key)
    stats.update(_storage_quota(request, stats))
    return JsonResponse({"ok": True, "deleted_ids": deleted_ids, "deleted_count": len(deleted_ids), "account_stats": stats})




@require_POST
def upload_music_project_asset(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)

    project = _music_project_queryset(owner_id, guest_key).filter(id=project_id).first()

    if not project:
        raise Http404("Music project not found")

    upload = _require_file(request, "file")

    if upload.size > settings.max_video_mb * 1024 * 1024:
        return _error_json(ValueError(f"File limit: {settings.max_video_mb} MB"), 400)

    asset_dir = _music_project_media_dir(project)
    asset_dir.mkdir(parents=True, exist_ok=True)

    upload_name = upload.name or "audio"
    suffix = Path(upload_name).suffix.lower()
    media_type = upload.content_type or mimetypes.guess_type(upload_name)[0] or "application/octet-stream"
    imported_assets: list[MusicEditorAsset] = []
    source_assets: list[MusicEditorAsset] = []
    skipped: list[str] = []

    if suffix == ".zip" or media_type in {"application/zip", "application/x-zip-compressed"}:
        try:
            archive_bytes = BytesIO()
            for chunk in upload.chunks():
                archive_bytes.write(chunk)
            archive_bytes.seek(0)
            with zipfile.ZipFile(archive_bytes) as archive:
                for member in archive.infolist():
                    if member.is_dir():
                        continue
                    member_path = Path(member.filename)
                    member_name = member_path.name
                    if not member_name or member.file_size <= 0:
                        continue
                    if member.file_size > settings.max_video_mb * 1024 * 1024:
                        skipped.append(member.filename)
                        continue
                    member_suffix = member_path.suffix.lower()
                    member_media_type = mimetypes.guess_type(member_name)[0] or "application/octet-stream"
                    if _music_import_is_audio(member_name, member_media_type):
                        imported_assets.append(_save_music_import_asset(project, asset_dir, member_name, member_media_type, archive.read(member), "audio"))
                    elif member_suffix in {".flp", ".mid", ".midi"}:
                        source_assets.append(_save_music_import_asset(project, asset_dir, member_name, member_media_type, archive.read(member), "source"))
                    else:
                        skipped.append(member.filename)
        except zipfile.BadZipFile:
            return _error_json(ValueError("ZIP archive is not readable"), 400)
        if not imported_assets and not source_assets:
            return _error_json(ValueError("ZIP does not contain supported audio, MIDI or FLP files"), 400)
    elif _music_import_is_audio(upload_name, media_type):
        imported_assets.append(_save_music_upload_asset(project, asset_dir, upload, upload_name, media_type, "audio"))
    elif suffix in {".flp", ".mid", ".midi"}:
        source_assets.append(_save_music_upload_asset(project, asset_dir, upload, upload_name, media_type, "source"))
    else:
        return _error_json(ValueError("Upload audio, FL Studio ZIP stems, MIDI or FLP source files"), 400)

    _update_music_project_metadata(project)
    project.storage_bytes = _music_project_storage_bytes(project)
    project.save(update_fields=["storage_bytes", "asset_count", "clip_count", "duration_seconds", "updated_at"])

    assets = imported_assets + source_assets
    return JsonResponse({
        "asset": _music_asset_payload(assets[0]) if len(assets) == 1 else None,
        "assets": [_music_asset_payload(asset) for asset in assets],
        "import": {
            "audio_count": len(imported_assets),
            "source_count": len(source_assets),
            "skipped": skipped[:20],
        },
        "project": _music_project_payload(project, owner_id=owner_id, guest_key=guest_key),
    })

@require_GET
def preview_music_project_asset(request: HttpRequest, project_id: int, asset_id: int) -> HttpResponse:
    owner_id, guest_key = _workspace_identity(request)

    project = _music_project_queryset(owner_id, guest_key).filter(id=project_id).first()

    if not project:
        raise Http404("Music project not found")

    asset = project.assets.filter(id=asset_id).first()

    if not asset:
        raise Http404("Music asset not found")

    path = Path(asset.file_path)

    if not path.exists() or not path.is_file():
        raise Http404("Music asset file not found")

    return _range_file_response(
        request,
        path,
        asset.media_type,
        asset.original_name or path.name,
    )

@require_POST
def delete_music_project_asset(request: HttpRequest, project_id: int, asset_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)

    project = _music_project_queryset(owner_id, guest_key).filter(id=project_id).first()

    if not project:
        raise Http404("Music project not found")

    asset = project.assets.filter(id=asset_id).first()

    if not asset:
        raise Http404("Music asset not found")

    asset_path = Path(asset.file_path)

    try:
        if asset_path.exists() and asset_path.is_file():
            asset_path.unlink()
    except Exception:
        pass

    asset.delete()

    state = project.state_json or {}

    if isinstance(state, dict):
        state["assets"] = [
            item for item in state.get("assets", [])
            if str(item.get("serverId", "")) != str(asset_id)
        ]

        state["clips"] = [
            clip for clip in state.get("clips", [])
            if str(clip.get("assetId", "")) != str(asset_id)
        ]

        project.state_json = state

    _update_music_project_metadata(project)
    project.storage_bytes = _music_project_storage_bytes(project)
    project.save(update_fields=["state_json", "storage_bytes", "asset_count", "clip_count", "duration_seconds", "updated_at"])

    return JsonResponse({
        "ok": True,
        "asset_id": asset_id,
        "project": _music_project_payload(project, owner_id=owner_id, guest_key=guest_key),
    })

@require_GET
@ensure_csrf_cookie
def video_project_list(request: HttpRequest):
    owner_id, guest_key = _workspace_identity(request)
    listing = _project_listing_payload(
        request,
        _video_project_queryset(owner_id, guest_key).prefetch_related("assets"),
        WorkspaceShare.RESOURCE_VIDEO,
        owner_id,
        guest_key,
        lambda project: _video_project_payload(project, include_state=False, owner_id=owner_id, guest_key=guest_key),
        per_page_default=120,
    )
    return render(
        request,
        "studio/video_projects.html",
        {
            "video_projects": listing["projects"],
            "project_listing": listing,
            "video_editor_url": reverse("studio:video_editor"),
            "video_projects_api_url": reverse("studio:video_projects"),
            "accent_color": _accent_color(request),
            "ui_accent_color": _ui_accent_color(request),
            "theme_mode": _theme_mode(request),
            "app_messages": app_messages(getattr(request, "interface_language", "en")),
        },
    )


@require_GET
def video_projects(request: HttpRequest) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    listing = _project_listing_payload(
        request,
        _video_project_queryset(owner_id, guest_key).prefetch_related("assets"),
        WorkspaceShare.RESOURCE_VIDEO,
        owner_id,
        guest_key,
        lambda project: _video_project_payload(project, include_state=False, owner_id=owner_id, guest_key=guest_key),
    )
    return JsonResponse(listing)


@require_POST
def create_video_project(request: HttpRequest) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    data = _json_body(request)
    state = data.get("state") if isinstance(data.get("state"), dict) else {}
    title = _clean_project_title(str(data.get("title") or state.get("title") or "New project"))
    project = VideoEditorProject.objects.create(
        owner=request.user if request.user.is_authenticated else None,
        guest_key="" if owner_id else guest_key,
        title=title,
        state_json=state,
        storage_bytes=_json_size(state),
    )
    _update_video_project_metadata(project, assets=[])
    project.storage_bytes = _video_project_storage_bytes(project)
    project.save(update_fields=["storage_bytes", "asset_count", "clip_count", "duration_seconds", "thumbnail_path", "updated_at"])
    return JsonResponse({"project": _video_project_payload(project, owner_id=owner_id, guest_key=guest_key)})


@require_GET
def video_project_detail(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _video_project_queryset(owner_id, guest_key).prefetch_related("assets").filter(id=project_id).first()
    if not project:
        raise Http404("Project not found")
    return JsonResponse({"project": _video_project_payload(project, owner_id=owner_id, guest_key=guest_key)})


@require_POST
def save_video_project(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _video_project_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Project not found")
    if not _resource_can_edit(WorkspaceShare.RESOURCE_VIDEO, project, owner_id, guest_key):
        raise Http404("Project not found")
    data = _json_body(request)
    state = data.get("state") if isinstance(data.get("state"), dict) else {}
    title = _clean_project_title(str(data.get("title") or state.get("title") or project.title))
    project.title = title
    project.state_json = state
    _update_video_project_metadata(project)
    project.storage_bytes = _video_project_storage_bytes(project)
    project.save(update_fields=["title", "state_json", "storage_bytes", "asset_count", "clip_count", "duration_seconds", "thumbnail_path", "updated_at"])
    return JsonResponse({"project": _video_project_payload(project, owner_id=owner_id, guest_key=guest_key)})


@require_POST
def rename_video_project(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _video_owner_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Project not found")
    data = _json_body(request)
    title = _clean_project_title(str(data.get("title") or project.title))
    project.title = title
    project.save(update_fields=["title", "updated_at"])
    return JsonResponse({"project": _video_project_payload(project, include_state=False, owner_id=owner_id, guest_key=guest_key)})


@require_POST
def start_video_project_export(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _video_project_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Project not found")
    if not _resource_can_edit(WorkspaceShare.RESOURCE_VIDEO, project, owner_id, guest_key):
        raise Http404("Project not found")
    data = _json_body(request)
    quality = "1080p" if data.get("quality") == "1080p" else "720p"
    preset = str(data.get("preset") or "").strip()
    job_id = uuid.uuid4().hex[:16]
    output_dir = _video_project_media_dir(project) / "exports"
    output_dir.mkdir(parents=True, exist_ok=True)
    output = output_dir / f"{clean_base_name(project.title, 'video-project')}_{quality}_{job_id}.mp4"
    job_record = JobRecord.objects.create(
        owner=project.owner,
        guest_key=project.guest_key,
        job_id=job_id,
        kind="video_export",
        title=f"Export {project.title}",
        status="queued",
        progress=2,
        message="Queued",
        params_json=json.dumps({"project_id": project.id, "path": str(output), "quality": quality, "preset": preset}, ensure_ascii=False),
    )
    JobEventRecord.objects.create(job=job_record, status="queued", progress=2, message="Queued")
    with _video_export_lock:
        _video_export_jobs[job_id] = {
            "id": job_id,
            "project_id": project.id,
            "owner_id": owner_id,
            "guest_key": guest_key,
            "status": "queued",
            "progress": 2,
            "message": "Queued",
            "error": "",
            "path": str(output),
            "quality": quality,
            "preset": preset,
            "created_at": time.time(),
        }
    if not settings.persistent_job_queue:
        _video_export_executor.submit(_run_video_project_export, job_id, project.id, quality, output)
    return JsonResponse({"job": _video_export_payload(request, project, job_id)})


@require_GET
def list_video_project_exports(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _video_project_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Project not found")
    jobs = JobRecord.objects.filter(kind__in=["video_export", "video_cover"], params_json__contains=f'"project_id": {project.id}').prefetch_related("outputs").order_by("-created_at")[:12]
    return JsonResponse({"jobs": [_video_export_record_payload(request, project, job) for job in jobs]})


@require_GET
def video_project_export_status(request: HttpRequest, project_id: int, job_id: str) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _video_project_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Project not found")
    if not _video_export_access(job_id, project, owner_id, guest_key):
        raise Http404("Export not found")
    return JsonResponse({"job": _video_export_payload(request, project, job_id)})


@require_POST
def cancel_video_project_export(request: HttpRequest, project_id: int, job_id: str) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _video_project_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Project not found")
    if not _video_export_access(job_id, project, owner_id, guest_key):
        raise Http404("Export not found")
    if not _resource_can_edit(WorkspaceShare.RESOURCE_VIDEO, project, owner_id, guest_key):
        raise Http404("Export not found")
    with _video_export_lock:
        process = _video_export_processes.get(job_id)
    if process and process.poll() is None:
        try:
            process.terminate()
        except Exception:
            pass
    _set_video_export_job(job_id, status="cancelled", progress=100, message="Cancelled")
    return JsonResponse({"job": _video_export_payload(request, project, job_id)})


@require_GET
def download_video_project_export(request: HttpRequest, project_id: int, job_id: str) -> FileResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _video_project_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Project not found")
    if not _video_export_access(job_id, project, owner_id, guest_key):
        raise Http404("Export not found")
    record = JobRecord.objects.filter(job_id=job_id).prefetch_related("outputs").first()
    output = record.outputs.first() if record else None
    if output:
        path = Path(output.path)
        ready = record.status in {"completed", "done"}
        media_type = output.media_type
    else:
        job = _video_export_jobs.get(job_id) or {}
        path = Path(str(job.get("path") or ""))
        ready = job.get("status") == "done"
        media_type = "video/mp4"
    if not ready or not path.exists() or not path.is_file():
        raise Http404("Export not ready")
    return FileResponse(path.open("rb"), as_attachment=True, filename=path.name, content_type=media_type)


@require_POST
def export_video_project_cover(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _video_project_queryset(owner_id, guest_key).prefetch_related("assets").filter(id=project_id).first()
    if not project:
        raise Http404("Project not found")
    if not _resource_can_edit(WorkspaceShare.RESOURCE_VIDEO, project, owner_id, guest_key):
        raise Http404("Project not found")
    data = _json_body(request)
    time_seconds = max(0, float(data.get("time") or 0))
    job_id = uuid.uuid4().hex[:16]
    output_dir = _video_project_media_dir(project) / "exports"
    output_dir.mkdir(parents=True, exist_ok=True)
    output = output_dir / f"{clean_base_name(project.title, 'cover')}_{job_id}.jpg"
    job_record = JobRecord.objects.create(
        owner=project.owner,
        guest_key=project.guest_key,
        job_id=job_id,
        kind="video_cover",
        title=f"Cover {project.title}",
        status="queued",
        progress=2,
        message="Queued",
        params_json=json.dumps({"project_id": project.id, "path": str(output), "time": time_seconds}, ensure_ascii=False),
    )
    JobEventRecord.objects.create(job=job_record, status="queued", progress=2, message="Queued")
    with _video_export_lock:
        _video_export_jobs[job_id] = {
            "id": job_id,
            "project_id": project.id,
            "owner_id": owner_id,
            "guest_key": guest_key,
            "kind": "video_cover",
            "status": "queued",
            "progress": 2,
            "message": "Queued",
            "error": "",
            "path": str(output),
            "created_at": time.time(),
        }
    if not settings.persistent_job_queue:
        _video_export_executor.submit(_run_video_project_cover, job_id, project.id, output, time_seconds)
    return JsonResponse({"job": _video_export_payload(request, project, job_id)})


@require_POST
def import_video_project_subtitles(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _video_project_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Project not found")
    if not _resource_can_edit(WorkspaceShare.RESOURCE_VIDEO, project, owner_id, guest_key):
        raise Http404("Project not found")
    upload = request.FILES.get("file")
    if not upload:
        return JsonResponse({"error": "Subtitle file is required"}, status=400)
    raw = upload.read()
    text = _decode_subtitle_bytes(raw)
    cues = _parse_subtitle_cues(text, upload.name or "")
    return JsonResponse({"cues": cues, "count": len(cues)})


@require_POST
def auto_video_project_subtitles(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _video_project_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Project not found")
    if not _resource_can_edit(WorkspaceShare.RESOURCE_VIDEO, project, owner_id, guest_key):
        raise Http404("Project not found")

    data = _json_body(request)
    try:
        asset_id = int(data.get("asset_id") or 0)
    except (TypeError, ValueError):
        asset_id = 0
    asset_query = VideoEditorAsset.objects.filter(project=project, kind__in=["audio", "video"])
    asset = asset_query.filter(id=asset_id).first() if asset_id else asset_query.first()
    if not asset:
        return JsonResponse({"error": "Select a video or audio clip first"}, status=400)

    source = Path(asset.file_path)
    if not source.exists() or not source.is_file():
        return JsonResponse({"error": "Source media file is missing"}, status=404)

    allowed_languages = {key for key, _label in actions.SUBTITLE_LANGUAGE_CHOICES if key != "auto"}
    language = normalize_subtitle_language(str(data.get("language") or "auto"))
    language = language if language in allowed_languages else None
    timeline_start = _subtitle_seconds(data.get("timeline_start"))
    source_start = _subtitle_seconds(data.get("source_start"))
    source_end = _subtitle_seconds(data.get("source_end"))
    clip_duration = _subtitle_seconds(data.get("clip_duration"))
    if source_end <= source_start and clip_duration > 0:
        source_end = source_start + clip_duration
    clip_end = timeline_start + max(0.0, source_end - source_start) if source_end > source_start else None

    job_id = uuid.uuid4().hex[:16]
    params = {
        "project_id": project.id,
        "asset_id": asset.id,
        "asset_name": asset.original_name,
        "language": language or "auto",
        "timeline_start": timeline_start,
        "source_start": source_start,
        "source_end": source_end,
        "clip_duration": clip_duration,
        "clip_end": clip_end,
    }
    job_record = JobRecord.objects.create(
        owner=project.owner,
        guest_key=project.guest_key,
        job_id=job_id,
        kind="video_subtitles",
        title=f"Auto subtitles {project.title}",
        status="queued",
        progress=2,
        message="Queued",
        params_json=json.dumps(params, ensure_ascii=False),
    )
    JobEventRecord.objects.create(job=job_record, status="queued", progress=2, message="Queued")
    with _video_export_lock:
        _video_export_jobs[job_id] = {
            "id": job_id,
            "project_id": project.id,
            "owner_id": owner_id,
            "guest_key": guest_key,
            "kind": "video_subtitles",
            "status": "queued",
            "progress": 2,
            "message": "Queued",
            "error": "",
            "created_at": time.time(),
        }
    if not settings.persistent_job_queue:
        _video_export_executor.submit(_run_video_project_subtitles, job_id, project.id, asset.id, params)
    return JsonResponse({"job": _video_export_payload(request, project, job_id)})


@require_GET
def export_video_project_subtitles(request: HttpRequest, project_id: int) -> HttpResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _video_project_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Project not found")
    if not _resource_can_view(WorkspaceShare.RESOURCE_VIDEO, project, owner_id, guest_key):
        raise Http404("Project not found")
    export_format = str(request.GET.get("format") or "srt").strip().lower()
    if export_format not in {"srt", "vtt", "ass", "json"}:
        export_format = "srt"
    cues = _video_project_caption_cues(project.state_json or {}, rich=export_format in {"ass", "json"})
    if export_format == "vtt":
        content = _render_vtt(cues)
        media_type = "text/vtt; charset=utf-8"
    elif export_format == "ass":
        content = _render_ass(cues, project.state_json or {})
        media_type = "text/x-ssa; charset=utf-8"
    elif export_format == "json":
        content = json.dumps(
            {
                "project_id": project.id,
                "title": project.title,
                "subtitleWorkflow": (project.state_json or {}).get("subtitleWorkflow", {}) if isinstance(project.state_json, dict) else {},
                "cues": cues,
            },
            ensure_ascii=False,
            indent=2,
        )
        media_type = "application/json; charset=utf-8"
    else:
        content = _render_srt(cues)
        media_type = "application/x-subrip; charset=utf-8"
    filename = f"{clean_base_name(project.title, 'subtitles')}.{export_format}"
    response = HttpResponse(content, content_type=media_type)
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


@require_POST
def delete_video_project(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _video_owner_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Project not found")
    _delete_video_project_media(project)
    deleted, _ = project.delete()
    if not deleted:
        raise Http404("Project not found")
    stats = actions.get_account_stats(owner_id, guest_key)
    stats.update(_storage_quota(request, stats))
    return JsonResponse({"ok": True, "account_stats": stats})


@require_POST
def delete_video_projects(request: HttpRequest) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    data = _json_body(request)
    raw_ids = data.get("ids") if isinstance(data.get("ids"), list) else []
    project_ids: list[int] = []
    for raw_id in raw_ids:
        try:
            project_id = int(raw_id)
        except (TypeError, ValueError):
            continue
        if project_id > 0 and project_id not in project_ids:
            project_ids.append(project_id)
    if not project_ids:
        stats = actions.get_account_stats(owner_id, guest_key)
        stats.update(_storage_quota(request, stats))
        return JsonResponse({"ok": True, "deleted_ids": [], "deleted_count": 0, "account_stats": stats})

    projects = list(_video_owner_queryset(owner_id, guest_key).filter(id__in=project_ids))
    deleted_ids: list[int] = []
    for project in projects:
        _delete_video_project_media(project)
        deleted_ids.append(project.id)
    if deleted_ids:
        VideoEditorProject.objects.filter(id__in=deleted_ids).delete()
    stats = actions.get_account_stats(owner_id, guest_key)
    stats.update(_storage_quota(request, stats))
    return JsonResponse({"ok": True, "deleted_ids": deleted_ids, "deleted_count": len(deleted_ids), "account_stats": stats})


@require_POST
def upload_video_project_asset(request: HttpRequest, project_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _video_project_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Project not found")
    if not _resource_can_edit(WorkspaceShare.RESOURCE_VIDEO, project, owner_id, guest_key):
        raise Http404("Project not found")
    upload = _require_file(request, "file")
    media_type = upload.content_type or mimetypes.guess_type(upload.name or "")[0] or "application/octet-stream"
    kind = (request.POST.get("kind") or "").strip().lower()
    if kind not in {"video", "audio", "image"}:
        if media_type.startswith("video/"):
            kind = "video"
        elif media_type.startswith("audio/"):
            kind = "audio"
        elif media_type.startswith("image/") or _is_visual_document_type(media_type, upload.name or ""):
            kind = "image"
        else:
            return _error_json(ValueError("Unsupported asset type"), 400)
    if kind == "video" and upload.size > settings.max_video_mb * 1024 * 1024:
        return _error_json(ValueError(f"Video limit: {settings.max_video_mb} MB"), 400)
    if kind == "image" and upload.size > settings.max_image_mb * 1024 * 1024:
        return _error_json(ValueError(f"Image limit: {settings.max_image_mb} MB"), 400)
    asset_dir = _video_project_media_dir(project)
    asset_dir.mkdir(parents=True, exist_ok=True)
    suffix = Path(upload.name or "").suffix[:16] or mimetypes.guess_extension(media_type) or ".bin"
    base = clean_base_name(upload.name or kind, kind)
    if kind == "image" and media_type.startswith("image/"):
        path, media_type = _save_optimized_editor_image(upload, asset_dir, base)
    else:
        path = asset_dir / f"{uuid.uuid4().hex[:12]}_{base}{suffix}"
        with path.open("wb") as destination:
            for chunk in upload.chunks():
                destination.write(chunk)
    thumbnail_path = ""
    thumbnail_data = request.POST.get("thumbnail", "")
    if thumbnail_data.startswith("data:image/"):
        thumbnail_path = str(_save_project_thumbnail(thumbnail_data, asset_dir))
    asset = VideoEditorAsset.objects.create(
        project=project,
        kind=kind,
        file_path=str(path),
        media_type=media_type,
        size=path.stat().st_size,
        original_name=(upload.name or kind)[:240],
        thumbnail_path=thumbnail_path,
        duration=float(request.POST.get("duration") or 0),
    )
    _update_video_project_metadata(project)
    project.storage_bytes = _video_project_storage_bytes(project)
    project.save(update_fields=["storage_bytes", "asset_count", "clip_count", "duration_seconds", "thumbnail_path", "updated_at"])
    return JsonResponse({"asset": _video_asset_payload(asset), "project": _video_project_payload(project, owner_id=owner_id, guest_key=guest_key)})


@require_GET
def preview_video_project_asset(request: HttpRequest, project_id: int, asset_id: int) -> HttpResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _video_project_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Project not found")
    if not _resource_can_view(WorkspaceShare.RESOURCE_VIDEO, project, owner_id, guest_key):
        raise Http404("Project not found")
    asset = project.assets.filter(id=asset_id).first()
    if not asset:
        raise Http404("Asset not found")
    path = Path(asset.file_path)
    if not path.exists() or not path.is_file():
        raise Http404("Asset file not found")
    if asset.kind in {"video", "audio"}:
        return _range_file_response(request, path, asset.media_type, asset.original_name or path.name)
    return FileResponse(path.open("rb"), as_attachment=False, filename=asset.original_name or path.name, content_type=asset.media_type)


@require_GET
def thumbnail_video_project_asset(request: HttpRequest, project_id: int, asset_id: int) -> HttpResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _video_project_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Project not found")
    if not _resource_can_view(WorkspaceShare.RESOURCE_VIDEO, project, owner_id, guest_key):
        raise Http404("Project not found")
    asset = project.assets.filter(id=asset_id).first()
    if not asset or not asset.thumbnail_path:
        raise Http404("Thumbnail not found")
    path = Path(asset.thumbnail_path)
    if not path.exists() or not path.is_file():
        raise Http404("Thumbnail not found")
    return FileResponse(path.open("rb"), as_attachment=False, filename=path.name, content_type="image/jpeg")


@require_http_methods(["GET", "POST"])
def video_project_asset_waveform(request: HttpRequest, project_id: int, asset_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _video_project_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Project not found")
    can_edit = _resource_can_edit(WorkspaceShare.RESOURCE_VIDEO, project, owner_id, guest_key)
    if request.method == "POST" and not can_edit:
        raise Http404("Waveform asset not found")
    asset = project.assets.filter(id=asset_id).first()
    if not asset or asset.kind not in {"audio", "video"}:
        raise Http404("Waveform asset not found")
    path = _video_asset_waveform_path(asset)
    if request.method == "POST" or (can_edit and not path.exists()):
        path.parent.mkdir(parents=True, exist_ok=True)
        samples = _build_compact_waveform(Path(asset.file_path))
        path.write_text(json.dumps({"samples": samples, "count": len(samples)}, separators=(",", ":")), encoding="utf-8")
    if not path.exists():
        return JsonResponse({"samples": [], "count": 0})
    return JsonResponse(json.loads(path.read_text(encoding="utf-8")))


@require_POST
def rename_video_project_asset(request: HttpRequest, project_id: int, asset_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _video_project_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Project not found")
    if not _resource_can_edit(WorkspaceShare.RESOURCE_VIDEO, project, owner_id, guest_key):
        raise Http404("Project not found")
    asset = project.assets.filter(id=asset_id).first()
    if not asset:
        raise Http404("Asset not found")
    data = _json_body(request)
    name = str(data.get("name") or asset.original_name or "").strip()
    if not name:
        return _error_json(ValueError("Asset name is required"), 400)
    asset.original_name = name[:240]
    asset.save(update_fields=["original_name"])
    project.save(update_fields=["updated_at"])
    return JsonResponse({"asset": _video_asset_payload(asset), "project": _video_project_payload(project, owner_id=owner_id, guest_key=guest_key)})


@require_POST
def delete_video_project_asset(request: HttpRequest, project_id: int, asset_id: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    project = _video_project_queryset(owner_id, guest_key).filter(id=project_id).first()
    if not project:
        raise Http404("Project not found")
    if not _resource_can_edit(WorkspaceShare.RESOURCE_VIDEO, project, owner_id, guest_key):
        raise Http404("Project not found")
    asset = project.assets.filter(id=asset_id).first()
    if not asset:
        raise Http404("Asset not found")
    _delete_video_asset_files(asset)
    asset.delete()
    state = project.state_json or {}
    if isinstance(state, dict):
        state["clips"] = [clip for clip in state.get("clips", []) if str(clip.get("assetId", "")) != str(asset_id)]
        project.state_json = state
    _update_video_project_metadata(project)
    project.storage_bytes = _video_project_storage_bytes(project)
    project.save(update_fields=["state_json", "storage_bytes", "asset_count", "clip_count", "duration_seconds", "thumbnail_path", "updated_at"])
    return JsonResponse({"ok": True, "project": _video_project_payload(project, owner_id=owner_id, guest_key=guest_key)})

def _dashboard_output_queryset(owner_id: int | None, guest_key: str, query: str = "", file_type: str = ""):
    queryset = JobOutputRecord.objects.select_related("job").only(
        "id",
        "job_id",
        "label",
        "path",
        "media_type",
        "size",
        "created_at",
        "job__job_id",
        "job__kind",
        "job__title",
        "job__params_json",
        "job__created_at",
    )
    if owner_id is not None:
        queryset = queryset.filter(job__owner_id=owner_id)
    else:
        queryset = queryset.filter(job__owner__isnull=True, job__guest_key=guest_key)
    if query:
        queryset = queryset.filter(Q(label__icontains=query) | Q(path__icontains=query) | Q(job__title__icontains=query))
    if file_type in {"image", "video", "pdf", "text", "audio", "other"}:
        image_q = Q(media_type__istartswith="image/")
        video_q = Q(media_type__istartswith="video/")
        pdf_q = Q(media_type="application/pdf") | Q(path__iendswith=".pdf")
        text_q = Q(media_type__istartswith="text/")
        for extension in (".txt", ".ass", ".srt", ".json", ".csv", ".rtf"):
            text_q |= Q(path__iendswith=extension)
        audio_q = Q(media_type__istartswith="audio/")
        if file_type == "image":
            queryset = queryset.filter(image_q)
        elif file_type == "video":
            queryset = queryset.filter(video_q)
        elif file_type == "pdf":
            queryset = queryset.filter(pdf_q)
        elif file_type == "text":
            queryset = queryset.filter(text_q)
        elif file_type == "audio":
            queryset = queryset.filter(audio_q)
        else:
            queryset = queryset.exclude(image_q | video_q | pdf_q | text_q | audio_q)
    return queryset.order_by("-job__created_at", "-created_at", "-id")


def _dashboard_output_indexes(outputs: list[JobOutputRecord]) -> dict[int, int]:
    job_ids = [output.job_id for output in outputs]
    if not job_ids:
        return {}
    indexes: dict[int, int] = {}
    current_job_id = None
    current_index = -1
    queryset = JobOutputRecord.objects.filter(job_id__in=job_ids).only("id", "job_id").order_by("job_id", "id")
    for output in queryset:
        if output.job_id != current_job_id:
            current_job_id = output.job_id
            current_index = 0
        else:
            current_index += 1
        indexes[output.id] = current_index
    return indexes


def _dashboard_output_payloads(outputs: list[JobOutputRecord]) -> list[dict[str, object]]:
    indexes = _dashboard_output_indexes(outputs)
    cache = _job_url_cache([output.job.job_id for output in outputs])
    payloads: list[dict[str, object]] = []
    for output in outputs:
        job = output.job
        job_id = job.job_id
        index = indexes.get(output.id, 0)
        name = Path(output.path).name or output.label
        item: dict[str, object] = {
            "index": index,
            "label": output.label,
            "name": name,
            "media_type": output.media_type,
            "size": output.size,
            "size_text": human_size(output.size or _path_size(output.path)),
            "job_title": job.title,
            "job_id": job_id,
            "detail_url": reverse("studio:job_detail", args=[job_id]),
            "url": reverse("studio:download_output", args=[job_id, index]),
            "preview_url": reverse("studio:preview_output", args=[job_id, index]),
        }
        item["preview_kind"] = _preview_kind(item)
        job_payload = {"id": job_id, "kind": job.kind}
        item["can_edit_design"] = _output_can_edit_design(item, job_payload)
        cached = cache.get(job_id, {})
        output_key = str(Path(output.path).resolve())
        if item["can_edit_design"]:
            item["edit_design_url"] = reverse("studio:edit_output_design", args=[job_id, index])
            design_map = cached.get("design_projects") if isinstance(cached.get("design_projects"), dict) else {}
            design_project_id = design_map.get(output_key) if output_key else None
            item["design_project_url"] = f"{reverse('studio:designer')}?{urlencode({'project': design_project_id})}" if design_project_id else ""
        item["can_edit_video"] = _output_can_edit_video(item, job_payload)
        if item["can_edit_video"]:
            item["edit_video_url"] = reverse("studio:edit_output_video", args=[job_id, index])
            video_map = cached.get("video_projects") if isinstance(cached.get("video_projects"), dict) else {}
            video_project_id = video_map.get(output_key) if output_key else None
            item["video_project_url"] = f"{reverse('studio:video_editor')}?{urlencode({'project': video_project_id})}" if video_project_id else ""
        payloads.append(item)
    return payloads


def _dashboard_page_range(page: int, pages: int) -> list[int | None]:
    if pages > 9:
        if page <= 5:
            return list(range(1, 7)) + [None, pages]
        if page >= pages - 4:
            return [1, None] + list(range(pages - 5, pages + 1))
        return [1, None] + list(range(page - 2, page + 3)) + [None, pages]
    return list(range(1, pages + 1))


def _dashboard_job_queryset(owner_id: int | None, guest_key: str, section: str):
    queryset = _owned_job_records(owner_id, guest_key).prefetch_related("outputs")
    if section == "active":
        queryset = queryset.filter(status__in=["queued", "running", "processing", "paused"])
    elif section == "completed":
        queryset = queryset.filter(status="completed")
    return queryset.order_by("-created_at", "-id")


def _dashboard_job_payloads(records: list[JobRecord], request: HttpRequest) -> list[dict]:
    language = getattr(request, "interface_language", "en")
    jobs = _attach_output_urls_many([actions._serialize_job_record(record) for record in records])
    return [_localize_job(job, language) for job in jobs]


@require_GET
def dashboard_detail(request: HttpRequest, section: str):
    owner_id, guest_key = _workspace_identity(request)
    stats = actions.get_account_stats(owner_id, guest_key)
    stats.update(_storage_quota(request, stats))
    normalized = section if section in {"all", "active", "completed", "files", "storage"} else "all"
    visible_jobs: list[dict] = []
    total_jobs = 0

    query = str(request.GET.get("q") or "").strip()
    file_type = str(request.GET.get("type") or "").strip().lower()
    try:
        page = int(request.GET.get("page") or 1)
    except (TypeError, ValueError):
        page = 1
    page = max(1, page)
    outputs = []
    total_outputs = 0
    pages = 1
    page_range = [1]
    per_page = 20
    has_access = user_has_active_access(request.user)

    if normalized in {"files", "storage"}:
        output_queryset = _dashboard_output_queryset(owner_id, guest_key, query, file_type)
        total_outputs = output_queryset.count()
        pages = max(1, (total_outputs + per_page - 1) // per_page)
        if page > pages:
            page = pages
        start = (page - 1) * per_page
        outputs = _dashboard_output_payloads(list(output_queryset[start : start + per_page]))

        page_range = _dashboard_page_range(page, pages)
    else:
        job_queryset = _dashboard_job_queryset(owner_id, guest_key, normalized)
        total_jobs = job_queryset.count()
        pages = max(1, (total_jobs + per_page - 1) // per_page)
        if page > pages:
            page = pages
        start = (page - 1) * per_page
        visible_jobs = _dashboard_job_payloads(list(job_queryset[start : start + per_page]), request)
        page_range = _dashboard_page_range(page, pages)

    context = {
        "section": normalized,
        "jobs": visible_jobs,
        "total_jobs": total_jobs,
        "outputs": outputs,
        "total_outputs": total_outputs,
        "page": page,
        "pages": pages,
        "page_range": page_range,
        "query": query,
        "file_type": file_type,
        "show_file_filters": normalized in {"files", "storage"},
        "has_access": has_access,
        "account_stats": stats,
        "is_guest": not request.user.is_authenticated,
        "checkout_url": _checkout_url(request),
        "pricing_url": reverse("billing:pricing"),
        "login_url": reverse("studio:login"),
        "settings_url": reverse("studio:account_settings"),
        "avatar_url": _avatar_url(request),
        "accent_color": _accent_color(request),
        "ui_accent_color": _ui_accent_color(request),
        "theme_mode": _theme_mode(request),
    }

    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        template = "studio/dashboard_files_section.html" if normalized in {"files", "storage"} else "studio/dashboard_tasks_section.html"
        html = render_to_string(template, context, request=request)
        return JsonResponse({"html": html, "total_outputs": total_outputs, "page": page, "pages": pages, "query": query})

    return render(request, "studio/dashboard_detail.html", context)


@require_POST
def start_convert(request: HttpRequest) -> JsonResponse:
    try:
        owner_id, guest_key = _workspace_identity(request)
        upload = _require_file(request, "file")
        target_format = request.POST.get("target_format", "webp")
        _validate_convert_target(upload, target_format, getattr(request, "interface_language", "en"))
        source = _save_upload(upload, "convert")
        job = actions.start_conversion_job(
            source=source,
            original_name=upload.name or "upload",
            content_type=getattr(upload, "content_type", ""),
            target_format=target_format,
            output_name=request.POST.get("output_name", "converted"),
            image_mode=request.POST.get("image_mode", "balanced"),
            owner_id=owner_id,
            guest_key=guest_key,
        )
        return _job_json(job)
    except Exception as exc:
        return _error_json(exc)


@require_POST
def start_youtube(request: HttpRequest) -> JsonResponse:
    try:
        owner_id, guest_key = _workspace_identity(request)
        job = actions.start_youtube_job(
            url=request.POST.get("url", ""),
            mode=request.POST.get("mode", "regular"),
            owner_id=owner_id,
            guest_key=guest_key,
            ai_improve=_post_bool(request, "ai_improve"),
            clip_count=int(request.POST.get("clip_count") or 10),
            processing_speed=request.POST.get("processing_speed", "auto"),
        )
        return _job_json(job)
    except Exception as exc:
        return _error_json(exc)


@require_POST
def start_cover(request: HttpRequest) -> JsonResponse:
    try:
        owner_id, guest_key = _workspace_identity(request)
        upload = _require_file(request, "file")
        source = _save_upload(upload, "cover")
        variants = int(request.POST.get("variants", "1") or 1)
        job = actions.start_cover_job(
            source=source,
            original_name=upload.name or "video",
            title=request.POST.get("title", ""),
            variants=variants,
            owner_id=owner_id,
            guest_key=guest_key,
            ai_cover=_post_bool(request, "ai_cover"),
        )
        return _job_json(job)
    except Exception as exc:
        return _error_json(exc)


@require_POST
def start_subtitles(request: HttpRequest) -> JsonResponse:
    try:
        owner_id, guest_key = _workspace_identity(request)
        upload = _require_file(request, "file")
        source = _save_upload(upload, "subtitles")
        job = actions.start_subtitle_job(
            source=source,
            original_name=upload.name or "video",
            title=request.POST.get("title", ""),
            style=request.POST.get("style", "pop"),
            language=request.POST.get("language", "auto"),
            owner_id=owner_id,
            guest_key=guest_key,
            ai_transcription=_post_bool(request, "ai_transcription"),
        )
        return _job_json(job)
    except Exception as exc:
        return _error_json(exc)


@require_POST
def start_package(request: HttpRequest) -> JsonResponse:
    try:
        owner_id, guest_key = _workspace_identity(request)
        upload = _require_file(request, "file")
        source = _save_upload(upload, "package")
        job = actions.start_package_job(
            source=source,
            original_name=upload.name or "video",
            title=request.POST.get("title", ""),
            style=request.POST.get("style", "pop"),
            language=request.POST.get("language", "auto"),
            owner_id=owner_id,
            guest_key=guest_key,
            ai_transcription=_post_bool(request, "ai_transcription"),
            ai_cover=_post_bool(request, "ai_cover"),
        )
        return _job_json(job)
    except Exception as exc:
        return _error_json(exc)


@require_POST
def start_resume(request: HttpRequest) -> JsonResponse:
    try:
        owner_id, guest_key = _workspace_identity(request)
        data = {field: request.POST.get(field, "").strip() for field in RESUME_FIELDS}
        data["lang"] = getattr(request, "interface_language", "en")
        photo = request.FILES.get("photo")
        if photo:
            data["photo_path"] = str(_save_upload(photo, "resume_photo"))
            data["photo_crop"] = request.POST.get("photo_crop", "").strip()
        elif request.POST.get("use_account_avatar") == "1" and request.user.is_authenticated:
            try:
                avatar_path = request.user.studio_profile.avatar_path
                if avatar_path and Path(avatar_path).exists():
                    data["photo_path"] = avatar_path
            except AccountProfile.DoesNotExist:
                pass
        job = actions.start_resume_job(data, request.POST.get("template", "1"), owner_id, guest_key)
        return _job_json(job)
    except Exception as exc:
        return _error_json(exc)


@require_POST
def resume_ai_rewrite(request: HttpRequest) -> JsonResponse:
    language = getattr(request, "interface_language", "en")
    payload = _resume_ai_payload(request, language)
    try:
        result = openai_ai.rewrite_resume_block(payload)
        if not result.get("text"):
            raise ValueError("OpenAI returned empty text")
        return JsonResponse({"ok": True, "used_ai": True, **result})
    except Exception as exc:
        result = _local_resume_rewrite(payload)
        return JsonResponse({"ok": True, "used_ai": False, "fallback_reason": str(exc)[:300], **result})


@require_POST
def resume_ai_match(request: HttpRequest) -> JsonResponse:
    language = getattr(request, "interface_language", "en")
    payload = _resume_ai_payload(request, language)
    try:
        result = openai_ai.analyze_resume_match(payload)
        return JsonResponse({"ok": True, "used_ai": True, **result})
    except Exception as exc:
        result = _local_resume_match(payload)
        return JsonResponse({"ok": True, "used_ai": False, "fallback_reason": str(exc)[:300], **result})


@require_POST
def resume_ai_cover_letter(request: HttpRequest) -> JsonResponse:
    language = getattr(request, "interface_language", "en")
    payload = _resume_ai_payload(request, language)
    try:
        result = openai_ai.generate_resume_cover_letter(payload)
        if not result.get("letter"):
            raise ValueError("OpenAI returned empty letter")
        return JsonResponse({"ok": True, "used_ai": True, **result})
    except Exception as exc:
        result = _local_resume_cover_letter(payload)
        return JsonResponse({"ok": True, "used_ai": False, "fallback_reason": str(exc)[:300], **result})


@require_POST
def start_originality(request: HttpRequest) -> JsonResponse:
    try:
        owner_id, guest_key = _workspace_identity(request)
        language = getattr(request, "interface_language", "en")
        mode_code = _originality_mode_code(request.POST.get("mode"))
        mode = ORIGINALITY_MODES[mode_code]
        text = str(request.POST.get("text", "") or "")
        upload = request.FILES.get("file")
        source_name = ""
        if upload:
            source_name = upload.name or "document"
            if int(getattr(upload, "size", 0) or 0) > ORIGINALITY_MAX_UPLOAD_BYTES:
                raise ValueError(_originality_runtime_text("file_too_large", language))
            extracted = _extract_originality_upload(upload, language)
            text = f"{text}\n\n{extracted}" if text.strip() else extracted
        text = _clean_originality_text(text)
        if len(text) < 40:
            raise ValueError(translate("originality_empty", language))
        max_chars = int(mode["max_chars"])
        truncated = len(text) > max_chars
        analysis_text = text[:max_chars]
        analysis = _analyze_originality_text(analysis_text, language, source_name, truncated)
        analysis["check"] = _originality_check_metadata(mode_code, analysis_text)
        record = _create_originality_record(analysis, analysis_text, owner_id, guest_key, language, source_name)
        return JsonResponse({"analysis": analysis, "job": _attach_output_urls(actions._serialize_job_record(record))})
    except Exception as exc:
        return _error_json(exc)


@require_POST
def originality_document_preview(request: HttpRequest) -> JsonResponse:
    try:
        language = getattr(request, "interface_language", "en")
        upload = request.FILES.get("file")
        if not upload:
            raise ValueError(translate("file_missing", language))
        if int(getattr(upload, "size", 0) or 0) > ORIGINALITY_MAX_UPLOAD_BYTES:
            raise ValueError(_originality_runtime_text("file_too_large", language))
        extracted = _extract_originality_upload(upload, language)
        text = _clean_originality_text(extracted)
        preview_text = re.sub(r"\n{3,}", "\n\n", str(extracted or "").strip())
        words = len(_originality_words(text))
        return JsonResponse(
            {
                "ok": True,
                "name": upload.name or "document",
                "size": int(getattr(upload, "size", 0) or 0),
                "text": (preview_text or text)[:12_000],
                "truncated": len(preview_text or text) > 12_000,
                "words": words,
            }
        )
    except Exception as exc:
        return _error_json(exc)


@require_GET
def job_detail(request: HttpRequest, job_id: str):
    owner_id, guest_key = _workspace_identity(request)
    job = actions.get_job(job_id, owner_id, guest_key)
    if not job:
        raise Http404("Job not found")
    language = getattr(request, "interface_language", "en")
    prepared_job = _localize_job(_attach_output_urls(job), language)
    originality_context = _originality_detail_context(request, prepared_job, owner_id, guest_key, language)
    display_outputs = originality_context.get("display_outputs")
    if not isinstance(display_outputs, list):
        display_outputs = prepared_job.get("outputs", [])
    return render(
        request,
        "studio/job_detail.html",
        {
            "job": prepared_job,
            "display_outputs": display_outputs,
            "events": _localize_events(actions.get_job_events(job_id, owner_id, guest_key), language),
            "has_access": user_has_active_access(request.user),
            "active_until": active_access_until(request.user),
            "is_guest": not request.user.is_authenticated,
            "display_name": _display_name(request),
            "checkout_url": _checkout_url(request),
            "pricing_url": reverse("billing:pricing"),
            "login_url": reverse("studio:login"),
            "settings_url": reverse("studio:account_settings"),
            "avatar_url": _avatar_url(request),
            "accent_color": _accent_color(request),
            "ui_accent_color": _ui_accent_color(request),
            "theme_mode": _theme_mode(request),
            **originality_context,
        },
    )


@require_GET
def job_status(request: HttpRequest, job_id: str) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    job = actions.get_job(job_id, owner_id, guest_key)
    if not job:
        raise Http404("Job not found")
    return _job_json(job)


@require_POST
def send_originality_report(request: HttpRequest, job_id: str) -> JsonResponse:
    try:
        owner_id, guest_key = _workspace_identity(request)
        language = getattr(request, "interface_language", "en")
        record = _job_record_for_workspace(job_id, owner_id, guest_key)
        if not record or record.kind != "originality":
            raise Http404("Report not found")
        email = _clean_email(request.POST.get("email") or (request.user.email if request.user.is_authenticated else ""))
        try:
            validate_email(email)
        except ValidationError as exc:
            raise ValueError(translate("checkout_email_invalid", language)) from exc
        output = _originality_html_output(record)
        if not output:
            raise ValueError(translate("file_missing", language))
        analysis = _load_originality_analysis(record)
        if analysis:
            _refresh_originality_html_report(record, analysis, language, output)
        _send_originality_email(request, record, output, email, language)
        return JsonResponse({"ok": True, "message": translate("originality_email_sent", language)})
    except Http404:
        raise
    except Exception as exc:
        return _error_json(exc)


@require_GET
def originality_shared_report(request: HttpRequest, token: str):
    record = _originality_record_by_share_token(token)
    if not record:
        raise Http404("Report not found")
    language = getattr(request, "interface_language", "en")
    analysis = _load_originality_analysis(record)
    if analysis:
        _refresh_originality_html_report(record, analysis, language)
    job = _attach_output_urls(actions._serialize_job_record(record))
    html_output = _originality_payload_output(job, "html")
    return render(
        request,
        "studio/originality_shared.html",
        {
            "record": record,
            "job": job,
            "analysis": analysis,
            "overall": analysis.get("overall", {}) if isinstance(analysis.get("overall"), dict) else {},
            "source": analysis.get("source", {}) if isinstance(analysis.get("source"), dict) else {},
            "metrics": analysis.get("metrics", []) if isinstance(analysis.get("metrics"), list) else [],
            "check": analysis.get("check", {}) if isinstance(analysis.get("check"), dict) else {},
            "score_degrees": round(max(0, min(100, int((analysis.get("overall", {}) if isinstance(analysis.get("overall"), dict) else {}).get("score") or 0))) * 3.6, 1),
            "report_url": html_output.get("preview_url", "") if html_output else "",
            "accent_color": _accent_color(request),
            "ui_accent_color": _ui_accent_color(request),
            "theme_mode": _theme_mode(request),
        },
    )


@require_POST
def edit_output_design(request: HttpRequest, job_id: str, index: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    job = actions.get_job(job_id, owner_id, guest_key)
    output = actions.get_output(job_id, index, owner_id, guest_key)
    if not job or not output:
        raise Http404("Cover output not found")
    if not _output_can_edit_design({"label": output.label, "name": output.name, "media_type": output.media_type}, job):
        raise Http404("Cover output not found")
    record = _job_record_for_workspace(job_id, owner_id, guest_key)
    if not record:
        raise Http404("Job not found")

    output_key = str(output.path.resolve())
    params = _job_record_params(record)
    design_map = params.get("design_projects") if isinstance(params.get("design_projects"), dict) else {}
    existing_id = design_map.get(output_key) if isinstance(design_map, dict) else None
    if existing_id:
        existing = _design_project_queryset(owner_id, guest_key).filter(id=existing_id).first()
        if existing:
            return JsonResponse(_design_open_payload(existing, owner_id, guest_key))

    project = _create_design_project_from_output(request, record, output, job)
    design_map = dict(design_map) if isinstance(design_map, dict) else {}
    design_map[output_key] = project.id
    params["design_projects"] = design_map
    record.params_json = json.dumps(params, ensure_ascii=False, default=str)
    record.save(update_fields=["params_json", "updated_at"])
    return JsonResponse(_design_open_payload(project, owner_id, guest_key))


@require_POST
def edit_output_video(request: HttpRequest, job_id: str, index: int) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    job = actions.get_job(job_id, owner_id, guest_key)
    output = actions.get_output(job_id, index, owner_id, guest_key)
    if not job or not output:
        raise Http404("Video output not found")
    if not _output_can_edit_video({"label": output.label, "name": output.name, "media_type": output.media_type}, job):
        raise Http404("Video output not found")
    record = _job_record_for_workspace(job_id, owner_id, guest_key)
    if not record:
        raise Http404("Job not found")

    output_key = str(output.path.resolve())
    params = _job_record_params(record)
    video_map = params.get("video_projects") if isinstance(params.get("video_projects"), dict) else {}
    existing_id = video_map.get(output_key) if isinstance(video_map, dict) else None
    if existing_id:
        existing = _video_project_queryset(owner_id, guest_key).filter(id=existing_id).first()
        if existing:
            _ensure_project_uses_editable_subtitle_source(existing, record, output.path)
            _append_job_output_subtitles_to_video_project(existing, record, output.path)
            return JsonResponse(_video_open_payload(existing, owner_id, guest_key))

    project = _create_video_project_from_output(request, record, output, job)
    video_map = dict(video_map) if isinstance(video_map, dict) else {}
    video_map[output_key] = project.id
    params["video_projects"] = video_map
    record.params_json = json.dumps(params, ensure_ascii=False, default=str)
    record.save(update_fields=["params_json", "updated_at"])
    return JsonResponse(_video_open_payload(project, owner_id, guest_key))


@require_POST
def repeat_job(request: HttpRequest, job_id: str) -> JsonResponse:
    try:
        owner_id, guest_key = _workspace_identity(request)
        job = actions.repeat_job(job_id, owner_id, guest_key)
    except Exception as exc:
        return _error_json(exc)
    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        return _job_json(job)
    return redirect("studio:job_detail", job_id=job["id"])


@require_POST
def cancel_job(request: HttpRequest, job_id: str):
    try:
        owner_id, guest_key = _workspace_identity(request)
        job = actions.cancel_job(job_id, owner_id, guest_key)
    except Exception as exc:
        return _error_json(exc)
    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        return _job_json(job)
    return redirect("studio:dashboard_detail", section="active")


@require_POST
def pause_job(request: HttpRequest, job_id: str):
    try:
        owner_id, guest_key = _workspace_identity(request)
        job = actions.pause_job(job_id, owner_id, guest_key)
    except Exception as exc:
        return _error_json(exc)
    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        return _job_json(job)
    return redirect("studio:dashboard_detail", section="active")


@require_POST
def resume_job(request: HttpRequest, job_id: str):
    try:
        owner_id, guest_key = _workspace_identity(request)
        job = actions.resume_job(job_id, owner_id, guest_key)
    except Exception as exc:
        return _error_json(exc)
    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        return _job_json(job)
    return redirect("studio:dashboard_detail", section="active")


@require_POST
def delete_job(request: HttpRequest, job_id: str):
    try:
        owner_id, guest_key = _workspace_identity(request)
        deleted = actions.delete_job_and_media(job_id, owner_id, guest_key)
        if not deleted:
            raise Http404("Job not found")
    except Exception as exc:
        return _error_json(exc)
    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        stats = actions.get_account_stats(owner_id, guest_key)
        stats.update(_storage_quota(request, stats))
        return JsonResponse({"ok": True, "account_stats": stats})
    return redirect("studio:index")


@require_GET
def preview_output(request: HttpRequest, job_id: str, index: int) -> HttpResponse:
    owner_id, guest_key = _workspace_identity(request)
    output = actions.get_output(job_id, index, owner_id, guest_key)
    if not output:
        raise Http404("Output not found")
    if str(output.media_type).startswith("video/"):
        return _range_file_response(request, output.path, output.media_type, output.name)
    if _is_subtitle_output(output.name, output.media_type):
        return _subtitle_preview_response(output.path, output.name)
    return FileResponse(
        output.path.open("rb"),
        as_attachment=False,
        filename=output.name,
        content_type=output.media_type,
    )


def _is_subtitle_output(name: str, media_type: str = "") -> bool:
    suffix = Path(str(name or "")).suffix.lower()
    return suffix in {".ass", ".srt", ".vtt"} or str(media_type or "").lower() in {"text/x-ssa", "text/x-ass"}


def _subtitle_preview_response(path: Path, name: str) -> HttpResponse:
    try:
        raw = path.read_bytes()
    except OSError as exc:
        raise Http404("Subtitle file not found") from exc
    text = _decode_subtitle_bytes(raw)
    cues = _parse_subtitle_cues(text, name)
    rows = []
    for index, cue in enumerate(cues[:500], start=1):
        start = _format_vtt_time(float(cue.get("start") or 0))
        end = _format_vtt_time(float(cue.get("end") or 0))
        caption = html.escape(str(cue.get("text") or "")).replace("\n", "<br>")
        rows.append(
            f'<article><span>{index:02d}</span><time>{html.escape(start)} - {html.escape(end)}</time><p>{caption}</p></article>'
        )
    if not rows:
        rows.append("<article><span>--</span><time>0:00</time><p>No readable subtitle cues found.</p></article>")
    raw_link = html.escape(path.name)
    page = f"""<!doctype html>
<html lang="uk">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(name)}</title>
  <style>
    :root {{ color-scheme: light; font-family: Inter, ui-sans-serif, system-ui, -apple-system, "Segoe UI", sans-serif; }}
    body {{ margin: 0; background: #f7fafc; color: #111827; }}
    main {{ display: grid; gap: 14px; padding: 18px; }}
    header {{ display: flex; align-items: end; justify-content: space-between; gap: 14px; border-bottom: 1px solid #dbe7f4; padding-bottom: 12px; }}
    h1 {{ margin: 0; font-size: 18px; line-height: 1.2; overflow-wrap: anywhere; }}
    header span {{ color: #64748b; font-size: 12px; font-weight: 800; }}
    section {{ display: grid; gap: 9px; }}
    article {{ display: grid; grid-template-columns: 44px 150px minmax(0, 1fr); gap: 12px; align-items: start; border: 1px solid #d8e4f0; border-radius: 14px; padding: 12px; background: #fff; box-shadow: 0 10px 26px rgba(15,23,42,.05); }}
    article span {{ display: inline-grid; width: 32px; height: 32px; place-items: center; border-radius: 999px; background: #eef6ff; color: #2563eb; font-size: 12px; font-weight: 900; }}
    time {{ color: #516176; font-size: 12px; font-weight: 850; white-space: nowrap; }}
    p {{ margin: 0; color: #0f172a; font-size: 15px; line-height: 1.5; overflow-wrap: anywhere; }}
    @media (max-width: 640px) {{ main {{ padding: 12px; }} article {{ grid-template-columns: 36px minmax(0, 1fr); }} time {{ grid-column: 2; }} p {{ grid-column: 1 / -1; }} }}
  </style>
</head>
<body>
  <main>
    <header><div><span>Subtitle preview</span><h1>{html.escape(name)}</h1></div><span>{len(cues)} cues</span></header>
    <section>{"".join(rows)}</section>
  </main>
</body>
</html>"""
    response = HttpResponse(page, content_type="text/html; charset=utf-8")
    response["X-Subtitle-Source"] = raw_link
    return response


def _range_file_response(request: HttpRequest, path: Path, content_type: str, filename: str) -> HttpResponse:
    file_size = path.stat().st_size
    range_header = request.headers.get("Range", "").strip()
    if not range_header.startswith("bytes="):
        response = FileResponse(path.open("rb"), as_attachment=False, filename=filename, content_type=content_type)
        response["Accept-Ranges"] = "bytes"
        response["Content-Length"] = str(file_size)
        return response

    start_text, _, end_text = range_header.removeprefix("bytes=").partition("-")
    try:
        if start_text:
            start = int(start_text)
            end = int(end_text) if end_text else file_size - 1
        else:
            suffix_length = int(end_text)
            start = max(0, file_size - suffix_length)
            end = file_size - 1
    except ValueError:
        start, end = 0, file_size - 1

    start = max(0, min(start, file_size - 1))
    end = max(start, min(end, file_size - 1))
    length = end - start + 1

    def stream():
        with path.open("rb") as handle:
            handle.seek(start)
            remaining = length
            while remaining > 0:
                chunk = handle.read(min(1024 * 1024, remaining))
                if not chunk:
                    break
                remaining -= len(chunk)
                yield chunk

    response = StreamingHttpResponse(stream(), status=206, content_type=content_type)
    response["Accept-Ranges"] = "bytes"
    response["Content-Range"] = f"bytes {start}-{end}/{file_size}"
    response["Content-Length"] = str(length)
    response["Content-Disposition"] = f'inline; filename="{filename}"'
    return response


@require_GET
def download_output(request: HttpRequest, job_id: str, index: int) -> FileResponse:
    owner_id, guest_key = _workspace_identity(request)
    output = actions.get_output(job_id, index, owner_id, guest_key)
    if not output:
        raise Http404("Output not found")
    if not user_has_active_access(request.user):
        return redirect(_checkout_url(request))
    return FileResponse(
        output.path.open("rb"),
        as_attachment=True,
        filename=output.name,
        content_type=output.media_type,
    )


@require_GET
def download_all_outputs(request: HttpRequest, job_id: str) -> FileResponse:
    owner_id, guest_key = _workspace_identity(request)
    if not actions.get_job(job_id, owner_id, guest_key):
        raise Http404("Job not found")
    if not user_has_active_access(request.user):
        return redirect(_checkout_url(request))
    try:
        archive = actions.zip_job_outputs(job_id, owner_id, guest_key)
    except Exception as exc:
        raise Http404(str(exc)) from exc
    return FileResponse(
        archive.open("rb"),
        as_attachment=True,
        filename=archive.name,
        content_type="application/zip",
    )


@require_GET
def health(request: HttpRequest) -> JsonResponse:
    owner_id, guest_key = _workspace_identity(request)
    return JsonResponse({
        "status": "ok",
        "native_acceleration": actions.native_status(),
        "persistent_job_queue": bool(settings.persistent_job_queue),
        "queue": actions.queue_status(owner_id, guest_key),
    })


def _workspace_identity(request: HttpRequest) -> tuple[int | None, str]:
    if request.user.is_authenticated:
        return request.user.id, ""
    return None, _guest_key(request)


def _guest_key(request: HttpRequest) -> str:
    if not request.session.session_key:
        request.session.save()
    return request.session.session_key or ""


SHARE_VIEW_ROLES = {WorkspaceShare.ROLE_VIEWER, WorkspaceShare.ROLE_EDITOR}
SHARE_EDIT_ROLES = {WorkspaceShare.ROLE_EDITOR}
SHARE_RESOURCE_TYPES = {WorkspaceShare.RESOURCE_DESIGN, WorkspaceShare.RESOURCE_VIDEO, WorkspaceShare.RESOURCE_MUSIC}


def _active_share_filter() -> Q:
    return Q(status=WorkspaceShare.STATUS_ACCEPTED, expires_at__gt=timezone.now())


def _shared_resource_ids(resource_type: str, owner_id: int | None) -> list[int]:
    if not owner_id:
        return []
    return list(
        WorkspaceShare.objects.filter(_active_share_filter(), resource_type=resource_type, invited_user_id=owner_id)
        .values_list("resource_id", flat=True)
    )


def _video_owner_queryset(owner_id: int | None, guest_key: str):
    queryset = VideoEditorProject.objects.all()
    if owner_id:
        return queryset.filter(owner_id=owner_id)
    return queryset.filter(owner__isnull=True, guest_key=guest_key)


def _design_owner_queryset(owner_id: int | None, guest_key: str):
    queryset = DesignerProject.objects.all()
    if owner_id:
        return queryset.filter(owner_id=owner_id)
    return queryset.filter(owner__isnull=True, guest_key=guest_key)


def _video_project_queryset(owner_id: int | None, guest_key: str):
    queryset = VideoEditorProject.objects.all()
    if owner_id:
        return queryset.filter(Q(owner_id=owner_id) | Q(id__in=_shared_resource_ids(WorkspaceShare.RESOURCE_VIDEO, owner_id))).distinct()
    return queryset.filter(owner__isnull=True, guest_key=guest_key)


def _design_project_queryset(owner_id: int | None, guest_key: str):
    queryset = DesignerProject.objects.all()
    if owner_id:
        return queryset.filter(Q(owner_id=owner_id) | Q(id__in=_shared_resource_ids(WorkspaceShare.RESOURCE_DESIGN, owner_id))).distinct()
    return queryset.filter(owner__isnull=True, guest_key=guest_key)


def _music_owner_queryset(owner_id: int | None, guest_key: str):
    queryset = MusicEditorProject.objects.all()
    if owner_id:
        return queryset.filter(owner_id=owner_id)
    return queryset.filter(owner__isnull=True, guest_key=guest_key)


def _music_project_queryset(owner_id: int | None, guest_key: str):
    queryset = MusicEditorProject.objects.all()
    if owner_id:
        return queryset.filter(Q(owner_id=owner_id) | Q(id__in=_shared_resource_ids(WorkspaceShare.RESOURCE_MUSIC, owner_id))).distinct()
    return queryset.filter(owner__isnull=True, guest_key=guest_key)


def _resource_access_role(resource_type: str, project, owner_id: int | None, guest_key: str = "") -> str:
    if not project:
        return ""
    cached_role = getattr(project, "_access_role", "")
    if cached_role:
        return cached_role
    if owner_id and project.owner_id == owner_id:
        return "owner"
    if not owner_id and project.owner_id is None and project.guest_key == guest_key:
        return "owner"
    if owner_id:
        share = WorkspaceShare.objects.filter(
            _active_share_filter(),
            resource_type=resource_type,
            resource_id=project.id,
            invited_user_id=owner_id,
        ).first()
        if share:
            return share.role
    return ""


def _attach_access_roles(projects: list, resource_type: str, owner_id: int | None, guest_key: str = "") -> list:
    if not projects:
        return projects
    for project in projects:
        if owner_id and project.owner_id == owner_id:
            project._access_role = "owner"
        elif not owner_id and project.owner_id is None and project.guest_key == guest_key:
            project._access_role = "owner"
        else:
            project._access_role = ""
    if owner_id:
        shared_ids = [project.id for project in projects if getattr(project, "_access_role", "") != "owner"]
        if shared_ids:
            shares = WorkspaceShare.objects.filter(
                _active_share_filter(),
                resource_type=resource_type,
                resource_id__in=shared_ids,
                invited_user_id=owner_id,
            )
            roles = {share.resource_id: share.role for share in shares}
            for project in projects:
                if getattr(project, "_access_role", "") != "owner":
                    project._access_role = roles.get(project.id, "")
    return projects


def _resource_can_view(resource_type: str, project, owner_id: int | None, guest_key: str = "") -> bool:
    role = _resource_access_role(resource_type, project, owner_id, guest_key)
    return role == "owner" or role in SHARE_VIEW_ROLES


def _resource_can_edit(resource_type: str, project, owner_id: int | None, guest_key: str = "") -> bool:
    role = _resource_access_role(resource_type, project, owner_id, guest_key)
    return role == "owner" or role in SHARE_EDIT_ROLES


def _resource_is_owner(project, owner_id: int | None, guest_key: str = "") -> bool:
    if not project:
        return False
    if owner_id:
        return project.owner_id == owner_id
    return project.owner_id is None and project.guest_key == guest_key


def _clean_email(value: str) -> str:
    return " ".join(value.strip().lower().split())


def _share_resource(resource_type: str, resource_id: int):
    if resource_type == WorkspaceShare.RESOURCE_DESIGN:
        return DesignerProject.objects.prefetch_related("assets").filter(id=resource_id).first()
    if resource_type == WorkspaceShare.RESOURCE_VIDEO:
        return VideoEditorProject.objects.prefetch_related("assets").filter(id=resource_id).first()
    if resource_type == WorkspaceShare.RESOURCE_MUSIC:
        return MusicEditorProject.objects.prefetch_related("assets").filter(id=resource_id).first()
    return None


def _share_resource_url(resource_type: str, resource_id: int) -> str:
    if resource_type == WorkspaceShare.RESOURCE_DESIGN:
        return f"{reverse('studio:designer')}?{urlencode({'project': resource_id})}"
    if resource_type == WorkspaceShare.RESOURCE_VIDEO:
        return f"{reverse('studio:video_editor')}?{urlencode({'project': resource_id})}"
    if resource_type == WorkspaceShare.RESOURCE_MUSIC:
        return f"{reverse('studio:music_editor')}?{urlencode({'project': resource_id})}"
    return reverse("studio:index")


def _share_resource_label(resource_type: str, language: str | None = None) -> str:
    return {
        WorkspaceShare.RESOURCE_DESIGN: translate("designer_project", language),
        WorkspaceShare.RESOURCE_VIDEO: translate("video_editor_nav", language),
        WorkspaceShare.RESOURCE_MUSIC: translate("music_editor_nav", language),
    }.get(resource_type, translate("project_invite", language))


def _share_resource_preview(resource_type: str, resource) -> str:
    if resource_type == WorkspaceShare.RESOURCE_DESIGN and getattr(resource, "preview_path", ""):
        return reverse("studio:design_project_preview", args=[resource.id])
    if resource_type == WorkspaceShare.RESOURCE_VIDEO:
        state = resource.state_json or {}
        first_thumb = next((asset for asset in _prefetched_assets(resource) if asset.thumbnail_path), None)
        if first_thumb:
            return reverse("studio:video_project_asset_thumbnail", args=[resource.id, first_thumb.id])
        if isinstance(state, dict):
            return str(state.get("thumbnail") or "")
    return ""


def _workspace_share_payload(request: HttpRequest, share: WorkspaceShare) -> dict[str, object]:
    expires_delta = share.expires_at - timezone.now()
    days_left = max(0, expires_delta.days + (1 if expires_delta.seconds else 0))
    language = getattr(request, "interface_language", "en")
    return {
        "id": share.id,
        "resource_type": share.resource_type,
        "resource_id": share.resource_id,
        "email": share.email,
        "role": share.role,
        "status": share.status,
        "invite_url": request.build_absolute_uri(reverse("studio:workspace_invite", args=[share.token])),
        "expires_at": share.expires_at.isoformat(),
        "expires_label": f"{days_left} {translate('days', language)}" if days_left else translate("expires_today", language),
        "created_at": share.created_at.isoformat(),
        "updated_at": share.updated_at.isoformat(),
    }


def _workspace_share_resource_payload(request: HttpRequest, resource_type: str, resource) -> dict[str, object]:
    language = getattr(request, "interface_language", "en")
    preview_url = _share_resource_preview(resource_type, resource)
    if preview_url and preview_url.startswith("/"):
        preview_url = request.build_absolute_uri(preview_url)
    elif preview_url and not preview_url.startswith(("http://", "https://")):
        preview_url = ""
    return {
        "title": getattr(resource, "title", "Shared project"),
        "label": _share_resource_label(resource_type, language),
        "preview_url": preview_url,
    }


def _invite_context(request: HttpRequest, share: WorkspaceShare | None, status: str) -> dict[str, object]:
    language = getattr(request, "interface_language", "en")
    resource = _share_resource(share.resource_type, share.resource_id) if share else None
    next_url = reverse("studio:workspace_invite", args=[share.token]) if share else reverse("studio:index")
    return {
        "share": share,
        "resource": resource,
        "resource_title": getattr(resource, "title", "Shared project") if resource else "Shared project",
        "resource_label": _share_resource_label(share.resource_type, language) if share else translate("project_invite", language),
        "preview_url": _share_resource_preview(share.resource_type, resource) if share and resource else "",
        "owner_display": (share.owner.first_name or share.owner.email) if share else "CherryX user",
        "role": share.role if share else "",
        "status": status,
        "login_url": f"{reverse('studio:login')}?{urlencode({'next': next_url})}",
        "register_url": f"{reverse('studio:register')}?{urlencode({'next': next_url})}",
        "accent_color": _accent_color(request),
        "ui_accent_color": _ui_accent_color(request),
        "theme_mode": _theme_mode(request),
    }


def _send_workspace_invite(request: HttpRequest, share: WorkspaceShare) -> None:
    resource = _share_resource(share.resource_type, share.resource_id)
    if not resource:
        return
    preview_url = _share_resource_preview(share.resource_type, resource)
    if preview_url and preview_url.startswith("/"):
        preview_url = request.build_absolute_uri(preview_url)
    elif preview_url and not preview_url.startswith(("http://", "https://")):
        preview_url = ""
    context = {
        "share": share,
        "resource_title": getattr(resource, "title", "Shared project"),
        "resource_label": _share_resource_label(share.resource_type, getattr(request, "interface_language", "en")),
        "owner_display": share.owner.first_name or share.owner.email,
        "invite_url": request.build_absolute_uri(reverse("studio:workspace_invite", args=[share.token])),
        "preview_url": preview_url,
        "role": share.role,
    }
    html = render_to_string("studio/email/invite.html", context)
    message = EmailMultiAlternatives(
        subject=f"{context['owner_display']} invited you to CherryX",
        body=strip_tags(html),
        from_email=None,
        to=[share.email],
    )
    message.attach_alternative(html, "text/html")
    message.send(fail_silently=True)


def _accept_pending_workspace_shares(user) -> None:
    email = _clean_email(getattr(user, "email", "") or "")
    if not email:
        return
    WorkspaceShare.objects.filter(
        email=email,
        status=WorkspaceShare.STATUS_PENDING,
        expires_at__gt=timezone.now(),
    ).update(invited_user=user, status=WorkspaceShare.STATUS_ACCEPTED, updated_at=timezone.now())


def _design_project_payload(project: DesignerProject | None, include_state: bool = True, owner_id: int | None = None, guest_key: str = "") -> dict[str, object]:
    if not project:
        return {}
    state = project.state_json or {}
    assets = _prefetched_assets(project)
    objects = _state_list(state, "objects")
    vectors = _state_list(state, "vectors")
    frames = [item for item in objects if isinstance(item, dict) and item.get("type") == "frame"]
    object_count = len(objects) if objects else int(getattr(project, "object_count", 0) or 0)
    vector_count = len(vectors)
    layer_count = object_count + vector_count
    has_preview_content = bool(project.preview_path or object_count or vector_count)
    preview_url = ""
    if has_preview_content:
        preview_url = f"{reverse('studio:design_project_preview', args=[project.id])}?v={int(project.updated_at.timestamp())}"
    access_role = _resource_access_role(WorkspaceShare.RESOURCE_DESIGN, project, owner_id, guest_key) if owner_id is not None or guest_key else "owner"
    is_owner = access_role == "owner"
    payload: dict[str, object] = {
        "id": project.id,
        "title": project.title,
        "preview_url": preview_url,
        "preview_focus": _design_project_preview_focus(state),
        "object_count": object_count,
        "frame_count": len(frames),
        "vector_count": vector_count,
        "layer_count": layer_count,
        "asset_count": int(getattr(project, "asset_count", 0) or len(assets)),
        "storage_bytes": project.storage_bytes,
        "storage_text": human_size(project.storage_bytes),
        "created_at": project.created_at.isoformat(),
        "updated_at": project.updated_at.isoformat(),
        "assets": [_design_asset_payload(asset) for asset in assets],
        "access_role": access_role,
        "is_owner": is_owner,
        "can_edit": is_owner or access_role == WorkspaceShare.ROLE_EDITOR,
        "can_share": is_owner and owner_id is not None,
        "publish_url": f"{reverse('studio:community_publish')}?{urlencode({'source': 'design_project', 'id': project.id})}" if is_owner and owner_id is not None else "",
    }
    if include_state:
        payload["state"] = project.state_json or {}
    return payload


def _design_project_preview_focus(state: dict[str, object]) -> dict[str, float | int]:
    design_width = 9000.0
    design_height = 6400.0
    candidates: list[tuple[float, float, float, float, float, float, float]] = []

    def as_float(value: object, default: float = 0) -> float:
        try:
            return float(value)
        except (TypeError, ValueError):
            return default

    def clamp_percent(value: float) -> int:
        return int(round(max(8, min(92, value))))

    def add_box(min_x: float, min_y: float, max_x: float, max_y: float, weight: float) -> None:
        width = max_x - min_x
        height = max_y - min_y
        if width <= 0 or height <= 0 or weight <= 0:
            return
        center_x = min_x + width / 2
        center_y = min_y + height / 2
        candidates.append((center_x, center_y, weight, min_x, min_y, max_x, max_y))

    if not isinstance(state, dict):
        return {"x": 50, "y": 50, "scale": 1.18, "hover_scale": 1.34}

    objects = state.get("objects") if isinstance(state.get("objects"), list) else []
    for item in objects:
        if not isinstance(item, dict):
            continue
        if item.get("hidden") or item.get("isHidden"):
            continue
        x = as_float(item.get("x"))
        y = as_float(item.get("y"))
        width = as_float(item.get("w") or item.get("width"), 80)
        height = as_float(item.get("h") or item.get("height"), 80)
        item_type = str(item.get("type") or "")
        area_weight = max(14.0, min(260.0, (max(1.0, width * height) ** 0.5) / 2.0))
        if item_type == "frame":
            area_weight *= 0.28
        elif item_type == "text":
            text = str(item.get("text") or item.get("content") or "")
            area_weight *= 1.25 + min(0.65, len(text) / 140)
        elif item_type in {"image", "photo"} or item.get("src") or item.get("assetId"):
            area_weight *= 1.35
        add_box(x, y, x + width, y + height, area_weight)

    for key in ("vectors", "strokes"):
        strokes = state.get(key) if isinstance(state.get(key), list) else []
        for stroke in strokes:
            if not isinstance(stroke, dict):
                continue
            points = stroke.get("points") if isinstance(stroke.get("points"), list) else []
            coords = [(as_float(point.get("x")), as_float(point.get("y"))) for point in points if isinstance(point, dict)]
            if not coords:
                continue
            xs = [point[0] for point in coords]
            ys = [point[1] for point in coords]
            width = max(2.0, max(xs) - min(xs))
            height = max(2.0, max(ys) - min(ys))
            stroke_weight = max(18.0, min(320.0, len(coords) * 5.0 + ((width * height) ** 0.5) / 7.0))
            add_box(min(xs), min(ys), max(xs), max(ys), stroke_weight)

    if not candidates:
        return {"x": 50, "y": 50, "scale": 1.18, "hover_scale": 1.34}

    def distance(left: tuple[float, float, float, float, float, float, float], right: tuple[float, float, float, float, float, float, float]) -> float:
        return ((left[0] - right[0]) ** 2 + (left[1] - right[1]) ** 2) ** 0.5

    best = max(
        candidates,
        key=lambda candidate: sum(other[2] / (1 + distance(candidate, other) / 900.0) for other in candidates),
    )
    cluster = [candidate for candidate in candidates if distance(best, candidate) <= 1800.0] or [best]
    total_weight = sum(candidate[2] for candidate in cluster) or 1.0
    x = sum(candidate[0] * candidate[2] for candidate in cluster) / total_weight
    y = sum(candidate[1] * candidate[2] for candidate in cluster) / total_weight
    min_x = min(candidate[3] for candidate in cluster)
    min_y = min(candidate[4] for candidate in cluster)
    max_x = max(candidate[5] for candidate in cluster)
    max_y = max(candidate[6] for candidate in cluster)
    content_ratio = max((max_x - min_x) / design_width, (max_y - min_y) / design_height)
    if content_ratio < 0.08:
        scale = 2.35
    elif content_ratio < 0.16:
        scale = 1.95
    elif content_ratio < 0.28:
        scale = 1.55
    else:
        scale = 1.22
    scale = round(max(1.16, min(2.65, scale)), 2)

    return {
        "x": clamp_percent((x / design_width) * 100),
        "y": clamp_percent((y / design_height) * 100),
        "scale": scale,
        "hover_scale": round(min(2.9, scale + 0.18), 2),
    }


def _render_design_project_preview(project: DesignerProject) -> BytesIO | None:
    state = project.state_json if isinstance(project.state_json, dict) else {}
    objects = [item for item in _state_list(state, "objects") if isinstance(item, dict) and not item.get("hidden") and not item.get("isHidden")]
    vectors = [item for item in _state_list(state, "vectors") if isinstance(item, dict) and not item.get("hidden") and isinstance(item.get("points"), list) and item.get("points")]
    if not objects and not vectors:
        return None

    def as_float(value: object, default: float = 0.0) -> float:
        try:
            return float(value)
        except (TypeError, ValueError):
            return default

    boxes: list[tuple[float, float, float, float]] = []
    for item in objects:
        x = as_float(item.get("x"))
        y = as_float(item.get("y"))
        width = max(1.0, as_float(item.get("w") or item.get("width"), 80.0))
        height = max(1.0, as_float(item.get("h") or item.get("height"), 80.0))
        boxes.append((x, y, x + width, y + height))
    for vector in vectors:
        coords = [(as_float(point.get("x")), as_float(point.get("y"))) for point in vector.get("points", []) if isinstance(point, dict)]
        if not coords:
            continue
        xs = [point[0] for point in coords]
        ys = [point[1] for point in coords]
        pad = max(8.0, as_float(vector.get("width"), 5.0) * 2)
        boxes.append((min(xs) - pad, min(ys) - pad, max(xs) + pad, max(ys) + pad))
    if not boxes:
        return None

    min_x = min(box[0] for box in boxes)
    min_y = min(box[1] for box in boxes)
    max_x = max(box[2] for box in boxes)
    max_y = max(box[3] for box in boxes)
    content_w = max(1.0, max_x - min_x)
    content_h = max(1.0, max_y - min_y)
    pad = max(80.0, min(520.0, max(content_w, content_h) * 0.1))
    min_x = max(0.0, min_x - pad)
    min_y = max(0.0, min_y - pad)
    content_w += pad * 2
    content_h += pad * 2
    scale = min(720.0 / content_w, 480.0 / content_h)
    width = max(240, int(round(content_w * scale)))
    height = max(160, int(round(content_h * scale)))
    scale = min(width / content_w, height / content_h)

    def px(value: float) -> int:
        return int(round((value - min_x) * scale))

    def py(value: float) -> int:
        return int(round((value - min_y) * scale))

    def color(value: object, fallback: str = "#2563eb", opacity: float = 1.0) -> tuple[int, int, int, int]:
        raw = str(value or fallback).strip()
        if raw.lower() == "transparent":
            return (0, 0, 0, 0)
        if raw.startswith("#"):
            raw = raw[1:]
            if len(raw) == 3:
                raw = "".join(part * 2 for part in raw)
            if len(raw) >= 6:
                try:
                    return (
                        int(raw[0:2], 16),
                        int(raw[2:4], 16),
                        int(raw[4:6], 16),
                        int(max(0, min(255, round(255 * opacity)))),
                    )
                except ValueError:
                    pass
        return color(fallback, "#2563eb", opacity)

    image = Image.new("RGBA", (width, height), (248, 251, 255, 255))
    draw = ImageDraw.Draw(image, "RGBA")
    grid = max(12, int(round(48 * scale)))
    for grid_x in range(0, width, grid):
        draw.line((grid_x, 0, grid_x, height), fill=(37, 99, 235, 18), width=1)
    for grid_y in range(0, height, grid):
        draw.line((0, grid_y, width, grid_y), fill=(37, 99, 235, 18), width=1)

    try:
        font = ImageFont.truetype("arial.ttf", max(12, int(round(24 * scale))))
    except Exception:
        font = ImageFont.load_default()

    for item in objects:
        opacity = max(0.0, min(1.0, as_float(item.get("opacity"), 1.0)))
        x = as_float(item.get("x"))
        y = as_float(item.get("y"))
        width_value = max(1.0, as_float(item.get("w") or item.get("width"), 80.0))
        height_value = max(1.0, as_float(item.get("h") or item.get("height"), 80.0))
        box = (px(x), py(y), px(x + width_value), py(y + height_value))
        item_type = str(item.get("type") or "")
        fill = color(item.get("fill"), "#ffffff" if item_type == "frame" else "transparent", opacity)
        stroke = color(item.get("stroke"), "#2563eb", opacity)
        stroke_width = max(1, int(round(as_float(item.get("strokeWidth"), 2.0) * scale)))
        if item_type == "text":
            if fill[3]:
                draw.rounded_rectangle(box, radius=max(3, int(round(as_float(item.get("cornerRadius"), 10) * scale))), fill=fill)
            text = str(item.get("text") or item.get("name") or "Text")[:140]
            draw.text((box[0] + 8, box[1] + 8), text, fill=stroke, font=font)
        elif item_type == "shape" and item.get("shape") == "ellipse":
            draw.ellipse(box, fill=fill, outline=stroke, width=stroke_width)
        elif item_type == "shape" and item.get("shape") in {"line", "arrow"}:
            center_y = (box[1] + box[3]) // 2
            draw.line((box[0] + 6, center_y, box[2] - 6, center_y), fill=stroke, width=max(2, stroke_width))
        else:
            radius = max(0, int(round(as_float(item.get("cornerRadius"), 10 if item_type != "frame" else 0) * scale)))
            draw.rounded_rectangle(box, radius=radius, fill=fill, outline=stroke, width=stroke_width)

    for vector in vectors:
        points = [(px(as_float(point.get("x"))), py(as_float(point.get("y")))) for point in vector.get("points", []) if isinstance(point, dict)]
        if not points:
            continue
        stroke = color(vector.get("color"), "#2563eb", max(0.0, min(1.0, as_float(vector.get("opacity"), 1.0))))
        width_px = max(2, int(round(as_float(vector.get("width"), 5.0) * scale)))
        if len(points) == 1:
            x, y = points[0]
            draw.ellipse((x - width_px, y - width_px, x + width_px, y + width_px), fill=stroke)
        else:
            draw.line(points, fill=stroke, width=width_px, joint="curve")

    output = BytesIO()
    image.convert("RGB").save(output, "JPEG", quality=86, optimize=True)
    output.seek(0)
    return output


def _design_asset_payload(asset: DesignerAsset) -> dict[str, object]:
    return {
        "id": asset.id,
        "kind": asset.kind,
        "name": asset.original_name,
        "media_type": asset.media_type,
        "size": asset.size,
        "size_text": human_size(asset.size),
        "preview_url": reverse("studio:design_project_asset_preview", args=[asset.project_id, asset.id]),
    }


def _design_open_payload(project: DesignerProject, owner_id: int | None, guest_key: str) -> dict[str, object]:
    return {
        "project": _design_project_payload(project, owner_id=owner_id, guest_key=guest_key),
        "designer_url": f"{reverse('studio:designer')}?{urlencode({'project': project.id})}",
    }


def _video_open_payload(project: VideoEditorProject, owner_id: int | None, guest_key: str) -> dict[str, object]:
    return {
        "project": _video_project_payload(project, owner_id=owner_id, guest_key=guest_key),
        "video_editor_url": f"{reverse('studio:video_editor')}?{urlencode({'project': project.id})}",
    }


def _job_record_for_workspace(job_id: str, owner_id: int | None, guest_key: str):
    queryset = JobRecord.objects.filter(job_id=job_id)
    if owner_id is not None:
        queryset = queryset.filter(owner_id=owner_id)
    else:
        queryset = queryset.filter(owner__isnull=True, guest_key=guest_key)
    return queryset.first()


def _job_record_params(record: JobRecord) -> dict[str, object]:
    try:
        data = json.loads(record.params_json or "{}")
        return data if isinstance(data, dict) else {}
    except json.JSONDecodeError:
        return {}


def _refresh_job_output_summary(record: JobRecord) -> None:
    outputs = list(record.outputs.all())
    record.output_count = len(outputs)
    record.total_output_size = sum(int(output.size or 0) for output in outputs)
    record.primary_output_type = str(outputs[0].media_type or "").split("/", 1)[0][:32] if outputs else ""
    record.save(update_fields=["output_count", "total_output_size", "primary_output_type", "updated_at"])


def _create_design_project_from_output(request: HttpRequest, record: JobRecord, output, job: dict) -> DesignerProject:
    owner_id, guest_key = _workspace_identity(request)
    output_path = Path(output.path)
    metadata = _cover_design_metadata(output_path)
    title = _clean_project_title(f"Cover design: {Path(output_path).stem}")
    project = DesignerProject.objects.create(
        owner=request.user if request.user.is_authenticated else None,
        guest_key="" if owner_id else guest_key,
        title=title,
        state_json={},
    )
    media_dir = _design_project_media_dir(project)
    media_dir.mkdir(parents=True, exist_ok=True)
    background_source = Path(str(metadata.get("background_path") or "")) if metadata else output_path
    if not background_source.exists() or not background_source.is_file():
        background_source = output_path
    asset = _copy_output_to_design_asset(project, background_source, output_path.name)
    state = _cover_design_state(project, asset, output_path, metadata, record, job)
    project.state_json = state
    _update_design_project_metadata(project)
    project.storage_bytes = _design_project_storage_bytes(project)
    project.save(update_fields=["state_json", "storage_bytes", "asset_count", "object_count", "updated_at"])
    return project


def _copy_output_to_design_asset(project: DesignerProject, source: Path, original_name: str) -> DesignerAsset:
    suffix = source.suffix.lower()[:12] or ".png"
    target = _design_project_media_dir(project) / f"{uuid.uuid4().hex[:12]}_{clean_base_name(original_name, 'cover')}{suffix}"
    shutil.copy2(source, target)
    media_type = mimetypes.guess_type(target.name)[0] or "image/png"
    return DesignerAsset.objects.create(project=project, kind="image", file_path=str(target), media_type=media_type, size=target.stat().st_size, original_name=original_name[:240])


def _create_video_project_from_output(request: HttpRequest, record: JobRecord, output, job: dict) -> VideoEditorProject:
    owner_id, guest_key = _workspace_identity(request)
    output_path = Path(output.path)
    edit_source = _short_video_edit_metadata(output_path)
    asset_source = Path(str(edit_source.get("source_path") or "")) if edit_source else output_path
    asset_original_name = output_path.name
    subtitle_source = _job_original_video_source_for_edit(record, output_path)
    if subtitle_source:
        asset_source = subtitle_source
        asset_original_name = subtitle_source.name
        edit_source = {}
    if not asset_source.exists() or not asset_source.is_file():
        fallback_source = Path(str(edit_source.get("fallback_source_path") or "")) if edit_source else output_path
        asset_source = fallback_source if fallback_source.exists() and fallback_source.is_file() else output_path
        if asset_source == output_path:
            edit_source = {}
    title = _clean_project_title(f"Edit: {Path(output_path).stem}")
    project = VideoEditorProject.objects.create(
        owner=request.user if request.user.is_authenticated else None,
        guest_key="" if owner_id else guest_key,
        title=title,
        state_json={},
    )
    asset = _copy_output_to_video_asset(project, asset_source, asset_original_name)
    duration = asset.duration or _safe_video_duration(Path(asset.file_path))
    if duration and not asset.duration:
        asset.duration = duration
        asset.save(update_fields=["duration"])
    project.state_json = _video_editor_state_from_asset(project, asset, title, duration, record, job, edit_source)
    _update_video_project_metadata(project)
    project.storage_bytes = _video_project_storage_bytes(project)
    project.save(update_fields=["state_json", "storage_bytes", "asset_count", "clip_count", "duration_seconds", "thumbnail_path", "updated_at"])
    _append_job_output_subtitles_to_video_project(project, record, output_path)
    return project


def _job_original_video_source_for_edit(record: JobRecord, video_output_path: Path) -> Path | None:
    params = _job_record_params(record)
    if str(params.get("action") or record.kind) not in {"subtitles", "package"}:
        return None
    if not _subtitle_output_path_for_video_edit(record, video_output_path):
        return None
    source = Path(str(params.get("source") or ""))
    if source.exists() and source.is_file():
        return source
    return None


def _ensure_project_uses_editable_subtitle_source(project: VideoEditorProject, record: JobRecord, video_output_path: Path) -> None:
    source = _job_original_video_source_for_edit(record, video_output_path)
    if not source:
        return
    state = project.state_json if isinstance(project.state_json, dict) else {}
    clips = state.get("clips") if isinstance(state.get("clips"), list) else []
    video_clip = next((clip for clip in clips if isinstance(clip, dict) and clip.get("type") == "video"), None)
    if not video_clip:
        return
    current_asset = project.assets.filter(id=int(video_clip.get("assetId") or 0)).first()
    if current_asset:
        try:
            if Path(current_asset.file_path).resolve() == source.resolve():
                return
        except OSError:
            pass
        if current_asset.original_name == source.name:
            return
    asset = _copy_output_to_video_asset(project, source, source.name)
    duration = asset.duration or _safe_video_duration(Path(asset.file_path))
    video_clip["assetId"] = asset.id
    video_clip["duration"] = round(max(0.25, duration or float(video_clip.get("duration") or 0) or 12), 3)
    video_clip["sourceStart"] = 0
    video_clip["sourceEnd"] = video_clip["duration"]
    state["clipName"] = source.name
    state["clips"] = clips
    project.state_json = state
    _update_video_project_metadata(project)
    project.storage_bytes = _video_project_storage_bytes(project)
    project.save(update_fields=["state_json", "storage_bytes", "asset_count", "clip_count", "duration_seconds", "thumbnail_path", "updated_at"])


def _append_job_output_subtitles_to_video_project(project: VideoEditorProject, record: JobRecord, video_output_path: Path) -> None:
    state = project.state_json if isinstance(project.state_json, dict) else {}
    clips = state.get("clips") if isinstance(state.get("clips"), list) else []
    if any(isinstance(clip, dict) and clip.get("type") == "caption" for clip in clips):
        _mark_generated_subtitle_project(project, record, video_output_path)
        _apply_job_subtitle_style_to_caption_clips(project, record)
        return
    subtitle_path = _subtitle_output_path_for_video_edit(record, video_output_path)
    if not subtitle_path:
        return
    try:
        text = _decode_subtitle_bytes(subtitle_path.read_bytes())
        cues = _parse_subtitle_cues(text, subtitle_path.name)
    except Exception:
        return
    if cues:
        _append_video_project_caption_clips(project, cues)
        _mark_generated_subtitle_project(project, record, video_output_path, subtitle_path=subtitle_path)
        _apply_job_subtitle_style_to_caption_clips(project, record)


def _mark_generated_subtitle_project(
    project: VideoEditorProject,
    record: JobRecord,
    video_output_path: Path,
    *,
    subtitle_path: Path | None = None,
) -> None:
    params = _job_record_params(record)
    subtitle_path = subtitle_path or _subtitle_output_path_for_video_edit(record, video_output_path)
    state = project.state_json if isinstance(project.state_json, dict) else {}
    state = json.loads(json.dumps(state))
    workflow = state.get("subtitleWorkflow") if isinstance(state.get("subtitleWorkflow"), dict) else {}
    workflow.update(
        {
            "mode": "generated-editable",
            "sourceJobId": record.job_id,
            "sourceJobKind": record.kind,
            "style": str(params.get("style") or "").strip().lower(),
            "sourceVideo": str(params.get("source") or ""),
            "renderedVideo": str(video_output_path),
            "subtitlePath": str(subtitle_path or ""),
        }
    )
    state["subtitleWorkflow"] = workflow
    clips = state.get("clips") if isinstance(state.get("clips"), list) else []
    for clip in clips:
        if isinstance(clip, dict) and clip.get("type") == "caption":
            clip["subtitleSource"] = {
                "kind": "generated",
                "jobId": record.job_id,
                "style": workflow.get("style") or "",
            }
    project.state_json = state
    _update_video_project_metadata(project)
    project.storage_bytes = _video_project_storage_bytes(project)
    project.save(update_fields=["state_json", "storage_bytes", "asset_count", "clip_count", "duration_seconds", "thumbnail_path", "updated_at"])


def _apply_job_subtitle_style_to_caption_clips(project: VideoEditorProject, record: JobRecord) -> None:
    params = _job_record_params(record)
    style = str(params.get("style") or "").strip().lower()
    clip_style = _video_editor_caption_style(style)
    state = project.state_json if isinstance(project.state_json, dict) else {}
    clips = state.get("clips") if isinstance(state.get("clips"), list) else []
    changed = False
    for clip in clips:
        if not isinstance(clip, dict) or clip.get("type") != "caption":
            continue
        current = clip.get("style") if isinstance(clip.get("style"), dict) else {}
        next_style = dict(current)
        next_style.update(clip_style)
        clip["style"] = next_style
        source = clip.get("subtitleSource") if isinstance(clip.get("subtitleSource"), dict) else {}
        if source:
            source["style"] = style
            clip["subtitleSource"] = source
        changed = True
    if changed:
        project.state_json = state
        project.save(update_fields=["state_json", "updated_at"])


def _video_editor_caption_style(style: str) -> dict[str, object]:
    base = {
        "font": "system-ui, -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif",
        "size": 32,
        "fontWeight": 900,
        "color": "#ffffff",
        "stroke": "#000000",
        "strokeWidth": 2,
        "bg": "#000000",
        "bgAlpha": 0,
        "textShadow": "0 3px 0 rgba(0,0,0,.72), 0 10px 24px rgba(0,0,0,.28)",
        "animation": style or "none",
    }
    if style in {"kinetic", "bounce", "pop", "headline"}:
        base.update({"size": 36, "stroke": "#db2777", "strokeWidth": 2, "bgAlpha": 0, "animation": "kinetic"})
    elif style == "neon":
        base.update({"size": 34, "stroke": "#7c3aed", "strokeWidth": 2, "bgAlpha": 0, "textShadow": "0 0 10px #22d3ee, 0 0 22px #f0abfc, 0 3px 0 #020617", "animation": "neon"})
    elif style == "candy":
        base.update({"size": 34, "stroke": "#7e22ce", "strokeWidth": 2, "bgAlpha": 0, "textShadow": "2px 2px 0 #672c92, -1px -1px 0 #672c92, 0 0 14px rgba(255,118,216,.52)", "animation": "candy"})
    elif style in {"clean", "minimal"}:
        base.update({"size": 28, "fontWeight": 760, "stroke": "#111827", "strokeWidth": 1, "bg": "#000000", "bgAlpha": 24, "textShadow": "none", "animation": "none"})
    elif style in {"editorial", "luxury"}:
        base.update({"size": 30, "color": "#f8fafc", "stroke": "#1f2937", "strokeWidth": 1, "bgAlpha": 18, "animation": "soft"})
    elif style == "typewriter":
        base.update({"font": "Consolas, monospace", "size": 28, "fontWeight": 800, "strokeWidth": 1, "bgAlpha": 20, "textShadow": "none", "animation": "typewriter"})
    return base


def _subtitle_output_path_for_video_edit(record: JobRecord, video_output_path: Path) -> Path | None:
    try:
        video_parent = video_output_path.resolve().parent
    except OSError:
        video_parent = video_output_path.parent
    candidates: list[Path] = []
    for output in record.outputs.all():
        path = Path(output.path)
        if path == video_output_path or not _is_subtitle_output(path.name, output.media_type):
            continue
        if not path.exists() or not path.is_file():
            continue
        candidates.append(path)
    if not candidates:
        return None
    try:
        same_folder = [path for path in candidates if path.resolve().parent == video_parent]
    except OSError:
        same_folder = [path for path in candidates if path.parent == video_parent]
    return (same_folder or candidates)[0]


def _short_video_edit_metadata(output_path: Path) -> dict[str, object]:
    meta_path = output_path.with_suffix(".edit.json")
    if not meta_path.exists() or not meta_path.is_file():
        return {}
    try:
        data = json.loads(meta_path.read_text(encoding="utf-8"))
    except Exception:
        return {}
    if not isinstance(data, dict) or data.get("kind") != "youtube_short_source":
        return {}
    return data


def _copy_output_to_video_asset(project: VideoEditorProject, source: Path, original_name: str) -> VideoEditorAsset:
    if not source.exists() or not source.is_file():
        raise FileNotFoundError(original_name or "Video output")
    target_dir = _video_project_media_dir(project)
    target_dir.mkdir(parents=True, exist_ok=True)
    suffix = source.suffix.lower()[:16] or ".mp4"
    target = target_dir / f"{uuid.uuid4().hex[:12]}_{clean_base_name(original_name, 'short')}{suffix}"
    shutil.copy2(source, target)
    media_type = mimetypes.guess_type(target.name)[0] or "video/mp4"
    return VideoEditorAsset.objects.create(
        project=project,
        kind="video",
        file_path=str(target),
        media_type=media_type,
        size=target.stat().st_size,
        original_name=original_name[:240],
        duration=_safe_video_duration(target),
    )


def _video_editor_state_from_asset(
    project: VideoEditorProject,
    asset: VideoEditorAsset,
    title: str,
    duration: float,
    record: JobRecord,
    job: dict,
    edit_source: dict[str, object] | None = None,
) -> dict[str, object]:
    safe_duration = max(0.25, float(duration or asset.duration or 12))
    edit_source = edit_source or {}
    try:
        source_start = max(0.0, float(edit_source.get("source_start") or 0))
    except (TypeError, ValueError):
        source_start = 0.0
    if source_start >= safe_duration:
        source_start = 0.0
    try:
        clip_duration = float(edit_source.get("clip_duration") or 0)
    except (TypeError, ValueError):
        clip_duration = 0
    clip_duration = max(0.25, clip_duration or safe_duration)
    if source_start + clip_duration > safe_duration and source_start < safe_duration:
        clip_duration = max(0.25, safe_duration - source_start)
    source_end = source_start + clip_duration
    return {
        "title": title,
        "clipName": asset.original_name or title,
        "aspect": str(edit_source.get("aspect") or "9 / 16"),
        "background": "#020617",
        "backgroundMode": "solid",
        "backgroundValue": "#020617",
        "tracks": [
            {"id": "video-main", "type": "video", "name": "Video", "order": 0},
            {"id": "text-1", "type": "text", "name": "Text", "order": 1},
            {"id": "image-1", "type": "image", "name": "Image", "order": 2},
            {"id": "audio-1", "type": "audio", "name": "Audio", "order": 3},
        ],
        "clips": [
            {
                "id": f"video-output-{asset.id}",
                "type": "video",
                "trackId": "video-main",
                "assetId": asset.id,
                "start": 0,
                "duration": clip_duration,
                "sourceStart": source_start,
                "sourceEnd": source_end,
                "x": 50,
                "y": 50,
                "scale": 100,
                "style": {"fit": "cover", "speed": 1, "opacity": 100, "transition": "none", "fadeIn": 0, "fadeOut": 0},
                "text": "",
            }
        ],
        "sourceJob": {"id": record.job_id, "kind": record.kind, "title": job.get("title") or record.title},
    }


def _prefetched_assets(project) -> list:
    if not project:
        return []
    cache = getattr(project, "_prefetched_objects_cache", {}) or {}
    if "assets" in cache:
        return list(cache["assets"])
    return list(project.assets.all())


def _state_list(state: dict[str, object], key: str) -> list:
    value = state.get(key) if isinstance(state, dict) else []
    return value if isinstance(value, list) else []


def _timeline_duration_seconds(state: dict[str, object], fallback: float = 0) -> float:
    clips = _state_list(state, "clips")
    duration = max(
        (
            float(clip.get("start") or 0) + float(clip.get("duration") or 0)
            for clip in clips
            if isinstance(clip, dict)
        ),
        default=float(fallback or 0),
    )
    return max(0.0, duration)


def _update_video_project_metadata(project: VideoEditorProject, *, assets: list[VideoEditorAsset] | None = None) -> None:
    state = project.state_json or {}
    current_assets = assets if assets is not None else _prefetched_assets(project)
    project.asset_count = len(current_assets)
    project.clip_count = len(_state_list(state, "clips"))
    project.duration_seconds = _timeline_duration_seconds(state)
    first_thumb = next((asset.thumbnail_path for asset in current_assets if getattr(asset, "thumbnail_path", "")), "")
    project.thumbnail_path = first_thumb or (str(state.get("thumbnail") or "") if isinstance(state, dict) else "")


def _update_design_project_metadata(project: DesignerProject, *, assets: list[DesignerAsset] | None = None) -> None:
    state = project.state_json or {}
    project.asset_count = len(assets) if assets is not None else project.assets.count()
    project.object_count = len(_state_list(state, "objects"))


def _update_music_project_metadata(project: MusicEditorProject, *, assets: list[MusicEditorAsset] | None = None) -> None:
    state = project.state_json or {}
    current_assets = assets if assets is not None else _prefetched_assets(project)
    project.asset_count = len(current_assets)
    project.clip_count = len(_state_list(state, "clips"))
    project.duration_seconds = _timeline_duration_seconds(state, sum(float(getattr(asset, "duration", 0) or 0) for asset in current_assets))


def _safe_video_duration(path: Path) -> float:
    try:
        return float(inspect_video(path).duration_seconds or 0)
    except Exception:
        return 0.0


def _cover_design_metadata(output_path: Path) -> dict[str, object]:
    meta_path = output_path.with_suffix(".design.json")
    if not meta_path.exists():
        return {}
    try:
        data = json.loads(meta_path.read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else {}
    except (OSError, json.JSONDecodeError):
        return {}


def _cover_design_state(project: DesignerProject, asset: DesignerAsset, output_path: Path, metadata: dict[str, object], record: JobRecord, job: dict) -> dict[str, object]:
    copy = metadata.get("copy") if isinstance(metadata.get("copy"), dict) else {}
    palette = metadata.get("palette") if isinstance(metadata.get("palette"), dict) else {}
    layout = metadata.get("layout") if isinstance(metadata.get("layout"), dict) else {}
    style = metadata.get("style") if isinstance(metadata.get("style"), dict) else {}
    focus = metadata.get("focus_point") if isinstance(metadata.get("focus_point"), dict) else {}
    headline = str(copy.get("headline") or record.title or job.get("title") or output_path.stem)
    description = str(copy.get("description") or "")
    eyebrow = str(copy.get("eyebrow") or "COVER")
    accent = str(palette.get("accent") or "#facc15")
    accent2 = str(palette.get("accent2") or "#38bdf8")
    dark = str(palette.get("dark") or "#05070c")
    paper = str(palette.get("paper") or "#ffffff")
    panel_x = int(float(layout.get("panel_x") or 58))
    headline_y = int(float(layout.get("headline_y") or 135))
    max_text_width = int(float(layout.get("max_text_width") or 650))
    badge_x = int(float(layout.get("badge_x") or panel_x))
    badge_y = int(float(layout.get("badge_y") or 48))
    hook_x = int(float(layout.get("hook_x") or panel_x))
    hook_y = int(float(layout.get("hook_y") or 624))
    hook_text = str(style.get("badge") or "WATCH TO THE END")
    layout_name = str(style.get("layout") or "split")
    mood = str(style.get("mood") or "premium")
    text_left = bool(style.get("text_left", panel_x < 300))
    objects = [
        {"id": "frame-cover", "type": "frame", "name": "Cover 1280x720", "x": 0, "y": 0, "w": 1280, "h": 720, "fill": dark, "stroke": accent, "strokeWidth": 2, "cornerRadius": 0, "clipContent": True},
        {"id": "cover-bg", "type": "image", "name": "Background", "x": 0, "y": 0, "w": 1280, "h": 720, "parentId": "frame-cover", "assetId": asset.id, "src": reverse("studio:design_project_asset_preview", args=[project.id, asset.id]), "naturalW": 1280, "naturalH": 720, "imageFit": "fill", "imageCrop": {"x": 0, "y": 0, "scale": 1}},
        {"id": "cover-fade", "type": "shape", "shape": "rect", "name": "Contrast overlay", "x": max(0, panel_x - 42), "y": 34, "w": min(720, max_text_width + 74), "h": 632, "parentId": "frame-cover", "fill": "#000000", "stroke": "transparent", "strokeWidth": 0, "cornerRadius": 18, "opacity": 0.46},
        {"id": "cover-diagonal-accent", "type": "shape", "shape": "rect", "name": f"{layout_name} accent", "x": 728 if text_left else 188, "y": -58, "w": 22, "h": 850, "parentId": "frame-cover", "fill": accent2, "stroke": "transparent", "strokeWidth": 0, "cornerRadius": 12, "opacity": 0.36, "rotation": -10 if text_left else 10},
        {"id": "cover-eyebrow-bg", "type": "shape", "shape": "rect", "name": "Eyebrow badge", "x": badge_x, "y": badge_y, "w": min(max_text_width, max(210, len(eyebrow) * 22 + 42)), "h": 50, "parentId": "frame-cover", "fill": accent, "stroke": "transparent", "strokeWidth": 0, "cornerRadius": 8},
        {"id": "cover-eyebrow", "type": "text", "name": "Eyebrow", "text": eyebrow, "x": badge_x + 20, "y": badge_y + 9, "w": max(180, min(max_text_width - 24, len(eyebrow) * 24)), "h": 40, "parentId": "frame-cover", "fill": "transparent", "stroke": dark, "fontSize": 28, "fontWeight": 900, "fontFamily": "Arial, sans-serif", "lineHeight": 1.05},
        {"id": "cover-headline-panel", "type": "shape", "shape": "rect", "name": "Headline backing", "x": max(0, panel_x - 22), "y": max(0, headline_y - 16), "w": max_text_width + 48, "h": 286, "parentId": "frame-cover", "fill": "#000000", "stroke": accent, "strokeWidth": 1, "cornerRadius": 14, "opacity": 0.38},
        {"id": "cover-headline", "type": "text", "name": "Headline", "text": headline, "x": panel_x, "y": headline_y, "w": max_text_width, "h": 286, "parentId": "frame-cover", "fill": "transparent", "stroke": paper, "fontSize": 82, "fontWeight": 950, "fontFamily": "Arial, sans-serif", "lineHeight": 1.02, "textStroke": "#000000", "textStrokeWidth": 2},
        {"id": "cover-accent-line", "type": "shape", "shape": "rect", "name": "Accent line", "x": panel_x, "y": headline_y + 327, "w": max_text_width, "h": 8, "parentId": "frame-cover", "fill": accent2, "stroke": "transparent", "strokeWidth": 0, "cornerRadius": 4},
        {"id": "cover-description", "type": "text", "name": "Description", "text": description or "Edit subtitle text", "x": panel_x, "y": headline_y + 353, "w": max_text_width, "h": 110, "parentId": "frame-cover", "fill": "transparent", "stroke": "#e2e8f0", "fontSize": 30, "fontWeight": 800, "fontFamily": "Arial, sans-serif", "lineHeight": 1.12},
        {"id": "cover-hook-bg", "type": "shape", "shape": "rect", "name": "Hook badge", "x": hook_x, "y": hook_y, "w": min(560, max(260, len(hook_text) * 19 + 64)), "h": 56, "parentId": "frame-cover", "fill": "#000000", "stroke": accent, "strokeWidth": 2, "cornerRadius": 8, "opacity": 0.88},
        {"id": "cover-hook-strip", "type": "shape", "shape": "rect", "name": "Hook strip", "x": hook_x, "y": hook_y, "w": 12, "h": 56, "parentId": "frame-cover", "fill": accent2, "stroke": "transparent", "strokeWidth": 0, "cornerRadius": 4},
        {"id": "cover-hook", "type": "text", "name": "Hook", "text": hook_text, "x": hook_x + 28, "y": hook_y + 12, "w": min(500, max(220, len(hook_text) * 18)), "h": 40, "parentId": "frame-cover", "fill": "transparent", "stroke": paper, "fontSize": 30, "fontWeight": 900, "fontFamily": "Arial, sans-serif", "lineHeight": 1.05},
        {"id": "cover-mood-chip", "type": "text", "name": "Mood", "text": f"{mood.upper()} / {layout_name.upper()}", "x": 1016, "y": 48, "w": 210, "h": 34, "parentId": "frame-cover", "fill": "transparent", "stroke": accent2, "fontSize": 22, "fontWeight": 900, "fontFamily": "Arial, sans-serif", "lineHeight": 1.05},
    ]
    if focus:
        fx = int(float(focus.get("x") or 0))
        fy = int(float(focus.get("y") or 0))
        if 0 < fx < 1280 and 0 < fy < 720:
            objects.append({"id": "cover-focus-ring", "type": "shape", "shape": "ellipse", "name": "Focus marker", "x": max(0, fx - 84), "y": max(0, fy - 84), "w": 168, "h": 168, "parentId": "frame-cover", "fill": "transparent", "stroke": accent2, "strokeWidth": 5, "cornerRadius": 0, "opacity": 0.52})
    return {
        "version": 2,
        "title": f"Cover design: {output_path.stem}",
        "zoom": 1,
        "didFit": False,
        "vectors": [],
        "objects": objects,
    }


def _music_asset_payload(asset: MusicEditorAsset) -> dict[str, object]:
    return {
        "id": asset.id,
        "kind": asset.kind,
        "name": asset.original_name,
        "media_type": asset.media_type,
        "size": asset.size,
        "size_text": human_size(asset.size),
        "duration": asset.duration,
        "preview_url": reverse("studio:music_project_asset_preview", args=[asset.project_id, asset.id]),
    }


def _music_project_payload(project: MusicEditorProject | None, include_state: bool = True, owner_id: int | None = None, guest_key: str = "") -> dict[str, object]:
    if not project:
        return {}
    state = project.state_json or {}
    assets = _prefetched_assets(project)
    access_role = _resource_access_role(WorkspaceShare.RESOURCE_MUSIC, project, owner_id, guest_key) if owner_id is not None or guest_key else "owner"
    is_owner = access_role == "owner"
    total_duration = float(getattr(project, "duration_seconds", 0) or 0) or sum(asset.duration for asset in assets if asset.duration)
    payload: dict[str, object] = {
        "id": project.id,
        "title": project.title,
        "track_count": int(getattr(project, "asset_count", 0) or len(assets)),
        "clip_count": int(getattr(project, "clip_count", 0) or len(_state_list(state, "clips"))),
        "total_duration": total_duration,
        "duration_text": f"{int(total_duration // 60)}:{int(total_duration % 60):02d}",
        "storage_bytes": project.storage_bytes,
        "storage_text": human_size(project.storage_bytes),
        "created_at": project.created_at.isoformat(),
        "updated_at": project.updated_at.isoformat(),
        "assets": [_music_asset_payload(asset) for asset in assets],
        "access_role": access_role,
        "is_owner": is_owner,
        "can_edit": is_owner or access_role == WorkspaceShare.ROLE_EDITOR,
        "can_share": is_owner and owner_id is not None,
        "publish_url": f"{reverse('studio:community_publish')}?{urlencode({'source': 'music_project', 'id': project.id})}" if is_owner and owner_id is not None else "",
    }
    if include_state:
        payload["state"] = project.state_json or {}
    return payload


def _video_project_payload(project: VideoEditorProject, include_state: bool = True, owner_id: int | None = None, guest_key: str = "") -> dict[str, object]:
    state = project.state_json or {}
    layers = _state_list(state, "layers") if include_state else []
    assets = _prefetched_assets(project)
    thumbnail_assets = [asset for asset in assets if asset.thumbnail_path]
    first_thumb = thumbnail_assets[project.id % len(thumbnail_assets)] if thumbnail_assets else None
    preview_sources = [
        reverse("studio:video_project_asset_preview", args=[project.id, asset.id])
        for asset in assets
        if asset.kind == "video"
    ]
    clip_count = int(getattr(project, "clip_count", 0) or 0)
    if include_state or not clip_count:
        clip_count = len(_state_list(state, "clips") or layers)
    access_role = _resource_access_role(WorkspaceShare.RESOURCE_VIDEO, project, owner_id, guest_key) if owner_id is not None or guest_key else "owner"
    is_owner = access_role == "owner"
    payload: dict[str, object] = {
        "id": project.id,
        "title": project.title,
        "thumbnail": reverse("studio:video_project_asset_thumbnail", args=[project.id, first_thumb.id]) if first_thumb else (getattr(project, "thumbnail_path", "") or (state.get("thumbnail", "") if isinstance(state, dict) else "")),
        "preview_sources": preview_sources,
        "aspect": state.get("aspect", "9 / 16") if isinstance(state, dict) else "9 / 16",
        "clip_name": state.get("clipName", "") if isinstance(state, dict) else "",
        "layer_count": clip_count,
        "track_count": len(_state_list(state, "tracks")) if include_state else int(getattr(project, "asset_count", 0) or len(assets)),
        "asset_count": int(getattr(project, "asset_count", 0) or len(assets)),
        "duration_seconds": float(getattr(project, "duration_seconds", 0) or _timeline_duration_seconds(state)),
        "storage_bytes": project.storage_bytes,
        "storage_text": human_size(project.storage_bytes),
        "created_at": project.created_at.isoformat(),
        "updated_at": project.updated_at.isoformat(),
        "assets": [_video_asset_payload(asset) for asset in assets],
        "access_role": access_role,
        "is_owner": is_owner,
        "can_edit": is_owner or access_role == WorkspaceShare.ROLE_EDITOR,
        "can_share": is_owner and owner_id is not None,
        "publish_url": f"{reverse('studio:community_publish')}?{urlencode({'source': 'video_project', 'id': project.id})}" if is_owner and owner_id is not None else "",
    }
    if include_state:
        payload["state"] = project.state_json or {}
    return payload


def _video_export_payload(request: HttpRequest, project: VideoEditorProject, job_id: str) -> dict[str, object]:
    record = JobRecord.objects.filter(job_id=job_id).prefetch_related("outputs").first()
    if record:
        return _video_export_record_payload(request, project, record)
    with _video_export_lock:
        job = dict(_video_export_jobs.get(job_id) or {})
    if not job:
        raise Http404("Export not found")
    payload = {
        "id": job_id,
        "status": job.get("status", "missing"),
        "progress": job.get("progress", 0),
        "message": job.get("message", ""),
        "error": job.get("error", ""),
        "quality": job.get("quality", "720p"),
    }
    path = Path(str(job.get("path") or ""))
    if job.get("status") == "done" and path.exists():
        payload["download_url"] = reverse("studio:download_video_project_export", args=[project.id, job_id])
        payload["size_text"] = human_size(path.stat().st_size)
        payload["filename"] = path.name
    return payload


def _video_export_record_payload(request: HttpRequest, project: VideoEditorProject, record: JobRecord) -> dict[str, object]:
    output = record.outputs.first()
    status = "done" if record.status in {"completed", "done"} else record.status
    try:
        events = [{"status": event.status, "progress": event.progress, "message": event.message} for event in record.events.all().order_by("-id")[:6]]
    except Exception:
        events = []
    payload: dict[str, object] = {
        "id": record.job_id,
        "status": status,
        "progress": record.progress,
        "message": record.message,
        "error": record.error,
        "quality": _job_param(record, "quality", "720p"),
        "kind": record.kind,
        "created_at": record.created_at.isoformat(),
        "events": events,
    }
    if record.kind == "video_subtitles":
        result_cues = _job_param(record, "result_cues", [])
        payload["cue_count"] = int(_job_param(record, "result_count", len(result_cues) if isinstance(result_cues, list) else 0) or 0)
        if isinstance(result_cues, list):
            payload["cues"] = result_cues
        if status == "done":
            fresh_project = VideoEditorProject.objects.prefetch_related("assets").filter(id=project.id).first() or project
            owner_id, guest_key = _workspace_identity(request)
            payload["project"] = _video_project_payload(fresh_project, owner_id=owner_id, guest_key=guest_key)
    if output and Path(output.path).exists() and status == "done":
        payload["download_url"] = reverse("studio:download_video_project_export", args=[project.id, record.job_id])
        payload["size_text"] = human_size(output.size or Path(output.path).stat().st_size)
        payload["filename"] = Path(output.path).name
        payload["media_type"] = output.media_type
    return payload


def _video_export_access(job_id: str, project: VideoEditorProject, owner_id: int | None, guest_key: str) -> bool:
    if not _resource_can_view(WorkspaceShare.RESOURCE_VIDEO, project, owner_id, guest_key):
        return False
    record = JobRecord.objects.filter(job_id=job_id).first()
    if record:
        return int(_job_param(record, "project_id", 0) or 0) == project.id
    with _video_export_lock:
        job = _video_export_jobs.get(job_id)
    if not job or int(job.get("project_id") or 0) != project.id:
        return False
    return True


def _set_video_export_job(job_id: str, **values: object) -> None:
    with _video_export_lock:
        job = _video_export_jobs.get(job_id)
        if not job:
            return
        job.update(values)
    try:
        record = JobRecord.objects.filter(job_id=job_id).first()
        if record:
            status = str(values.get("status", record.status))
            record.status = "completed" if status == "done" else status
            record.progress = int(values.get("progress", record.progress) or 0)
            record.message = str(values.get("message", record.message) or "")
            record.error = str(values.get("error", record.error) or "")
            record.save(update_fields=["status", "progress", "message", "error", "updated_at"])
            try:
                JobEventRecord.objects.create(job=record, status=record.status, progress=record.progress, message=record.error or record.message)
            except Exception:
                pass
    except Exception:
        pass


def _run_video_project_export(job_id: str, project_id: int, quality: str, output: Path) -> None:
    try:
        _set_video_export_job(job_id, status="running", progress=8, message="Preparing timeline")
        project = VideoEditorProject.objects.prefetch_related("assets").get(id=project_id)
        width, height = _video_export_size(project.state_json or {}, quality)
        duration = _video_export_duration(project.state_json or {})
        assets = {asset.id: asset for asset in project.assets.all()}
        clips = (project.state_json or {}).get("clips", []) if isinstance(project.state_json, dict) else []
        video_clip = next((clip for clip in clips if clip.get("type") == "video" and assets.get(int(clip.get("assetId") or 0))), None)
        _set_video_export_job(job_id, progress=22, message="Rendering MP4")
        output.parent.mkdir(parents=True, exist_ok=True)
        if video_clip:
            _render_video_project_from_clip(video_clip, assets, project.state_json or {}, output, width, height, duration)
        elif _video_export_visual_clips(project.state_json or {}, assets):
            _render_visual_card_project(project.state_json or {}, assets, output, width, height, duration)
        else:
            _render_empty_video_project(output, width, height, duration)
        if JobRecord.objects.filter(job_id=job_id, status="cancelled").exists():
            return
        _update_video_project_metadata(project, assets=list(project.assets.all()))
        project.storage_bytes = _video_project_storage_bytes(project)
        project.last_export_status = "completed"
        project.save(update_fields=["storage_bytes", "asset_count", "clip_count", "duration_seconds", "thumbnail_path", "last_export_status", "updated_at"])
        record = JobRecord.objects.filter(job_id=job_id).first()
        if record and output.exists():
            JobOutputRecord.objects.get_or_create(
                job=record,
                path=str(output),
                label="MP4 export",
                defaults={"media_type": "video/mp4", "size": output.stat().st_size},
            )
            _refresh_job_output_summary(record)
        _set_video_export_job(job_id, status="done", progress=100, message="Ready")
    except Exception as exc:
        _set_video_export_job(job_id, status="failed", progress=100, message="Export failed", error=str(exc))


def _run_video_project_cover(job_id: str, project_id: int, output: Path, time_seconds: float) -> None:
    try:
        _set_video_export_job(job_id, status="running", progress=12, message="Rendering cover")
        project = VideoEditorProject.objects.prefetch_related("assets").get(id=project_id)
        _render_video_project_cover(project, output, time_seconds)
        record = JobRecord.objects.filter(job_id=job_id).first()
        if record and output.exists():
            JobOutputRecord.objects.get_or_create(
                job=record,
                path=str(output),
                label="Cover frame",
                defaults={"media_type": "image/jpeg", "size": output.stat().st_size},
            )
            _refresh_job_output_summary(record)
        _set_video_export_job(job_id, status="done", progress=100, message="Ready")
    except Exception as exc:
        _set_video_export_job(job_id, status="failed", progress=100, message="Cover failed", error=str(exc))


def _run_video_project_subtitles(job_id: str, project_id: int, asset_id: int, params: dict[str, object]) -> None:
    try:
        _set_video_export_job(job_id, status="running", progress=10, message="Transcribing audio")
        project = VideoEditorProject.objects.prefetch_related("assets").get(id=project_id)
        asset = project.assets.filter(id=asset_id, kind__in=["audio", "video"]).first()
        if not asset:
            raise FileNotFoundError("Source media is missing")
        source = Path(asset.file_path)
        if not source.exists() or not source.is_file():
            raise FileNotFoundError("Source media file is missing")

        language = str(params.get("language") or "auto").strip().lower()
        language = None if language == "auto" else language
        transcribed = transcribe_subtitle_cues(source, settings.subtitle_model, language)
        _set_video_export_job(job_id, progress=70, message="Saving captions")

        cues = _normalize_auto_subtitle_cues(transcribed, params)
        project = _append_video_project_caption_clips(project, cues)
        _merge_job_params(job_id, {"result_cues": cues, "result_count": len(cues)})
        _set_video_export_job(
            job_id,
            status="done",
            progress=100,
            message=f"Subtitles added: {len(cues)}" if cues else "No speech detected",
        )
    except SubtitleUnavailableError as exc:
        _set_video_export_job(job_id, status="failed", progress=100, message="Auto subtitles unavailable", error=str(exc))
    except Exception as exc:
        _set_video_export_job(job_id, status="failed", progress=100, message="Auto subtitles failed", error=str(exc))


def _video_export_size(state: dict[str, object], quality: str) -> tuple[int, int]:
    aspect = str(state.get("aspect") or "9 / 16").replace(" ", "")
    long_side = 1080 if quality == "1080p" else 720
    try:
        left, right = (max(1, int(part)) for part in aspect.split("/", 1))
    except (TypeError, ValueError):
        left, right = 9, 16
    if left == right:
        return long_side, long_side
    if left > right:
        height = long_side
        width = int(round(height * left / right))
    else:
        width = long_side
        height = int(round(width * right / left))
    return max(2, width + width % 2), max(2, height + height % 2)


def _video_export_duration(state: dict[str, object]) -> float:
    clips = state.get("clips") if isinstance(state.get("clips"), list) else []
    duration = max((float(clip.get("start") or 0) + float(clip.get("duration") or 0) for clip in clips if isinstance(clip, dict)), default=5.0)
    return max(0.25, min(duration, 60 * 30))


def _video_clip_base_filter(clip: dict[str, object], width: int, height: int) -> str:
    style = clip.get("style") if isinstance(clip.get("style"), dict) else {}
    fit = str(style.get("fit") or "contain")
    x_pct = max(0.0, min(1.0, float(clip.get("x") or 50) / 100))
    y_pct = max(0.0, min(1.0, float(clip.get("y") or 50) / 100))
    if fit in {"cover", "crop"}:
        zoom = max(1.0, min(3.0, float(clip.get("scale") or 100) / 100))
        target_width = max(width, int(round(width * zoom)))
        target_height = max(height, int(round(height * zoom)))
        target_width += target_width % 2
        target_height += target_height % 2
        return (
            f"scale={target_width}:{target_height}:force_original_aspect_ratio=increase,"
            f"crop={width}:{height}:(iw-ow)*{x_pct:.6f}:(ih-oh)*{y_pct:.6f},setsar=1"
        )
    zoom = max(0.08, min(1.0, float(clip.get("scale") or 100) / 100))
    target_width = max(2, int(round(width * zoom)))
    target_height = max(2, int(round(height * zoom)))
    target_width += target_width % 2
    target_height += target_height % 2
    return (
        f"scale={target_width}:{target_height}:force_original_aspect_ratio=decrease,"
        f"pad={width}:{height}:(ow-iw)*{x_pct:.6f}:(oh-ih)*{y_pct:.6f}:color=black,setsar=1"
    )


def _render_video_project_from_clip(clip: dict[str, object], assets: dict[int, VideoEditorAsset], state: dict[str, object], output: Path, width: int, height: int, duration: float) -> None:
    asset = assets[int(clip.get("assetId") or 0)]
    source = Path(asset.file_path)
    if not source.exists():
        raise FileNotFoundError(asset.original_name or "Video asset not found")
    source_start = max(0, float(clip.get("sourceStart") or 0))
    clip_duration = max(0.25, min(float(clip.get("duration") or duration), duration))
    base_filter = _video_clip_base_filter(clip, width, height)
    args = [
        ffmpeg_path(),
        "-y",
        "-ss",
        str(source_start),
        "-t",
        str(clip_duration),
        "-i",
        str(source),
    ]
    overlay_inputs: list[tuple[dict[str, object], VideoEditorAsset]] = []
    for visual_clip in _video_export_visual_clips(state, assets):
        if visual_clip is clip or visual_clip.get("type") != "image":
            continue
        overlay_asset = assets.get(int(visual_clip.get("assetId") or 0))
        if overlay_asset and _asset_is_previewable_image(overlay_asset):
            overlay_inputs.append((visual_clip, overlay_asset))
            args += ["-loop", "1", "-t", str(duration), "-i", str(overlay_asset.file_path)]
    filters = [f"[0:v]{base_filter}[v0]"]
    current_label = "v0"
    for index, (visual_clip, overlay_asset) in enumerate(overlay_inputs, start=1):
        scale_pct = max(8, min(120, float(visual_clip.get("scale") or 42))) / 100
        overlay_width = max(24, min(width, int(width * scale_pct)))
        x_pct = max(0, min(100, float(visual_clip.get("x") or 50))) / 100
        y_pct = max(0, min(100, float(visual_clip.get("y") or 50))) / 100
        start = max(0, float(visual_clip.get("start") or 0))
        end = min(duration, start + max(0.25, float(visual_clip.get("duration") or 5)))
        img_label = f"img{index}"
        out_label = f"v{index}"
        filters.append(f"[{index}:v]scale={overlay_width}:-1[{img_label}]")
        filters.append(
            f"[{current_label}][{img_label}]overlay="
            f"x=(W-w)*{x_pct:.4f}:y=(H-h)*{y_pct:.4f}:"
            f"enable='between(t,{start:.3f},{end:.3f})'[{out_label}]"
        )
        current_label = out_label
    text_index = 0
    for visual_clip in _video_export_visual_clips(state, assets):
        if visual_clip.get("type") != "image":
            continue
        visual_asset = assets.get(int(visual_clip.get("assetId") or 0))
        if not visual_asset or _asset_is_previewable_image(visual_asset):
            continue
        text = _ffmpeg_drawtext_escape(f"{_asset_visual_kind(visual_asset).upper()}  {visual_asset.original_name or 'File'}")
        start = max(0, float(visual_clip.get("start") or 0))
        end = min(duration, start + max(0.25, float(visual_clip.get("duration") or 5)))
        y = int(height * (float(visual_clip.get("y") or 50) / 100))
        next_label = f"doc{text_index}"
        filters.append(
            f"[{current_label}]drawtext="
            f"text='{text}':"
            f"x=(w-text_w)/2:y={max(40, min(height - 90, y))}:fontsize={max(18, int(width * 0.045))}:fontcolor=white:"
            "box=1:boxcolor=black@0.55:boxborderw=22:"
            f"enable='between(t,{start:.3f},{end:.3f})'[{next_label}]"
        )
        current_label = next_label
        text_index += 1
    for text_clip in _video_export_text_clips(state):
        text = _ffmpeg_drawtext_escape(str(text_clip.get("text") or "Text"))
        style = text_clip.get("style") if isinstance(text_clip.get("style"), dict) else {}
        size = max(12, min(96, int(float(style.get("size") or 32))))
        start = max(0, float(text_clip.get("start") or 0))
        end = min(duration, start + max(0.25, float(text_clip.get("duration") or 4)))
        x_expr, y_expr = _ffmpeg_drawtext_xy(text_clip, width, height)
        draw_style = _ffmpeg_drawtext_style(style)
        filters.append(
            f"[{current_label}]drawtext="
            f"text='{text}':"
            f"x={x_expr}:y={y_expr}:fontsize={size}:{draw_style}:"
            f"enable='between(t,{start:.3f},{end:.3f})'[txt{text_index}]"
        )
        current_label = f"txt{text_index}"
        text_index += 1
    args += [
        "-filter_complex",
        ";".join(filters),
        "-map",
        f"[{current_label}]",
        "-map",
        "0:a?",
        "-c:v",
        "libx264",
        "-preset",
        "veryfast",
        "-crf",
        "22",
        "-pix_fmt",
        "yuv420p",
        "-c:a",
        "aac",
        "-b:a",
        "160k",
        "-movflags",
        "+faststart",
        str(output),
    ]
    completed = subprocess.run(args, capture_output=True, text=True, timeout=settings.video_timeout_seconds)
    if completed.returncode != 0:
        raise RuntimeError((completed.stderr.strip().splitlines() or ["FFmpeg export failed"])[-1])


def _render_empty_video_project(output: Path, width: int, height: int, duration: float) -> None:
    args = [
        ffmpeg_path(),
        "-y",
        "-f",
        "lavfi",
        "-i",
        f"color=c=black:s={width}x{height}:d={duration}:r=30",
        "-c:v",
        "libx264",
        "-preset",
        "veryfast",
        "-crf",
        "22",
        "-pix_fmt",
        "yuv420p",
        str(output),
    ]
    completed = subprocess.run(args, capture_output=True, text=True, timeout=settings.video_timeout_seconds)
    if completed.returncode != 0:
        raise RuntimeError((completed.stderr.strip().splitlines() or ["FFmpeg export failed"])[-1])


def _render_visual_card_project(state: dict[str, object], assets: dict[int, VideoEditorAsset], output: Path, width: int, height: int, duration: float) -> None:
    clips = _video_export_visual_clips(state, assets)
    first_image = next((clip for clip in clips if _asset_is_previewable_image(assets.get(int(clip.get("assetId") or 0)))), None)
    first_asset = assets.get(int(first_image.get("assetId") or 0)) if first_image else None
    filters: list[str] = []
    args = [ffmpeg_path(), "-y"]
    if first_asset:
        args += ["-loop", "1", "-t", str(duration), "-i", str(first_asset.file_path)]
        filters.append(f"scale={width}:{height}:force_original_aspect_ratio=decrease,pad={width}:{height}:(ow-iw)/2:(oh-ih)/2:color=black,setsar=1")
    else:
        args += ["-f", "lavfi", "-i", f"color=c=black:s={width}x{height}:d={duration}:r=30"]
    for visual_clip in clips:
        asset = assets.get(int(visual_clip.get("assetId") or 0))
        if not asset or (first_asset and asset.id == first_asset.id and visual_clip is first_image):
            continue
        if _asset_is_previewable_image(asset):
            text = _ffmpeg_drawtext_escape(asset.original_name or "Image")
        else:
            text = _ffmpeg_drawtext_escape(f"{_asset_visual_kind(asset).upper()}  {asset.original_name or 'File'}")
        start = max(0, float(visual_clip.get("start") or 0))
        end = min(duration, start + max(0.25, float(visual_clip.get("duration") or 5)))
        y = int(height * (float(visual_clip.get("y") or 50) / 100))
        filters.append(
            "drawtext="
            f"text='{text}':"
            f"x=(w-text_w)/2:y={max(40, min(height - 90, y))}:fontsize={max(18, int(width * 0.045))}:fontcolor=white:"
            "box=1:boxcolor=black@0.55:boxborderw=22:"
            f"enable='between(t,{start:.3f},{end:.3f})'"
        )
    for text_clip in _video_export_text_clips(state):
        text = _ffmpeg_drawtext_escape(str(text_clip.get("text") or "Text"))
        style = text_clip.get("style") if isinstance(text_clip.get("style"), dict) else {}
        size = max(12, min(96, int(float(style.get("size") or 32))))
        start = max(0, float(text_clip.get("start") or 0))
        end = min(duration, start + max(0.25, float(text_clip.get("duration") or 4)))
        x_expr, y_expr = _ffmpeg_drawtext_xy(text_clip, width, height)
        draw_style = _ffmpeg_drawtext_style(style)
        filters.append(
            "drawtext="
            f"text='{text}':"
            f"x={x_expr}:y={y_expr}:fontsize={size}:{draw_style}:"
            f"enable='between(t,{start:.3f},{end:.3f})'"
        )
    if filters:
        args += ["-vf", ",".join(filters)]
    args += [
        "-t",
        str(duration),
        "-c:v",
        "libx264",
        "-preset",
        "veryfast",
        "-crf",
        "22",
        "-pix_fmt",
        "yuv420p",
        "-movflags",
        "+faststart",
        str(output),
    ]
    completed = subprocess.run(args, capture_output=True, text=True, timeout=settings.video_timeout_seconds)
    if completed.returncode != 0:
        raise RuntimeError((completed.stderr.strip().splitlines() or ["FFmpeg export failed"])[-1])


def _video_export_text_clips(state: dict[str, object]) -> list[dict[str, object]]:
    clips = state.get("clips") if isinstance(state.get("clips"), list) else []
    return [clip for clip in clips if isinstance(clip, dict) and clip.get("type") in {"text", "caption"}]


def _video_export_visual_clips(state: dict[str, object], assets: dict[int, VideoEditorAsset]) -> list[dict[str, object]]:
    clips = state.get("clips") if isinstance(state.get("clips"), list) else []
    visual: list[dict[str, object]] = []
    for clip in clips:
        if not isinstance(clip, dict) or clip.get("type") not in {"video", "image"}:
            continue
        asset = assets.get(int(clip.get("assetId") or 0))
        if asset and asset.kind in {"video", "image"}:
            visual.append(clip)
    return sorted(visual, key=lambda item: float(item.get("start") or 0))


def _ffmpeg_drawtext_style(style: dict[str, object]) -> str:
    color = _ffmpeg_color(str(style.get("color") or "#ffffff"), fallback="white")
    stroke = _ffmpeg_color(str(style.get("stroke") or "#000000"), fallback="black")
    stroke_width = max(0, min(8, int(float(style.get("strokeWidth") or 0))))
    bg = _ffmpeg_color(str(style.get("bg") or "#000000"), fallback="black")
    bg_alpha = max(0, min(100, int(float(style.get("bgAlpha") or 0)))) / 100
    parts = [f"fontcolor={color}"]
    if stroke_width:
        parts.extend([f"borderw={stroke_width}", f"bordercolor={stroke}"])
    if bg_alpha > 0:
        parts.extend(["box=1", f"boxcolor={bg}@{bg_alpha:.2f}", "boxborderw=18"])
    else:
        parts.extend(["box=0", "boxborderw=0"])
    if str(style.get("textShadow") or "").strip().lower() != "none":
        parts.extend(["shadowcolor=black@0.55", "shadowx=2", "shadowy=3"])
    return ":".join(parts)


def _ffmpeg_drawtext_xy(clip: dict[str, object], width: int, height: int) -> tuple[str, str]:
    x_pct = max(-140.0, min(240.0, float(clip.get("x") or 50)))
    y_pct = max(-140.0, min(240.0, float(clip.get("y") or 78)))
    x = int(width * (x_pct / 100))
    y = int(height * (y_pct / 100))
    x_expr = f"{x}-text_w/2"
    y_expr = f"{y}-text_h/2"
    return x_expr, y_expr


def _ffmpeg_color(value: str, *, fallback: str) -> str:
    raw = str(value or "").strip()
    if re.fullmatch(r"#[0-9a-fA-F]{6}", raw):
        return f"0x{raw[1:]}"
    if re.fullmatch(r"[A-Za-z]+", raw):
        return raw.lower()
    return fallback


def _ffmpeg_drawtext_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace(":", "\\:").replace("'", "\\'").replace("\n", " ")[:500]


def _decode_subtitle_bytes(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "cp1251", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _parse_subtitle_cues(text: str, filename: str = "") -> list[dict[str, object]]:
    source = str(text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not source:
        return []
    if filename.lower().endswith(".ass") or "[events]" in source.lower():
        cues = _parse_ass_subtitles(source)
    elif filename.lower().endswith(".json") or source[:1] in {"[", "{"}:
        cues = _parse_json_subtitles(source)
    else:
        cues = _parse_srt_vtt_subtitles(source)
    return _normalize_subtitle_cues(cues)


def _parse_json_subtitles(text: str) -> list[dict[str, object]]:
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        return []
    raw_cues = data.get("cues") if isinstance(data, dict) else data
    if not isinstance(raw_cues, list):
        return []
    cues: list[dict[str, object]] = []
    for item in raw_cues:
        if not isinstance(item, dict):
            continue
        cue: dict[str, object] = {"start": item.get("start"), "end": item.get("end"), "text": item.get("text")}
        for key in ("x", "y", "scale", "boxWidth", "rotation", "style", "source"):
            if key in item:
                cue[key] = item[key]
        cues.append(cue)
    return cues


def _parse_ass_subtitles(text: str) -> list[dict[str, object]]:
    cues: list[dict[str, object]] = []
    for line in text.splitlines():
        if not line.lower().startswith("dialogue:"):
            continue
        payload = line.split(":", 1)[1].strip()
        parts = payload.split(",", 9)
        if len(parts) < 10:
            continue
        cues.append({"start": parts[1], "end": parts[2], "text": parts[9].replace("\\N", "\n").strip()})
    return cues


def _parse_srt_vtt_subtitles(text: str) -> list[dict[str, object]]:
    cleaned = re.sub(r"^\ufeff?WEBVTT[^\n]*\n+", "", text.strip(), flags=re.IGNORECASE)
    blocks = re.split(r"\n{2,}", cleaned)
    cues: list[dict[str, object]] = []
    for block in blocks:
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if not lines or lines[0].upper().startswith(("NOTE", "STYLE", "REGION")):
            continue
        time_index = next((index for index, line in enumerate(lines) if "-->" in line), -1)
        if time_index < 0:
            continue
        start_raw, end_raw = lines[time_index].split("-->", 1)
        end_raw = end_raw.strip().split()[0]
        caption_text = "\n".join(lines[time_index + 1 :]).strip()
        cues.append({"start": start_raw.strip().split()[0], "end": end_raw, "text": _clean_subtitle_text(caption_text)})
    return cues


def _normalize_subtitle_cues(raw_cues: list[dict[str, object]]) -> list[dict[str, object]]:
    cues: list[dict[str, object]] = []
    for item in raw_cues:
        start = _subtitle_seconds(item.get("start"))
        end = _subtitle_seconds(item.get("end"))
        text = _clean_subtitle_text(str(item.get("text") or ""))
        if end <= start:
            end = start + 2.0
        if text:
            cue: dict[str, object] = {"start": round(start, 3), "end": round(end, 3), "text": text[:800]}
            for key in ("x", "y", "scale", "boxWidth", "rotation", "style", "source"):
                if key in item:
                    cue[key] = item[key]
            cues.append(cue)
    cues.sort(key=lambda cue: (float(cue["start"]), float(cue["end"])))
    return cues[:5000]


def _normalize_auto_subtitle_cues(transcribed, params: dict[str, object]) -> list[dict[str, object]]:
    timeline_start = _subtitle_seconds(params.get("timeline_start"))
    source_start = _subtitle_seconds(params.get("source_start"))
    source_end = _subtitle_seconds(params.get("source_end"))
    clip_duration = _subtitle_seconds(params.get("clip_duration"))
    if source_end <= source_start and clip_duration > 0:
        source_end = source_start + clip_duration
    clip_end = timeline_start + max(0.0, source_end - source_start) if source_end > source_start else None

    raw_cues: list[dict[str, object]] = []
    for cue in transcribed:
        start = max(0.0, float(cue.start))
        end = max(start + 0.05, float(cue.end))
        if end <= source_start:
            continue
        if source_end > source_start and start >= source_end:
            continue
        timeline_cue_start = timeline_start + max(0.0, start - source_start)
        timeline_cue_end = timeline_start + max(0.05, end - source_start)
        if clip_end is not None:
            timeline_cue_end = min(timeline_cue_end, clip_end)
        if timeline_cue_end <= timeline_cue_start:
            continue
        raw_cues.append({"start": timeline_cue_start, "end": timeline_cue_end, "text": cue.text})
    return _normalize_subtitle_cues(raw_cues)


def _append_video_project_caption_clips(project: VideoEditorProject, cues: list[dict[str, object]]) -> VideoEditorProject:
    state = project.state_json if isinstance(project.state_json, dict) else {}
    state = json.loads(json.dumps(state))
    tracks = state.get("tracks") if isinstance(state.get("tracks"), list) else []
    text_track = next((track for track in tracks if isinstance(track, dict) and track.get("type") == "text"), None)
    if not text_track:
        text_track = {"id": "text-1", "type": "text", "name": "Text", "order": 1}
        tracks.append(text_track)
    state["tracks"] = tracks
    clips = state.get("clips") if isinstance(state.get("clips"), list) else []
    for cue in cues:
        start = max(0.0, float(cue.get("start") or 0))
        end = max(start + 0.2, float(cue.get("end") or start + 2))
        cue_style = cue.get("style") if isinstance(cue.get("style"), dict) else {}
        default_style = {
            "font": "system-ui, -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif",
            "size": 24,
            "color": "#ffffff",
            "stroke": "#000000",
            "strokeWidth": 1,
            "bg": "#000000",
            "bgAlpha": 42,
        }
        default_style.update(cue_style)
        source = cue.get("source") if isinstance(cue.get("source"), dict) else {}
        clips.append(
            {
                "id": f"caption-{uuid.uuid4().hex[:10]}",
                "type": "caption",
                "trackId": text_track.get("id") or "text-1",
                "start": round(start, 3),
                "duration": round(max(0.2, end - start), 3),
                "x": max(0.0, min(100.0, float(cue.get("x") or 50))),
                "y": max(0.0, min(100.0, float(cue.get("y") or 84))),
                "scale": max(10.0, min(300.0, float(cue.get("scale") or 100))),
                "boxWidth": max(18.0, min(86.0, float(cue.get("boxWidth") or 42))),
                "rotation": float(cue.get("rotation") or 0),
                "text": str(cue.get("text") or "Caption"),
                "style": default_style,
                "subtitleSource": source,
            }
        )
    state["clips"] = clips
    project.state_json = state
    _update_video_project_metadata(project)
    project.storage_bytes = _video_project_storage_bytes(project)
    project.save(update_fields=["state_json", "storage_bytes", "asset_count", "clip_count", "duration_seconds", "thumbnail_path", "updated_at"])
    return project


def _video_project_caption_cues(state: dict[str, object], *, rich: bool = False) -> list[dict[str, object]]:
    raw_clips = state.get("clips") if isinstance(state.get("clips"), list) else []
    cues: list[dict[str, object]] = []
    for clip in raw_clips:
        if not isinstance(clip, dict) or clip.get("type") not in {"caption", "text"}:
            continue
        start = max(0.0, float(clip.get("start") or 0))
        duration = max(0.1, float(clip.get("duration") or 0))
        text = _clean_subtitle_text(str(clip.get("text") or ""))
        if text:
            cue: dict[str, object] = {"start": round(start, 3), "end": round(start + duration, 3), "text": text[:800]}
            if rich:
                style = clip.get("style") if isinstance(clip.get("style"), dict) else {}
                cue.update(
                    {
                        "x": max(0.0, min(100.0, float(clip.get("x") or 50))),
                        "y": max(0.0, min(100.0, float(clip.get("y") or 84))),
                        "scale": max(10.0, min(300.0, float(clip.get("scale") or 100))),
                        "boxWidth": max(18.0, min(86.0, float(clip.get("boxWidth") or 42))),
                        "rotation": float(clip.get("rotation") or 0),
                        "style": style,
                        "source": clip.get("subtitleSource") if isinstance(clip.get("subtitleSource"), dict) else {},
                    }
                )
            cues.append(cue)
    return _normalize_subtitle_cues(cues)


def _subtitle_seconds(value: object) -> float:
    if isinstance(value, (int, float)):
        return max(0.0, float(value))
    raw = str(value or "").strip()
    if not raw:
        return 0.0
    raw = raw.replace(",", ".")
    match = re.search(r"(?:(\d+):)?(\d{1,2}):(\d{1,2}(?:\.\d+)?)", raw)
    if match:
        hours = float(match.group(1) or 0)
        minutes = float(match.group(2) or 0)
        seconds = float(match.group(3) or 0)
        return max(0.0, hours * 3600 + minutes * 60 + seconds)
    try:
        return max(0.0, float(raw))
    except ValueError:
        return 0.0


def _clean_subtitle_text(value: str) -> str:
    text = re.sub(r"<[^>]+>", "", str(value or ""))
    text = re.sub(r"\{\\[^}]+\}", "", text)
    return "\n".join(_polish_subtitle_line(line) for line in text.splitlines() if line.strip()).strip()


def _polish_subtitle_line(value: str) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    text = re.sub(r"\s+([,.;:!?…])", r"\1", text)
    text = re.sub(r"([,.;:!?…])([^\s,.;:!?…])", r"\1 \2", text)
    text = re.sub(r"\.{2,}", "…", text)
    text = re.sub(r"([!?]){3,}", r"\1", text)
    text = re.sub(r"\b(\w{2,})\s+\1\b", r"\1", text, flags=re.IGNORECASE)
    if text:
        text = text[0].upper() + text[1:]
    return text


def _render_srt(cues: list[dict[str, object]]) -> str:
    blocks = []
    for index, cue in enumerate(cues, start=1):
        blocks.append(f"{index}\n{_format_srt_time(float(cue['start']))} --> {_format_srt_time(float(cue['end']))}\n{cue['text']}")
    return "\n\n".join(blocks) + ("\n" if blocks else "")


def _render_vtt(cues: list[dict[str, object]]) -> str:
    blocks = ["WEBVTT"]
    for cue in cues:
        blocks.append(f"{_format_vtt_time(float(cue['start']))} --> {_format_vtt_time(float(cue['end']))}\n{cue['text']}")
    return "\n\n".join(blocks) + "\n"


def _render_ass(cues: list[dict[str, object]], state: dict[str, object]) -> str:
    width, height = _video_export_size(state if isinstance(state, dict) else {}, "1080p")
    workflow = state.get("subtitleWorkflow") if isinstance(state, dict) and isinstance(state.get("subtitleWorkflow"), dict) else {}
    workflow_style = str(workflow.get("style") or "").strip().lower()
    default_style = _video_editor_caption_style(workflow_style)
    lines = [
        "[Script Info]",
        "ScriptType: v4.00+",
        f"PlayResX: {width}",
        f"PlayResY: {height}",
        "WrapStyle: 0",
        "ScaledBorderAndShadow: yes",
        "",
        "[V4+ Styles]",
        "Format: Name, Fontname, Fontsize, PrimaryColour, SecondaryColour, OutlineColour, BackColour, Bold, Italic, Underline, StrikeOut, ScaleX, ScaleY, Spacing, Angle, BorderStyle, Outline, Shadow, Alignment, MarginL, MarginR, MarginV, Encoding",
        _ass_style_line("Default", default_style),
        "",
        "[Events]",
        "Format: Layer, Start, End, Style, Name, MarginL, MarginR, MarginV, Effect, Text",
    ]
    for cue in cues:
        style = cue.get("style") if isinstance(cue.get("style"), dict) else {}
        merged_style = dict(default_style)
        merged_style.update(style)
        x = int(width * (max(0.0, min(100.0, float(cue.get("x") or 50))) / 100))
        y = int(height * (max(0.0, min(100.0, float(cue.get("y") or 84))) / 100))
        scale = max(10.0, min(300.0, float(cue.get("scale") or 100)))
        rotation = float(cue.get("rotation") or 0)
        start = float(cue["start"])
        end = float(cue["end"])
        duration_ms = max(100, int((end - start) * 1000))
        tags = [r"\an2", fr"\pos({x},{y})"]
        if scale != 100:
            tags.append(fr"\fscx{scale:.0f}\fscy{scale:.0f}")
        if abs(rotation) >= 0.5:
            tags.append(fr"\frz{rotation:.1f}")
        tags.extend(_ass_inline_style_tags(merged_style))
        source = cue.get("source") if isinstance(cue.get("source"), dict) else {}
        cue_style = str(source.get("style") or workflow_style or "").strip().lower()
        if cue_style in {"kinetic", "bounce", "pop", "headline"}:
            tags.append(r"\fad(80,120)")
            tags.append(r"\t(0,160,\fscx108\fscy108)")
            tags.append(r"\t(160,320,\fscx100\fscy100)")
        elif cue_style in {"neon", "candy"}:
            tags.append(r"\fad(70,140)")
            tags.append(fr"\t(0,{min(duration_ms, 260)},\blur1)")
        text = _ass_escape_text(str(cue["text"]))
        lines.append(f"Dialogue: 0,{_format_ass_time(start)},{_format_ass_time(end)},Default,,0,0,0,,{{{''.join(tags)}}}{text}")
    return "\n".join(lines) + "\n"


def _ass_style_line(name: str, style: dict[str, object]) -> str:
    font = _ass_font_name(str(style.get("font") or "Arial"))
    size = max(12, min(140, int(float(style.get("size") or 52))))
    color = _hex_to_ass_color(str(style.get("color") or "#ffffff"), alpha=0)
    stroke = _hex_to_ass_color(str(style.get("stroke") or "#000000"), alpha=0)
    bg_alpha = max(0, min(100, int(float(style.get("bgAlpha") or 0))))
    bg = _hex_to_ass_color(str(style.get("bg") or "#000000"), alpha=255 - int(bg_alpha * 2.55))
    outline = max(0, min(12, int(float(style.get("strokeWidth") or 2))))
    bold = -1 if int(float(style.get("fontWeight") or 800)) >= 700 else 0
    return (
        f"Style: {name},{font},{size},{color},&H000000FF,{stroke},{bg},"
        f"{bold},0,0,0,100,100,0,0,1,{outline},2,2,54,54,120,1"
    )


def _ass_inline_style_tags(style: dict[str, object]) -> list[str]:
    tags: list[str] = []
    if style.get("color"):
        tags.append(fr"\c{_hex_to_ass_color(str(style.get('color')), alpha=0)}")
    if style.get("stroke"):
        tags.append(fr"\3c{_hex_to_ass_color(str(style.get('stroke')), alpha=0)}")
    if style.get("strokeWidth") is not None:
        tags.append(fr"\bord{max(0, min(12, int(float(style.get('strokeWidth') or 0))))}")
    if style.get("size"):
        tags.append(fr"\fs{max(12, min(140, int(float(style.get('size') or 52))))}")
    if int(float(style.get("fontWeight") or 800)) >= 700:
        tags.append(r"\b1")
    return tags


def _ass_font_name(value: str) -> str:
    font = value.split(",", 1)[0].strip().strip("'\"") or "Arial"
    aliases = {"system-ui": "Segoe UI", "-apple-system": "Segoe UI", "blinkmacsystemfont": "Segoe UI"}
    return aliases.get(font.lower(), font)


def _hex_to_ass_color(value: str, *, alpha: int = 0) -> str:
    raw = str(value or "").strip()
    match = re.fullmatch(r"#?([0-9a-fA-F]{6})", raw)
    if not match:
        raw = "#ffffff"
        match = re.fullmatch(r"#?([0-9a-fA-F]{6})", raw)
    hex_value = match.group(1)
    r = int(hex_value[0:2], 16)
    g = int(hex_value[2:4], 16)
    b = int(hex_value[4:6], 16)
    return f"&H{max(0, min(255, alpha)):02X}{b:02X}{g:02X}{r:02X}"


def _ass_escape_text(value: str) -> str:
    return str(value).replace("{", "(").replace("}", ")").replace("\n", "\\N")


def _format_srt_time(seconds: float) -> str:
    hours, minutes, whole, millis = _split_subtitle_time(seconds)
    return f"{hours:02}:{minutes:02}:{whole:02},{millis:03}"


def _format_vtt_time(seconds: float) -> str:
    hours, minutes, whole, millis = _split_subtitle_time(seconds)
    return f"{hours:02}:{minutes:02}:{whole:02}.{millis:03}"


def _format_ass_time(seconds: float) -> str:
    hours, minutes, whole, millis = _split_subtitle_time(seconds)
    centis = min(99, int(round(millis / 10)))
    return f"{hours}:{minutes:02}:{whole:02}.{centis:02}"


def _split_subtitle_time(seconds: float) -> tuple[int, int, int, int]:
    total_millis = max(0, int(round(float(seconds) * 1000)))
    hours, remainder = divmod(total_millis, 3600_000)
    minutes, remainder = divmod(remainder, 60_000)
    whole, millis = divmod(remainder, 1000)
    return hours, minutes, whole, millis


def _json_body(request: HttpRequest) -> dict[str, object]:
    if not request.body:
        return {}
    try:
        value = json.loads(request.body.decode("utf-8"))
    except Exception:
        return {}
    return value if isinstance(value, dict) else {}


def _project_listing_payload(
    request: HttpRequest,
    queryset,
    resource_type: str,
    owner_id: int | None,
    guest_key: str,
    payload_builder,
    *,
    per_page_default: int = 24,
) -> dict[str, object]:
    query = " ".join(str(request.GET.get("q") or request.GET.get("query") or "").split())
    sort = str(request.GET.get("sort") or "updated").strip().lower()
    try:
        page = max(1, int(request.GET.get("page") or 1))
    except (TypeError, ValueError):
        page = 1
    try:
        per_page = int(request.GET.get("per_page") or per_page_default)
    except (TypeError, ValueError):
        per_page = per_page_default
    per_page = max(1, min(80, per_page))

    if query:
        queryset = queryset.filter(title__icontains=query)
    ordering = {
        "created": "-created_at",
        "title": "title",
        "name": "title",
        "oldest": "updated_at",
        "updated": "-updated_at",
        "recent": "-updated_at",
    }.get(sort, "-updated_at")
    queryset = queryset.order_by(ordering, "-id" if ordering != "-id" else "id")

    total = queryset.count()
    start = (page - 1) * per_page
    rows = list(queryset[start : start + per_page + 1])
    has_more = len(rows) > per_page
    rows = rows[:per_page]
    rows = _attach_access_roles(rows, resource_type, owner_id, guest_key)
    return {
        "projects": [payload_builder(project) for project in rows],
        "page": page,
        "per_page": per_page,
        "has_more": has_more,
        "total": total,
        "query": query,
        "sort": sort,
    }


def _post_bool(request: HttpRequest, name: str) -> bool:
    return str(request.POST.get(name, "")).strip().lower() in {"1", "true", "yes", "on"}


def _json_size(value: object) -> int:
    return len(json.dumps(value, ensure_ascii=False, separators=(",", ":")).encode("utf-8"))


def _video_asset_payload(asset: VideoEditorAsset) -> dict[str, object]:
    extension = Path(asset.original_name or asset.file_path).suffix.lower()
    payload = {
        "id": asset.id,
        "kind": asset.kind,
        "name": asset.original_name,
        "media_type": asset.media_type,
        "visual_kind": _asset_visual_kind(asset),
        "is_previewable_image": _asset_is_previewable_image(asset),
        "extension": extension,
        "size": asset.size,
        "size_text": human_size(asset.size),
        "duration": asset.duration,
        "preview_url": reverse("studio:video_project_asset_preview", args=[asset.project_id, asset.id]),
        "waveform_url": reverse("studio:video_project_asset_waveform", args=[asset.project_id, asset.id]) if asset.kind in {"audio", "video"} else "",
        "rename_url": reverse("studio:rename_video_project_asset", args=[asset.project_id, asset.id]),
        "delete_url": reverse("studio:delete_video_project_asset", args=[asset.project_id, asset.id]),
        "thumbnail_url": "",
    }
    if asset.thumbnail_path:
        payload["thumbnail_url"] = reverse("studio:video_project_asset_thumbnail", args=[asset.project_id, asset.id])
    return payload


def _asset_visual_kind(asset: VideoEditorAsset) -> str:
    if asset.kind != "image":
        return asset.kind
    media_type = asset.media_type or ""
    extension = Path(asset.original_name or asset.file_path).suffix.lower()
    if media_type.startswith("image/"):
        return "image"
    if media_type == "application/pdf" or extension == ".pdf":
        return "pdf"
    if media_type.startswith("text/") or extension in {".doc", ".docx", ".ppt", ".pptx", ".xls", ".xlsx", ".txt", ".csv", ".json", ".rtf"}:
        return "document"
    return "file"


def _asset_is_previewable_image(asset: VideoEditorAsset | None) -> bool:
    return bool(asset and asset.media_type.startswith("image/") and Path(asset.file_path).exists())


def _is_visual_document_type(media_type: str, filename: str) -> bool:
    extension = Path(filename or "").suffix.lower()
    return (
        media_type == "application/pdf"
        or media_type.startswith("text/")
        or extension in {".pdf", ".doc", ".docx", ".ppt", ".pptx", ".xls", ".xlsx", ".txt", ".csv", ".json", ".rtf"}
    )


def _save_optimized_editor_image(upload, target_dir: Path, base: str) -> tuple[Path, str]:
    try:
        image = Image.open(upload)
        image = ImageOps.exif_transpose(image)
        image.thumbnail((2560, 2560), Image.Resampling.LANCZOS)
        has_alpha = image.mode in {"RGBA", "LA"} or (image.mode == "P" and "transparency" in image.info)
        suffix = ".webp" if has_alpha else ".jpg"
        path = target_dir / f"{uuid.uuid4().hex[:12]}_{base}{suffix}"
        if has_alpha:
            image.save(path, "WEBP", quality=88, method=6)
            return path, "image/webp"
        image.convert("RGB").save(path, "JPEG", quality=88, optimize=True, progressive=True)
        return path, "image/jpeg"
    except Exception:
        upload.seek(0)
        media_type = upload.content_type or mimetypes.guess_type(upload.name or "")[0] or "image/jpeg"
        suffix = Path(upload.name or "").suffix[:16] or mimetypes.guess_extension(media_type) or ".img"
        path = target_dir / f"{uuid.uuid4().hex[:12]}_{base}{suffix}"
        with path.open("wb") as destination:
            for chunk in upload.chunks():
                destination.write(chunk)
        return path, media_type


def _video_project_storage_bytes(project: VideoEditorProject) -> int:
    return _json_size(project.state_json or {}) + sum(asset.size + _path_size(asset.thumbnail_path) for asset in project.assets.all())


def _design_project_storage_bytes(project: DesignerProject) -> int:
    return _json_size(project.state_json or {}) + _path_size(project.preview_path) + sum(asset.size for asset in project.assets.all())


def _music_project_storage_bytes(project: MusicEditorProject) -> int:
    return _json_size(project.state_json or {}) + sum(asset.size for asset in project.assets.all())


def _music_import_is_audio(name: str, media_type: str = "") -> bool:
    suffix = Path(name or "").suffix.lower()
    return (media_type or "").startswith("audio/") or suffix in {".wav", ".mp3", ".ogg", ".oga", ".flac", ".m4a", ".aac", ".aiff", ".aif", ".webm"}


def _save_music_import_asset(project: MusicEditorProject, asset_dir: Path, original_name: str, media_type: str, data: bytes, kind: str = "audio") -> MusicEditorAsset:
    suffix = Path(original_name or "").suffix[:16] or mimetypes.guess_extension(media_type) or (".wav" if kind == "audio" else ".bin")
    base = clean_base_name(Path(original_name or kind).stem or kind, kind)
    path = asset_dir / f"{uuid.uuid4().hex[:12]}_{base}{suffix}"
    path.write_bytes(data)
    duration = _get_audio_duration(path) if kind == "audio" else 0.0
    return MusicEditorAsset.objects.create(
        project=project,
        kind=kind,
        file_path=str(path),
        media_type=media_type or "application/octet-stream",
        size=path.stat().st_size,
        original_name=(original_name or kind)[:240],
        duration=duration,
    )


def _save_music_upload_asset(project: MusicEditorProject, asset_dir: Path, upload, original_name: str, media_type: str, kind: str = "audio") -> MusicEditorAsset:
    suffix = Path(original_name or "").suffix[:16] or mimetypes.guess_extension(media_type) or (".wav" if kind == "audio" else ".bin")
    base = clean_base_name(Path(original_name or kind).stem or kind, kind)
    path = asset_dir / f"{uuid.uuid4().hex[:12]}_{base}{suffix}"
    with path.open("wb") as destination:
        for chunk in upload.chunks():
            destination.write(chunk)
    duration = _get_audio_duration(path) if kind == "audio" else 0.0
    return MusicEditorAsset.objects.create(
        project=project,
        kind=kind,
        file_path=str(path),
        media_type=media_type or "application/octet-stream",
        size=path.stat().st_size,
        original_name=(original_name or kind)[:240],
        duration=duration,
    )


def _owned_job_records(owner_id: int | None, guest_key: str = ""):
    queryset = JobRecord.objects.all()
    if owner_id is not None:
        return queryset.filter(owner_id=owner_id)
    return queryset.filter(owner__isnull=True, guest_key=guest_key)


def _job_param(record: JobRecord, key: str, default=None):
    try:
        params = json.loads(record.params_json or "{}")
    except Exception:
        params = {}
    return params.get(key, default) if isinstance(params, dict) else default


def _merge_job_params(job_id: str, values: dict[str, object]) -> None:
    try:
        record = JobRecord.objects.filter(job_id=job_id).first()
        if not record:
            return
        try:
            params = json.loads(record.params_json or "{}")
        except Exception:
            params = {}
        if not isinstance(params, dict):
            params = {}
        params.update(values)
        record.params_json = json.dumps(params, ensure_ascii=False, default=str)
        record.save(update_fields=["params_json", "updated_at"])
    except Exception:
        return


def _video_asset_waveform_path(asset: VideoEditorAsset) -> Path:
    return _video_project_media_dir(asset.project) / "waveforms" / f"{asset.id}.json"


def _build_compact_waveform(path: Path, count: int = 96) -> list[float]:
    if not path.exists():
        return [0.0] * count
    data = path.read_bytes()
    if not data:
        return [0.0] * count
    step = max(1, len(data) // count)
    samples: list[float] = []
    for index in range(count):
        chunk = data[index * step : (index + 1) * step]
        if not chunk:
            samples.append(0.0)
            continue
        value = sum(abs(byte - 128) for byte in chunk) / (len(chunk) * 128)
        samples.append(round(max(0.04, min(1.0, value)), 3))
    return samples


def _render_video_project_cover(project: VideoEditorProject, output: Path, time_seconds: float) -> None:
    state = project.state_json or {}
    width, height = _video_export_size(state, "720p")
    assets = {asset.id: asset for asset in project.assets.all()}
    clips = state.get("clips") if isinstance(state.get("clips"), list) else []
    active = [clip for clip in clips if isinstance(clip, dict) and float(clip.get("start") or 0) <= time_seconds <= float(clip.get("start") or 0) + float(clip.get("duration") or 0)]
    image = Image.new("RGB", (width, height), _state_background_color(state))
    visual = next((clip for clip in active if clip.get("type") in {"image", "video"} and assets.get(int(clip.get("assetId") or 0))), None)
    if visual:
        asset = assets[int(visual.get("assetId") or 0)]
        source = Path(asset.thumbnail_path or asset.file_path)
        if source.exists() and asset.media_type.startswith("image/"):
            overlay = Image.open(source).convert("RGB")
            overlay.thumbnail((width, height), Image.Resampling.LANCZOS)
            image.paste(overlay, ((width - overlay.width) // 2, (height - overlay.height) // 2))
    output.parent.mkdir(parents=True, exist_ok=True)
    image.save(output, "JPEG", quality=90, optimize=True)


def _state_background_color(state: dict[str, object]) -> str:
    value = str(state.get("backgroundValue") or state.get("background") or "#000000")
    if value.startswith("#") and len(value) in {4, 7}:
        return value
    return "#000000"


def _path_size(path_text: str) -> int:
    if not path_text:
        return 0
    path = Path(path_text)
    return path.stat().st_size if path.exists() and path.is_file() else 0


def _get_audio_duration(path: Path) -> float:
    """Get audio duration in seconds using ffprobe."""
    try:
        result = subprocess.run(
            [
                str(ffmpeg_path("ffprobe")),
                "-v",
                "error",
                "-show_entries",
                "format=duration",
                "-of",
                "default=noprint_wrappers=1:nokey=1:nokey=1",
                str(path),
            ],
            capture_output=True,
            text=True,
            timeout=10,
        )
        return float(result.stdout.strip()) if result.stdout.strip() else 0.0
    except Exception:
        return 0.0


def _video_project_media_dir(project: VideoEditorProject) -> Path:
    owner = f"user_{project.owner_id}" if project.owner_id else f"guest_{project.guest_key or 'anon'}"
    return settings.storage_dir / "django_video_projects" / owner / str(project.id)


def _design_project_media_dir(project: DesignerProject) -> Path:
    owner = f"user_{project.owner_id}" if project.owner_id else f"guest_{project.guest_key or 'anon'}"
    return settings.storage_dir / "django_design_projects" / owner / str(project.id)


def _music_project_media_dir(project: MusicEditorProject) -> Path:
    owner = f"user_{project.owner_id}" if project.owner_id else f"guest_{project.guest_key or 'anon'}"
    return settings.storage_dir / "django_music_projects" / owner / str(project.id)


def _save_design_project_preview(data_url: str, target_dir: Path) -> Path:
    target_dir.mkdir(parents=True, exist_ok=True)
    raw = data_url.split(",", 1)[1] if "," in data_url else data_url
    image = Image.open(BytesIO(b64decode(raw))).convert("RGB")
    image.thumbnail((900, 640), Image.Resampling.LANCZOS)
    path = target_dir / f"preview_{uuid.uuid4().hex[:12]}.jpg"
    image.save(path, "JPEG", quality=82, optimize=True)
    return path


def _save_project_thumbnail(data_url: str, target_dir: Path) -> Path:
    raw = data_url.split(",", 1)[1] if "," in data_url else data_url
    image = Image.open(BytesIO(b64decode(raw))).convert("RGB")
    image.thumbnail((960, 540))
    path = target_dir / f"thumb_{uuid.uuid4().hex[:12]}.jpg"
    image.save(path, "JPEG", quality=82, optimize=True)
    return path


def _delete_video_asset_files(asset: VideoEditorAsset) -> None:
    for path_text in (asset.file_path, asset.thumbnail_path):
        if not path_text:
            continue
        path = Path(path_text)
        try:
            if path.resolve().is_relative_to(settings.storage_dir.resolve()) and path.exists() and path.is_file():
                path.unlink()
        except Exception:
            pass


def _delete_video_project_media(project: VideoEditorProject) -> None:
    for asset in project.assets.all():
        _delete_video_asset_files(asset)
    directory = _video_project_media_dir(project)
    try:
        if directory.resolve().is_relative_to(settings.storage_dir.resolve()) and directory.exists():
            shutil.rmtree(directory)
    except Exception:
        pass


def _delete_design_project_media(project: DesignerProject) -> None:
    for asset in project.assets.all():
        path = Path(asset.file_path)
        try:
            if path.resolve().is_relative_to(settings.storage_dir.resolve()) and path.exists() and path.is_file():
                path.unlink()
        except Exception:
            pass
    if project.preview_path:
        path = Path(project.preview_path)
        try:
            if path.resolve().is_relative_to(settings.storage_dir.resolve()) and path.exists() and path.is_file():
                path.unlink()
        except Exception:
            pass
    directory = _design_project_media_dir(project)
    try:
        if directory.resolve().is_relative_to(settings.storage_dir.resolve()) and directory.exists():
            shutil.rmtree(directory)
    except Exception:
        pass


def _delete_music_project_media(project: MusicEditorProject) -> None:
    for asset in project.assets.all():
        path = Path(asset.file_path)
        try:
            if path.resolve().is_relative_to(settings.storage_dir.resolve()) and path.exists() and path.is_file():
                path.unlink()
        except Exception:
            pass
    directory = _music_project_media_dir(project)
    try:
        if directory.resolve().is_relative_to(settings.storage_dir.resolve()) and directory.exists():
            shutil.rmtree(directory)
    except Exception:
        pass


def _clean_project_title(value: str) -> str:
    title = " ".join(repair_mojibake(value).strip().split())
    return (title or "New project")[:180]


def _transfer_guest_video_projects(guest_key: str, user) -> None:
    if not guest_key:
        return
    VideoEditorProject.objects.filter(owner__isnull=True, guest_key=guest_key).update(owner=user, guest_key="")


def _transfer_guest_design_projects(guest_key: str, user) -> None:
    if not guest_key:
        return
    DesignerProject.objects.filter(owner__isnull=True, guest_key=guest_key).update(owner=user, guest_key="")


def _checkout_url(request: HttpRequest) -> str:
    return f"{reverse('billing:checkout')}?{urlencode({'next': request.get_full_path()})}"


def _display_name(request: HttpRequest) -> str:
    if request.user.is_authenticated:
        return request.user.first_name or request.user.email or "CherryX user"
    return "Guest workspace"


def _resume_prefill(request: HttpRequest) -> dict[str, object]:
    if not request.user.is_authenticated:
        return {"is_guest": True, "name": "", "email": "", "avatar_url": _avatar_url(request), "can_use_avatar": False}
    user = request.user
    display = user.get_full_name() or user.first_name or ""
    if not display and user.email:
        display = user.email.split("@", 1)[0].replace(".", " ").replace("_", " ").title()
    can_use_avatar = False
    try:
        avatar_path = user.studio_profile.avatar_path
        can_use_avatar = bool(avatar_path and Path(avatar_path).exists())
    except AccountProfile.DoesNotExist:
        pass
    return {
        "is_guest": False,
        "name": display,
        "email": user.email or "",
        "avatar_url": _avatar_url(request),
        "can_use_avatar": can_use_avatar,
    }


def _cherryx_balance(request: HttpRequest) -> int:
    if request.user.is_authenticated:
        try:
            return int(request.user.studio_profile.cherryx_balance or 0)
        except AccountProfile.DoesNotExist:
            return 0
    seed = len(_guest_key(request))
    return 1200 + (seed % 7) * 150


def _avatar_url(request: HttpRequest) -> str:
    if request.user.is_authenticated:
        try:
            if request.user.studio_profile.avatar_path:
                return reverse("studio:account_avatar")
            if request.user.studio_profile.avatar_url:
                return request.user.studio_profile.avatar_url
        except AccountProfile.DoesNotExist:
            pass
    return "https://api.iconify.design/lucide/circle-user-round.svg?color=%232563eb&width=80&height=80"


def _accent_color(request: HttpRequest) -> str:
    if request.user.is_authenticated:
        try:
            return request.user.studio_profile.accent_color or "#2563eb"
        except AccountProfile.DoesNotExist:
            pass
    return "#2563eb"


def _theme_mode(request: HttpRequest) -> str:
    if request.user.is_authenticated:
        try:
            mode = request.user.studio_profile.theme_mode or "light"
            return mode if mode in {"light", "soft", "dark"} else "light"
        except AccountProfile.DoesNotExist:
            pass
    return "light"


def _ui_accent_color(request: HttpRequest) -> str:
    accent = _accent_color(request)
    if _theme_mode(request) == "dark" and accent.lower() in {"#111827", "#000000", "black"}:
        return "#7aa2ff"
    return accent


def _auth_context(request: HttpRequest) -> dict[str, str]:
    return {
        "accent_color": _accent_color(request),
        "ui_accent_color": _ui_accent_color(request),
        "theme_mode": _theme_mode(request),
        "next_url": _safe_next_url(request),
    }


def _safe_next_url(request: HttpRequest) -> str:
    value = (request.POST.get("next") or request.GET.get("next") or "").strip()
    if value and url_has_allowed_host_and_scheme(value, allowed_hosts={request.get_host()}, require_https=request.is_secure()):
        return value
    return ""


def _subscription_meter(request: HttpRequest) -> dict[str, object]:
    if not request.user.is_authenticated:
        return {"percent": 0, "days_left": 0}
    try:
        access = request.user.billing_access
    except Exception:
        return {"percent": 0, "days_left": 0}
    now = timezone.now()
    if access.active_until <= now:
        return {"percent": 0, "days_left": 0}
    total_seconds = max((access.active_until - access.created_at).total_seconds(), 1)
    left_seconds = max((access.active_until - now).total_seconds(), 0)
    percent = max(0, min(100, round((left_seconds / total_seconds) * 100)))
    days_left = max(1, int((left_seconds + 86399) // 86400))
    return {"percent": percent, "days_left": days_left}


def _subscription_panel(request: HttpRequest, account_stats: dict[str, object], language: str) -> dict[str, object]:
    now = timezone.now()
    access = None
    if request.user.is_authenticated:
        try:
            access = request.user.billing_access
        except Exception:
            access = None
    current_plan = get_plan(access.plan_code if access else "starter")
    active_until = access.active_until if access and access.active_until > now else None
    left_seconds = max((active_until - now).total_seconds(), 0) if active_until else 0
    days_left = int((left_seconds + 86399) // 86400) if left_seconds else 0
    period_seconds = max(current_plan.period_days * 86400, 1)
    used_seconds = max(0, period_seconds - left_seconds) if active_until else 0
    days_used = min(current_plan.period_days, int(used_seconds // 86400)) if active_until else 0
    remaining_value_cents = 0
    if active_until and current_plan.period_days:
        remaining_value_cents = round(current_plan.price_cents * min(1, left_seconds / period_seconds))
    plan_payload = []
    for plan in PLANS:
        localized = localized_plan(plan, language)
        due_cents = prorated_due_cents(current_plan if active_until else None, plan, left_seconds)
        credit_cents = 0 if plan.code == current_plan.code else min(plan.price_cents, remaining_value_cents)
        plan_payload.append(
            {
                **vars(localized),
                "checkout_url": f"{reverse('billing:checkout')}?{urlencode({'plan': plan.code, 'due': due_cents, 'next': request.get_full_path()})}",
                "due_cents": due_cents,
                "due_display": _money_display(due_cents),
                "credit_cents": credit_cents,
                "credit_display": _money_display(credit_cents),
                "is_current": plan.code == current_plan.code,
            }
        )
    return {
        "current_plan": localized_plan(current_plan, language),
        "plans": plan_payload,
        "active_until": active_until,
        "days_left": days_left,
        "days_used": days_used,
        "remaining_value_cents": remaining_value_cents,
        "remaining_value_display": _money_display(remaining_value_cents),
        "storage_used_text": account_stats.get("total_output_size_text", ""),
        "storage_available_text": account_stats.get("storage_available_text", ""),
        "storage_percent": account_stats.get("storage_percent", 0),
    }


def _money_display(cents: int) -> str:
    return f"{cents // 100}$" if cents % 100 == 0 else f"{cents / 100:.2f}$"


def _storage_quota(request: HttpRequest, account_stats: dict[str, object]) -> dict[str, object]:
    used = int(account_stats.get("total_output_size") or 0)
    owner_id, guest_key = _workspace_identity(request)
    project_bytes = sum(_video_owner_queryset(owner_id, guest_key).values_list("storage_bytes", flat=True))
    design_bytes = sum(_design_owner_queryset(owner_id, guest_key).values_list("storage_bytes", flat=True))
    used += int(project_bytes or 0)
    used += int(design_bytes or 0)
    plan_code = "starter"
    if request.user.is_authenticated:
        try:
            plan_code = request.user.billing_access.plan_code
        except Exception:
            plan_code = "starter"
    limit = get_plan(plan_code).storage_bytes
    percent = 0 if limit <= 0 else max(0, min(100, round((used / limit) * 100)))
    return {
        "storage_limit": limit,
        "storage_limit_text": human_size(limit),
        "storage_available": max(0, limit - used),
        "storage_available_text": human_size(max(0, limit - used)),
        "storage_percent": percent,
        "total_output_size_text": f"{human_size(used)} / {human_size(limit)}",
    }


def _language_options_for_form(form: AccountSettingsForm) -> list[dict[str, object]]:
    current = form["interface_language"].value() or "en"
    return [
        {
            "code": value,
            "native": label,
            "flag": {
                "en": "https://flagcdn.com/gb.svg",
                "ru": "https://flagcdn.com/ru.svg",
                "uk": "https://flagcdn.com/ua.svg",
                "fr": "https://flagcdn.com/fr.svg",
                "de": "https://flagcdn.com/de.svg",
                "es": "https://flagcdn.com/es.svg",
                "ka": "https://flagcdn.com/ge.svg",
                "hy": "https://flagcdn.com/am.svg",
                "it": "https://flagcdn.com/it.svg",
            }.get(value, "https://flagcdn.com/gb.svg"),
            "active": value == current,
        }
        for value, label in form.fields["interface_language"].choices
    ]


def _localized_image_modes(language: str) -> list[tuple[str, str]]:
    labels = {
        "light": {"en": "Light file", "ru": "Лёгкий файл", "uk": "Легкий файл", "fr": "Fichier léger", "de": "Leichte Datei", "es": "Archivo ligero", "ka": "მსუბუქი ფაილი", "hy": "Թեթև ֆայլ", "it": "File leggero"},
        "balanced": {"en": "Balanced", "ru": "Баланс", "uk": "Баланс", "fr": "Équilibre", "de": "Ausgewogen", "es": "Equilibrado", "ka": "ბალანსი", "hy": "Բալանս", "it": "Bilanciato"},
        "quality": {"en": "Quality", "ru": "Качество", "uk": "Якість", "fr": "Qualité", "de": "Qualität", "es": "Calidad", "ka": "ხარისხი", "hy": "Որակ", "it": "Qualità"},
    }
    return [(value, labels.get(value, {}).get(language, labels.get(value, {}).get("en", label))) for value, label in actions.IMAGE_MODE_CHOICES]


def _localized_youtube_modes(language: str) -> list[dict[str, str]]:
    labels = {
        "regular": {"en": "Shorts classic", "ru": "Shorts классика", "uk": "Shorts класика", "fr": "Shorts classique", "de": "Shorts klassisch", "es": "Shorts clásico", "ka": "Shorts კლასიკა", "hy": "Shorts դասական", "it": "Shorts classici"},
        "dynamic": {"en": "Shorts dynamic", "ru": "Shorts динамика", "uk": "Shorts динаміка", "fr": "Shorts dynamique", "de": "Shorts dynamisch", "es": "Shorts dinámico", "ka": "Shorts დინამიკა", "hy": "Shorts դինամիկ", "it": "Shorts dinamici"},
        "podcast": {"en": "Shorts podcast", "ru": "Shorts подкаст", "uk": "Shorts подкаст", "fr": "Shorts podcast", "de": "Shorts Podcast", "es": "Shorts podcast", "ka": "Shorts პოდკასტი", "hy": "Shorts փոդքաստ", "it": "Shorts podcast"},
        "calm": {"en": "Shorts calm", "ru": "Shorts спокойно", "uk": "Shorts спокійно", "fr": "Shorts calme", "de": "Shorts ruhig", "es": "Shorts calmado", "ka": "Shorts მშვიდი", "hy": "Shorts հանգիստ", "it": "Shorts morbidi"},
        "backstage30": {"en": "Preview 30s", "ru": "Preview 30 сек", "uk": "Preview 30 с", "fr": "Preview 30 s", "de": "Preview 30 s", "es": "Preview 30 s", "ka": "Preview 30 წმ", "hy": "Preview 30 վրկ", "it": "Preview 30 s"},
        "backstage60": {"en": "Preview 60s", "ru": "Preview 60 сек", "uk": "Preview 60 с", "fr": "Preview 60 s", "de": "Preview 60 s", "es": "Preview 60 s", "ka": "Preview 60 წმ", "hy": "Preview 60 վրկ", "it": "Preview 60 s"},
        "backstage90": {"en": "Preview 90s", "ru": "Preview 90 сек", "uk": "Preview 90 с", "fr": "Preview 90 s", "de": "Preview 90 s", "es": "Preview 90 s", "ka": "Preview 90 წმ", "hy": "Preview 90 վրկ", "it": "Preview 90 s"},
        "download": {"en": "Download MP4", "ru": "Скачать MP4", "uk": "Завантажити MP4", "fr": "Télécharger MP4", "de": "MP4 herunterladen", "es": "Descargar MP4", "ka": "MP4 ჩამოტვირთვა", "hy": "Ներբեռնել MP4", "it": "Scarica MP4"},
        "cover": {"en": "PNG cover", "ru": "PNG-обложка", "uk": "PNG-обкладинка", "fr": "Couverture PNG", "de": "PNG-Cover", "es": "Portada PNG", "ka": "PNG ყდა", "hy": "PNG շապիկ", "it": "Copertina PNG"},
    }
    hint_keys = {
        "regular": "youtube_mode_regular_hint",
        "dynamic": "youtube_mode_dynamic_hint",
        "podcast": "youtube_mode_podcast_hint",
        "calm": "youtube_mode_calm_hint",
        "backstage30": "youtube_mode_backstage_hint",
        "backstage60": "youtube_mode_backstage_hint",
        "backstage90": "youtube_mode_backstage_hint",
        "download": "youtube_mode_download_hint",
        "cover": "youtube_mode_cover_hint",
    }
    icons = {
        "regular": "classic",
        "dynamic": "dynamic",
        "podcast": "podcast",
        "calm": "calm",
        "backstage30": "backstage",
        "backstage60": "backstage",
        "backstage90": "backstage",
        "download": "download",
        "cover": "cover",
    }
    return [
        {
            "value": value,
            "label": labels.get(value, {}).get(language, labels.get(value, {}).get("en", label)),
            "hint": translate(hint_keys.get(value, "youtube_mode_regular_hint"), language),
            "icon": icons.get(value, "classic"),
        }
        for value, label in actions.YOUTUBE_MODE_CHOICES
    ]


def _localized_subtitle_styles(request: HttpRequest) -> list[dict[str, str]]:
    language = getattr(request, "interface_language", "en")
    current_path = request.get_full_path()
    items: list[dict[str, str]] = []
    for value, fallback in actions.SUBTITLE_STYLE_CHOICES:
        detail = _subtitle_style_detail(value, language)
        if not detail:
            detail = {
                "code": value,
                "label": fallback,
                "short": fallback,
                "description": fallback,
                "look": fallback,
                "best_for": fallback,
                "preview_text": fallback,
                "preview_class": value,
            }
        items.append(
            {
                **detail,
                "detail_url": f"{reverse('studio:subtitle_style_detail', args=[value])}?{urlencode({'next': current_path})}",
            }
        )
    return items


def _subtitle_style_detail(style: str, language: str) -> dict[str, str] | None:
    code = clean_base_name(style, "pop").lower()
    styles = _subtitle_style_catalog()
    data = styles.get(code)
    if not data:
        return None
    lang = clean_language(language)
    detail = {
        "code": code,
        "label": data["label"].get(lang) or data["label"]["en"],
        "short": data["short"].get(lang) or data["short"]["en"],
        "description": data["description"].get(lang) or data["description"]["en"],
        "look": data["look"].get(lang) or data["look"]["en"],
        "best_for": data["best_for"].get(lang) or data["best_for"]["en"],
        "preview_text": data["preview_text"].get(lang) or data["preview_text"]["en"],
        "preview_class": code,
    }
    if code == "candy":
        candy_text = {
            "short": _style_text(
                "Bright outline captions with candy-colored word accents.",
                "Яркие outline-субтитры с candy-акцентами на словах.",
                "Яскраві outline-субтитри з candy-акцентами на словах.",
                "Sous-titres outline avec accents candy colorés.",
                "Helle Outline-Untertitel mit Candy-Farbakzenten.",
                "Subtítulos outline con acentos candy de color.",
                "კაშკაშა outline სუბტიტრები candy ფერის აქცენტებით.",
                "Վառ outline ենթագրեր candy գունային շեշտերով։",
                "Sottotitoli outline con accenti candy colorati.",
            ),
            "description": _style_text(
                "Candy renders as large mobile-first captions: white base text, purple outline and rotating pink, yellow and green highlights on key words. It stays playful without covering the video with a colored block.",
                "Candy рендерится как крупные mobile-first субтитры: белая основа, фиолетовая обводка и розовые, жёлтые, зелёные акценты на ключевых словах. Стиль игривый, но не закрывает видео цветной плашкой.",
                "Candy рендериться як великі mobile-first субтитри: біла основа, фіолетова обводка та рожеві, жовті, зелені акценти на ключових словах. Стиль грайливий, але не закриває відео кольоровою плашкою.",
                "Candy affiche de grands sous-titres mobile-first : base blanche, contour violet et accents rose, jaune, vert sur les mots clés, sans bloc coloré massif.",
                "Candy rendert große Mobile-Untertitel: weiße Basis, violette Kontur und pinke, gelbe, grüne Akzente auf Schlüsselwörtern, ohne das Video mit einer Farbfläche zu verdecken.",
                "Candy se renderiza como subtítulos grandes para móvil: texto blanco, contorno violeta y acentos rosa, amarillo y verde en palabras clave, sin tapar el vídeo con un bloque de color.",
                "Candy არის დიდი mobile-first სუბტიტრები: თეთრი ტექსტი, იისფერი კონტური და ვარდისფერი, ყვითელი, მწვანე აქცენტები მთავარ სიტყვებზე.",
                "Candy-ն մեծ mobile-first ենթագրեր է՝ սպիտակ հիմք, մանուշակագույն եզրագիծ և վարդագույն, դեղին, կանաչ շեշտեր կարևոր բառերի վրա։",
                "Candy rende sottotitoli grandi mobile-first: testo bianco, contorno viola e accenti rosa, gialli e verdi sulle parole chiave, senza coprire il video.",
            ),
            "look": _style_text(
                "Uppercase, thick outline, soft pop-in motion, colored key words.",
                "Uppercase, плотная обводка, мягкий pop-in, цветные ключевые слова.",
                "Uppercase, щільна обводка, м'який pop-in, кольорові ключові слова.",
                "Uppercase, contour épais, pop-in doux, mots clés colorés.",
                "Uppercase, starke Kontur, weiches Pop-in, farbige Schlüsselwörter.",
                "Uppercase, contorno grueso, pop-in suave, palabras clave de color.",
                "Uppercase, მკვრივი კონტური, რბილი pop-in, ფერადი მთავარი სიტყვები.",
                "Uppercase, հաստ եզրագիծ, մեղմ pop-in, գունավոր կարևոր բառեր։",
                "Uppercase, contorno spesso, pop-in morbido, parole chiave colorate.",
            ),
        }
        for key, values in candy_text.items():
            detail[key] = values.get(lang) or values["en"]
    return detail


def _style_text(en: str, ru: str, uk: str, fr: str, de: str, es: str, ka: str, hy: str, it: str) -> dict[str, str]:
    return {"en": en, "ru": ru, "uk": uk, "fr": fr, "de": de, "es": es, "ka": ka, "hy": hy, "it": it}


def _subtitle_style_catalog() -> dict[str, dict[str, dict[str, str]]]:
    shared_best = _style_text(
        "Short clips, reels, lessons and creator videos.",
        "Шорты, Reels, уроки и авторские видео.",
        "Shorts, Reels, уроки та авторські відео.",
        "Shorts, reels, tutoriels et vidéos de créateurs.",
        "Shorts, Reels, Tutorials und Creator-Videos.",
        "Shorts, reels, tutoriales y videos de creadores.",
        "Shorts, Reels, გაკვეთილები და creator ვიდეოები.",
        "Shorts, Reels, դասեր և ստեղծագործական տեսանյութեր։",
        "Short, reels, tutorial e video creator.",
    )
    return {
        "pop": {
            "label": _style_text("Pop", "Pop", "Pop", "Pop", "Pop", "Pop", "Pop", "Pop", "Pop"),
            "short": _style_text("Bright, punchy and easy to read.", "Яркий, ударный и легко читается.", "Яскравий, виразний і легко читається.", "Vif, percutant et lisible.", "Hell, kräftig und gut lesbar.", "Brillante, potente y legible.", "კაშკაშა, ენერგიული და ადვილად წასაკითხი.", "Վառ, ազդեցիկ և հեշտ ընթեռնելի։", "Vivace, incisivo e leggibile."),
            "description": _style_text("Large subtitles with a strong outline and energetic color accents. They feel modern and work well on mobile screens.", "Крупные субтитры с плотной обводкой и энергичными цветными акцентами. Смотрятся современно и хорошо читаются на телефоне.", "Великі субтитри з щільною обводкою та енергійними кольоровими акцентами. Сучасні й добре читаються на телефоні.", "Sous-titres larges avec contour marqué et accents colorés. Modernes et très lisibles sur mobile.", "Große Untertitel mit starker Kontur und farbigen Akzenten. Modern und gut auf Mobilgeräten lesbar.", "Subtítulos grandes con contorno fuerte y acentos de color. Modernos y claros en móvil.", "დიდი სუბტიტრები მკვეთრი კონტურით და ფერადი აქცენტებით. კარგად იკითხება მობილურზე.", "Մեծ ենթագրեր հստակ եզրագծով և գունային շեշտերով։ Լավ ընթեռնելի է հեռախոսում։", "Sottotitoli grandi con contorno deciso e accenti colorati. Moderni e leggibili su mobile."),
            "look": _style_text("Bold white text, dark outline, warm highlights.", "Белый жирный текст, тёмная обводка, тёплые акценты.", "Білий жирний текст, темна обводка, теплі акценти.", "Texte blanc épais, contour sombre, accents chauds.", "Fetter weißer Text, dunkle Kontur, warme Akzente.", "Texto blanco grueso, contorno oscuro, acentos cálidos.", "მსხვილი თეთრი ტექსტი, მუქი კონტური, თბილი აქცენტები.", "Հաստ սպիտակ տեքստ, մուգ եզրագիծ, տաք շեշտեր։", "Testo bianco bold, contorno scuro, accenti caldi."),
            "best_for": shared_best,
            "preview_text": _style_text("Big moment!", "Главный момент!", "Головний момент!", "Moment fort !", "Starker Moment!", "¡Momento clave!", "მთავარი მომენტი!", "Գլխավոր պահը։", "Momento chiave!"),
        },
        "neon": {
            "label": _style_text("Neon", "Neon", "Neon", "Néon", "Neon", "Neón", "Neon", "Neon", "Neon"),
            "short": _style_text("Glowing, nightlife, high energy.", "Свечение, ночной вайб, высокая энергия.", "Світіння, нічний вайб, висока енергія.", "Lumineux, nocturne, énergique.", "Leuchtend, Nachtgefühl, energiegeladen.", "Brillante, nocturno y enérgico.", "ნათება, ღამის განწყობა, მაღალი ენერგია.", "Փայլող, գիշերային, էներգիկ։", "Luminoso, notturno, energico."),
            "description": _style_text("Neon adds glow and punch. It is best when the video has music, nightlife, gaming or a tech mood.", "Neon добавляет свечение и драйв. Лучше всего подходит для музыки, ночных сцен, игр и технологичного вайба.", "Neon додає світіння й драйв. Найкраще для музики, нічних сцен, ігор та tech-настрою.", "Néon ajoute de la lumière et du rythme. Idéal pour musique, nuit, gaming et tech.", "Neon bringt Glow und Tempo. Ideal für Musik, Nacht, Gaming und Tech.", "Neón añade brillo y energía. Ideal para música, noche, gaming y tecnología.", "Neon მატებს ნათებას და ენერგიას. კარგია მუსიკის, ღამის, gaming და tech ვიდეოებისთვის.", "Neon-ը ավելացնում է փայլ և դինամիկա։ Լավ է երաժշտության, գիշերային, gaming և tech տեսանյութերի համար։", "Neon aggiunge luce ed energia. Ideale per musica, notte, gaming e tech."),
            "look": _style_text("Bright glow with saturated color edges.", "Яркое свечение с насыщенными цветными краями.", "Яскраве світіння з насиченими кольоровими краями.", "Lueur vive avec bords colorés saturés.", "Heller Glow mit kräftigen Farbkanten.", "Brillo intenso con bordes saturados.", "კაშკაშა ნათება ფერადი კიდეებით.", "Վառ փայլ հագեցած գունային եզրերով։", "Bagliore intenso con bordi saturi."),
            "best_for": shared_best,
            "preview_text": _style_text("Feel the glow", "Светится!", "Світиться!", "Ça brille", "Es leuchtet", "Brilla", "ანათებს", "Փայլում է", "Si illumina"),
        },
        "candy": {
            "label": _style_text("Candy", "Candy", "Candy", "Candy", "Candy", "Candy", "Candy", "Candy", "Candy"),
            "short": _style_text("Sweet, colorful and playful.", "Сочный, цветной и игривый.", "Соковитий, кольоровий та грайливий.", "Coloré, doux et ludique.", "Süß, bunt und verspielt.", "Dulce, colorido y juguetón.", "ტკბილი, ფერადი და მხიარული.", "Քաղցր, գունավոր և խաղային։", "Dolce, colorato e giocoso."),
            "description": _style_text("Candy uses soft bright colors and a friendly look. It is great for lifestyle, beauty, food and upbeat clips.", "Candy даёт мягкие яркие цвета и дружелюбный вид. Хорош для lifestyle, beauty, еды и позитивных роликов.", "Candy дає м'які яскраві кольори та дружній вигляд. Добре для lifestyle, beauty, їжі й позитивних роликів.", "Candy utilise des couleurs douces et vives. Parfait pour lifestyle, beauté, food et clips positifs.", "Candy nutzt weiche helle Farben. Gut für Lifestyle, Beauty, Food und positive Clips.", "Candy usa colores suaves y vivos. Perfecto para lifestyle, belleza, comida y clips positivos.", "Candy იყენებს რბილ ნათელ ფერებს. კარგია lifestyle, beauty, food და პოზიტიური ვიდეოებისთვის.", "Candy-ն օգտագործում է մեղմ վառ գույներ։ Հարմար է lifestyle, beauty, food և դրական տեսանյութերի համար։", "Candy usa colori morbidi e vivaci. Ottimo per lifestyle, beauty, food e clip positive."),
            "look": _style_text("Rounded, bright accents, friendly contrast.", "Округло, яркие акценты, мягкий контраст.", "Округло, яскраві акценти, м'який контраст.", "Arrondi, accents vifs, contraste doux.", "Rund, helle Akzente, weicher Kontrast.", "Redondeado, acentos vivos, contraste suave.", "მომრგვალებული, ნათელი აქცენტები, რბილი კონტრასტი.", "Կլորացված, վառ շեշտեր, մեղմ կոնտրաստ։", "Arrotondato, accenti vivaci, contrasto morbido."),
            "best_for": shared_best,
            "preview_text": _style_text("Sweet cut", "Сочно!", "Соковито!", "Très doux", "Sweet cut", "Muy dulce", "ტკბილი კადრი", "Քաղցր կադր", "Dolce taglio"),
        },
        "kinetic": {
            "label": _style_text("Kinetic", "Kinetic", "Kinetic", "Cinétique", "Kinetic", "Cinético", "Kinetic", "Kinetic", "Kinetic"),
            "short": _style_text("Moving captions with impact.", "Движущиеся субтитры с акцентом.", "Рухомі субтитри з акцентом.", "Sous-titres animés avec impact.", "Bewegte Untertitel mit Wirkung.", "Subtítulos en movimiento con impacto.", "მოძრავი სუბტიტრები ძლიერი აქცენტით.", "Շարժվող ենթագրեր ուժեղ շեշտով։", "Sottotitoli in movimento con impatto."),
            "description": _style_text("Kinetic adds subtle movement and scale. It makes speech feel dynamic without needing manual animation.", "Kinetic добавляет движение и масштаб. Речь становится динамичной без ручной анимации.", "Kinetic додає рух і масштаб. Мова виглядає динамічно без ручної анімації.", "Cinétique ajoute mouvement et échelle. La parole paraît dynamique sans animation manuelle.", "Kinetic fügt Bewegung und Skalierung hinzu. Sprache wirkt dynamisch ohne manuelle Animation.", "Cinético añade movimiento y escala. La voz se siente dinámica sin animación manual.", "Kinetic ამატებს მოძრაობას და მასშტაბს. მეტყველება დინამიკური ჩანს ხელით ანიმაციის გარეშე.", "Kinetic-ը ավելացնում է շարժում և մասշտաբ։ Խոսքը դինամիկ է թվում առանց ձեռքով անիմացիայի։", "Kinetic aggiunge movimento e scala. Il parlato sembra dinamico senza animazione manuale."),
            "look": _style_text("Small motion, bold text, punchy entrance.", "Лёгкое движение, жирный текст, резкий вход.", "Легкий рух, жирний текст, різкий вхід.", "Léger mouvement, texte fort, entrée vive.", "Leichte Bewegung, fetter Text, starker Einstieg.", "Movimiento leve, texto fuerte, entrada marcada.", "მსუბუქი მოძრაობა, მსხვილი ტექსტი, მკვეთრი შესვლა.", "Թեթև շարժում, հաստ տեքստ, հստակ մուտք։", "Movimento leggero, testo bold, ingresso deciso."),
            "best_for": shared_best,
            "preview_text": _style_text("Move fast", "Двигаемся", "Рухаємось", "Ça bouge", "Bewegung", "Muévete", "მოძრაობა", "Շարժում", "In movimento"),
        },
        "bounce": {
            "label": _style_text("Bounce", "Bounce", "Bounce", "Bounce", "Bounce", "Bounce", "Bounce", "Bounce", "Bounce"),
            "short": _style_text("Bouncy, friendly motion.", "Пружинит и выглядит дружелюбно.", "Пружинить і виглядає дружньо.", "Rebondissant et convivial.", "Federnd und freundlich.", "Rebota y se ve amable.", "ხტუნავს და მეგობრულად ჩანს.", "Ցատկոտ և ընկերական տեսք։", "Rimbalzante e amichevole."),
            "description": _style_text("Bounce makes each phrase pop in with a soft jump. It is playful but still readable.", "Bounce выводит фразы мягким прыжком. Игриво, но всё ещё читаемо.", "Bounce виводить фрази м'яким стрибком. Грайливо, але читабельно.", "Bounce fait apparaître les phrases avec un petit saut. Ludique mais lisible.", "Bounce lässt Sätze weich einspringen. Verspielt, aber lesbar.", "Bounce muestra frases con un salto suave. Juguetón pero legible.", "Bounce ფრაზებს რბილი ნახტომით აჩენს. მხიარულია, მაგრამ იკითხება.", "Bounce-ը արտահայտությունները մեղմ ցատկով է ցույց տալիս։ Խաղային է, բայց ընթեռնելի։", "Bounce fa entrare le frasi con un piccolo salto. Giocoso ma leggibile."),
            "look": _style_text("Soft jump, bright outline, upbeat rhythm.", "Мягкий прыжок, яркая обводка, бодрый ритм.", "М'який стрибок, яскрава обводка, бадьорий ритм.", "Petit saut, contour vif, rythme positif.", "Weicher Sprung, helle Kontur, flotter Rhythmus.", "Salto suave, contorno vivo, ritmo alegre.", "რბილი ნახტომი, ნათელი კონტური, ენერგიული რიტმი.", "Մեղմ ցատկ, վառ եզրագիծ, ուրախ ռիթմ։", "Salto morbido, contorno vivo, ritmo allegro."),
            "best_for": shared_best,
            "preview_text": _style_text("Pop in", "Прыг!", "Стриб!", "Hop !", "Plopp", "Salta", "ხტომა", "Ցատկ", "Salta"),
        },
        "comic": {
            "label": _style_text("Comic", "Comic", "Comic", "Comic", "Comic", "Comic", "Comic", "Comic", "Comic"),
            "short": _style_text("Bold meme/comic energy.", "Жирная мемная энергия.", "Жирна мемна енергія.", "Énergie meme et BD.", "Meme- und Comic-Energie.", "Energía meme/cómic.", "მემისა და კომიქსის ენერგია.", "Մեմ/կոմիքս էներգիա։", "Energia meme/fumetto."),
            "description": _style_text("Comic feels loud, fun and expressive. Use it when the video should feel humorous or dramatic.", "Comic громкий, весёлый и выразительный. Подходит для юмора, реакций и драматичных моментов.", "Comic гучний, веселий і виразний. Підходить для гумору, реакцій і драматичних моментів.", "Comic est fort, drôle et expressif. Idéal pour humour, réactions et moments dramatiques.", "Comic ist laut, lustig und ausdrucksstark. Gut für Humor, Reactions und Drama.", "Comic es fuerte, divertido y expresivo. Ideal para humor, reacciones y drama.", "Comic ხმამაღალი, მხიარული და გამომხატველია. კარგია იუმორისა და რეაქციებისთვის.", "Comic-ը բարձր, զվարճալի և արտահայտիչ է։ Հարմար է հումորի և ռեակցիաների համար։", "Comic è forte, divertente ed espressivo. Ideale per umorismo, reaction e drama."),
            "look": _style_text("Chunky letters, strong outline, playful pop.", "Крупные буквы, мощная обводка, playful-pop.", "Великі літери, сильна обводка, playful-pop.", "Lettres épaisses, contour fort, pop ludique.", "Massive Buchstaben, starke Kontur, spielerischer Pop.", "Letras gruesas, contorno fuerte, pop divertido.", "მსხვილი ასოები, ძლიერი კონტური, მხიარული pop.", "Խոշոր տառեր, ուժեղ եզրագիծ, խաղային pop։", "Lettere grosse, contorno forte, pop giocoso."),
            "best_for": shared_best,
            "preview_text": _style_text("No way!", "Да ладно!", "Та ну!", "Incroyable !", "Nicht wahr!", "¡No puede ser!", "არ მჯერა!", "Չի կարող պատահել։", "Ma dai!"),
        },
        "clean": {
            "label": _style_text("Clean", "Clean", "Clean", "Clean", "Clean", "Clean", "Clean", "Clean", "Clean"),
            "short": _style_text("Simple and professional.", "Просто и профессионально.", "Просто й професійно.", "Simple et professionnel.", "Einfach und professionell.", "Simple y profesional.", "მარტივი და პროფესიული.", "Պարզ և պրոֆեսիոնալ։", "Semplice e professionale."),
            "description": _style_text("Clean keeps attention on the message. It is readable, neutral and safe for business or education.", "Clean удерживает внимание на смысле. Нейтральный, читаемый и безопасный для бизнеса или обучения.", "Clean тримає увагу на змісті. Нейтральний, читабельний і безпечний для бізнесу чи навчання.", "Clean garde l'attention sur le message. Neutre, lisible, sûr pour business et éducation.", "Clean hält den Fokus auf der Botschaft. Neutral, lesbar, gut für Business und Bildung.", "Clean centra la atención en el mensaje. Neutral, legible, seguro para negocio o educación.", "Clean ყურადღებას აზრზე ტოვებს. ნეიტრალური და კარგი ბიზნესისთვის ან სწავლისთვის.", "Clean-ը ուշադրությունը պահում է մտքի վրա։ Չեզոք և հարմար է բիզնեսի կամ ուսուցման համար։", "Clean mantiene il focus sul messaggio. Neutro, leggibile, adatto a business e formazione."),
            "look": _style_text("White text, subtle outline, calm layout.", "Белый текст, лёгкая обводка, спокойная раскладка.", "Білий текст, легка обводка, спокійна композиція.", "Texte blanc, contour léger, mise en page calme.", "Weißer Text, dezente Kontur, ruhiges Layout.", "Texto blanco, contorno sutil, diseño tranquilo.", "თეთრი ტექსტი, მსუბუქი კონტური, მშვიდი განლაგება.", "Սպիտակ տեքստ, մեղմ եզրագիծ, հանգիստ դասավորություն։", "Testo bianco, contorno leggero, layout calmo."),
            "best_for": shared_best,
            "preview_text": _style_text("Clear message", "Ясная мысль", "Чітка думка", "Message clair", "Klare Aussage", "Mensaje claro", "გასაგები აზრი", "Հստակ միտք", "Messaggio chiaro"),
        },
        "minimal": {
            "label": _style_text("Minimal", "Minimal", "Minimal", "Minimal", "Minimal", "Minimal", "Minimal", "Minimal", "Minimal"),
            "short": _style_text("Quiet, thin and elegant.", "Тихий, тонкий и элегантный.", "Тихий, тонкий та елегантний.", "Discret, fin et élégant.", "Ruhig, fein und elegant.", "Discreto, fino y elegante.", "მშვიდი, თხელი და ელეგანტური.", "Հանգիստ, նուրբ և էլեգանտ։", "Sobrio, sottile ed elegante."),
            "description": _style_text("Minimal is restrained and refined. Use it when subtitles should not overpower the visuals.", "Minimal сдержанный и аккуратный. Используй, когда субтитры не должны перебивать картинку.", "Minimal стриманий і акуратний. Використовуй, коли субтитри не мають перебивати картинку.", "Minimal est sobre et raffiné. À utiliser quand les sous-titres ne doivent pas dominer l'image.", "Minimal ist zurückhaltend und fein. Wenn Untertitel das Bild nicht überlagern sollen.", "Minimal es sobrio y refinado. Úsalo cuando los subtítulos no deben dominar la imagen.", "Minimal თავშეკავებული და დახვეწილია. როცა სუბტიტრებმა კადრი არ უნდა გადაფაროს.", "Minimal-ը զուսպ և նուրբ է։ Երբ ենթագրերը չպետք է գերիշխեն կադրին։", "Minimal è sobrio e raffinato. Quando i sottotitoli non devono dominare l'immagine."),
            "look": _style_text("Small clean text with light shadow.", "Небольшой чистый текст с лёгкой тенью.", "Невеликий чистий текст з легкою тінню.", "Petit texte propre avec ombre légère.", "Kleiner sauberer Text mit leichtem Schatten.", "Texto pequeño y limpio con sombra ligera.", "პატარა სუფთა ტექსტი მსუბუქი ჩრდილით.", "Փոքր մաքուր տեքստ թեթև ստվերով։", "Testo piccolo e pulito con ombra leggera."),
            "best_for": shared_best,
            "preview_text": _style_text("Less is more", "Меньше — лучше", "Менше — краще", "Moins, mieux", "Weniger ist mehr", "Menos es más", "ნაკლები უკეთესია", "Քիչն ավելի լավ է", "Meno è meglio"),
        },
        "editorial": {
            "label": _style_text("Editorial", "Editorial", "Editorial", "Éditorial", "Editorial", "Editorial", "Editorial", "Editorial", "Editorial"),
            "short": _style_text("Magazine-like and polished.", "Как журнал: аккуратно и дорого.", "Як журнал: акуратно й дорого.", "Style magazine, soigné.", "Magazinartig und hochwertig.", "Tipo revista y pulido.", "ჟურნალის სტილი, დახვეწილი.", "Ամսագրային և հղկված։", "Stile magazine, curato."),
            "description": _style_text("Editorial feels curated and premium. It works well for interviews, brand stories and polished explainers.", "Editorial выглядит собранно и премиально. Хорош для интервью, брендов и аккуратных объясняющих видео.", "Editorial виглядає зібрано й преміально. Добре для інтерв'ю, брендів і пояснювальних відео.", "Éditorial paraît premium et maîtrisé. Bien pour interviews, marques et vidéos explicatives.", "Editorial wirkt hochwertig und kuratiert. Gut für Interviews, Marken und Erklärvideos.", "Editorial se siente premium y cuidado. Bueno para entrevistas, marcas y explicativos.", "Editorial პრემიუმ და მოწესრიგებულია. კარგია ინტერვიუებისა და ბრენდის ისტორიებისთვის.", "Editorial-ը պրեմիում և մշակված տեսք ունի։ Հարմար է հարցազրույցների և բրենդի պատմությունների համար։", "Editorial è premium e curato. Ottimo per interviste, brand story e spiegazioni."),
            "look": _style_text("Balanced typography, refined spacing.", "Сбалансированная типографика, аккуратные интервалы.", "Збалансована типографіка, акуратні інтервали.", "Typographie équilibrée, espacements soignés.", "Ausgewogene Typografie, feine Abstände.", "Tipografía equilibrada, espaciado cuidado.", "დაბალანსებული ტიპოგრაფიკა და დახვეწილი ინტერვალები.", "Հավասարակշռված տպագրություն և նուրբ տարածություններ։", "Tipografia bilanciata, spaziatura curata."),
            "best_for": shared_best,
            "preview_text": _style_text("In focus", "В фокусе", "У фокусі", "En focus", "Im Fokus", "En foco", "ფოკუსში", "Ֆոկուսում", "In focus"),
        },
        "typewriter": {
            "label": _style_text("Typewriter", "Typewriter", "Typewriter", "Machine", "Typewriter", "Máquina", "Typewriter", "Typewriter", "Typewriter"),
            "short": _style_text("Documentary and typed feel.", "Документальный печатный вайб.", "Документальний друкований вайб.", "Style tapé et documentaire.", "Getippt, dokumentarisch.", "Estilo escrito/documental.", "დოკუმენტური, საბეჭდი სტილი.", "Փաստագրական, տպագրական զգացողություն։", "Effetto scritto/documentario."),
            "description": _style_text("Typewriter is calm, focused and slightly retro. It suits diaries, documentary edits and thoughtful narration.", "Typewriter спокойный, фокусный и немного ретро. Подходит для дневников, документалок и вдумчивого рассказа.", "Typewriter спокійний, фокусний і трохи ретро. Для щоденників, документальних монтажів і вдумливої оповіді.", "Machine est calme, focalisé et un peu rétro. Pour journaux, documentaires et narration posée.", "Typewriter ist ruhig, fokussiert und leicht retro. Für Tagebücher, Dokus und nachdenkliche Erzählung.", "Máquina es calmado, enfocado y retro. Para diarios, documentales y narración reflexiva.", "Typewriter მშვიდი, ფოკუსირებული და ოდნავ რეტროა. კარგია დღიურებისა და დოკუმენტური ნარატივისთვის.", "Typewriter-ը հանգիստ, կենտրոնացած և մի փոքր ռետրո է։ Հարմար է օրագրերի և փաստագրական պատմության համար։", "Typewriter è calmo, focalizzato e un po' retro. Per diari, documentari e narrazione riflessiva."),
            "look": _style_text("Monospace rhythm and calm contrast.", "Моноширинный ритм и спокойный контраст.", "Моноширинний ритм і спокійний контраст.", "Rythme monospace et contraste calme.", "Monospace-Rhythmus und ruhiger Kontrast.", "Ritmo monoespaciado y contraste calmado.", "მონოსივრცული რიტმი და მშვიდი კონტრასტი.", "Մոնոսփեյս ռիթմ և հանգիստ կոնտրաստ։", "Ritmo monospazio e contrasto calmo."),
            "best_for": shared_best,
            "preview_text": _style_text("Typed note", "Печатная заметка", "Друкована нотатка", "Note tapée", "Getippte Notiz", "Nota escrita", "დაბეჭდილი ნოტა", "Տպված նշում", "Nota scritta"),
        },
        "headline": {
            "label": _style_text("Headline", "Headline", "Headline", "Titre", "Headline", "Titular", "Headline", "Headline", "Headline"),
            "short": _style_text("Big statement captions.", "Большие заголовочные фразы.", "Великі заголовкові фрази.", "Grandes phrases titre.", "Große Headline-Sätze.", "Frases grandes tipo titular.", "დიდი სათაურის ფრაზები.", "Մեծ վերնագրային արտահայտություններ։", "Frasi grandi da titolo."),
            "description": _style_text("Headline makes speech feel like a strong title. Use for hooks, bold claims and key moments.", "Headline превращает речь в сильный заголовок. Для хуков, заявлений и ключевых моментов.", "Headline перетворює мову на сильний заголовок. Для хуків, заяв і ключових моментів.", "Titre transforme la parole en titre fort. Pour hooks, annonces et moments clés.", "Headline macht Sprache zur starken Überschrift. Für Hooks, Claims und Kernmomente.", "Titular convierte la voz en un gran titular. Para hooks, frases fuertes y momentos clave.", "Headline მეტყველებას ძლიერ სათაურად აქცევს. კარგია hooks და მთავარი მომენტებისთვის.", "Headline-ը խոսքը դարձնում է ուժեղ վերնագիր։ Հարմար է hooks և գլխավոր պահերի համար։", "Headline trasforma il parlato in un titolo forte. Per hook, claim e momenti chiave."),
            "look": _style_text("Large uppercase feel, very strong contrast.", "Крупное заголовочное ощущение, очень сильный контраст.", "Велике заголовкове відчуття, дуже сильний контраст.", "Grand format titre, contraste très fort.", "Großes Headline-Gefühl, sehr starker Kontrast.", "Sensación de titular grande, contraste fuerte.", "დიდი სათაურის შეგრძნება, ძლიერი კონტრასტი.", "Մեծ վերնագրի զգացողություն, ուժեղ կոնտրաստ։", "Effetto titolo grande, contrasto forte."),
            "best_for": shared_best,
            "preview_text": _style_text("THE HOOK", "ГЛАВНЫЙ ХУК", "ГОЛОВНИЙ ХУК", "LE HOOK", "DER HOOK", "EL HOOK", "HOOK", "HOOK", "IL HOOK"),
        },
        "luxury": {
            "label": _style_text("Luxury", "Luxury", "Luxury", "Luxury", "Luxury", "Luxury", "Luxury", "Luxury", "Luxury"),
            "short": _style_text("Premium, elegant, soft.", "Премиально, элегантно, мягко.", "Преміально, елегантно, м'яко.", "Premium, élégant, doux.", "Premium, elegant, weich.", "Premium, elegante, suave.", "პრემიუმი, ელეგანტური, რბილი.", "Պրեմիում, էլեգանտ, մեղմ։", "Premium, elegante, morbido."),
            "description": _style_text("Luxury is polished and calm with a premium mood. Use it for fashion, real estate, products and refined stories.", "Luxury спокойный и дорогой по ощущению. Для fashion, недвижимости, продуктов и утончённых историй.", "Luxury спокійний і дорогий за відчуттям. Для fashion, нерухомості, продуктів і витончених історій.", "Luxury est calme et premium. Pour mode, immobilier, produits et récits raffinés.", "Luxury wirkt ruhig und hochwertig. Für Fashion, Immobilien, Produkte und feine Stories.", "Luxury se siente premium y calmado. Para moda, inmuebles, productos e historias refinadas.", "Luxury მშვიდი და პრემიუმ განწყობისაა. მოდის, უძრავი ქონებისა და პროდუქტებისთვის.", "Luxury-ը հանգիստ և պրեմիում զգացողություն ունի։ Հարմար է fashion, անշարժ գույք և ապրանքների համար։", "Luxury è calmo e premium. Per moda, immobiliare, prodotti e storie raffinate."),
            "look": _style_text("Elegant type, soft shadow, restrained color.", "Элегантный шрифт, мягкая тень, сдержанный цвет.", "Елегантний шрифт, м'яка тінь, стриманий колір.", "Typo élégante, ombre douce, couleur sobre.", "Elegante Schrift, weicher Schatten, dezente Farbe.", "Tipo elegante, sombra suave, color sobrio.", "ელეგანტური შრიფტი, რბილი ჩრდილი, თავშეკავებული ფერი.", "Էլեգանտ տառատեսակ, մեղմ ստվեր, զուսպ գույն։", "Font elegante, ombra morbida, colore sobrio."),
            "best_for": shared_best,
            "preview_text": _style_text("Premium detail", "Премиальная деталь", "Преміальна деталь", "Détail premium", "Premium Detail", "Detalle premium", "პრემიუმ დეტალი", "Պրեմիում դետալ", "Dettaglio premium"),
        },
        "mono": {
            "label": _style_text("Mono", "Mono", "Mono", "Mono", "Mono", "Mono", "Mono", "Mono", "Mono"),
            "short": _style_text("Tech, code and structured.", "Техно, код и структура.", "Техно, код і структура.", "Tech, code et structuré.", "Tech, Code und Struktur.", "Tech, código y estructura.", "ტექნო, კოდი და სტრუქტურა.", "Տեխնո, կոդ և կառուցվածք։", "Tech, codice e struttura."),
            "description": _style_text("Mono uses a technical monospace feel. It is great for software, finance, analytics and explainers.", "Mono даёт технический моноширинный характер. Хорош для софта, финансов, аналитики и объяснений.", "Mono дає технічний моноширинний характер. Добре для софту, фінансів, аналітики й пояснень.", "Mono donne une sensation technique monospace. Idéal pour logiciel, finance, analyse et explications.", "Mono nutzt technische Monospace-Optik. Gut für Software, Finanzen, Analyse und Erklärvideos.", "Mono usa estilo monoespaciado técnico. Bueno para software, finanzas, análisis y explicativos.", "Mono ტექნიკურ monospace სტილს იყენებს. კარგია software, finance, analytics და ახსნებისთვის.", "Mono-ն տեխնիկական monospace տեսք ունի։ Հարմար է software, finance, analytics և բացատրությունների համար։", "Mono usa un look monospazio tecnico. Ottimo per software, finanza, analisi e spiegazioni."),
            "look": _style_text("Monospace, green accents, precise rhythm.", "Моноширинный, зелёные акценты, точный ритм.", "Моноширинний, зелені акценти, точний ритм.", "Monospace, accents verts, rythme précis.", "Monospace, grüne Akzente, präziser Rhythmus.", "Monoespaciado, acentos verdes, ritmo preciso.", "monospace, მწვანე აქცენტები, ზუსტი რიტმი.", "Monospace, կանաչ շեշտեր, ճշգրիտ ռիթմ։", "Monospazio, accenti verdi, ritmo preciso."),
            "best_for": shared_best,
            "preview_text": _style_text("System ready", "Система готова", "Система готова", "Système prêt", "System bereit", "Sistema listo", "სისტემა მზადაა", "Համակարգը պատրաստ է", "Sistema pronto"),
        },
        "soft": {
            "label": _style_text("Soft", "Soft", "Soft", "Soft", "Soft", "Soft", "Soft", "Soft", "Soft"),
            "short": _style_text("Gentle, warm and calm.", "Нежно, тепло и спокойно.", "Ніжно, тепло й спокійно.", "Doux, chaleureux et calme.", "Sanft, warm und ruhig.", "Suave, cálido y tranquilo.", "ნაზი, თბილი და მშვიდი.", "Նուրբ, տաք և հանգիստ։", "Delicato, caldo e calmo."),
            "description": _style_text("Soft is gentle and human. It is good for calm narration, family, wellness and emotional stories.", "Soft мягкий и человечный. Хорош для спокойного рассказа, семьи, wellness и эмоциональных историй.", "Soft м'який і людяний. Добре для спокійної оповіді, сім'ї, wellness та емоційних історій.", "Soft est doux et humain. Pour narration calme, famille, bien-être et histoires émotionnelles.", "Soft ist sanft und menschlich. Für ruhige Erzählung, Familie, Wellness und Emotion.", "Soft es suave y humano. Para narración tranquila, familia, bienestar e historias emocionales.", "Soft ნაზი და ადამიანურია. კარგია მშვიდი თხრობის, ოჯახის, wellness და ემოციური ისტორიებისთვის.", "Soft-ը նուրբ և մարդկային է։ Հարմար է հանգիստ պատմության, ընտանիքի, wellness և զգացմունքային պատմությունների համար։", "Soft è delicato e umano. Per narrazione calma, famiglia, wellness e storie emotive."),
            "look": _style_text("Soft shadow, warm contrast, calm placement.", "Мягкая тень, тёплый контраст, спокойное положение.", "М'яка тінь, теплий контраст, спокійне розташування.", "Ombre douce, contraste chaud, placement calme.", "Weicher Schatten, warmer Kontrast, ruhige Position.", "Sombra suave, contraste cálido, posición tranquila.", "რბილი ჩრდილი, თბილი კონტრასტი, მშვიდი განლაგება.", "Մեղմ ստվեր, տաք կոնտրաստ, հանգիստ տեղադրում։", "Ombra morbida, contrasto caldo, posizione calma."),
            "best_for": shared_best,
            "preview_text": _style_text("A softer note", "Мягкая мысль", "М'яка думка", "Note douce", "Sanfte Note", "Nota suave", "რბილი აზრი", "Մեղմ միտք", "Nota morbida"),
        },
    }


def _localized_subtitle_languages(language: str) -> list[dict[str, str]]:
    labels = {
        "auto": {"en": "Auto", "ru": "Авто", "uk": "Авто", "fr": "Auto", "de": "Auto", "es": "Auto", "ka": "ავტო", "hy": "Ավտո", "it": "Auto"},
    }
    subtitle_language_meta = {
        "en": {"native": "English", "flag": "https://flagcdn.com/gb.svg"},
        "ru": {"native": "Русский", "flag": "https://flagcdn.com/ru.svg"},
        "uk": {"native": "Українська", "flag": "https://flagcdn.com/ua.svg"},
        "fr": {"native": "Français", "flag": "https://flagcdn.com/fr.svg"},
        "de": {"native": "Deutsch", "flag": "https://flagcdn.com/de.svg"},
        "es": {"native": "Español", "flag": "https://flagcdn.com/es.svg"},
        "it": {"native": "Italiano", "flag": "https://flagcdn.com/it.svg"},
        "pt": {"native": "Português", "flag": "https://flagcdn.com/pt.svg"},
        "pl": {"native": "Polski", "flag": "https://flagcdn.com/pl.svg"},
        "tr": {"native": "Türkçe", "flag": "https://flagcdn.com/tr.svg"},
        "nl": {"native": "Nederlands", "flag": "https://flagcdn.com/nl.svg"},
        "sv": {"native": "Svenska", "flag": "https://flagcdn.com/se.svg"},
        "ar": {"native": "العربية", "flag": "https://flagcdn.com/sa.svg"},
        "hi": {"native": "हिन्दी", "flag": "https://flagcdn.com/in.svg"},
        "ja": {"native": "日本語", "flag": "https://flagcdn.com/jp.svg"},
        "ko": {"native": "한국어", "flag": "https://flagcdn.com/kr.svg"},
        "zh": {"native": "中文", "flag": "https://flagcdn.com/cn.svg"},
        "ka": {"native": "ქართული", "flag": "https://flagcdn.com/ge.svg"},
        "hy": {"native": "Հայերեն", "flag": "https://flagcdn.com/am.svg"},
    }
    available = {value for value, _label in actions.SUBTITLE_LANGUAGE_CHOICES}
    options = [
        {
            "code": "auto",
            "label": labels["auto"].get(language, labels["auto"]["en"]),
            "flag": "",
        }
    ]
    language_lookup = {item["code"]: item for item in LANGUAGE_OPTIONS}
    options.extend(
        {
            "code": value,
            "label": subtitle_language_meta.get(value, {}).get("native") or language_lookup.get(value, {}).get("native") or label,
            "flag": subtitle_language_meta.get(value, {}).get("flag") or language_lookup.get(value, {}).get("flag", ""),
        }
        for value, label in actions.SUBTITLE_LANGUAGE_CHOICES
        if value != "auto" and value in available
    )
    return options


def _localized_resume_templates(language: str) -> list[tuple[str, str]]:
    labels = {
        "1": {"en": "Classic", "ru": "Классика", "uk": "Класика", "fr": "Classique", "de": "Klassisch", "es": "Clásico", "ka": "კლასიკური", "hy": "Դասական", "it": "Classico"},
        "2": {"en": "Executive", "ru": "Руководитель", "uk": "Керівник", "fr": "Direction", "de": "Executive", "es": "Ejecutivo", "ka": "აღმასრულებელი", "hy": "Ղեկավար", "it": "Executive"},
        "3": {"en": "Creative", "ru": "Креатив", "uk": "Креатив", "fr": "Créatif", "de": "Kreativ", "es": "Creativo", "ka": "კრეატიული", "hy": "Ստեղծարար", "it": "Creativo"},
        "4": {"en": "Modern", "ru": "Современный", "uk": "Сучасний", "fr": "Moderne", "de": "Modern", "es": "Moderno", "ka": "თანამედროვე", "hy": "Ժամանակակից", "it": "Moderno"},
        "5": {"en": "Tech", "ru": "Техно", "uk": "Техно", "fr": "Tech", "de": "Tech", "es": "Tech", "ka": "ტექ", "hy": "Տեխ", "it": "Tech"},
        "6": {"en": "Minimal", "ru": "Минимал", "uk": "Мінімал", "fr": "Minimal", "de": "Minimal", "es": "Minimal", "ka": "მინიმალი", "hy": "Մինիմալ", "it": "Minimal"},
        "7": {"en": "Premium", "ru": "Премиум", "uk": "Преміум", "fr": "Premium", "de": "Premium", "es": "Premium", "ka": "პრემიუმი", "hy": "Պրեմիում", "it": "Premium"},
        "8": {"en": "Focus", "ru": "Фокус", "uk": "Фокус", "fr": "Focus", "de": "Fokus", "es": "Enfoque", "ka": "ფოკუსი", "hy": "Ֆոկուս", "it": "Focus"},
        "9": {"en": "Nordic", "ru": "Сканди", "uk": "Сканди", "fr": "Nordique", "de": "Nordisch", "es": "Nórdico", "ka": "ნორდიკული", "hy": "Նորդիկ", "it": "Nordico"},
        "10": {"en": "Legal", "ru": "Юридический", "uk": "Юридичний", "fr": "Juridique", "de": "Legal", "es": "Legal", "ka": "იურიდიული", "hy": "Իրավական", "it": "Legale"},
        "11": {"en": "Startup", "ru": "Стартап", "uk": "Стартап", "fr": "Startup", "de": "Startup", "es": "Startup", "ka": "სტარტაპი", "hy": "Ստարտափ", "it": "Startup"},
        "12": {"en": "Finance", "ru": "Финансы", "uk": "Фінанси", "fr": "Finance", "de": "Finanzen", "es": "Finanzas", "ka": "ფინანსები", "hy": "Ֆինանսներ", "it": "Finanza"},
        "13": {"en": "Academic", "ru": "Академический", "uk": "Академічний", "fr": "Académique", "de": "Akademisch", "es": "Académico", "ka": "აკადემიური", "hy": "Ակադեմիական", "it": "Accademico"},
        "14": {"en": "Compact", "ru": "Компактный", "uk": "Компактний", "fr": "Compact", "de": "Kompakt", "es": "Compacto", "ka": "კომპაქტური", "hy": "Կոմպակտ", "it": "Compatto"},
    }
    return [(value, labels.get(value, {}).get(language, labels.get(value, {}).get("en", label))) for value, label in actions.RESUME_TEMPLATE_CHOICES]


def _localized_value(key: str, language: str, fallback: str) -> str:
    values = {
        "queued": {"en": "Queued", "ru": "Очередь", "uk": "Черга", "fr": "En file", "de": "Warteschlange", "es": "En cola", "ka": "რიგშია", "hy": "Հերթում", "it": "In coda"},
        "running": {"en": "Running", "ru": "В работе", "uk": "В роботі", "fr": "En cours", "de": "Läuft", "es": "En curso", "ka": "მიმდინარეობს", "hy": "Ընթացքում", "it": "In corso"},
        "completed": {"en": "Completed", "ru": "Готово", "uk": "Готово", "fr": "Terminé", "de": "Abgeschlossen", "es": "Completado", "ka": "დასრულდა", "hy": "Ավարտված", "it": "Completato"},
        "failed": {"en": "Failed", "ru": "Ошибка", "uk": "Помилка", "fr": "Échec", "de": "Fehler", "es": "Error", "ka": "შეცდომა", "hy": "Սխալ", "it": "Errore"},
        "convert": {"en": "Convert", "ru": "Конвертер", "uk": "Конвертер", "fr": "Convertir", "de": "Konvertieren", "es": "Convertir", "ka": "კონვერტაცია", "hy": "Փոխարկել", "it": "Converti"},
        "image": {"en": "Image", "ru": "Изображение", "uk": "Зображення", "fr": "Image", "de": "Bild", "es": "Imagen", "ka": "სურათი", "hy": "Պատկեր", "it": "Immagine"},
        "video": {"en": "Video", "ru": "Видео", "uk": "Відео", "fr": "Vidéo", "de": "Video", "es": "Vídeo", "ka": "ვიდეო", "hy": "Տեսանյութ", "it": "Video"},
        "ready": {"en": "Done", "ru": "Готово", "uk": "Готово", "fr": "Terminé", "de": "Fertig", "es": "Listo", "ka": "მზადაა", "hy": "Պատրաստ", "it": "Pronto"},
        "queued_msg": {"en": "Queued", "ru": "В очереди", "uk": "У черзі", "fr": "En file", "de": "In der Warteschlange", "es": "En cola", "ka": "რიგშია", "hy": "Հերթում է", "it": "In coda"},
        "starting": {"en": "Starting task", "ru": "Стартую задачу", "uk": "Запускаю задачу", "fr": "Démarrage de la tâche", "de": "Aufgabe startet", "es": "Iniciando tarea", "ka": "ამოცანა იწყება", "hy": "Առաջադրանքը սկսվում է", "it": "Avvio attività"},
        "convert_image": {"en": "Converting image", "ru": "Конвертирую изображение", "uk": "Конвертую зображення", "fr": "Conversion de l'image", "de": "Bild wird konvertiert", "es": "Convirtiendo imagen", "ka": "სურათი კონვერტირდება", "hy": "Պատկերը փոխարկվում է", "it": "Conversione immagine"},
        "read_video": {"en": "Reading video parameters", "ru": "Читаю параметры видео", "uk": "Читаю параметри відео", "fr": "Lecture des paramètres vidéo", "de": "Videoparameter werden gelesen", "es": "Leyendo parámetros de vídeo", "ka": "ვიდეო პარამეტრების კითხვა", "hy": "Տեսանյութի պարամետրերի ընթերցում", "it": "Lettura parametri video"},
        "ready_file": {"en": "File ready", "ru": "Файл готов", "uk": "Файл готовий", "fr": "Fichier prêt", "de": "Datei bereit", "es": "Archivo listo", "ka": "ფაილი მზადაა", "hy": "Ֆայլը պատրաստ է", "it": "File pronto"},
        "interrupted": {"en": "Task interrupted", "ru": "Задача прервана", "uk": "Задачу перервано", "fr": "Tâche interrompue", "de": "Aufgabe unterbrochen", "es": "Tarea interrumpida", "ka": "ამოცანა შეწყდა", "hy": "Առաջադրանքը ընդհատվեց", "it": "Attività interrotta"},
        "server_restarted": {"en": "The server restarted before the task finished.", "ru": "Сервер был перезапущен до завершения задачи.", "uk": "Сервер було перезапущено до завершення задачі.", "fr": "Le serveur a redémarré avant la fin de la tâche.", "de": "Der Server wurde vor Abschluss der Aufgabe neu gestartet.", "es": "El servidor se reinició antes de finalizar la tarea.", "ka": "სერვერი ამოცანის დასრულებამდე გადაიტვირთა.", "hy": "Սերվերը վերագործարկվեց մինչև առաջադրանքի ավարտը։", "it": "Il server si è riavviato prima della fine dell'attività."},
    }
    values["cancelled"] = {"en": "Cancelled", "ru": "Отменено", "uk": "Скасовано", "fr": "Annulé", "de": "Abgebrochen", "es": "Cancelado", "ka": "გაუქმდა", "hy": "Չեղարկված", "it": "Annullato"}
    values["paused"] = {"en": "Paused", "ru": "На паузе", "uk": "На паузі", "fr": "En pause", "de": "Pausiert", "es": "En pausa", "ka": "Paused", "hy": "Paused", "it": "In pausa"}
    values["processing"] = {"en": "Processing", "ru": "В обработке", "uk": "В обробці", "fr": "Traitement", "de": "Verarbeitung", "es": "Procesando", "ka": "Processing", "hy": "Processing", "it": "In elaborazione"}
    return values.get(key, {}).get(language, values.get(key, {}).get("en", fallback))


def _localize_runtime_text(text: str, language: str) -> str:
    if not text:
        return text
    replacements = {
        "В очереди": _localized_value("queued_msg", language, "Queued"),
        "Стартую задачу": _localized_value("starting", language, "Starting task"),
        "Конвертирую изображение": _localized_value("convert_image", language, "Converting image"),
        "Читаю параметры видео": _localized_value("read_video", language, "Reading video parameters"),
        "Задача прервана": _localized_value("interrupted", language, "Task interrupted"),
        "Сервер был перезапущен до завершения задачи.": _localized_value("server_restarted", language, "The server restarted before the task finished."),
        "Ошибка": _localized_value("failed", language, "Failed"),
    }
    value = str(text)
    value = value.replace("Cancelled", _localized_value("cancelled", language, "Cancelled"))
    for source, target in replacements.items():
        value = value.replace(source, target)
    value = value.replace("Файл готов", _localized_value("ready_file", language, "File ready"))
    value = value.replace("Готово", _localized_value("ready", language, "Done"))
    value = value.replace("Изображение", _localized_value("image", language, "Image"))
    value = value.replace("Видео", _localized_value("video", language, "Video"))
    value = value.replace("Конвертация", _localized_value("convert", language, "Convert"))
    return value


def _ai_summary_items(ai_meta: object, language: str) -> list[dict[str, str]]:
    if not isinstance(ai_meta, dict):
        return []
    items: list[dict[str, str]] = []
    for name, meta in ai_meta.items():
        if not isinstance(meta, dict):
            continue
        key = str(name or "").strip()
        status = str(meta.get("status") or "unknown").strip().lower() or "unknown"
        label_key = "ai_used" if status == "used" else "ai_fallback" if status == "fallback" else "ai_unknown"
        reason = _friendly_ai_reason(str(meta.get("fallback_reason") or meta.get("reason") or "").strip(), language)
        model = str(meta.get("model") or "").strip()
        selected_outputs = meta.get("selected_outputs")
        output_count = len(selected_outputs) if isinstance(selected_outputs, list) else 0
        scenario = translate(f"ai_{key}", language) if key else translate("ai_unknown", language)
        details = [scenario]
        if model:
            details.append(model)
        if reason and status != "fallback":
            details.append(reason)
        if output_count:
            details.append(f"{output_count} {translate('ai_outputs', language)}")
        items.append(
            {
                "key": key,
                "status": status,
                "label": translate(label_key, language),
                "detail": " · ".join(part for part in details if part),
                "reason": reason,
                "model": model,
            }
        )
    return items


def _friendly_ai_reason(reason: str, language: str) -> str:
    value = _localize_runtime_text(str(reason or "").strip(), language)
    if not value:
        return ""
    lowered = value.lower()
    localized = {
        "quota": {
            "en": "OpenAI quota is exhausted, local fallback used",
            "ru": "Лимит OpenAI исчерпан, использован локальный fallback",
            "uk": "Ліміт OpenAI вичерпано, використано локальний fallback",
        },
        "not_configured": {
            "en": "OpenAI is not configured, local fallback used",
            "ru": "OpenAI не настроен, использован локальный fallback",
            "uk": "OpenAI не налаштовано, використано локальний fallback",
        },
        "no_usable": {
            "en": "OpenAI returned no usable result, local fallback used",
            "ru": "OpenAI не вернул подходящий результат, использован локальный fallback",
            "uk": "OpenAI не повернув придатний результат, використано локальний fallback",
        },
        "generic": {
            "en": "AI unavailable, local fallback used",
            "ru": "AI недоступен, использован локальный fallback",
            "uk": "AI недоступний, використано локальний fallback",
        },
    }
    lang = language if language in {"en", "ru", "uk"} else "en"
    if "insufficient_quota" in lowered or "insufficient quota" in lowered or "exceeded your current quota" in lowered or "error code: 429" in lowered:
        return localized["quota"][lang]
    if "openai is not configured" in lowered:
        return localized["not_configured"][lang]
    if "no usable" in lowered or "returned no cues" in lowered:
        return localized["no_usable"][lang]
    if "error code:" in lowered or "traceback" in lowered or "{'error'" in lowered or '"error"' in lowered:
        return localized["generic"][lang]
    value = re.sub(r"https?://\S+", "", value)
    value = re.sub(r"\s+", " ", value).strip(" -:;·")
    return value[:140]


def _localize_job(job: dict, language: str) -> dict:
    prepared = dict(job)
    prepared["title"] = _localize_runtime_text(str(prepared.get("title") or ""), language)
    prepared["message"] = _localize_runtime_text(str(prepared.get("message") or ""), language)
    prepared["error"] = _localize_runtime_text(str(prepared.get("error") or ""), language)
    prepared["status_label"] = _localized_value(str(prepared.get("status") or ""), language, str(prepared.get("status") or ""))
    prepared["kind_label"] = _localized_value(str(prepared.get("kind") or ""), language, str(prepared.get("kind") or ""))
    outputs = []
    for output in prepared.get("outputs", []):
        item = dict(output)
        item["label"] = _localize_runtime_text(str(item.get("label") or ""), language)
        outputs.append(item)
    prepared["outputs"] = outputs
    prepared["ai_summary"] = _ai_summary_items(prepared.get("ai"), language)
    return prepared


def _localize_events(events: list[dict], language: str) -> list[dict]:
    localized = []
    for event in events:
        item = dict(event)
        item["status_label"] = _localized_value(str(item.get("status") or ""), language, str(item.get("status") or ""))
        item["message"] = _localize_runtime_text(str(item.get("message") or ""), language)
        localized.append(item)
    return localized


def _resume_ai_payload(request: HttpRequest, language: str) -> dict[str, object]:
    fields = {field: request.POST.get(field, "").strip() for field in RESUME_FIELDS}
    resume_text = "\n".join(
        value
        for key, value in fields.items()
        if key not in {"vacancy_text", "cover_letter"} and value
    )
    return {
        **fields,
        "field": request.POST.get("field", "").strip(),
        "mode": request.POST.get("mode", "stronger").strip(),
        "text": request.POST.get("text", "").strip(),
        "vacancy": request.POST.get("vacancy_text", "").strip(),
        "resume": resume_text,
        "language": language,
    }


def _resume_tokens(text: str) -> list[str]:
    stop = {
        "and", "the", "for", "with", "you", "your", "are", "this", "that", "from", "will", "have", "has", "our",
        "для", "или", "как", "что", "это", "при", "на", "по", "над", "под", "про", "без", "або", "що", "цей",
    }
    words = re.findall(r"[\w+#.-]{3,}", str(text or "").lower(), flags=re.UNICODE)
    return [word for word in words if word not in stop and not word.isdigit()]


def _local_resume_rewrite(payload: dict[str, object]) -> dict[str, object]:
    text = str(payload.get("text") or "").strip()
    mode = str(payload.get("mode") or "stronger")
    field = str(payload.get("field") or "")
    if mode == "shorter":
        result = "\n".join(line.strip() for line in text.splitlines() if line.strip())[:900]
    elif mode == "english":
        result = text if re.search(r"[A-Za-z]", text) else "Product-minded specialist focused on clear execution, measurable outcomes and practical teamwork."
    elif mode == "formal":
        result = f"{text}\nFocus: structured communication, ownership, measurable delivery and clear business value.".strip()
    else:
        prefix = "- " if field in {"experience", "achievements"} else ""
        result = f"{text}\n{prefix}Strengthened outcomes through clearer priorities, practical execution and measurable improvements.".strip()
    return {"text": result[:7000], "notes": ["Local fallback used", "Add real metrics where possible"], "model": "local"}


def _local_resume_match(payload: dict[str, object]) -> dict[str, object]:
    vacancy_words = list(dict.fromkeys(_resume_tokens(str(payload.get("vacancy") or ""))))[:28]
    resume_set = set(_resume_tokens(str(payload.get("resume") or "")))
    matched = [word for word in vacancy_words if word in resume_set]
    missing = [word for word in vacancy_words if word not in resume_set]
    base = round((len(matched) / len(vacancy_words)) * 70) if vacancy_words else 0
    resume = str(payload.get("resume") or "")
    bonus = 0
    bonus += 6 if str(payload.get("name") or "").strip() else 0
    bonus += 6 if str(payload.get("contact") or "").strip() else 0
    bonus += 8 if len(_resume_tokens(str(payload.get("summary") or ""))) >= 12 else 0
    bonus += 8 if re.search(r"\d", resume) else 0
    bonus += 8 if len(_resume_tokens(str(payload.get("skills") or ""))) >= 5 else 0
    suggestions = []
    if missing:
        suggestions.append("Add true missing keywords to summary, skills or experience.")
    if not re.search(r"\d", resume):
        suggestions.append("Add measurable results: %, count, budget, team size or deadline.")
    if len(_resume_tokens(str(payload.get("summary") or ""))) < 12:
        suggestions.append("Make summary more specific: role, domain, strength, outcome.")
    return {
        "score": max(0, min(100, base + bonus)),
        "matched_keywords": matched[:16],
        "missing_keywords": missing[:16],
        "suggestions": suggestions[:8],
        "summary": "Local keyword and ATS checklist analysis.",
        "model": "local",
    }


def _local_resume_cover_letter(payload: dict[str, object]) -> dict[str, object]:
    name = str(payload.get("name") or "").strip() or "Candidate"
    role = str(payload.get("target_role") or payload.get("position") or "").strip() or "this role"
    value = str(payload.get("value_offer") or payload.get("summary") or "").strip() or "I can bring structured execution, clear communication and measurable results."
    skills = str(payload.get("skills") or "").strip() or "relevant tools, teamwork and delivery discipline"
    return {
        "subject": f"Application for {role}",
        "letter": (
            f"Hello,\n\nI am interested in the {role} position. {value}\n\n"
            f"My relevant strengths include {skills}. I would be glad to discuss how my experience can help your team move faster and produce stronger results.\n\n"
            f"Best regards,\n{name}"
        )[:6000],
        "model": "local",
    }


def _extract_originality_upload(upload, language: str) -> str:
    suffix = Path(upload.name or "").suffix.lower()
    data = b"".join(upload.chunks())
    if suffix == ".doc":
        raise ValueError(_originality_runtime_text("legacy_doc", language))
    if suffix == ".docx":
        return _extract_docx_text(data, language)
    if suffix == ".pdf":
        return _extract_pdf_text(data, language)
    if suffix in {".txt", ".md", ".csv", ".json", ".rtf", ".html", ".htm", ""} or str(getattr(upload, "content_type", "") or "").startswith("text/"):
        text = _decode_text_bytes(data)
        if suffix in {".html", ".htm"}:
            text = html.unescape(strip_tags(text))
        if suffix == ".rtf":
            text = _strip_rtf_text(text)
        return text
    raise ValueError(_originality_runtime_text("unsupported", language))


def _extract_docx_text(data: bytes, language: str) -> str:
    try:
        document = Document(BytesIO(data))
    except Exception as exc:
        raise ValueError(_originality_runtime_text("docx_failed", language)) from exc
    parts: list[str] = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)


def _extract_pdf_text(data: bytes, language: str) -> str:
    try:
        reader = PdfReader(BytesIO(data))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                pass
        pages = []
        for page in reader.pages[:80]:
            pages.append(page.extract_text() or "")
    except Exception as exc:
        raise ValueError(_originality_runtime_text("pdf_failed", language)) from exc
    text = "\n\n".join(page.strip() for page in pages if page.strip())
    if not text.strip():
        raise ValueError(_originality_runtime_text("pdf_empty", language))
    return _repair_extracted_text(text)


def _decode_text_bytes(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "cp1251", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _strip_rtf_text(text: str) -> str:
    text = re.sub(r"\\'[0-9a-fA-F]{2}", " ", text)
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", text)
    text = text.replace("{", " ").replace("}", " ")
    return re.sub(r"\s+", " ", text)


def _clean_originality_text(text: str) -> str:
    text = html.unescape(str(text or ""))
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\u00a0", " ")
    text = re.sub(r"[ \t\f\v]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return _repair_extracted_text(text.strip())


def _repair_extracted_text(text: str) -> str:
    text = _repair_mojibake_cyrillic(str(text or "")).replace("\r\n", "\n").replace("\r", "\n")
    raw_lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    lines = [line for line in raw_lines if line and not re.fullmatch(r"[-–—_ ]*\d{1,4}[-–—_ ]*", line)]
    if not lines:
        return text.strip()
    word_counts = [len(_originality_words(line)) for line in lines]
    short_ratio = sum(1 for count in word_counts if count <= 6) / max(1, len(word_counts))
    median_words = statistics.median(word_counts) if word_counts else 0
    if len(lines) < 18 or (short_ratio < 0.48 and median_words > 7):
        return _repair_pdf_spacing_artifacts(re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip())

    paragraphs: list[str] = []
    current = ""
    for line in lines:
        if _is_noise_pdf_line(line):
            continue
        line = _normalize_pdf_line(line)
        if not line:
            continue
        is_heading = _looks_like_section_heading(line)
        should_flush = bool(current) and (is_heading or _ends_complete_sentence(current) and len(_originality_words(current)) >= 18)
        if should_flush:
            paragraphs.append(current.strip())
            current = ""
        current = f"{current} {line}".strip() if current else line
        if is_heading and len(_originality_words(current)) <= 12:
            paragraphs.append(current.strip())
            current = ""
    if current.strip():
        paragraphs.append(current.strip())
    repaired = "\n\n".join(_dedupe_repeated_short_lines(paragraphs))
    repaired = re.sub(r"\s+([,;:!?])", r"\1", repaired).strip()
    return _repair_pdf_spacing_artifacts(repaired)


def _repair_mojibake_cyrillic(text: str) -> str:
    text = str(text or "")
    candidates = [text]
    for encoding in ("cp1251", "latin-1", "cp1252"):
        try:
            repaired = text.encode(encoding, errors="ignore").decode("utf-8", errors="ignore")
        except Exception:
            continue
        if repaired.strip():
            candidates.append(repaired)
    return max(candidates, key=_cyrillic_text_quality)


def _cyrillic_text_quality(text: str) -> int:
    text = str(text or "")
    cyrillic = len(re.findall(r"[А-Яа-яЁёІіЇїЄєҐґ]", text))
    common_lower = len(re.findall(r"[а-яёіїєґ]", text))
    mojibake_pairs = len(re.findall(r"[РС][\u00a0-\u00bf\u0400-\u045f\u2010-\u203a\u20ac\u2116]", text))
    replacements = text.count("\ufffd") + text.count("?")
    return cyrillic + common_lower * 2 - mojibake_pairs * 8 - replacements * 12


def _normalize_pdf_line(line: str) -> str:
    line = re.sub(r"(?<=\d)\s+\.\s*(?=\d)", ".", line)
    line = _repair_link_spacing_artifacts(line)
    line = re.sub(r"\b([A-Za-zА-Яа-яЁёІіЇїЄєҐґ])\s+\.\s+", r"\1. ", line)
    line = _repair_pdf_spacing_artifacts(line)
    line = re.sub(r"\s{2,}", " ", line)
    return line.strip()


def _repair_link_spacing_artifacts(text: str) -> str:
    text = str(text or "")
    text = re.sub(r"\b(https?):\s*/\s*/", r"\1://", text, flags=re.IGNORECASE)
    text = re.sub(r"\b(www)\s+\.\s+", r"\1.", text, flags=re.IGNORECASE)
    text = re.sub(r"([A-Za-z0-9._%+-]+@[A-Za-z0-9._%+-]+)\s*\.\s+([A-Za-z]{2,})(?=(?:[\s/]|$))", r"\1.\2", text)
    text = re.sub(r"((?:https?://|www\.)[A-Za-z0-9.-]+)\s*\.\s+([A-Za-z]{2,})(?=(?:[\s/]|$))", r"\1.\2", text, flags=re.IGNORECASE)
    text = re.sub(r"\b([A-Za-z0-9-]{3,}\d[A-Za-z0-9-]*)\s*\.\s+([A-Za-z]{2,})(?=(?:[\s/]|$))", r"\1.\2", text)
    text = re.sub(r"\b(https?://)\s+", r"\1", text, flags=re.IGNORECASE)
    text = re.sub(r"/\s+/", "//", text)
    return text


def _repair_pdf_spacing_artifacts(text: str) -> str:
    text = str(text or "")
    replacements = (
        (r"\bPy\s+thon\b", "Python"),
        (r"\bD\s+j\s+ango\b", "Django"),
        (r"\bD\s+jango\b", "Django"),
        (r"\bF\s+igma\b", "Figma"),
        (r"\bG\s+it\s+H\s+u\s+b\b", "GitHub"),
        (r"\bV\s+S\s+C\s+o\s+d\s+e\b", "VS Code"),
        (r"\bX\s+co\s+d\s+e\b", "Xcode"),
        (r"\bP\s+ostgre\s+S\s+Q\s+L\b", "PostgreSQL"),
        (r"\bR\s+E\s+S\s+T\b", "REST"),
        (r"\bC\s+R\s+M\b", "CRM"),
        (r"\bA\s+I\b", "AI"),
        (r"\bU\s+I\s*/\s*U\s+X\b", "UI/UX"),
        (r"\bi\s+O\s+S\b", "iOS"),
        (r"\bDe\s+v\s+elopment\b", "Development"),
        (r"\bde\s+v\s+elopment\b", "development"),
        (r"\blea\s+d\s+ership\b", "leadership"),
        (r"\bhan\s+d\s+s-on\b", "hands-on"),
        (r"\bPro\s+j\s+ect\b", "Project"),
        (r"\bpro\s+j\s+ects\b", "projects"),
        (r"\bFullstack\b", "Fullstack"),
        (r"\bF\s+ullstack\b", "Fullstack"),
        (r"\btra\s+f\s+fi\s+c\b", "traffic"),
        (r"\bplat\s+f\s+orm\b", "platform"),
    )
    for pattern, value in replacements:
        text = re.sub(pattern, value, text, flags=re.IGNORECASE)
    text = re.sub(r"\b([a-z]{3,})\s+([b-hj-z])\s+([a-z]{2,})(?=[\s,.;:!?)]|$)", r"\1\2\3", text)
    text = re.sub(r"\b([A-Z][a-z]{2,})\s+([b-hj-z])\s+([a-z]{2,})(?=[\s,.;:!?)]|$)", r"\1\2\3", text)
    text = re.sub(r"\b(?![AI]\s)([A-Z])\s+([a-z]{3,})(?=[\s,.;:!?)]|$)", r"\1\2", text)
    return _repair_link_spacing_artifacts(text)


def _is_noise_pdf_line(line: str) -> bool:
    if len(line) <= 2:
        return True
    alpha_count = len(re.findall(r"[A-Za-zА-Яа-яЁёІіЇїЄєҐґ]", line))
    if alpha_count == 0 and len(line) < 12:
        return True
    if len(_originality_words(line)) <= 1 and re.fullmatch(r"[\d\s./№()_-]+", line):
        return True
    return False


def _looks_like_section_heading(line: str) -> bool:
    clean = line.strip()
    words = _originality_words(clean)
    if not (2 <= len(words) <= 12):
        return False
    if re.search(r"[;:!?…]$", clean):
        return False
    if re.fullmatch(r"\d+(?:\.\d+)*\.?", clean):
        return False
    alpha_count = len(re.findall(r"[A-Za-zА-Яа-яЁёІіЇїЄєҐґ]", clean))
    if alpha_count < 6:
        return False
    digit_ratio = len(re.findall(r"\d", clean)) / max(1, len(clean))
    if digit_ratio > 0.32:
        return False
    return bool(re.match(r"^(?:\d+(?:\.\d+)*\.?\s+)?[A-ZА-ЯІЇЄҐ]", clean))


def _ends_complete_sentence(text: str) -> bool:
    return bool(re.search(r"[.!?…][\"')\]]?$", text.strip()))


def _dedupe_repeated_short_lines(lines: list[str]) -> list[str]:
    counts = Counter(line for line in lines if len(line) <= 90)
    seen: Counter[str] = Counter()
    result = []
    for line in lines:
        if counts[line] >= 4:
            seen[line] += 1
            if seen[line] > 1:
                continue
        result.append(line)
    return result


def _analyze_originality_text(text: str, language: str, source_name: str = "", truncated: bool = False) -> dict[str, object]:
    sentences = _originality_sentences(text)
    words = _originality_words(text)
    paragraphs = _analysis_paragraphs(text)
    document_kind = _originality_document_kind(text)
    word_count = len(words)
    sentence_count = len(sentences)
    paragraph_count = len(paragraphs)
    lengths = [len(item["words"]) for item in sentences if item["words"]]
    avg_sentence = statistics.mean(lengths) if lengths else 0
    sentence_stdev = statistics.pstdev(lengths) if len(lengths) > 1 else 0
    lexical_diversity = len(set(words)) / max(1, word_count)

    shingle_counts = Counter(_word_shingles(words, 8))
    shingle_total = sum(shingle_counts.values())
    duplicate_shingles = sum(count - 1 for count in shingle_counts.values() if count > 1)
    duplicate_ratio = duplicate_shingles / max(1, shingle_total)
    normalized_sentences = [item["norm"] for item in sentences if len(item["norm"]) > 34]
    sentence_counts = Counter(normalized_sentences)
    repeated_sentence_norms = {value for value, count in sentence_counts.items() if count > 1}
    repeated_sentence_ratio = sum(count - 1 for count in sentence_counts.values() if count > 1) / max(1, len(normalized_sentences))

    long_ratio = sum(1 for length in lengths if length >= 36) / max(1, len(lengths))
    short_ratio = sum(1 for length in lengths if length <= 6) / max(1, len(lengths))
    ai_marker_hits = _ai_marker_hits(text)
    source_count = _source_marker_count(text)
    source_claim_count = _source_claim_count(text)
    heading_count = min(_heading_count(text), max(0, paragraph_count + 4))
    bullet_count = len(re.findall(r"(?m)^\s*(?:[-*•]|\d+[.)])\s+", text))

    uniqueness_score = _clamp_score(100 - duplicate_ratio * 310 - repeated_sentence_ratio * 90)
    repetition_score = _clamp_score(100 - duplicate_ratio * 360 - repeated_sentence_ratio * 110)
    readability_score = _clamp_score(96 - long_ratio * 54 - short_ratio * 18 - max(0, avg_sentence - 28) * 1.4 - max(0, 10 - avg_sentence) * 1.2)
    structure_score = _clamp_score(62 + min(22, paragraph_count * 2.2) + min(10, heading_count * 3) + min(6, bullet_count * 1.2) - (18 if paragraph_count <= 1 and word_count > 280 else 0))
    sources_score = _source_support_score(word_count, source_count, source_claim_count, document_kind)
    ai_risk = _ai_risk_score(
        word_count=word_count,
        sentence_count=sentence_count,
        sentence_stdev=sentence_stdev,
        avg_sentence=avg_sentence,
        marker_hits=ai_marker_hits,
        lexical_diversity=lexical_diversity,
        long_ratio=long_ratio,
        source_count=source_count,
    )
    overall_score = _clamp_score(
        uniqueness_score * 0.32
        + (100 - ai_risk) * 0.22
        + readability_score * 0.18
        + repetition_score * 0.13
        + sources_score * 0.08
        + structure_score * 0.07
    )

    metrics = [
        _originality_metric("uniqueness", translate("metric_uniqueness", language), uniqueness_score, _score_tone(uniqueness_score), _metric_detail("uniqueness", language, round(duplicate_ratio * 100), len(repeated_sentence_norms))),
        _originality_metric("ai_risk", translate("metric_ai_risk", language), ai_risk, _risk_tone(ai_risk), _metric_detail("ai_risk", language, ai_marker_hits, round(sentence_stdev, 1))),
        _originality_metric("readability", translate("metric_readability", language), readability_score, _score_tone(readability_score), _metric_detail("readability", language, round(avg_sentence, 1), round(long_ratio * 100))),
        _originality_metric("repetition", translate("metric_repetition", language), repetition_score, _score_tone(repetition_score), _metric_detail("repetition", language, duplicate_shingles, round(repeated_sentence_ratio * 100))),
        _originality_metric("sources", translate("metric_sources", language), sources_score, _score_tone(sources_score), _metric_detail("sources", language, source_count, source_claim_count)),
        _originality_metric("structure", translate("metric_structure", language), structure_score, _score_tone(structure_score), _metric_detail("structure", language, paragraph_count, heading_count)),
    ]
    segments = _originality_segments(sentences, repeated_sentence_norms, source_count, language, document_kind)
    return {
        "overall": {"score": overall_score, "tone": _score_tone(overall_score), "label": translate("originality_score", language)},
        "source": {"name": source_name, "words": word_count, "sentences": sentence_count, "paragraphs": paragraph_count, "kind": document_kind, "kind_label": _document_kind_label(document_kind, language), "truncated": truncated},
        "metrics": metrics,
        "segments": segments,
        "highlights": _originality_highlight_summary(segments, language),
    }


def _originality_sentences(text: str) -> list[dict[str, object]]:
    units = []
    protected = _protect_originality_sentence_boundaries(text)
    for match in ORIGINALITY_SENTENCE_RE.finditer(protected):
        stripped = _restore_originality_sentence_boundaries(match.group(0).strip())
        if not stripped:
            continue
        words = _originality_words(stripped)
        units.append({"text": stripped, "words": words, "norm": _normalize_originality_sentence(stripped)})
    if not units and text.strip():
        units.append({"text": text.strip(), "words": _originality_words(text), "norm": _normalize_originality_sentence(text)})
    return units


def _analysis_paragraphs(text: str) -> list[str]:
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n+", text) if part.strip()]
    if not paragraphs:
        return [text.strip()] if text.strip() else []
    words = _originality_words(text)
    if len(paragraphs) > max(8, len(words) // 24):
        merged: list[str] = []
        current = ""
        for part in paragraphs:
            if current and (len(_originality_words(current)) >= 55 or _looks_like_section_heading(part)):
                merged.append(current.strip())
                current = ""
            current = f"{current} {part}".strip() if current else part
        if current:
            merged.append(current.strip())
        return merged
    return paragraphs


def _protect_originality_sentence_boundaries(text: str) -> str:
    protected = text
    dot_token = "__ORIGINALITY_DOT__"

    def protect_dots(match: re.Match[str]) -> str:
        return match.group(0).replace(".", dot_token)

    protected = re.sub(r"\b[\w.%+-]+@[\w.-]+\.[A-Za-z]{2,}\b", protect_dots, protected, flags=re.UNICODE)
    protected = re.sub(r"\bhttps?://[^\s]+", protect_dots, protected, flags=re.IGNORECASE)
    protected = re.sub(r"\bwww\.[^\s]+", protect_dots, protected, flags=re.IGNORECASE)
    protected = re.sub(r"\b[A-Za-z0-9-]{2,}\.[A-Za-z]{2,}(?:\.[A-Za-z]{2,})?\b", protect_dots, protected)
    protected = re.sub(r"(?<=\d)\.(?=\d)", "∯", protected)
    protected = re.sub(r"\b([A-Za-zА-Яа-яЁёІіЇїЄєҐґ])\.(?=\s)", r"\1∯", protected)
    protected = re.sub(r"\b(м|р|п|ст|кв|вул|ім|т|гр|Mr|Mrs|Ms|Dr|Prof)\.(?=\s)", lambda match: match.group(0).replace(".", "∯"), protected, flags=re.IGNORECASE)
    protected = re.sub(r"\b(\d{1,3})\.(?=\s+(?:[a-zа-яіїєґ]|[0-9]))", r"\1∯", protected, flags=re.IGNORECASE)
    protected = re.sub(r"\b(\d{1,3})\.(?=\s+[A-ZА-ЯІЇЄҐ][a-zа-яіїєґ])", r"\1∯", protected)
    return protected


def _restore_originality_sentence_boundaries(text: str) -> str:
    text = text.replace("__ORIGINALITY_DOT__", ".")
    return text.replace("∯", ".")


def _originality_words(text: str) -> list[str]:
    return [word.strip("’'-").lower() for word in ORIGINALITY_WORD_RE.findall(text.lower()) if len(word.strip("’'-")) > 1]


def _normalize_originality_sentence(text: str) -> str:
    return re.sub(r"\s+", " ", re.sub(r"[^\wА-Яа-яЁёІіЇїЄєҐґ]+", " ", text.lower(), flags=re.UNICODE)).strip()


def _word_shingles(words: list[str], size: int) -> list[str]:
    if len(words) < size:
        return []
    return [" ".join(words[index : index + size]) for index in range(0, len(words) - size + 1)]


def _originality_mode_code(value: object) -> str:
    code = str(value or "local").strip().lower()
    return code if code in ORIGINALITY_MODES else "local"


def _originality_check_metadata(mode_code: str, text: str) -> dict[str, object]:
    mode = ORIGINALITY_MODES[_originality_mode_code(mode_code)]
    probes = _originality_web_probe_queries(text, int(mode["max_web_queries"]))
    return {
        "mode": mode_code,
        "mode_label": mode["label"],
        "price_cherryx": mode["price"],
        "max_chars": mode["max_chars"],
        "web_queries_limit": mode["max_web_queries"],
        "web_status": "planned" if probes else "not_requested",
        "web_provider": "not_connected",
        "web_probes": probes,
        "share_enabled": True,
    }


def _originality_web_probe_queries(text: str, limit: int) -> list[dict[str, object]]:
    if limit <= 0:
        return []
    sentences = [item for item in _originality_sentences(text) if len(item.get("words", [])) >= 8]
    scored: list[tuple[int, str]] = []
    seen: set[str] = set()
    for item in sentences:
        raw = re.sub(r"\s+", " ", str(item.get("text") or "")).strip()
        words = _originality_words(raw)
        if len(words) < 8:
            continue
        quote = " ".join(raw.split()[: min(14, max(8, len(raw.split())))])
        key = _normalize_originality_sentence(quote)[:120]
        if not key or key in seen:
            continue
        seen.add(key)
        score = len(set(words)) + min(24, len(words))
        scored.append((score, quote))
    scored.sort(reverse=True)
    return [
        {"query": f'"{quote}"', "status": "queued_stub", "matches": []}
        for _, quote in scored[:limit]
    ]


def _ai_marker_hits(text: str) -> int:
    lower = text.lower()
    markers = (
        "it is important to note", "it should be noted", "in conclusion", "moreover", "furthermore", "comprehensive analysis",
        "multifaceted", "delves into", "plays a crucial role", "realm of", "significant impact",
        "важно отметить", "следует отметить", "необходимо подчеркнуть", "таким образом", "в заключение", "данная работа",
        "актуальность данной", "в современном мире", "комплексный подход", "многогранный", "играет важную роль",
        "варто зазначити", "слід зазначити", "слід підкреслити", "таким чином", "у сучасному світі", "актуальність цієї",
        "комплексний підхід", "відіграє важливу роль",
    )
    return sum(lower.count(marker) for marker in markers)


def _source_marker_count(text: str) -> int:
    return len(re.findall(r"https?://|www\.|\[[0-9]{1,3}\]|\((?:[^()]{2,40}),\s*(?:19|20)\d{2}[a-z]?\)|\bdoi\s*:|\bisbn\s*:", text, flags=re.IGNORECASE))


def _source_claim_count(text: str) -> int:
    lower = text.lower()
    markers = (
        "according to", "research shows", "studies show", "statistics", "data shows", "survey",
        "согласно", "по данным", "исследования показывают", "статистика", "опрос", "учёные", "ученые",
        "згідно з", "за даними", "дослідження показують", "статистика", "опитування",
    )
    return sum(lower.count(marker) for marker in markers)


def _heading_count(text: str) -> int:
    count = 0
    for line in text.splitlines():
        clean = line.strip()
        words = _originality_words(clean)
        alpha_count = len(re.findall(r"[A-Za-zА-Яа-яЁёІіЇїЄєҐґ]", clean))
        if (
            8 <= len(clean) <= 96
            and 2 <= len(words) <= 12
            and alpha_count >= 6
            and not re.search(r"[.!?…;:]$", clean)
            and not re.fullmatch(r"[\d\s./№()_-]+", clean)
        ):
            count += 1
    return count


def _originality_document_kind(text: str) -> str:
    lower = text.lower()
    legal_hits = sum(lower.count(marker) for marker in ("договір", "договору", "сторін", "замовник", "розробник", "виконавець", "форс-мажор", "рнокпп", "підпис", "відповідальність сторін"))
    academic_hits = sum(lower.count(marker) for marker in ("курсова", "диплом", "вступ", "висновки", "список використаних джерел", "актуальність", "мета дослідження", "об'єкт дослідження", "предмет дослідження"))
    if legal_hits >= 4 and legal_hits >= academic_hits:
        return "legal"
    if academic_hits >= 3:
        return "academic"
    return "general"


def _source_support_score(word_count: int, source_count: int, source_claim_count: int, document_kind: str = "general") -> int:
    if document_kind == "legal":
        return _clamp_score(86 + min(8, source_count * 2) - max(0, source_claim_count - source_count) * 2)
    if word_count < 160:
        return 72 if source_count == 0 else 88
    if source_count:
        return _clamp_score(74 + min(22, source_count * 4) - max(0, source_claim_count - source_count) * 2)
    if source_claim_count:
        return 46
    return 64 if word_count > 700 else 76


def _ai_risk_score(*, word_count: int, sentence_count: int, sentence_stdev: float, avg_sentence: float, marker_hits: int, lexical_diversity: float, long_ratio: float, source_count: int) -> int:
    if sentence_count <= 1:
        return 22
    uniformity = max(0.0, min(1.0, 1 - (sentence_stdev / max(avg_sentence, 1))))
    marker_pressure = min(1.0, marker_hits / max(1.0, sentence_count / 7))
    diversity_pressure = 0.0
    if word_count > 180:
        if lexical_diversity < 0.36:
            diversity_pressure = min(1.0, (0.36 - lexical_diversity) * 5)
        elif 0.42 <= lexical_diversity <= 0.58:
            diversity_pressure = 0.28
    risk = 14 + uniformity * 26 + marker_pressure * 34 + diversity_pressure * 18 + long_ratio * 9 + (5 if source_count == 0 and word_count > 450 else 0)
    if word_count < 140:
        risk -= 10
    return _clamp_score(risk)


def _originality_segments(sentences: list[dict[str, object]], repeated_norms: set[str], source_count: int, language: str, document_kind: str = "general") -> list[dict[str, object]]:
    segments: list[dict[str, object]] = []
    for item in sentences[:ORIGINALITY_VISIBLE_SEGMENTS]:
        text = str(item["text"])
        lower = text.lower()
        words = item["words"] if isinstance(item.get("words"), list) else []
        issue_keys: list[str] = []
        if item.get("norm") in repeated_norms:
            issue_keys.append("issue_repetition")
        if _ai_marker_hits(text):
            issue_keys.append("issue_ai_pattern")
        if len(words) >= 38:
            issue_keys.append("issue_long_sentence")
        if document_kind != "legal" and source_count == 0 and _source_claim_count(text):
            issue_keys.append("issue_sources")
        if any(marker in lower for marker in ("данная работа", "в современном мире", "it is important to note", "in conclusion", "таким образом", "слід зазначити")):
            issue_keys.append("issue_boilerplate")
        issue_keys = list(dict.fromkeys(issue_keys))
        issues = [translate(key, language) for key in issue_keys]
        severity = "none"
        if issues:
            severity = "high" if "issue_repetition" in issue_keys or len(issue_keys) >= 3 else "medium"
        segments.append({"text": text, "severity": severity, "issues": issues})
    if len(sentences) > ORIGINALITY_VISIBLE_SEGMENTS:
        segments.append({"text": "...", "severity": "none", "issues": []})
    return segments


def _originality_highlight_summary(segments: list[dict[str, object]], language: str) -> list[dict[str, object]]:
    counts: Counter[str] = Counter()
    for segment in segments:
        for issue in segment.get("issues", []):
            counts[str(issue)] += 1
    return [{"label": label, "count": count, "tone": "bad" if count >= 4 else "warn"} for label, count in counts.most_common(6)] or [{"label": translate("originality_no_issues", language), "count": 0, "tone": "good"}]


def _originality_detail_context(request: HttpRequest, job: dict, owner_id: int | None, guest_key: str, language: str) -> dict[str, object]:
    if str(job.get("kind") or "") != "originality":
        return {"is_originality": False}
    record = _job_record_for_workspace(str(job.get("id") or ""), owner_id, guest_key)
    analysis = _load_originality_analysis(record) if record else {}
    if record and analysis:
        _refresh_originality_html_report(record, analysis, language)
    html_output = _originality_payload_output(job, "html")
    json_output = _originality_payload_output(job, "json")
    display_outputs = [html_output] if html_output else []
    next_url = request.get_full_path()
    overall = analysis.get("overall", {}) if isinstance(analysis.get("overall"), dict) else {}
    score = max(0, min(100, int(overall.get("score") or 0)))
    params = _job_record_params(record) if record else {}
    originality_params = params.get("originality") if isinstance(params.get("originality"), dict) else {}
    share_token = str(originality_params.get("share_token") or "")
    return {
        "is_originality": True,
        "display_outputs": display_outputs,
        "originality_report": analysis,
        "originality_overall": overall,
        "originality_score_degrees": round(score * 3.6, 1),
        "originality_source": analysis.get("source", {}) if isinstance(analysis.get("source"), dict) else {},
        "originality_metrics": analysis.get("metrics", []) if isinstance(analysis.get("metrics"), list) else [],
        "originality_highlights": analysis.get("highlights", []) if isinstance(analysis.get("highlights"), list) else [],
        "originality_report_output": html_output,
        "originality_json_output": json_output,
        "originality_report_url": html_output.get("preview_url", "") if html_output else "",
        "originality_json_url": json_output.get("preview_url", "") if json_output else "",
        "originality_send_url": reverse("studio:send_originality_report", args=[job.get("id")]),
        "originality_result_link": request.build_absolute_uri(reverse("studio:job_detail", args=[job.get("id")])),
        "originality_share_link": request.build_absolute_uri(reverse("studio:originality_shared_report", args=[share_token])) if share_token else "",
        "originality_check": analysis.get("check", {}) if isinstance(analysis.get("check"), dict) else {},
        "originality_recent_reports": _recent_originality_reports(owner_id, guest_key, exclude_job_id=str(job.get("id") or "")),
        "originality_email_default": request.user.email if request.user.is_authenticated else "",
        "originality_account_url": f"{reverse('studio:login')}?{urlencode({'next': next_url})}",
    }


def _recent_originality_reports(owner_id: int | None, guest_key: str, exclude_job_id: str = "") -> list[dict[str, object]]:
    queryset = JobRecord.objects.filter(kind="originality").order_by("-created_at")
    if owner_id is not None:
        queryset = queryset.filter(owner_id=owner_id)
    else:
        queryset = queryset.filter(owner__isnull=True, guest_key=guest_key)
    reports = []
    for record in queryset[:8]:
        if record.job_id == exclude_job_id:
            continue
        params = _job_record_params(record)
        originality = params.get("originality") if isinstance(params.get("originality"), dict) else {}
        check = originality.get("check") if isinstance(originality.get("check"), dict) else {}
        reports.append(
            {
                "title": record.title,
                "url": reverse("studio:job_detail", args=[record.job_id]),
                "created_at": record.created_at,
                "mode_label": check.get("mode_label") or "Local",
                "price": check.get("price_cherryx") or 0,
                "status": record.status,
            }
        )
        if len(reports) >= 4:
            break
    return reports


def _load_originality_analysis(record: JobRecord | None) -> dict[str, object]:
    if not record:
        return {}
    for output in record.outputs.all():
        path = Path(output.path)
        if path.suffix.lower() != ".json" and "json" not in str(output.media_type).lower():
            continue
        try:
            payload = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            continue
        analysis = payload.get("analysis") if isinstance(payload, dict) else None
        if isinstance(analysis, dict):
            return analysis
    return {}


def _originality_record_by_share_token(token: str) -> JobRecord | None:
    clean = str(token or "").strip()
    if len(clean) < 12:
        return None
    queryset = JobRecord.objects.filter(kind="originality").prefetch_related("outputs").order_by("-created_at")
    for record in queryset[:500]:
        params = _job_record_params(record)
        originality = params.get("originality") if isinstance(params.get("originality"), dict) else {}
        if originality.get("share_public", True) and originality.get("share_token") == clean:
            return record
    return None


def _originality_payload_output(job: dict, kind: str) -> dict[str, object]:
    for output in job.get("outputs", []):
        name = str(output.get("name") or "").lower()
        media_type = str(output.get("media_type") or "").lower()
        if kind == "html" and (name.endswith(".html") or media_type.startswith("text/html")):
            return output
        if kind == "json" and (name.endswith(".json") or "json" in media_type):
            return output
    return {}


def _originality_html_output(record: JobRecord) -> tuple[int, JobOutputRecord] | None:
    for index, output in enumerate(record.outputs.all()):
        name = Path(output.path).name.lower()
        media_type = str(output.media_type or "").lower()
        if name.endswith(".html") or media_type.startswith("text/html"):
            return index, output
    return None


def _refresh_originality_html_report(
    record: JobRecord,
    analysis: dict[str, object],
    language: str,
    output_pair: tuple[int, JobOutputRecord] | None = None,
) -> None:
    output_pair = output_pair or _originality_html_output(record)
    if not output_pair or not analysis:
        return
    _, output = output_pair
    path = Path(output.path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(_render_originality_report_html_v2(analysis, language), encoding="utf-8")
    size = path.stat().st_size
    if output.size != size or not str(output.media_type or "").lower().startswith("text/html"):
        output.size = size
        output.media_type = "text/html; charset=utf-8"
        output.save(update_fields=["size", "media_type"])
    try:
        record.total_output_size = sum(int(item.size or 0) for item in record.outputs.all())
        record.save(update_fields=["total_output_size", "updated_at"])
    except Exception:
        pass


def _send_originality_email(request: HttpRequest, record: JobRecord, output_pair: tuple[int, JobOutputRecord], email: str, language: str) -> None:
    index, output = output_pair
    path = Path(output.path)
    if not path.exists():
        raise ValueError(translate("file_missing", language))
    detail_url = request.build_absolute_uri(reverse("studio:job_detail", args=[record.job_id]))
    preview_url = request.build_absolute_uri(reverse("studio:preview_output", args=[record.job_id, index]))
    subject = f"CherryX · {translate('originality_title', language)}"
    text_body = (
        f"{translate('originality_email_intro', language)}\n\n"
        f"{record.message}\n\n"
        f"{translate('originality_full_report', language)}: {preview_url}\n"
        f"{translate('copy_result_link', language)}: {detail_url}\n"
    )
    html_body = (
        f"<p>{html.escape(translate('originality_email_intro', language))}</p>"
        f"<p><strong>{html.escape(record.message)}</strong></p>"
        f'<p><a href="{html.escape(preview_url)}">{html.escape(translate("originality_full_report", language))}</a></p>'
        f'<p><a href="{html.escape(detail_url)}">{html.escape(translate("copy_result_link", language))}</a></p>'
    )
    message = EmailMultiAlternatives(subject=subject, body=text_body, from_email=None, to=[email])
    message.attach_alternative(html_body, "text/html")
    message.attach(path.name, path.read_bytes(), "text/html")
    message.send(fail_silently=False)


def _create_originality_record(analysis: dict[str, object], text: str, owner_id: int | None, guest_key: str, language: str, source_name: str) -> JobRecord:
    job_id = uuid.uuid4().hex[:16]
    target_dir = settings.storage_dir / "originality_reports" / job_id
    target_dir.mkdir(parents=True, exist_ok=True)
    json_path = target_dir / "originality_report.json"
    html_path = target_dir / "originality_report.html"
    json_payload = {"analysis": analysis, "text_excerpt": text[:8000]}
    json_path.write_text(json.dumps(json_payload, ensure_ascii=False, indent=2), encoding="utf-8")
    html_path.write_text(_render_originality_report_html_v2(analysis, language), encoding="utf-8")

    source = analysis.get("source") if isinstance(analysis.get("source"), dict) else {}
    overall = analysis.get("overall") if isinstance(analysis.get("overall"), dict) else {}
    metrics = analysis.get("metrics") if isinstance(analysis.get("metrics"), list) else []
    ai_metric = next((metric for metric in metrics if isinstance(metric, dict) and metric.get("key") == "ai_risk"), {})
    score = int(overall.get("score") or 0)
    ai_score = int(ai_metric.get("score") or 0) if isinstance(ai_metric, dict) else 0
    words = int(source.get("words") or 0)
    title_name = Path(source_name).name if source_name else translate("originality_title", language)
    record = JobRecord.objects.create(
        owner_id=owner_id,
        guest_key=guest_key if owner_id is None else "",
        job_id=job_id,
        kind="originality",
        title=f"{translate('originality_check', language)}: {title_name}"[:240],
        status="completed",
        progress=100,
        message=f"{translate('originality_score', language)}: {score}/100 · {translate('metric_ai_risk', language)}: {ai_score}/100 · {words} {translate('originality_words', language)}",
        params_json="",
        output_count=2,
        total_output_size=html_path.stat().st_size + json_path.stat().st_size,
        primary_output_type="text",
    )
    record.params_json = json.dumps(
        {
            "originality": {
                "share_token": secrets.token_urlsafe(18),
                "share_public": True,
                "share_summary_only": False,
                "created_language": language,
                "check": analysis.get("check") if isinstance(analysis.get("check"), dict) else {},
                "source_name": source_name,
            }
        },
        ensure_ascii=False,
        default=str,
    )
    record.save(update_fields=["params_json", "updated_at"])
    JobEventRecord.objects.create(job=record, status="queued", progress=12, message=_originality_process_message("received", language))
    JobEventRecord.objects.create(job=record, status="running", progress=44, message=_originality_process_message("extracted", language, words))
    JobEventRecord.objects.create(job=record, status="running", progress=78, message=_originality_process_message("metrics", language))
    JobEventRecord.objects.create(job=record, status="completed", progress=100, message=_originality_process_message("ready", language, score))
    JobOutputRecord.objects.create(job=record, label=_originality_output_label("html", language), path=str(html_path), media_type="text/html; charset=utf-8", size=html_path.stat().st_size)
    JobOutputRecord.objects.create(job=record, label=_originality_output_label("json", language), path=str(json_path), media_type="application/json; charset=utf-8", size=json_path.stat().st_size)
    record._prefetched_objects_cache = {"outputs": list(record.outputs.all())}
    return record


def _render_originality_report_html_v2(analysis: dict[str, object], language: str) -> str:
    overall = analysis.get("overall") if isinstance(analysis.get("overall"), dict) else {}
    source = analysis.get("source") if isinstance(analysis.get("source"), dict) else {}
    metrics = analysis.get("metrics") if isinstance(analysis.get("metrics"), list) else []
    highlights = analysis.get("highlights") if isinstance(analysis.get("highlights"), list) else []
    segments = analysis.get("segments") if isinstance(analysis.get("segments"), list) else []
    score = max(0, min(100, int(overall.get("score") or 0)))
    tone = str(overall.get("tone") or "none")
    if tone not in {"good", "warn", "bad"}:
        tone = "none"
    problem_segments = [
        segment
        for segment in segments
        if isinstance(segment, dict)
        and str(segment.get("severity") or "none") in {"medium", "high"}
        and isinstance(segment.get("issues"), list)
        and segment.get("issues")
    ]
    metric_markup = "\n".join(
        _render_originality_report_metric(metric)
        for metric in metrics
        if isinstance(metric, dict)
    )
    highlight_markup = "\n".join(
        _render_originality_report_highlight(item)
        for item in highlights
        if isinstance(item, dict)
    )
    segment_markup = "\n".join(_render_originality_report_problem_segment(segment) for segment in problem_segments)
    if not segment_markup:
        segment_markup = (
            f'<div class="map-empty"><b>{html.escape(translate("originality_no_issues", language))}</b>'
            f'<span>{html.escape(translate("originality_report_ready", language))}</span></div>'
        )
    title = str(source.get("name") or translate("originality_title", language))
    kind = str(source.get("kind_label") or "")
    meta_items = [
        title,
        kind,
        f'{int(source.get("words") or 0)} {translate("originality_words", language)}',
        f'{int(source.get("sentences") or 0)} {translate("originality_sentences", language)}',
    ]
    meta_markup = "\n".join(f"<span>{html.escape(item)}</span>" for item in meta_items if item)
    score_label = html.escape(str(overall.get("label") or translate("originality_score", language)))
    page_title = html.escape(translate("originality_title", language))
    metrics_title = html.escape(translate("originality_metrics", language))
    highlights_title = html.escape(translate("originality_highlights", language))
    score_degrees = score * 3.6
    return f"""<!doctype html>
<html lang="{html.escape(clean_language(language))}">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{page_title}</title>
  <style>
    :root {{ color-scheme: light; --blue:#2563eb; --ink:#101827; --muted:#64748b; --line:#d8e3f0; --surface:#ffffff; --good:#10b981; --warn:#f59e0b; --bad:#ef4444; }}
    * {{ box-sizing:border-box; }}
    body {{ margin:0; padding:24px; background:linear-gradient(180deg,#eef6ff 0,#f8fbff 42%,#ffffff 100%); color:var(--ink); font-family:Inter,Segoe UI,Arial,sans-serif; }}
    main {{ max-width:1120px; margin:0 auto; display:grid; gap:16px; }}
    .hero {{ position:relative; overflow:hidden; display:grid; grid-template-columns:auto minmax(0,1fr); gap:20px; align-items:center; padding:22px; border:1px solid color-mix(in srgb,var(--blue) 18%,var(--line)); border-radius:18px; background:linear-gradient(135deg,rgba(255,255,255,.98),rgba(255,255,255,.78)); box-shadow:0 24px 70px rgba(15,23,42,.12); }}
    .hero::before {{ content:""; position:absolute; inset:0; background:radial-gradient(circle at 8% 0%,color-mix(in srgb,var(--blue) 14%,transparent),transparent 34%),repeating-linear-gradient(135deg,color-mix(in srgb,var(--blue) 8%,transparent) 0 1px,transparent 1px 20px); pointer-events:none; }}
    .hero > * {{ position:relative; }}
    .score {{ --score-size:122px; --score-ring:10px; --score-deg:{score_degrees}deg; position:relative; isolation:isolate; display:flex; flex-direction:column; align-items:center; justify-content:center; gap:1px; width:var(--score-size); height:var(--score-size); border-radius:50%; color:var(--blue); background:conic-gradient(from -90deg,currentColor 0 var(--score-deg),color-mix(in srgb,currentColor 11%,#e8eef8) var(--score-deg) 360deg); box-shadow:0 18px 38px color-mix(in srgb,currentColor 16%,transparent),inset 0 0 0 1px color-mix(in srgb,currentColor 22%,transparent); animation:score-breathe 3.8s ease-in-out infinite; }}
    .score::before {{ content:""; position:absolute; inset:calc(var(--score-ring) + 1px); z-index:0; border-radius:inherit; background:radial-gradient(circle at 32% 22%,rgba(255,255,255,.98),transparent 34%),linear-gradient(180deg,#fff,color-mix(in srgb,currentColor 8%,#f7fbff)); box-shadow:inset 0 1px 0 rgba(255,255,255,.96),inset 0 -12px 26px color-mix(in srgb,currentColor 7%,transparent),0 0 0 1px color-mix(in srgb,currentColor 12%,transparent); }}
    .score::after {{ content:""; position:absolute; inset:-3px; z-index:1; border-radius:inherit; background:conic-gradient(from 0deg,transparent 0 68%,color-mix(in srgb,currentColor 24%,transparent) 74%,transparent 82%); mask:radial-gradient(circle,transparent calc(50% - var(--score-ring) - 4px),#000 calc(50% - var(--score-ring) - 3px),#000 calc(50% + 3px),transparent calc(50% + 5px)); -webkit-mask:radial-gradient(circle,transparent calc(50% - var(--score-ring) - 4px),#000 calc(50% - var(--score-ring) - 3px),#000 calc(50% + 3px),transparent calc(50% + 5px)); opacity:.48; animation:score-orbit 5.8s linear infinite; pointer-events:none; }}
    .score b,.score span {{ position:relative; z-index:2; }}
    .score b {{ font-size:40px; font-variant-numeric:tabular-nums; letter-spacing:0; line-height:.88; text-shadow:0 8px 18px color-mix(in srgb,currentColor 18%,transparent); }}
    .score span {{ color:color-mix(in srgb,var(--ink) 72%,var(--muted)); font-size:13px; font-weight:900; line-height:1; }}
    .score.is-good {{ color:var(--good); }} .score.is-warn {{ color:var(--warn); }} .score.is-bad {{ color:var(--bad); }}
    h1 {{ margin:0 0 8px; font-size:clamp(26px,4vw,42px); line-height:1; letter-spacing:0; }}
    p {{ margin:0; color:var(--muted); line-height:1.45; }}
    .meta {{ display:flex; flex-wrap:wrap; gap:8px; margin-top:12px; }}
    .meta span {{ display:inline-flex; align-items:center; min-height:30px; padding:6px 10px; border:1px solid color-mix(in srgb,var(--blue) 15%,var(--line)); border-radius:999px; background:rgba(255,255,255,.76); color:#334155; font-size:12px; font-weight:800; }}
    .metrics {{ display:grid; grid-template-columns:repeat(3,minmax(0,1fr)); gap:12px; }}
    .metric {{ display:grid; grid-template-columns:minmax(0,1fr) auto; gap:10px 12px; align-items:center; padding:15px; border:1px solid var(--line); border-radius:14px; background:var(--surface); box-shadow:0 12px 34px rgba(15,23,42,.06); }}
    .metric span {{ display:block; color:var(--ink); font-weight:900; }}
    .metric b {{ color:var(--blue); font-size:30px; line-height:1; }}
    .metric small {{ display:block; margin-top:4px; color:var(--muted); font-size:12px; line-height:1.3; }}
    .metric i {{ grid-column:1 / -1; height:8px; overflow:hidden; border-radius:999px; background:#e8edf6; }}
    .metric em {{ display:block; height:100%; border-radius:inherit; background:var(--blue); }}
    .metric.is-good em {{ background:var(--good); }} .metric.is-good b {{ color:var(--good); }}
    .metric.is-warn em {{ background:var(--warn); }} .metric.is-warn b {{ color:var(--warn); }}
    .metric.is-bad em {{ background:var(--bad); }} .metric.is-bad b {{ color:var(--bad); }}
    .grid {{ display:grid; grid-template-columns:minmax(250px,.34fr) minmax(0,1fr); gap:14px; }}
    section {{ padding:17px; border:1px solid var(--line); border-radius:16px; background:rgba(255,255,255,.92); box-shadow:0 14px 40px rgba(15,23,42,.06); }}
    h2 {{ margin:0 0 12px; font-size:18px; }}
    ul {{ display:grid; gap:8px; list-style:none; margin:0; padding:0; }}
    li {{ display:flex; align-items:center; justify-content:space-between; gap:10px; min-height:42px; padding:9px 0; border-top:1px solid #edf1f7; font-weight:800; }}
    li:first-child {{ border-top:0; }}
    li span {{ color:#334155; }}
    li b {{ min-width:30px; height:28px; display:grid; place-items:center; border-radius:999px; background:var(--blue); color:#fff; font-size:12px; }}
    li.is-good b {{ background:var(--good); }} li.is-warn b {{ background:var(--warn); }} li.is-bad b {{ background:var(--bad); }}
    .map {{ max-height:640px; overflow:auto; padding-right:6px; color:#334155; font-size:14px; line-height:1.75; scrollbar-width:thin; scrollbar-color:color-mix(in srgb,var(--blue) 42%,#9db7ff) transparent; }}
    .map::-webkit-scrollbar {{ width:8px; }} .map::-webkit-scrollbar-track {{ background:transparent; }} .map::-webkit-scrollbar-thumb {{ border-radius:999px; background:color-mix(in srgb,var(--blue) 42%,#9db7ff); }}
    .map-empty {{ display:grid; gap:6px; min-height:160px; place-content:center; text-align:center; border:1px dashed color-mix(in srgb,var(--good) 36%,var(--line)); border-radius:14px; background:linear-gradient(180deg,color-mix(in srgb,var(--good) 8%,#fff),#fff); }}
    .map-empty b {{ color:#047857; font-size:16px; }}
    .map-empty span {{ color:var(--muted); font-size:13px; font-weight:750; }}
    mark {{ display:inline; margin:0 1px; padding:2px 4px; border-radius:6px; color:inherit; }}
    mark.medium {{ background:rgba(245,158,11,.22); }} mark.high {{ background:rgba(239,68,68,.18); }}
    mark small {{ display:inline-flex; margin-left:4px; padding:2px 5px; border-radius:999px; background:rgba(15,23,42,.78); color:#fff; font-size:10px; font-weight:800; }}
    @keyframes score-breathe {{ 0%,100% {{ filter:saturate(1); }} 50% {{ filter:saturate(1.08) brightness(1.02); }} }}
    @keyframes score-orbit {{ to {{ transform:rotate(360deg); }} }}
    @media (prefers-reduced-motion:reduce) {{ .score,.score::after {{ animation:none; }} }}
    @media (max-width:760px) {{ body {{ padding:12px; }} .hero,.grid,.metrics {{ grid-template-columns:1fr; }} .score {{ --score-size:104px; --score-ring:9px; }} }}
  </style>
</head>
<body>
  <main>
    <header class="hero">
      <div class="score is-{html.escape(tone)}"><b>{score}</b><span>/100</span></div>
      <div>
        <h1>{page_title}</h1>
        <p>{score_label}</p>
        <div class="meta">{meta_markup}</div>
      </div>
    </header>
    <div class="metrics">{metric_markup}</div>
    <div class="grid">
      <section><h2>{metrics_title}</h2><ul>{highlight_markup}</ul></section>
      <section><h2>{highlights_title}</h2><div class="map">{segment_markup}</div></section>
    </div>
  </main>
</body>
</html>"""


def _render_originality_report_metric(metric: dict[str, object]) -> str:
    score = max(0, min(100, int(metric.get("score") or 0)))
    tone = str(metric.get("tone") or "none")
    if tone not in {"good", "warn", "bad"}:
        tone = "none"
    return (
        f'<article class="metric is-{html.escape(tone)}">'
        f'<div><span>{html.escape(str(metric.get("label") or ""))}</span>'
        f'<small>{html.escape(str(metric.get("detail") or ""))}</small></div>'
        f"<b>{score}</b><i><em style=\"width:{score}%\"></em></i></article>"
    )


def _render_originality_report_highlight(item: dict[str, object]) -> str:
    tone = str(item.get("tone") or "none")
    if tone not in {"good", "warn", "bad"}:
        tone = "none"
    return (
        f'<li class="is-{html.escape(tone)}">'
        f'<span>{html.escape(str(item.get("label") or ""))}</span>'
        f'<b>{int(item.get("count") or 0)}</b></li>'
    )


def _render_originality_report_problem_segment(segment: dict[str, object]) -> str:
    text = html.escape(str(segment.get("text") or ""))
    severity = str(segment.get("severity") or "none")
    if severity not in {"medium", "high"}:
        severity = "medium"
    issues = segment.get("issues") if isinstance(segment.get("issues"), list) else []
    issue_text = " / ".join(str(issue) for issue in issues if issue)
    if not text or not issue_text:
        return ""
    return f'<mark class="{html.escape(severity)}">{text}<small>{html.escape(issue_text)}</small></mark>'


def _render_originality_report_html(analysis: dict[str, object], language: str) -> str:
    overall = analysis.get("overall") if isinstance(analysis.get("overall"), dict) else {}
    source = analysis.get("source") if isinstance(analysis.get("source"), dict) else {}
    metrics = analysis.get("metrics") if isinstance(analysis.get("metrics"), list) else []
    highlights = analysis.get("highlights") if isinstance(analysis.get("highlights"), list) else []
    segments = analysis.get("segments") if isinstance(analysis.get("segments"), list) else []
    score = int(overall.get("score") or 0)
    metric_markup = "\n".join(
        f'<article class="metric {html.escape(str(metric.get("tone") or ""))}"><span>{html.escape(str(metric.get("label") or ""))}</span><b>{int(metric.get("score") or 0)}</b><small>{html.escape(str(metric.get("detail") or ""))}</small><i><em style="width:{max(0, min(100, int(metric.get("score") or 0)))}%"></em></i></article>'
        for metric in metrics
        if isinstance(metric, dict)
    )
    highlight_markup = "\n".join(
        f'<li><span>{html.escape(str(item.get("label") or ""))}</span><b>{int(item.get("count") or 0)}</b></li>'
        for item in highlights
        if isinstance(item, dict)
    )
    segment_markup = "\n".join(_render_originality_report_segment(segment) for segment in segments if isinstance(segment, dict))
    title = html.escape(str(source.get("name") or translate("originality_title", language)))
    kind = html.escape(str(source.get("kind_label") or ""))
    return f"""<!doctype html>
<html lang="{html.escape(clean_language(language))}">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(translate("originality_title", language))}</title>
  <style>
    :root {{ color-scheme: light; --blue:#2563eb; --ink:#111827; --muted:#64748b; --line:#d7e1ec; }}
    * {{ box-sizing:border-box; }}
    body {{ margin:0; padding:24px; background:#f5f8fc; color:var(--ink); font-family:Inter,Segoe UI,Arial,sans-serif; }}
    main {{ max-width:1100px; margin:0 auto; display:grid; gap:18px; }}
    header {{ display:grid; grid-template-columns:auto minmax(0,1fr); gap:18px; align-items:center; padding:18px; border:1px solid var(--line); border-radius:12px; background:#fff; box-shadow:0 18px 44px rgba(15,23,42,.08); }}
    .score {{ display:grid; place-items:center; width:104px; height:104px; border-radius:50%; color:var(--blue); background:radial-gradient(circle,color-mix(in srgb,var(--blue) 12%,transparent),transparent 62%),#fff; box-shadow:inset 0 0 0 10px color-mix(in srgb,var(--blue) 10%,transparent); border:1px solid color-mix(in srgb,var(--blue) 24%,var(--line)); }}
    .score b {{ font-size:34px; line-height:.9; }}
    .score span {{ font-size:13px; color:var(--muted); font-weight:800; }}
    h1 {{ margin:0 0 8px; font-size:clamp(24px,4vw,38px); line-height:1; }}
    p {{ margin:0; color:var(--muted); line-height:1.45; }}
    .metrics {{ display:grid; grid-template-columns:repeat(3,minmax(0,1fr)); gap:12px; }}
    .metric {{ display:grid; gap:7px; padding:14px; border:1px solid var(--line); border-radius:10px; background:#fff; }}
    .metric span {{ font-weight:850; }}
    .metric b {{ font-size:28px; color:var(--blue); }}
    .metric small {{ min-height:30px; color:var(--muted); font-size:12px; line-height:1.25; }}
    .metric i {{ height:7px; overflow:hidden; border-radius:999px; background:#e8edf6; }}
    .metric em {{ display:block; height:100%; border-radius:inherit; background:var(--blue); }}
    .metric.good em {{ background:#10b981; }} .metric.warn em {{ background:#f59e0b; }} .metric.bad em {{ background:#ef4444; }}
    .grid {{ display:grid; grid-template-columns:minmax(220px,.32fr) minmax(0,1fr); gap:14px; }}
    section {{ padding:16px; border:1px solid var(--line); border-radius:12px; background:#fff; }}
    h2 {{ margin:0 0 12px; font-size:17px; }}
    ul {{ display:grid; gap:9px; list-style:none; margin:0; padding:0; }}
    li {{ display:flex; align-items:center; justify-content:space-between; gap:10px; padding-bottom:9px; border-bottom:1px solid #edf1f7; font-weight:760; }}
    li b {{ min-width:30px; height:26px; display:grid; place-items:center; border-radius:999px; background:#2563eb; color:#fff; font-size:12px; }}
    .map {{ max-height:640px; overflow:auto; padding-right:6px; font-size:14px; line-height:1.75; }}
    mark {{ display:inline; margin:0 1px; padding:2px 4px; border-radius:6px; color:inherit; }}
    mark.medium {{ background:rgba(245,158,11,.22); }} mark.high {{ background:rgba(239,68,68,.18); }}
    mark small {{ display:inline-flex; margin-left:4px; padding:2px 5px; border-radius:999px; background:rgba(15,23,42,.78); color:#fff; font-size:10px; font-weight:800; }}
    @media (max-width:760px) {{ body {{ padding:12px; }} header,.grid,.metrics {{ grid-template-columns:1fr; }} }}
  </style>
</head>
<body>
  <main>
    <header>
      <div class="score"><b>{score}</b><span>/100</span></div>
      <div>
        <h1>{html.escape(translate("originality_title", language))}</h1>
        <p>{title} · {kind} · {int(source.get("words") or 0)} {html.escape(translate("originality_words", language))} · {int(source.get("sentences") or 0)} {html.escape(translate("originality_sentences", language))}</p>
      </div>
    </header>
    <div class="metrics">{metric_markup}</div>
    <div class="grid">
      <section><h2>{html.escape(translate("originality_metrics", language))}</h2><ul>{highlight_markup}</ul></section>
      <section><h2>{html.escape(translate("originality_highlights", language))}</h2><div class="map">{segment_markup}</div></section>
    </div>
  </main>
</body>
</html>"""


def _render_originality_report_segment(segment: dict[str, object]) -> str:
    text = html.escape(str(segment.get("text") or ""))
    severity = str(segment.get("severity") or "none")
    issues = segment.get("issues") if isinstance(segment.get("issues"), list) else []
    issue_text = " · ".join(str(issue) for issue in issues if issue)
    if severity in {"medium", "high"} and issue_text:
        return f'<mark class="{html.escape(severity)}">{text}<small>{html.escape(issue_text)}</small></mark>'
    return f"<span>{text}</span>"


def _document_kind_label(kind: str, language: str) -> str:
    values = {
        "legal": {"en": "Legal document", "ru": "Юридический документ", "uk": "Юридичний документ", "fr": "Document juridique", "de": "Rechtsdokument", "es": "Documento legal", "ka": "იურიდიული დოკუმენტი", "hy": "Իրավական փաստաթուղթ", "it": "Documento legale"},
        "academic": {"en": "Academic work", "ru": "Академическая работа", "uk": "Академічна робота", "fr": "Travail académique", "de": "Akademische Arbeit", "es": "Trabajo académico", "ka": "აკადემიური ნაშრომი", "hy": "Ակադեմիական աշխատանք", "it": "Lavoro accademico"},
        "general": {"en": "Text document", "ru": "Текстовый документ", "uk": "Текстовий документ", "fr": "Document texte", "de": "Textdokument", "es": "Documento de texto", "ka": "ტექსტური დოკუმენტი", "hy": "Տեքստային փաստաթուղթ", "it": "Documento di testo"},
    }
    language = clean_language(language)
    item = values.get(kind, values["general"])
    return item.get(language) or item["en"]


def _originality_process_message(key: str, language: str, value: int | None = None) -> str:
    values = {
        "received": {"en": "Document received", "ru": "Документ получен", "uk": "Документ отримано", "fr": "Document reçu", "de": "Dokument empfangen", "es": "Documento recibido", "ka": "დოკუმენტი მიღებულია", "hy": "Փաստաթուղթը ստացվել է", "it": "Documento ricevuto"},
        "extracted": {"en": f"Text extracted: {value or 0} words", "ru": f"Текст извлечён: {value or 0} слов", "uk": f"Текст витягнуто: {value or 0} слів", "fr": f"Texte extrait : {value or 0} mots", "de": f"Text extrahiert: {value or 0} Wörter", "es": f"Texto extraído: {value or 0} palabras", "ka": f"ტექსტი ამოღებულია: {value or 0} სიტყვა", "hy": f"Տեքստը հանվել է՝ {value or 0} բառ", "it": f"Testo estratto: {value or 0} parole"},
        "metrics": {"en": "Metrics calculated and fragments marked", "ru": "Метрики рассчитаны, фрагменты отмечены", "uk": "Метрики розраховано, фрагменти позначено", "fr": "Mesures calculées et fragments marqués", "de": "Metriken berechnet und Stellen markiert", "es": "Métricas calculadas y fragmentos marcados", "ka": "მეტრიკები დათვლილია და ფრაგმენტები მონიშნულია", "hy": "Ցուցիչները հաշվարկվել են, հատվածները նշվել են", "it": "Metriche calcolate e frammenti segnati"},
        "ready": {"en": f"Report ready: {value or 0}/100", "ru": f"Отчёт готов: {value or 0}/100", "uk": f"Звіт готовий: {value or 0}/100", "fr": f"Rapport prêt : {value or 0}/100", "de": f"Bericht fertig: {value or 0}/100", "es": f"Informe listo: {value or 0}/100", "ka": f"ანგარიში მზადაა: {value or 0}/100", "hy": f"Հաշվետվությունը պատրաստ է՝ {value or 0}/100", "it": f"Report pronto: {value or 0}/100"},
    }
    item = values.get(key, {})
    language = clean_language(language)
    return item.get(language) or item.get("en") or key


def _originality_output_label(key: str, language: str) -> str:
    values = {
        "html": {"en": "HTML originality report", "ru": "HTML-отчёт уникальности", "uk": "HTML-звіт унікальності", "fr": "Rapport originalité HTML", "de": "HTML-Originalitätsbericht", "es": "Informe HTML de originalidad", "ka": "HTML ორიგინალურობის ანგარიში", "hy": "HTML ինքնատիպության հաշվետվություն", "it": "Report originalità HTML"},
        "json": {"en": "JSON metrics", "ru": "JSON-метрики", "uk": "JSON-метрики", "fr": "Métriques JSON", "de": "JSON-Metriken", "es": "Métricas JSON", "ka": "JSON მეტრიკები", "hy": "JSON ցուցիչներ", "it": "Metriche JSON"},
    }
    item = values.get(key, {})
    language = clean_language(language)
    return item.get(language) or item.get("en") or key


def _originality_metric(key: str, label: str, score: float, tone: str, detail: str) -> dict[str, object]:
    return {"key": key, "label": label, "score": _clamp_score(score), "tone": tone, "detail": detail}


def _metric_detail(key: str, language: str, first: float | int, second: float | int) -> str:
    if key == "uniqueness":
        return f"{_originality_phrase('repeated_shingles', language)}: {first}% · {_originality_phrase('duplicate_sentences', language)}: {second}"
    if key == "ai_risk":
        return f"{_originality_phrase('ai_markers', language)}: {first} · {_originality_phrase('length_spread', language)}: {second}"
    if key == "readability":
        return f"{_originality_phrase('avg_sentence', language)}: {first} · {_originality_phrase('long_sentences', language)}: {second}%"
    if key == "repetition":
        return f"{_originality_phrase('repeated_shingles_count', language)}: {first} · {_originality_phrase('sentence_repeats', language)}: {second}%"
    if key == "sources":
        return f"{_originality_phrase('links_citations', language)}: {first} · {_originality_phrase('claims', language)}: {second}"
    if key == "structure":
        return f"{_originality_phrase('paragraphs', language)}: {first} · {_originality_phrase('headings', language)}: {second}"
    return ""


def _originality_phrase(key: str, language: str) -> str:
    phrases = {
        "repeated_shingles": {"en": "Repeated shingles", "ru": "Повторы шинглов", "uk": "Повтори шинглів", "fr": "Shingles répétés", "de": "Wiederholte Shingles", "es": "Shingles repetidos", "ka": "განმეორებული შინგლები", "hy": "Կրկնվող շինգլներ", "it": "Shingle ripetuti"},
        "duplicate_sentences": {"en": "duplicate sentences", "ru": "дубли фраз", "uk": "дублі фраз", "fr": "phrases doublées", "de": "doppelte Sätze", "es": "frases duplicadas", "ka": "დუბლირებული წინადადებები", "hy": "կրկնվող նախադասություններ", "it": "frasi duplicate"},
        "ai_markers": {"en": "AI markers", "ru": "AI-маркеры", "uk": "AI-маркери", "fr": "marqueurs IA", "de": "KI-Marker", "es": "marcadores IA", "ka": "AI ნიშნები", "hy": "AI նշաններ", "it": "marker IA"},
        "length_spread": {"en": "length spread", "ru": "разброс длины", "uk": "розкид довжини", "fr": "écart longueur", "de": "Längenstreuung", "es": "variación longitud", "ka": "სიგრძის გადახრა", "hy": "երկարության շեղում", "it": "scarto lunghezza"},
        "avg_sentence": {"en": "Avg sentence", "ru": "Средняя фраза", "uk": "Середня фраза", "fr": "Phrase moyenne", "de": "Durchschnittssatz", "es": "Frase media", "ka": "საშუალო წინადადება", "hy": "Միջին նախադասություն", "it": "Frase media"},
        "long_sentences": {"en": "long", "ru": "длинные", "uk": "довгі", "fr": "longues", "de": "lange", "es": "largas", "ka": "გრძელი", "hy": "երկար", "it": "lunghe"},
        "repeated_shingles_count": {"en": "Repeated shingles", "ru": "Повторных шинглов", "uk": "Повторних шинглів", "fr": "Shingles répétés", "de": "Wiederholte Shingles", "es": "Shingles repetidos", "ka": "განმეორებული შინგლები", "hy": "Կրկնվող շինգլներ", "it": "Shingle ripetuti"},
        "sentence_repeats": {"en": "sentence repeats", "ru": "повтор фраз", "uk": "повтор фраз", "fr": "répétitions phrases", "de": "Satzwiederholungen", "es": "repetición frases", "ka": "წინადადებების გამეორება", "hy": "նախադասության կրկնություն", "it": "ripetizioni frasi"},
        "links_citations": {"en": "Links/citations", "ru": "Ссылки/цитаты", "uk": "Посилання/цитати", "fr": "Liens/citations", "de": "Links/Zitate", "es": "Enlaces/citas", "ka": "ბმულები/ციტატები", "hy": "Հղումներ/մեջբերումներ", "it": "Link/citazioni"},
        "claims": {"en": "claims", "ru": "утверждения", "uk": "твердження", "fr": "affirmations", "de": "Aussagen", "es": "afirmaciones", "ka": "მტკიცებები", "hy": "պնդումներ", "it": "affermazioni"},
        "paragraphs": {"en": "Paragraphs", "ru": "Абзацы", "uk": "Абзаци", "fr": "Paragraphes", "de": "Absätze", "es": "Párrafos", "ka": "აბზაცები", "hy": "Պարբերություններ", "it": "Paragrafi"},
        "headings": {"en": "headings", "ru": "заголовки", "uk": "заголовки", "fr": "titres", "de": "Überschriften", "es": "títulos", "ka": "სათაურები", "hy": "վերնագրեր", "it": "titoli"},
    }
    values = phrases.get(key, {})
    language = clean_language(language)
    return values.get(language) or values.get("en") or key


def _score_tone(score: float) -> str:
    if score >= 76:
        return "good"
    if score >= 55:
        return "warn"
    return "bad"


def _risk_tone(score: float) -> str:
    if score <= 34:
        return "good"
    if score <= 64:
        return "warn"
    return "bad"


def _clamp_score(value: float) -> int:
    return int(round(max(0, min(100, value))))


def _originality_runtime_text(key: str, language: str) -> str:
    messages = {
        "file_too_large": {"en": "Document is larger than 8 MB.", "ru": "Документ больше 8 MB.", "uk": "Документ більший за 8 MB."},
        "legacy_doc": {"en": "Old .doc files are not supported. Save as DOCX, PDF or TXT.", "ru": "Старый .doc не поддерживается. Сохраните как DOCX, PDF или TXT.", "uk": "Старий .doc не підтримується. Збережіть як DOCX, PDF або TXT."},
        "unsupported": {"en": "Unsupported document format. Use DOCX, PDF, TXT, MD, RTF or HTML.", "ru": "Формат документа не поддерживается. Используйте DOCX, PDF, TXT, MD, RTF или HTML.", "uk": "Формат документа не підтримується. Використайте DOCX, PDF, TXT, MD, RTF або HTML."},
        "docx_failed": {"en": "Could not read DOCX text.", "ru": "Не получилось прочитать текст DOCX.", "uk": "Не вдалося прочитати текст DOCX."},
        "pdf_failed": {"en": "Could not read PDF text.", "ru": "Не получилось прочитать текст PDF.", "uk": "Не вдалося прочитати текст PDF."},
        "pdf_empty": {"en": "PDF has no readable text layer.", "ru": "В PDF нет читаемого текстового слоя.", "uk": "У PDF немає читабельного текстового шару."},
    }
    values = messages.get(key, {})
    language = clean_language(language)
    return values.get(language) or values.get("en") or key


def _require_file(request: HttpRequest, field: str):
    upload = request.FILES.get(field)
    if not upload:
        raise ValueError("Файл не выбран")
    return upload


def _validate_convert_target(upload, target_format: str, language: str) -> None:
    image_formats = {value for value, _ in actions.IMAGE_FORMAT_CHOICES}
    video_formats = {value for value, _ in actions.VIDEO_FORMAT_CHOICES}
    kind = _upload_kind(upload)
    if kind == "image" and target_format in video_formats:
        raise ValueError(_localized_format_error("image", language))
    if kind == "video" and target_format in image_formats:
        raise ValueError(_localized_format_error("video", language))


def _upload_kind(upload) -> str:
    content_type = str(getattr(upload, "content_type", "") or "").lower()
    if content_type.startswith("video/"):
        return "video"
    if content_type.startswith("image/"):
        return "image"
    suffix = Path(getattr(upload, "name", "") or "").suffix.lower()
    if suffix in {".mp4", ".webm", ".mov", ".m4v", ".mkv", ".avi", ".wmv", ".flv"}:
        return "video"
    return "image"


def _localized_format_error(kind: str, language: str) -> str:
    messages = {
        "image": {
            "en": "Images can only be converted to image formats.",
            "ru": "Фото можно конвертировать только в форматы изображений.",
            "uk": "Фото можна конвертувати тільки у формати зображень.",
            "fr": "Les images ne peuvent être converties qu'en formats image.",
            "de": "Bilder können nur in Bildformate konvertiert werden.",
            "es": "Las imágenes solo se pueden convertir a formatos de imagen.",
            "ka": "სურათების კონვერტაცია მხოლოდ სურათის ფორმატებში შეიძლება.",
            "hy": "Պատկերները կարելի է փոխարկել միայն պատկերի ձևաչափերի։",
            "it": "Le immagini possono essere convertite solo in formati immagine.",
        },
        "video": {
            "en": "Videos can only be converted to video formats.",
            "ru": "Видео можно конвертировать только в видеоформаты.",
            "uk": "Відео можна конвертувати тільки у відеоформати.",
            "fr": "Les vidéos ne peuvent être converties qu'en formats vidéo.",
            "de": "Videos können nur in Videoformate konvertiert werden.",
            "es": "Los vídeos solo se pueden convertir a formatos de vídeo.",
            "ka": "ვიდეოს კონვერტაცია მხოლოდ ვიდეო ფორმატებში შეიძლება.",
            "hy": "Տեսանյութերը կարելի է փոխարկել միայն տեսանյութի ձևաչափերի։",
            "it": "I video possono essere convertiti solo in formati video.",
        },
    }
    return messages[kind].get(clean_language(language), messages[kind]["en"])


def _save_upload(upload, group: str) -> Path:
    base = clean_base_name(upload.name or "upload", "upload")
    suffix = Path(upload.name or "").suffix[:16]
    target_dir = settings.storage_dir / "django_uploads" / group
    target_dir.mkdir(parents=True, exist_ok=True)
    path = target_dir / f"{uuid.uuid4().hex[:12]}_{base}{suffix}"
    with path.open("wb") as destination:
        for chunk in upload.chunks():
            destination.write(chunk)
    return path


def _save_avatar_upload(upload, user_id: int) -> Path:
    suffix = Path(upload.name or "").suffix.lower()[:8] or ".jpg"
    target_dir = settings.storage_dir / "django_uploads" / "avatars" / str(user_id)
    target_dir.mkdir(parents=True, exist_ok=True)
    path = target_dir / f"avatar_{uuid.uuid4().hex[:12]}{suffix}"
    with path.open("wb") as destination:
        for chunk in upload.chunks():
            destination.write(chunk)
    return path


def _save_avatar_crop(data_url: str, user_id: int) -> Path:
    if "," in data_url:
        data_url = data_url.split(",", 1)[1]
    raw = b64decode(data_url)
    image = Image.open(BytesIO(raw)).convert("RGB")
    image.thumbnail((1200, 1200))
    target_dir = settings.storage_dir / "django_uploads" / "avatars" / str(user_id)
    target_dir.mkdir(parents=True, exist_ok=True)
    path = target_dir / f"avatar_{uuid.uuid4().hex[:12]}.jpg"
    image.save(path, format="JPEG", quality=92, optimize=True)
    return path


def _job_json(job: dict) -> JsonResponse:
    return JsonResponse({"job": _attach_output_urls(job)})


def _error_json(exc: Exception, status: int = 400) -> JsonResponse:
    return JsonResponse({"error": str(exc) or exc.__class__.__name__}, status=status)


def _attach_output_urls_many(jobs: list[dict]) -> list[dict]:
    cache = _job_url_cache([str(job.get("id") or "") for job in jobs])
    return [_attach_output_urls(job, cache) for job in jobs]


def _attach_output_urls(job: dict, url_cache: dict[str, dict[str, object]] | None = None) -> dict:
    prepared = dict(job)
    prepared["detail_url"] = reverse("studio:job_detail", args=[job["id"]])
    prepared["download_all_url"] = reverse("studio:download_all_outputs", args=[job["id"]])
    prepared["delete_url"] = reverse("studio:delete_job", args=[job["id"]])
    prepared["repeat_url"] = reverse("studio:repeat_job", args=[job["id"]])
    prepared["pause_url"] = reverse("studio:pause_job", args=[job["id"]])
    prepared["resume_url"] = reverse("studio:resume_job", args=[job["id"]])
    prepared["cancel_url"] = reverse("studio:cancel_job", args=[job["id"]])
    prepared["publish_url"] = f"{reverse('studio:community_publish')}?{urlencode({'source': 'job', 'id': job['id']})}"
    outputs = []
    job_id = str(job.get("id") or "")
    cached = (url_cache or {}).get(job_id, {})
    design_map = cached.get("design_projects") if isinstance(cached.get("design_projects"), dict) else _job_design_map(job_id)
    video_map = cached.get("video_projects") if isinstance(cached.get("video_projects"), dict) else _job_video_map(job_id)
    path_keys = cached.get("path_keys") if isinstance(cached.get("path_keys"), dict) else {}
    for output in job.get("outputs", []):
        item = dict(output)
        item["url"] = reverse("studio:download_output", args=[job["id"], item["index"]])
        item["preview_url"] = reverse("studio:preview_output", args=[job["id"], item["index"]])
        item["preview_kind"] = _preview_kind(item)
        item["can_edit_design"] = _output_can_edit_design(item, job)
        if item["can_edit_design"]:
            item["edit_design_url"] = reverse("studio:edit_output_design", args=[job["id"], item["index"]])
            output_key = str(path_keys.get(int(item.get("index") or 0)) or "") or _job_output_path_key(job_id, int(item.get("index") or 0))
            design_project_id = design_map.get(output_key) if output_key else None
            item["design_project_url"] = f"{reverse('studio:designer')}?{urlencode({'project': design_project_id})}" if design_project_id else ""
        item["can_edit_video"] = _output_can_edit_video(item, job)
        if item["can_edit_video"]:
            item["edit_video_url"] = reverse("studio:edit_output_video", args=[job["id"], item["index"]])
            output_key = str(path_keys.get(int(item.get("index") or 0)) or "") or _job_output_path_key(job_id, int(item.get("index") or 0))
            video_project_id = video_map.get(output_key) if output_key else None
            item["video_project_url"] = f"{reverse('studio:video_editor')}?{urlencode({'project': video_project_id})}" if video_project_id else ""
        outputs.append(item)
    prepared["outputs"] = outputs
    prepared["thumbnail"] = _job_thumbnail_payload(prepared)
    return prepared


def _job_thumbnail_kind_from_output(output: dict) -> str:
    media_type = str(output.get("media_type") or "").lower()
    name = str(output.get("name") or output.get("label") or "").lower()
    if media_type.startswith("image/") or name.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif")):
        return "image"
    if media_type.startswith("video/") or name.endswith((".mp4", ".webm", ".mov", ".m4v", ".mkv")):
        return "video"
    if media_type.startswith("audio/") or name.endswith((".mp3", ".wav", ".m4a", ".ogg", ".flac")):
        return "audio"
    if media_type == "application/pdf" or name.endswith(".pdf"):
        return "pdf"
    if media_type.startswith("text/") or name.endswith((".txt", ".ass", ".srt", ".json", ".csv", ".md")):
        return "text"
    if name.endswith((".zip", ".rar", ".7z")):
        return "archive"
    return "file"


def _job_thumbnail_kind_from_job(job: dict) -> str:
    kind = str(job.get("kind") or "").lower()
    title = str(job.get("title") or "").lower()
    if "cover" in kind or "image" in kind or "convert" in kind and any(ext in title for ext in (".png", ".jpg", ".jpeg", ".webp", ".jfif")):
        return "image"
    if "youtube" in kind or "video" in kind or "short" in title:
        return "video"
    if "music" in kind or "audio" in kind:
        return "audio"
    if "subtitle" in kind or "originality" in kind:
        return "text"
    if "package" in kind:
        return "archive"
    return "file"


def _job_thumbnail_payload(job: dict) -> dict[str, object]:
    outputs = list(job.get("outputs") or [])
    output = next((item for item in outputs if _job_thumbnail_kind_from_output(item) == "image"), None) or (outputs[0] if outputs else None)
    if output:
        thumb_kind = _job_thumbnail_kind_from_output(output)
        return {
            "kind": thumb_kind,
            "url": output.get("preview_url") if thumb_kind == "image" else "",
            "label": str(output.get("label") or output.get("name") or thumb_kind).split(".")[-1][:8].upper(),
        }
    thumb_kind = _job_thumbnail_kind_from_job(job)
    return {"kind": thumb_kind, "url": "", "label": thumb_kind.upper()}


def _job_url_cache(job_ids: list[str]) -> dict[str, dict[str, object]]:
    clean_ids = [job_id for job_id in dict.fromkeys(job_ids) if job_id]
    if not clean_ids:
        return {}
    records = JobRecord.objects.filter(job_id__in=clean_ids).only("job_id", "params_json").prefetch_related("outputs")
    cache: dict[str, dict[str, object]] = {}
    for record in records:
        params = _job_record_params(record)
        cache[record.job_id] = {
            "design_projects": _job_project_map_from_params(params, "design_projects"),
            "video_projects": _job_project_map_from_params(params, "video_projects"),
            "path_keys": {index: str(Path(output.path).resolve()) for index, output in enumerate(record.outputs.all())},
        }
    return cache


def _job_design_map(job_id: str) -> dict[str, int]:
    return _job_project_map(job_id, "design_projects")


def _job_video_map(job_id: str) -> dict[str, int]:
    return _job_project_map(job_id, "video_projects")


def _job_project_map(job_id: str, key: str) -> dict[str, int]:
    record = JobRecord.objects.filter(job_id=job_id).only("params_json").first()
    if not record:
        return {}
    return _job_project_map_from_params(_job_record_params(record), key)


def _job_project_map_from_params(params: dict[str, object], key: str) -> dict[str, int]:
    raw = params.get(key)
    if not isinstance(raw, dict):
        return {}
    result: dict[str, int] = {}
    for key, value in raw.items():
        try:
            result[str(key)] = int(value)
        except (TypeError, ValueError):
            continue
    return result


def _job_output_path_key(job_id: str, index: int) -> str:
    record = JobRecord.objects.filter(job_id=job_id).prefetch_related("outputs").first()
    if not record:
        return ""
    outputs = list(record.outputs.all())
    if index < 0 or index >= len(outputs):
        return ""
    return str(Path(outputs[index].path).resolve())


def _output_can_edit_design(output: dict, job: dict | None = None) -> bool:
    media_type = str(output.get("media_type") or "").lower()
    name = str(output.get("name") or "").lower()
    label = str(output.get("label") or "").lower()
    job_kind = str((job or {}).get("kind") or "").lower()
    if not (media_type.startswith("image/") or name.endswith((".png", ".jpg", ".jpeg", ".webp"))):
        return False
    if name.endswith(".zip"):
        return False
    cover_words = ("cover", "облож", "обл", "thumbnail", "png-cover")
    return job_kind in {"cover", "youtube_cover"} or any(word in name or word in label for word in cover_words)


def _output_can_edit_video(output: dict, job: dict | None = None) -> bool:
    media_type = str(output.get("media_type") or "").lower()
    name = str(output.get("name") or "").lower()
    label = str(output.get("label") or "").lower()
    job_kind = str((job or {}).get("kind") or "").lower()
    if name.endswith(".zip"):
        return False
    is_video = media_type.startswith("video/") or name.endswith((".mp4", ".webm", ".mov", ".m4v", ".mkv"))
    if not is_video:
        return False
    return job_kind in {"youtube", "download", "convert", "subtitles", "package", "video_export"} or any(word in label for word in ("short", "preview", "mp4", "video"))


def _preview_kind(output: dict) -> str:
    media_type = str(output.get("media_type") or "")
    name = str(output.get("name") or "").lower()
    if media_type.startswith("image/"):
        return "image"
    if media_type.startswith("video/"):
        return "video"
    if _is_subtitle_output(name, media_type):
        return "subtitle"
    if media_type == "application/pdf" or name.endswith(".pdf"):
        return "embed"
    if media_type.startswith("text/") or name.endswith((".txt", ".ass", ".srt", ".json", ".csv")):
        return "embed"
    return "download"
